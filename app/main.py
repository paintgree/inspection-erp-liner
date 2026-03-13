from __future__ import annotations

# ======================
# Standard Library
# ======================
import os
import json
import hashlib
import mimetypes
import tempfile
import traceback
import subprocess

from datetime import datetime, date, time as dtime, timedelta
from io import BytesIO
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from types import SimpleNamespace
from zoneinfo import ZoneInfo
from collections import defaultdict
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.platypus import Table, TableStyle
from passlib.context import CryptContext


# ======================
# Third Party
# ======================
import openpyxl
from openpyxl.drawing.image import Image as XLImage
from sqlalchemy import func
from pypdf import PdfWriter, PdfReader, Transformation

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.lib.units import inch

from fastapi import (
    FastAPI,
    Depends,
    Request,
    Form,
    HTTPException,
    UploadFile,
    File,
)

from fastapi.responses import (
    HTMLResponse,
    RedirectResponse,
    FileResponse,
    StreamingResponse,
    JSONResponse,
)

from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

from sqlmodel import Session, select

# ======================
# Local App Imports
# ======================
from .auth import hash_password, verify_password
from .db import create_db_and_tables, get_session
from .models import (
    User,
    ProductionRun,
    RunMachine,
    RunParameter,
    InspectionEntry,
    InspectionValue,
    InspectionValueAudit,
    MaterialLot,
    MaterialUseEvent,
    MrrDocument,
    MrrReceiving,
    MrrReceivingInspection,
    MrrInspection,
    MrrInspectionPhoto,
    BurstTestReport,
    BurstAttachment,
    BurstSample,
    BurstAuditLog,
    BurstReportRevision,
    HydroTestRecord,
    FinalInspectionReel,
    FinalInspectionBatchNote,
    FinalInspectionPhase,
    FinalInspectionReel,
    FinalInspectionBatchNote,
    RfiRecord,
    RfiAttachment,

)

SLOTS = [
    "00:00", "02:00", "04:00", "06:00",
    "08:00", "10:00", "12:00", "14:00",
    "16:00", "18:00", "20:00", "22:00"
]

# =========================
# Default run parameters per process
# =========================
# Used when creating a new ProductionRun to seed RunParameter rows.
PROCESS_PARAMS: Dict[str, List[Tuple[str, str, str]]] = {
    "LINER": [
        ("length_m", "Length", "m"),
        ("od_mm", "OD", "mm"),
        ("wall_thickness_mm", "Wall thickness", "mm"),
        ("cooling_water_c", "Cooling water temp", "°C"),
        ("line_speed_m_min", "Line speed", "m/min"),
        ("tractor_pressure_mpa", "Tractor pressure", "MPa"),
        ("body_temp_1", "Body temp 1", "°C"),
        ("body_temp_2", "Body temp 2", "°C"),
        ("body_temp_3", "Body temp 3", "°C"),
        ("body_temp_4", "Body temp 4", "°C"),
        ("body_temp_5", "Body temp 5", "°C"),
        ("noising_temp_1", "Noising temp 1", "°C"),
        ("noising_temp_2", "Noising temp 2", "°C"),
        ("noising_temp_3", "Noising temp 3", "°C"),
        ("noising_temp_4", "Noising temp 4", "°C"),
        ("noising_temp_5", "Noising temp 5", "°C"),
    ],
    "COVER": [
        ("length_m", "Length", "m"),
        ("od_mm", "OD", "mm"),
        ("wall_thickness_mm", "Wall thickness", "mm"),
        ("cooling_water_c", "Cooling water temp", "°C"),
        ("line_speed_m_min", "Line speed", "m/min"),
        ("tractor_pressure_mpa", "Tractor pressure", "MPa"),
        ("body_temp_1", "Body temp 1", "°C"),
        ("body_temp_2", "Body temp 2", "°C"),
        ("body_temp_3", "Body temp 3", "°C"),
        ("body_temp_4", "Body temp 4", "°C"),
        ("body_temp_5", "Body temp 5", "°C"),
        ("noising_temp_1", "Noising temp 1", "°C"),
        ("noising_temp_2", "Noising temp 2", "°C"),
        ("noising_temp_3", "Noising temp 3", "°C"),
        ("noising_temp_4", "Noising temp 4", "°C"),
        ("noising_temp_5", "Noising temp 5", "°C"),
    ],
    "REINFORCEMENT": [
        ("length_m", "Length", "m"),
        ("annular_od_1_mm", "Annular OD 1", "mm"),
        ("annular_od_2_mm", "Annular OD 2", "mm"),
        ("internal_tensile_od_mm", "Internal tensile OD", "mm"),
        ("external_tensile_od_mm", "External tensile OD", "mm"),
        ("core_mould_dia_mm", "Core mould dia", "mm"),
        ("annular_width_1_mm", "Annular width 1", "mm"),
        ("annular_width_2_mm", "Annular width 2", "mm"),
        ("screw_yarn_width_mm", "Screw yarn width", "mm"),
    ],
}

# =========================
# MRR helpers (units + report no)
# =========================

UNIT_MULTIPLIER = {
    "KG": 1.0,
    "T": 1000.0,   # 1 Ton = 1000 KG
}
def safe_json_loads(val):
    if not val:
        return {}
    if isinstance(val, dict):
        return val
    try:
        return json.loads(val)
    except Exception:
        return {}

def normalize_qty_to_kg(qty: float, unit: str) -> float:
    u = (unit or "KG").upper().strip()
    if u in ["T", "TON", "TONNE"]:
        return float(qty) * 1000.0
    if u in ["KG", "KGS"]:
        return float(qty)
    # For PCS we cannot convert to KG; keep as-is (handled by units_compatible)
    return float(qty)

def dn_doc_exists(session: Session, lot_id: int, dn_number: str) -> bool:
    dn_number = (dn_number or "").strip()
    if not dn_number:
        return False
    doc = session.exec(
        select(MrrDocument).where(
            (MrrDocument.ticket_id == lot_id) &
            (MrrDocument.doc_type == "DELIVERY_NOTE") &
            (MrrDocument.doc_number == dn_number)
        )
    ).first()
    return bool(doc)
    
def quality_doc_exists(session: Session, lot_id: int) -> bool:
    """
    True if at least ONE quality doc exists:
    COA OR MTC OR INSPECTION_REPORT
    """
    q = select(MrrDocument).where(
        (MrrDocument.ticket_id == lot_id) &
        (MrrDocument.doc_type.in_(["COA", "MTC", "INSPECTION_REPORT"]))
    )
    return session.exec(q).first() is not None


def get_submitted_shipment_by_dn(session: Session, lot_id: int, dn_number: str):
    dn_number = (dn_number or "").strip()
    if not dn_number:
        return None
    # DN is "consumed" ONLY if shipment was submitted (inspector_confirmed True) or approved
    return session.exec(
        select(MrrReceivingInspection).where(
            (MrrReceivingInspection.ticket_id == lot_id) &
            (MrrReceivingInspection.delivery_note_no == dn_number) &
            (
                (MrrReceivingInspection.inspector_confirmed == True) |
                (MrrReceivingInspection.manager_approved == True)
            )
        )
    ).first()

def get_draft_shipment_by_dn(session: Session, lot_id: int, dn_number: str):
    dn_number = (dn_number or "").strip()
    if not dn_number:
        return None
    # Draft = created but NOT submitted yet
    return session.exec(
        select(MrrReceivingInspection).where(
            (MrrReceivingInspection.ticket_id == lot_id) &
            (MrrReceivingInspection.delivery_note_no == dn_number) &
            (MrrReceivingInspection.inspector_confirmed == False) &
            (MrrReceivingInspection.manager_approved == False)
        )
    ).first()

def get_latest_draft_shipment(session: Session, lot_id: int):
    return session.exec(
        select(MrrReceivingInspection).where(
            (MrrReceivingInspection.ticket_id == lot_id) &
            (MrrReceivingInspection.inspector_confirmed == False) &
            (MrrReceivingInspection.manager_approved == False)
        ).order_by(MrrReceivingInspection.created_at.desc())
    ).first()


def units_compatible(po_unit: str, arrived_unit: str) -> bool:
    pu = (po_unit or "").upper().strip()
    au = (arrived_unit or "").upper().strip()
    # KG and T compatible, PCS only with PCS
    if pu == "PCS" or au == "PCS":
        return pu == au
    return True

def generate_report_no(ticket_id: int, seq: int) -> str:
    # MRR-YYYY-MM-<ticket>-<shipment>
    now = datetime.utcnow()
    return f"MRR-{now.year}-{now.month:02d}-{ticket_id:04d}-{seq:02d}"

# =========================
# MRR REPORT TEMPLATES PATHS
# =========================
import os
import io
import re
import json
import zipfile
import subprocess
import tempfile
import shutil
from docx import Document
from datetime import datetime, date

import openpyxl
from fastapi import HTTPException
from fastapi.responses import Response

import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MRR_TEMPLATE_DIR = os.path.join(BASE_DIR, "templates", "templates_xlsx")

MRR_TEMPLATE_XLSX_MAP = {
    "RAW": os.path.join(MRR_TEMPLATE_DIR, "QAP0600-F01.xlsx"),
    # later we can add other xlsx templates if needed
}

MRR_TEMPLATE_DOCX_MAP = {
    "OUTSOURCED": os.path.join(MRR_TEMPLATE_DIR, "QAP0600-F02.docx"),
}

HYDRO_TEMPLATE_DOCX = os.path.join(MRR_TEMPLATE_DIR, "New hydrotest template 2026.docx")

# =========================
# =========================

def _safe_upper(x: str | None) -> str:
    return (x or "").strip().upper()

def _to_float(x, default=0.0):
    try:
        if x is None or x == "":
            return default
        return float(x)
    except Exception:
        return default


def _normalize_key(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip()).lower()

def _xlsx_bytes_from_wb(wb) -> bytes:
    bio = io.BytesIO()
    wb.save(bio)
    return bio.getvalue()

from openpyxl.utils.cell import coordinate_to_tuple

def _ws_set_value_safe(ws, addr: str, value):
    """
    Safely set a value even if the address points inside a merged cell.
    Writes to the top-left cell of the merged range.
    """
    r, c = coordinate_to_tuple(addr)

    # If addr is inside a merged cell, redirect to the merged range start
    for mrange in ws.merged_cells.ranges:
        if mrange.min_row <= r <= mrange.max_row and mrange.min_col <= c <= mrange.max_col:
            ws.cell(mrange.min_row, mrange.min_col).value = value
            return

    ws[addr].value = value


def _apply_excel_print_fit_settings_f01(wb):
    """
    Force Excel print settings so LibreOffice exports F01 correctly:
    - Fit all columns on 1 page width
    - Allow multiple pages height
    - A4 portrait
    """
    from openpyxl.worksheet.page import PageMargins

    for ws in wb.worksheets:
        try:
            # Enable "Fit to page"
            ws.sheet_properties.pageSetUpPr.fitToPage = True

            # A4 paper
            ws.page_setup.paperSize = ws.PAPERSIZE_A4

            # Portrait (change to 'landscape' only if your template is landscape)
            ws.page_setup.orientation = "portrait"

            # ✅ This is the key:
            ws.page_setup.fitToWidth = 1
            ws.page_setup.fitToHeight = 0  # 0 = as many pages tall as needed

            # Reasonable margins (tune if needed)
            ws.page_margins = PageMargins(
                left=0.25, right=0.25, top=0.35, bottom=0.35,
                header=0.2, footer=0.2
            )
        except Exception:
            pass

def fill_mrr_f01_xlsx_bytes(
    *,
    lot,
    receiving,
    inspection,
    docs: list,
    photos_by_group: dict | None = None,
) -> bytes:
    """
    Fills QAP0600-F01.xlsx using:
    - lot
    - receiving
    - inspection (inspection_json tables + basic fields)
    - docs (uploaded docs list)
    """
    template_path = MRR_TEMPLATE_XLSX_MAP.get("RAW")
    if not template_path or not os.path.exists(template_path):
        raise HTTPException(
            500,
            f"RAW template missing. Put QAP0600-F01.xlsx in {MRR_TEMPLATE_DIR}",
        )

    wb = openpyxl.load_workbook(template_path)
    ws = wb.active

    # -------------------------
    # Load inspection JSON
    # -------------------------
    try:
        data = json.loads(getattr(inspection, "inspection_json", None) or "{}")
        if not isinstance(data, dict):
            data = {}
    except Exception:
        data = {}

    # Normalize family from UI into a small set your code understands
    fam_ui = (data.get("material_family") or data.get("material_fam") or data.get("material_type") or "").strip().upper()
    if fam_ui in ("PE_RT", "PE100", "PE"):
        fam_ui = "PE"
    elif fam_ui in ("POLYESTER_FIBER", "FIBER"):
        fam_ui = "FIBER"
    elif fam_ui == "OTHER":
        fam_ui = "OTHER"
    else:
        fam_ui = fam_ui or ""

    grade = (data.get("material_grade") or data.get("grade") or "").strip()

    # -------------------------
    # HEADER MAP (F01)
    # -------------------------
    _ws_set_value_safe(ws, "E1", getattr(inspection, "report_no", "") or "")
    _ws_set_value_safe(ws, "E2", _as_date_str(getattr(inspection, "created_at", None) or datetime.utcnow()))

    _ws_set_value_safe(ws, "B6", getattr(lot, "supplier_name", "") or "")
    _ws_set_value_safe(ws, "F6", getattr(lot, "po_number", "") or "")

    # Material type in header should match dropdown selection exactly (NOT grade)
    raw_type = (data.get("material_family") or data.get("material_type") or "").strip().upper()
    
    type_label = ""
    if raw_type in ("PE_RT", "POLYETHYLENE (PE-RT)", "POLYETHYLENE_PE_RT"):
        type_label = "Polyethylene (PE-RT)"
    elif raw_type in ("PE100", "POLYETHYLENE (PE100)", "POLYETHYLENE_PE100"):
        type_label = "Polyethylene (PE100)"
    elif raw_type in ("POLYESTER_FIBER", "POLYESTER FIBER", "FIBER"):
        type_label = "POLYESTER FIBER"
    else:
        # fallback: show whatever was saved
        type_label = data.get("material_family") or data.get("material_type") or ""
    
    _ws_set_value_safe(ws, "B7", type_label)

    # Batch / grade / DN
    _ws_set_value_safe(ws, "B8", grade)
    _ws_set_value_safe(ws, "F8", getattr(inspection, "delivery_note_no", "") or "")

    # Batch numbers might be a list or string
    bn = data.get("batch_numbers")
    if isinstance(bn, list):
        _ws_set_value_safe(ws, "B9", ", ".join([str(x).strip() for x in bn if str(x).strip()]))
    elif isinstance(bn, str):
        _ws_set_value_safe(ws, "B9", bn.strip())
    else:
        _ws_set_value_safe(ws, "B9", getattr(lot, "batch_no", "") or "")

    # Quantity received
    qty_arrived = getattr(inspection, "qty_arrived", None)
    qty_unit = getattr(inspection, "qty_unit", None) or ""
    _ws_set_value_safe(ws, "F9", f"{_to_float(qty_arrived, 0)} {qty_unit}".strip())

    # Comments / remarks
    _ws_set_value_safe(ws, "A46", (data.get("remarks") or "").strip())

    # -------------------------
    # TABLE TITLES
    # -------------------------
    pe_title = f"POLYETHYLENE ({grade})" if grade else "POLYETHYLENE"
    fb_title = f"POLYESTER FIBER ({grade})" if grade else "POLYESTER FIBER"

    if fam_ui == "PE":
        _ws_set_value_safe(ws, "A11", pe_title)
        _ws_set_value_safe(ws, "A24", "")   # clear fiber title
    elif fam_ui == "FIBER":
        _ws_set_value_safe(ws, "A24", fb_title)
        _ws_set_value_safe(ws, "A11", "")   # clear PE title
    else:
        _ws_set_value_safe(ws, "A11", "")
        _ws_set_value_safe(ws, "A24", "")

    # -------------------------
    # PROPERTIES TABLE FILL
    # -------------------------
    def _build_prop_map(items):
        m = {}
        if not isinstance(items, list):
            return m
        for it in items:
            if not isinstance(it, dict):
                continue
            name = _normalize_key(str(it.get("name") or it.get("property") or ""))
            if not name:
                continue
            m[name] = it
        return m

    def _write_cell_safe(sheet, row, col, value):
        """Works with merged cells: write to TOP-LEFT of the merged range."""
        cell = sheet.cell(row, col)
        if isinstance(cell, openpyxl.cell.cell.MergedCell):
            for mr in sheet.merged_cells.ranges:
                if mr.min_row <= row <= mr.max_row and mr.min_col <= col <= mr.max_col:
                    sheet.cell(mr.min_row, mr.min_col).value = value
                    return
            return
        cell.value = value

    # If UI didn’t send a ready "properties" list, we convert from form keys
    prop_items = data.get("properties")
    if not isinstance(prop_items, list) or len(prop_items) == 0:
        converted = []

        # Labels MUST match Excel Column A exactly
        pe_rows = [
            ("density", "Density"),
            ("mfr", "Melt Flow Rate (MFR) -190°C / 5kg"),
            ("flexural", "Flexural Modulus"),
            ("tensile", "Tensile Strength at Yield"),
            ("tensile_break", "Tensile Strength at Break"),
            ("elong", "Elongation at Break"),
            ("escr", "(ESCR)"),
            ("oits", "Oxidative Induction Time (OIT)"),
            ("hdb", "HDB (23C°) /  MRS (20C°)"),
            ("cb", "Carbon Black Content"),
            ("melt", "Melting Point"),
        ]

        fb_rows = [
            ("linear_density", "Linear Density"),
            ("breaking_strength", "Breaking Strength"),
            ("tenacity", "Tenacity"),
            ("elongation", "Elongation at Break"),
            ("melting_point", "Melting Point"),
        ]

        if fam_ui == "PE":
            for k, label in pe_rows:
                r = (data.get(f"pe_{k}_result") or "").strip()
                rm = (data.get(f"pe_{k}_remarks") or "").strip()
                if r or rm:
                    converted.append({"name": label, "result": r, "remarks": rm})

        elif fam_ui == "FIBER":
            for k, label in fb_rows:
                r = (data.get(f"fb_{k}_result") or "").strip()
                rm = (data.get(f"fb_{k}_remarks") or "").strip()
                if r or rm:
                    converted.append({"name": label, "result": r, "remarks": rm})

        prop_items = converted

    prop_map = _build_prop_map(prop_items)

    # Allowed rows by family (based on your template)
    if fam_ui == "PE":
        allowed_row_min, allowed_row_max = 13, 23
    elif fam_ui == "FIBER":
        allowed_row_min, allowed_row_max = 26, 30
    else:
        # If unknown, don’t write into table area
        allowed_row_min, allowed_row_max = 0, -1

    # Columns: results at G (7), remarks at I (9)
    RESULTS_COL = 7  # G
    REMARKS_COL = 9  # I

    if allowed_row_min <= allowed_row_max:
        for r in range(allowed_row_min, allowed_row_max + 1):
            label = ws.cell(r, 1).value  # Column A
            if not isinstance(label, str):
                continue
            key = _normalize_key(label)
            if key in prop_map:
                it = prop_map[key]
                _write_cell_safe(ws, r, RESULTS_COL, it.get("result") or it.get("value") or "")
                _write_cell_safe(ws, r, REMARKS_COL, it.get("remarks") or "")

    # -------------------------
    # FOOTER / PRINT SETTINGS
    # -------------------------
    try:
        ws.oddFooter.left.text = "QAP0600-F01"
        ws.oddFooter.left.size = 9
        ws.oddFooter.left.font = "Arial"
    except Exception:
        pass

    _apply_excel_print_fit_settings_f01(wb)
    return _xlsx_bytes_from_wb(wb)


    # ---- VISUAL + DOC REVIEW (use fixed row mapping; match Jinja keys exactly) ----

    def _slug_visual_jinja(s: str) -> str:
        # Matches your Jinja visual slug exactly:
        # item|replace(" ", "_")|replace("/", "_")|replace("(", "")|replace(")", "")|replace(".", "")|lower
        s = (s or "").strip().lower()
        s = s.replace(" ", "_")
        s = s.replace("/", "_")
        s = s.replace("(", "").replace(")", "")
        s = s.replace(".", "")
        # IMPORTANT: your Jinja does NOT remove commas, so we keep commas here
        return s

    def _slug_doc_jinja(s: str) -> str:
        # Matches your Jinja doc slug exactly:
        # item|replace(" ", "_")|replace("’","")|replace("'","")|replace("/", "_")|replace("(", "")|replace(")", "")|replace(".", "")|lower
        s = (s or "").strip().lower()
        s = s.replace("’", "").replace("'", "")
        s = s.replace(" ", "_")
        s = s.replace("/", "_")
        s = s.replace("(", "").replace(")", "")
        s = s.replace(".", "")
        # IMPORTANT: your Jinja does NOT remove commas, so we keep commas here
        return s

    def _get_any(d: dict, keys: list[str]) -> str:
        # return first non-empty value
        for k in keys:
            v = d.get(k)
            if v is not None and str(v).strip() != "":
                return str(v).strip()
        return ""

    # These MUST match your HTML list (inspection page)
    visual_items = [
        "Physical Condition of Material",
        "Identification/Marking as per specifications",
        "Confirm that the packaging is undamaged, sealed, and properly labeled.",
        "Ensure there are no signs of chemical exposure that might degrade the material.",
    ]
    # Excel rows for those 4 visual items
    visual_rows = [33, 34, 35, 36]

    for item, r in zip(visual_items, visual_rows):
        k1 = _slug_visual_jinja(item)          # current Jinja slug (keeps commas)
        k2 = k1.replace(",", "")               # fallback if you later remove commas in Jinja

        yn = _get_any(data, [f"vc_{k1}_yn", f"vc_{k2}_yn"]).upper()
        rm = _get_any(data, [f"vc_{k1}_remarks", f"vc_{k2}_remarks"])

        if yn in ["YES", "NO"]:
            _ws_set_value_safe(ws, f"G{r}", yn)   # G:H merged
        if rm:
            _ws_set_value_safe(ws, f"I{r}", rm)   # I:J merged

    doc_items = [
        "Ensure the material’s quantity, type, and specification match the Purchase Order (PO)",
        "Confirm the availability of Certificate of Analysis (COA).",
        "Review the Delivery Note to verify correct Delivery.",
    ]
    doc_rows = [39, 40, 41]

    for item, r in zip(doc_items, doc_rows):
        k1 = _slug_doc_jinja(item)              # current Jinja slug (keeps commas)
        k2 = k1.replace(",", "")               # fallback if you later remove commas in Jinja

        yn = _get_any(data, [f"dr_{k1}_yn", f"dr_{k2}_yn"]).upper()
        rm = _get_any(data, [f"dr_{k1}_remarks", f"dr_{k2}_remarks"])

        if yn in ["YES", "NO"]:
            _ws_set_value_safe(ws, f"G{r}", yn)
        if rm:
            _ws_set_value_safe(ws, f"I{r}", rm)

    # ---- APPROVAL STATUS (do NOT overwrite text; keep text + add ✓) ----
    status = (data.get("approval_status") or "").strip().upper()

    v_text = "Verified and Confirmed"
    h_text = "On Hold (Specify Reason Below)"
    n_text = "Non-Conformity"

    _ws_set_value_safe(ws, "A44", f"✓ {v_text}" if status == "VERIFIED" else v_text)
    _ws_set_value_safe(ws, "D44", f"✓ {h_text}" if status == "HOLD" else h_text)
    _ws_set_value_safe(ws, "G44", f"✓ {n_text}" if status == "NONCONFORM" else n_text)

    # ---- COMMENTS BOX (A46:J48 merged) ----
    remarks = (data.get("remarks") or "").strip()
    on_hold_reason = (data.get("on_hold_reason") or "").strip()

    lines = []
    if status == "HOLD" and on_hold_reason:
        lines.append(f"On Hold Reason: {on_hold_reason}")
    if remarks:
        lines.append(f"Remarks: {remarks}")

    _ws_set_value_safe(ws, "A46", "\n".join(lines).strip())

    # ---- SIGNATURES ----
    _ws_set_value_safe(ws, "B51", getattr(inspection, "inspector_name", "") or "")
    _ws_set_value_safe(ws, "B52", _as_date_str(datetime.utcnow()))

    if bool(getattr(inspection, "manager_approved", False)):
        _ws_set_value_safe(ws, "D51", "MANAGER")
        _ws_set_value_safe(ws, "D52", _as_date_str(datetime.utcnow()))




def _set_cell_text(cell, text: str):
    # Clear existing paragraph runs and set new text cleanly
    cell.text = ""
    cell.text = (text or "").strip()

def _find_cell_by_label(doc: Document, label: str):
    """
    Find the table cell that contains 'label' (case-insensitive).
    Returns (table, row_idx, col_idx) or (None, None, None).
    """
    target = (label or "").strip().lower()
    for t in doc.tables:
        for r_idx, row in enumerate(t.rows):
            for c_idx, cell in enumerate(row.cells):
                if target in (cell.text or "").strip().lower():
                    return t, r_idx, c_idx
    return None, None, None

def _set_value_next_to_label(doc: Document, label: str, value: str):
    """
    Finds a cell containing label and writes value in the cell to its right.
    """
    t, r, c = _find_cell_by_label(doc, label)
    if t is None:
        return False
    # write into next cell if exists, else same cell
    if c + 1 < len(t.rows[r].cells):
        _set_cell_text(t.rows[r].cells[c + 1], value)
    else:
        _set_cell_text(t.rows[r].cells[c], value)
    return True

def _find_row_index_with_headers(table, headers: list[str]) -> int | None:
    """
    Find the row index where all headers appear (case-insensitive).
    """
    hdrs = [h.strip().lower() for h in headers]
    for r_idx, row in enumerate(table.rows):
        row_text = " | ".join([(c.text or "").strip().lower() for c in row.cells])
        if all(h in row_text for h in hdrs):
            return r_idx
    return None

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def _iter_paragraphs_in_doc(doc):
    # paragraphs in body
    for p in doc.paragraphs:
        yield p
    # paragraphs in tables
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def _clear_paragraph_runs(paragraph):
    for r in paragraph.runs:
        r.text = ""

def _set_bookmark_text(doc, bookmark_name: str, text: str) -> bool:
    """
    Writes `text` at a Word bookmark position.
    This version clears the paragraph cleanly but keeps it inside the same cell/box.
    """
    text = "" if text is None else str(text)

    for p in _iter_paragraphs_in_doc(doc):
        for child in p._p.iter():
            if child.tag == qn("w:bookmarkStart") and child.get(qn("w:name")) == bookmark_name:
                # wipe paragraph
                for r in p.runs:
                    r.text = ""
                if not p.runs:
                    p.add_run(text)
                else:
                    p.runs[0].text = text
                return True
    return False

def _set_bookmark_signature_block(
    doc: Document,
    bookmark_name: str,
    signed_line: str,
    dt_line: str,
    font_pt: float = 8.5,
    rgb=(70, 70, 70),
) -> bool:
    """
    Writes 2 lines inside the bookmark paragraph:
      line1
      line2
    Darker gray for better contrast.
    """
    from docx.shared import Pt, RGBColor

    for p in _iter_paragraphs_in_doc(doc):
        for child in p._p.iter():
            if child.tag == qn("w:bookmarkStart") and child.get(qn("w:name")) == bookmark_name:
                # clear runs
                for r in p.runs:
                    r.text = ""

                # rebuild with formatting
                r1 = p.add_run((signed_line or "").strip())
                r1.font.size = Pt(font_pt)
                r1.font.color.rgb = RGBColor(*rgb)

                p.add_run().add_break()

                r2 = p.add_run((dt_line or "").strip())
                r2.font.size = Pt(font_pt)
                r2.font.color.rgb = RGBColor(*rgb)

                return True
    return False

def _find_cell_by_bookmark(doc, bookmark_name: str):
    """
    Find the table cell that contains a bookmark.
    Returns (table, row_index, col_index) or None.
    """
    for t in doc.tables:
        for r_i, row in enumerate(t.rows):
            for c_i, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    for child in p._p.iter():
                        if child.tag == qn("w:bookmarkStart") and child.get(qn("w:name")) == bookmark_name:
                            return (t, r_i, c_i)
    return None

def _apply_f02_pdf_layout_tweaks(doc: Document) -> None:
    """
    Make DOCX->PDF output closer to the Word template:
    - slightly smaller text
    - push content down to avoid logo overlap
    """
    from docx.shared import Pt, Inches

    # Push body down a bit (helps if header/logo area differs in PDF conversion)
    for sec in doc.sections:
        sec.top_margin = Inches(0.75)         # increase top margin
        sec.header_distance = Inches(0.25)   # distance between header and body

    # Default font size smaller
    try:
        normal = doc.styles["Normal"]
        normal.font.size = Pt(9)
    except Exception:
        pass

    # Force all table runs to same size (LibreOffice conversion behaves better this way)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        try:
                            r.font.size = Pt(9)
                        except Exception:
                            pass

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _insert_multiline_text_at_bookmark(doc: Document, bookmark_name: str, lines: list[str], rgb_hex: str = "333333") -> bool:
    """
    Insert text at a WORD bookmark (real bookmark, not placeholder text).
    Supports multiple lines using <w:br/>.
    Example bookmarks: BM_INSPECTED_BY, BM_REVIEWD_BY, BM_APPROVED_BY
    """
    if not lines:
        return False

    # Find the bookmarkStart node
    bm_nodes = doc._element.xpath(
        f'.//w:bookmarkStart[@w:name="{bookmark_name}"]',
        namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
    )
    if not bm_nodes:
        return False

    bm = bm_nodes[0]
    parent = bm.getparent()

    # Build a run with color + font size + line breaks
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), rgb_hex)  # darker gray (not black)
    rPr.append(color)

    # 9pt font => 18 half-points
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    rPr.append(sz)

    r.append(rPr)

    for i, line in enumerate(lines):
        if i > 0:
            br = OxmlElement("w:br")
            r.append(br)

        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = line
        r.append(t)

    # Insert immediately after bookmarkStart
    idx = parent.index(bm)
    parent.insert(idx + 1, r)
    return True


def _bookmark_exists(doc: Document, bookmark_name: str) -> bool:
    try:
        nodes = doc._element.xpath(
            f'.//w:bookmarkStart[@w:name="{bookmark_name}"]',
            namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
        )
        return bool(nodes)
    except Exception:
        return False


def _insert_multiline_at_bookmark(
    doc: Document,
    bookmark_name: str,
    lines: list[str],
    rgb_hex: str = "333333",   # darker gray (not black)
    font_half_points: str = "18",  # 9pt
) -> bool:
    """
    Inserts multiline text exactly at bookmark position (works inside tables).
    """
    if not lines:
        return False

    try:
        nodes = doc._element.xpath(
            f'.//w:bookmarkStart[@w:name="{bookmark_name}"]',
            namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
        )
        if not nodes:
            return False

        bm = nodes[0]
        parent = bm.getparent()

        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        color = OxmlElement("w:color")
        color.set(qn("w:val"), rgb_hex)
        rPr.append(color)

        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), font_half_points)
        rPr.append(sz)

        r.append(rPr)

        for i, line in enumerate(lines):
            if i > 0:
                br = OxmlElement("w:br")
                r.append(br)

            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = str(line)
            r.append(t)

        idx = parent.index(bm)
        parent.insert(idx + 1, r)
        return True

    except Exception:
        return False


def _apply_f02_bookmark_signatures(doc: Document, inspection: "MrrReceivingInspection") -> None:
    """
    Writes digital signature blocks INSIDE the Word bookmarks (so position stays correct even if multi-page).

    Bookmarks used in QAP0600-F02.docx:
      - BM_INSPECTED_BY
      - BM_REVIEWED_BY (also supports BM_REVIEWD_BY)
      - BM_APPROVED_BY

    This version is "FORCED":
      - If inspector_name exists, we write the inspector signature even if inspector_confirmed is False.
      - If manager_approved is True, we write manager signature even if manager_approved_by is missing (fallback: "MANAGER").
      - Includes date + time (YYYY-MM-DD HH:MM).
    """
    try:
        data = safe_json_loads(getattr(inspection, "inspection_json", None)) or {}
        if not isinstance(data, dict):
            data = {}
    except Exception:
        data = {}

    # ----- Names (with safe fallbacks) -----
    inspector_name = (getattr(inspection, "inspector_name", "") or "").strip() or (data.get("inspected_by") or "").strip()
    reviewer_name = (data.get("reviewed_by") or "").strip()
    manager_name = (data.get("manager_approved_by") or "").strip()

    # ----- Timestamps (date + time) -----
    # Prefer stored timestamps if present, otherwise use created_at
    submitted_at = data.get("submitted_at_utc") or data.get("inspected_at_utc") or getattr(inspection, "created_at", None) or datetime.utcnow()
    reviewed_at = data.get("reviewed_at_utc") or ""
    approved_at = data.get("manager_approved_at_utc") or ""

    inspector_dt = _as_datetime_str(submitted_at)
    reviewer_dt = _as_datetime_str(reviewed_at) if reviewer_name else ""
    manager_dt = _as_datetime_str(approved_at) if approved_at else _as_datetime_str(datetime.utcnow())

    # Style
    font_pt = 8.5
    rgb = (70, 70, 70)  # dark gray

    # Helper: write 2 lines into bookmark paragraph (cleans old runs => no duplicates)
    def _write_sig(bookmark: str, line1: str, line2: str) -> bool:
        ok = _set_bookmark_signature_block(
            doc,
            bookmark,
            line1,
            line2,
            font_pt=font_pt,
            rgb=rgb,
        )
        if ok:
            return True
        # Fallback to XML insert method (if paragraph-based method fails for any reason)
        return _insert_multiline_at_bookmark(
            doc,
            bookmark,
            [line1, line2],
            rgb_hex="333333",
            font_half_points="18",
        )

    # ----- INSPECTOR (FORCED if name exists) -----
    if inspector_name:
        _write_sig(
            "BM_INSPECTED_BY",
            f"Digitally signed by: {inspector_name}",
            f"Date: {inspector_dt}",
        )

    # ----- REVIEWER (optional) -----
    if reviewer_name:
        # Your system supports both spellings, keep both:
        ok = _write_sig(
            "BM_REVIEWD_BY",
            f"Digitally reviewed by: {reviewer_name}",
            f"Date: {reviewer_dt}",
        )
        if not ok:
            _write_sig(
                "BM_REVIEWED_BY",
                f"Digitally reviewed by: {reviewer_name}",
                f"Date: {reviewer_dt}",
            )

    # ----- MANAGER (if approved) -----
    if bool(getattr(inspection, "manager_approved", False)):
        if not manager_name:
            manager_name = "MANAGER"
        _write_sig(
            "BM_APPROVED_BY",
            f"Digitally approved by: {manager_name}",
            f"Date: {manager_dt}",
        )

def fill_mrr_f02_docx_bytes(*, lot, inspection, receiving, docs: list) -> bytes:
    template_path = MRR_TEMPLATE_DOCX_MAP.get("OUTSOURCED")
    if not template_path or not os.path.exists(template_path):
        raise HTTPException(500, f"OUTSOURCED template missing. Put QAP0600-F02.docx in {MRR_TEMPLATE_DIR}")

    doc = Document(template_path)

    # Load saved inspection JSON safely
    data = safe_json_loads(getattr(inspection, "inspection_json", None)) or {}
    if not isinstance(data, dict):
        data = {}

    # ---------- Header bookmarks ----------
    _set_bookmark_text(doc, "BM_REPORT_NO", getattr(inspection, "report_no", "") or "")
    _set_bookmark_text(doc, "BM_REPORT_DATE", _as_date_str(getattr(inspection, "created_at", None) or datetime.utcnow()))
    _set_bookmark_text(doc, "BM_DELIVERY_NOTE", getattr(inspection, "delivery_note_no", "") or "")
    _set_bookmark_text(doc, "BM_PO_NUMBER", getattr(lot, "po_number", "") or "")

    # ---------- Helper: resolve column index for bookmarks ----------
    def _bookmark_col_index(bookmark: str):
        hit = _find_cell_by_bookmark(doc, bookmark)
        if not hit:
            return None
        t0, r0, c0 = hit
        min_c = c0
        for _c in range(len(t0.rows[r0].cells)):
            cell = t0.rows[r0].cells[_c]
            for p in cell.paragraphs:
                for child in p._p.iter():
                    if child.tag == qn("w:bookmarkStart") and child.get(qn("w:name")) == bookmark:
                        min_c = min(min_c, _c)
        return (t0, r0, min_c)

    # ---------- Items table ----------
    items_item = data.get("items_item[]", [])
    items_desc = data.get("items_desc[]", [])
    items_size = data.get("items_size[]", [])
    items_type = data.get("items_type[]", [])
    items_pressure = data.get("items_pressure[]", [])
    items_qty = data.get("items_qty[]", [])
    items_mtc = data.get("items_mtc[]", [])

    max_items = max(
        len(items_item), len(items_desc), len(items_size), len(items_type),
        len(items_pressure), len(items_qty), len(items_mtc), 0
    )

    cols = {idx: _bookmark_col_index(f"BM_ITEMS_R1_C{idx}") for idx in range(1, 8)}
    if cols.get(1):
        t, start_r, _ = cols[1]
        col_indices = [(cols[idx][2] if cols.get(idx) else None) for idx in range(1, 8)]
        if any(c is None for c in col_indices):
            col_indices = list(range(0, 7))

        for i in range(max_items):
            while (start_r + i) >= len(t.rows):
                t.add_row()

            row_cells = t.rows[start_r + i].cells
            vals = [
                items_item[i] if i < len(items_item) else "",
                items_desc[i] if i < len(items_desc) else "",
                items_size[i] if i < len(items_size) else "",
                items_type[i] if i < len(items_type) else "",
                items_pressure[i] if i < len(items_pressure) else "",
                items_qty[i] if i < len(items_qty) else "",
                items_mtc[i] if i < len(items_mtc) else "",
            ]
            for j, v in enumerate(vals):
                cidx = col_indices[j]
                if cidx is None or cidx >= len(row_cells):
                    continue
                _set_cell_text(row_cells[cidx], str(v))

    # ---------- Visual table ----------
    vis_batch = data.get("vis_batch[]", [])
    vis_flange = data.get("vis_flange[]", [])
    vis_surface = data.get("vis_surface[]", [])
    vis_damage = data.get("vis_damage[]", [])
    vis_package = data.get("vis_package[]", [])
    vis_marking = data.get("vis_marking[]", [])
    vis_result = data.get("vis_result[]", [])

    max_vis = max(
        len(vis_batch), len(vis_flange), len(vis_surface), len(vis_damage),
        len(vis_package), len(vis_marking), len(vis_result), 0
    )

    vcols = {idx: _bookmark_col_index(f"BM_VIS_R1_C{idx}") for idx in range(1, 8)}
    if vcols.get(1):
        t, start_r, _ = vcols[1]
        v_col_indices = [(vcols[idx][2] if vcols.get(idx) else None) for idx in range(1, 8)]
        if any(c is None for c in v_col_indices):
            v_col_indices = list(range(0, 7))

        for i in range(max_vis):
            while (start_r + i) >= len(t.rows):
                t.add_row()

            row_cells = t.rows[start_r + i].cells
            vals = [
                vis_batch[i] if i < len(vis_batch) else "",
                vis_flange[i] if i < len(vis_flange) else "",
                vis_surface[i] if i < len(vis_surface) else "",
                vis_damage[i] if i < len(vis_damage) else "",
                vis_package[i] if i < len(vis_package) else "",
                vis_marking[i] if i < len(vis_marking) else "",
                vis_result[i] if i < len(vis_result) else "",
            ]
            for j, v in enumerate(vals):
                cidx = v_col_indices[j]
                if cidx is None or cidx >= len(row_cells):
                    continue
                _set_cell_text(row_cells[cidx], str(v))

    # ---------- Remarks ----------
    _set_bookmark_text(doc, "BM_REMARKS", (data.get("remarks") or "").strip())

    # ✅ ---------- Signatures using BOOKMARKS (with TIME) ----------
    # IMPORTANT: this MUST run regardless of manager_approved;
    # the helper itself decides what to write based on inspector_confirmed/manager_approved.
    try:
        _apply_f02_bookmark_signatures(doc, inspection)
    except Exception:
        pass

    # ---------- Layout tweaks ----------
    try:
        _apply_f02_pdf_layout_tweaks(doc)
    except Exception:
        pass

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def docx_bytes_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    import tempfile
    import subprocess
    import os

    with tempfile.TemporaryDirectory() as tmpdir:
        in_path = os.path.join(tmpdir, "hydro_export.docx")
        out_path = os.path.join(tmpdir, "hydro_export.pdf")

        with open(in_path, "wb") as f:
            f.write(docx_bytes)

        cmd = [
            "soffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", tmpdir,
            in_path,
        ]
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)

        if not os.path.exists(out_path):
            raise HTTPException(500, "PDF conversion failed")

        with open(out_path, "rb") as f:
            return f.read()

def fill_hydro_template_docx_bytes(record: HydroTestRecord) -> bytes:
    """
    Fill the hydrotest DOCX template using Word bookmarks, then return DOCX bytes.
    Template file location:
    app/templates/templates_xlsx/New hydrotest template 2026.docx
    """
    if not os.path.exists(HYDRO_TEMPLATE_DOCX):
        raise HTTPException(
            500,
            f"Hydro template missing. Put 'New hydrotest template 2026.docx' in {MRR_TEMPLATE_DIR}"
        )

    doc = Document(HYDRO_TEMPLATE_DOCX)

    def txt(v):
        return "" if v is None else str(v)

    def num1(v):
        try:
            return f"{float(v):.1f}"
        except Exception:
            return ""

    def num2(v):
        try:
            return f"{float(v):.2f}"
        except Exception:
            return ""

    test_date = record.tested_at.strftime("%Y-%m-%d") if record.tested_at else ""

    approved_dt = format_oman_dt(record.approved_at) if record.approved_at else ""
    approved_by = (record.approved_by_user_name or "").strip()
    qaqc_display = (record.qaqc_name or record.assigned_qaqc_display_name or "").strip()

    bookmark_values = {
        "BM_CLIENT_NAME": txt(record.client_name),
        "BM_PO": txt(record.client_po),
        "BM_TEST_DATE": test_date,
        "BM_REPORT_NO": txt(record.report_no or f"HT-{record.id:04d}"),
        "BM_PIPE_SPEC": txt(record.pipe_specification),
        "BM_REF_STANDARD": txt(record.reference_standard),
        "BM_BATCH": txt(record.batch_no),
        "BM_REF_PROC": txt(record.reference_dhtp_procedure),
        "BM_MACHINE": txt(record.machine_model),
        "BM_CALIBRATION": txt(record.calibration_status),
        "BM_STARTING_LENGTH": num1(record.start_length_m),
        "BM_ENDING_LENGTH": num1(record.end_length_m),
        "BM_TEST_MEDIUM": txt(record.test_medium),
        "BM_TOTAL_LEGNTH": num1(record.tested_length_m),
        "BM_TEST_PRESSURE": num2(record.hydrotest_pressure_mpa),
        "BM_HOLD_TIME": txt(record.pressure_holding_time_min),
        "BM_MAX_PRESS": num2(record.highest_pressure_recorded_mpa),
        "BM_MIN_PRESS": num2(record.lowest_pressure_recorded_mpa),
        "BM_QAQC": txt(qaqc_display),
        "BM_TECH": txt(record.testing_operator_name or record.technician_name),
    }

    for bm, value in bookmark_values.items():
        _set_bookmark_text(doc, bm, value)

    if (record.approval_status or "").upper() == "APPROVED" and approved_by and approved_dt:
        _set_bookmark_signature_block(
            doc,
            "BM_APPROVAL_SIGNATURE",
            f"Digitally signed by: {approved_by}",
            f"Date/Time: {approved_dt} (Oman)",
            font_pt=8.5,
            rgb=(70, 70, 70),
        )
    else:
        _set_bookmark_text(doc, "BM_APPROVAL_SIGNATURE", "")

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
    # -------------------------
    # LOGO
    # -------------------------
    # Do NOT insert logo into worksheet cells.
    # We stamp the logo into the PDF after conversion (true header behavior).


    # -------------------------
    # PDF EXPORT PAGE SETUP
    # -------------------------
    # IMPORTANT: make print area end at last real content row, otherwise it shrinks everything.
    try:
        ws.page_setup.scale = None
        ws.page_setup.orientation = ws.ORIENTATION_PORTRAIT
        ws.page_setup.paperSize = ws.PAPERSIZE_A4
    
        # Fit to 1 page
        ws.page_setup.fitToWidth = 1
        ws.page_setup.fitToHeight = 1
        ws.sheet_properties.pageSetUpPr.fitToPage = True
    
        # ---- AUTO PRINT AREA (A1 to L<last_used_row>) ----
        # Scan for last row that has any value in columns A..L
        last_used_row = 1
        for r in range(ws.max_row, 0, -1):
            row_has_data = False
            for c in range(1, 13):  # A..L
                v = ws.cell(r, c).value
                if v is not None and str(v).strip() != "":
                    row_has_data = True
                    break
            if row_has_data:
                last_used_row = r
                break
    
        # Add a small safety pad (in case borders/text are near the end)
        last_used_row = min(last_used_row + 2, ws.max_row)
    
        ws.print_area = f"A1:L{last_used_row}"
    
        # Remove manual page breaks embedded in template
        try:
            ws.row_breaks.brk = []
            ws.col_breaks.brk = []
        except Exception:
            pass
    
        # Margins (small = bigger content)
        ws.page_margins.left = 0.10
        ws.page_margins.right = 0.10
        ws.page_margins.top = 0.10
        ws.page_margins.bottom = 0.15
    except Exception:
        pass


    return _xlsx_bytes_from_wb(wb)


# =========================
# PDF post-processing helpers
# =========================
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from pypdf import PdfReader, PdfWriter, Transformation

import zipfile
import mimetypes
from io import BytesIO

from pypdf import PdfReader, PdfWriter  # you already use pypdf elsewhere


def _safe_filename(name: str) -> str:
    name = (name or "").strip()
    if not name:
        return "file"
    bad = ['\\', '/', ':', '*', '?', '"', '<', '>', '|', '\n', '\r', '\t']
    for ch in bad:
        name = name.replace(ch, "_")
    return name[:180]


def _read_file_bytes(path: str) -> bytes:
    with open(path, "rb") as f:
        return f.read()

def _docx_replace_all(doc: Document, needle: str, repl: str) -> None:
    """
    Replace text in paragraphs AND table cells (python-docx has no global replace).
    """
    # paragraphs
    for p in doc.paragraphs:
        if needle in p.text:
            for r in p.runs:
                if needle in r.text:
                    r.text = r.text.replace(needle, repl)

    # tables
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if needle in p.text:
                        for r in p.runs:
                            if needle in r.text:
                                r.text = r.text.replace(needle, repl)


def _image_path_to_pdf_bytes(image_path: str) -> bytes:
    """
    Convert one image file into a one-page PDF (fits into A4).
    """
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.utils import ImageReader

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    page_w, page_h = A4

    img = ImageReader(image_path)
    iw, ih = img.getSize()

    margin = 36  # 0.5 inch
    max_w = page_w - 2 * margin
    max_h = page_h - 2 * margin

    scale = min(max_w / iw, max_h / ih)
    w = iw * scale
    h = ih * scale

    x = (page_w - w) / 2
    y = (page_h - h) / 2

    c.drawImage(img, x, y, width=w, height=h, preserveAspectRatio=True, mask="auto")
    c.showPage()
    c.save()

    buf.seek(0)
    return buf.getvalue()


def _doc_path_to_pdf_bytes(path: str) -> bytes | None:
    """
    Convert an attachment into PDF bytes if possible.
    - PDF returns bytes
    - DOC/DOCX/XLS/XLSX converts via LibreOffice
    - Images convert to PDF
    - Other formats => None (kept as original in ZIP)
    """
    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        return _read_file_bytes(path)

    if ext in [".png", ".jpg", ".jpeg", ".webp"]:
        return _image_path_to_pdf_bytes(path)

    # Use LibreOffice for office files (most reliable)
    if ext in [".doc", ".docx", ".xls", ".xlsx"]:
        return _soffice_convert_file_to_pdf_bytes(path)

    return None


def _generate_shipment_report_pdf_bytes(*, lot, insp, receiving, docs) -> bytes:
    """
    Generate the standalone shipment REPORT PDF (F01 or F02 depending on template),
    including footer + digital signatures (same behavior as export/pdf endpoint).
    """
    tpl = _resolve_template_type(lot, insp)

    # Read inspection json (for timestamps / approved-by name)
    try:
        data = json.loads(getattr(insp, "inspection_json", None) or "{}")
    except Exception:
        data = {}

    inspector_name = (getattr(insp, "inspector_name", "") or "").strip()
    inspector_date = ""
    if getattr(insp, "inspector_confirmed", False):
        ts = data.get("submitted_at_utc") or getattr(insp, "created_at", None)
        inspector_date = _as_date_str(ts) if ts else _as_date_str(datetime.utcnow())

    manager_name = (data.get("manager_approved_by") or "").strip()
    manager_date = ""
    if getattr(insp, "manager_approved", False):
        ts2 = data.get("manager_approved_at_utc") or datetime.utcnow().isoformat()
        manager_date = _as_date_str(ts2)

    if tpl == "OUTSOURCED":
        docx_bytes = fill_mrr_f02_docx_bytes(
            lot=lot,
            inspection=insp,
            receiving=receiving,
            docs=docs,
        )
        pdf_bytes = docx_bytes_to_pdf_bytes(docx_bytes)

        try:
            pdf_bytes = stamp_footer_on_pdf(pdf_bytes, "QAP0600-F02")
        except Exception:
            pass

        return pdf_bytes

    # RAW => XLSX => PDF
    xlsx_bytes = fill_mrr_f01_xlsx_bytes(
        lot=lot,
        receiving=receiving,
        inspection=insp,
        docs=docs,
        photos_by_group=None,
    )
    pdf_bytes = _try_convert_xlsx_to_pdf_bytes(xlsx_bytes)

    try:
        pdf_bytes = stamp_footer_on_pdf(pdf_bytes, "QAP0600-F01")
    except Exception:
        pass

    try:
        pdf_bytes = stamp_signatures_on_pdf(
            pdf_bytes,
            inspector_name=inspector_name if getattr(insp, "inspector_confirmed", False) else "",
            inspector_date=inspector_date if getattr(insp, "inspector_confirmed", False) else "",
            manager_name=manager_name if getattr(insp, "manager_approved", False) else "",
            manager_date=manager_date if getattr(insp, "manager_approved", False) else "",
        )
    except Exception:
        pass

    return pdf_bytes

def _soffice_convert_file_to_pdf_bytes(input_path: str) -> bytes | None:
    """
    Convert a file to PDF using LibreOffice (soffice).
    Works for: .doc, .docx, .xls, .xlsx, and many others.
    Returns PDF bytes or None if conversion fails.
    """
    import subprocess
    import tempfile
    import shutil

    if not input_path or not os.path.exists(input_path):
        return None

    with tempfile.TemporaryDirectory() as tmpdir:
        # Copy to temp to avoid filename issues
        base = os.path.basename(input_path)
        safe_base = _safe_filename(base)
        tmp_in = os.path.join(tmpdir, safe_base)

        try:
            shutil.copy2(input_path, tmp_in)
        except Exception:
            tmp_in = input_path  # fallback

        cmd = [
            "soffice",
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to",
            "pdf",
            "--outdir",
            tmpdir,
            tmp_in,
        ]

        try:
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        except Exception:
            return None

        # LibreOffice outputs same filename with .pdf
        out_pdf = os.path.splitext(os.path.basename(tmp_in))[0] + ".pdf"
        out_path = os.path.join(tmpdir, out_pdf)

        if not os.path.exists(out_path):
            # Sometimes LO changes extension casing, do a scan
            for f in os.listdir(tmpdir):
                if f.lower().endswith(".pdf"):
                    out_path = os.path.join(tmpdir, f)
                    break

        if not os.path.exists(out_path):
            return None

        try:
            return _read_file_bytes(out_path)
        except Exception:
            return None

def fit_pdf_pages_to_a4(
    pdf_bytes: bytes,
    margin_left_right: float = 3.0,
    margin_bottom: float = 3.0,
    header_reserved: float = 78.0,
    zoom: float = 1.18,        # bigger content
    shift_x: float = 40.0,     # move LEFT/RIGHT
    shift_y: float = -80.0,    # move UP/DOWN
) -> bytes:
    """
    Force pages onto A4, reserve header space, then zoom and shift.
    """
    # IMPORTANT: use a local alias so it can NEVER break from duplicate imports
    _B = BytesIO

    reader = PdfReader(_B(pdf_bytes))
    writer = PdfWriter()

    a4_w, a4_h = A4
    usable_w = a4_w - 2 * margin_left_right
    usable_h = a4_h - margin_bottom - header_reserved

    for page in reader.pages:
        src_w = float(page.mediabox.width)
        src_h = float(page.mediabox.height)

        base_scale = min(usable_w / src_w, usable_h / src_h)
        scale = base_scale * zoom

        new_page = writer.add_blank_page(width=a4_w, height=a4_h)

        content_w = src_w * scale
        content_h = src_h * scale

        tx = (a4_w - content_w) / 2.0 + shift_x
        ty = margin_bottom + (usable_h - content_h) / 2.0 + shift_y

        new_page.merge_transformed_page(
            page,
            Transformation().scale(scale).translate(tx, ty)
        )

    out = _B()
    writer.write(out)
    out.seek(0)
    return out.read()


def make_logo_stamp_pdf(page_w: float, page_h: float, logo_path: str) -> bytes:
    """
    Create a transparent 1-page PDF with a centered logo at top.
    """
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=(page_w, page_h))

    img = ImageReader(logo_path)
    iw, ih = img.getSize()

    target_w = page_w * 0.32
    scale = target_w / float(iw)
    target_h = float(ih) * scale

    top_margin = 16
    x = (page_w - target_w) / 2.0
    y = page_h - top_margin - target_h

    c.drawImage(img, x, y, width=target_w, height=target_h, mask="auto")
    c.showPage()
    c.save()

    buf.seek(0)
    return buf.getvalue()


def stamp_logo_on_pdf(pdf_bytes: bytes, logo_path: str) -> bytes:
    """
    Overlay the logo stamp onto every page (top-center).
    """
    if not logo_path or not os.path.exists(logo_path):
        return pdf_bytes

    reader = PdfReader(BytesIO(pdf_bytes))
    writer = PdfWriter()

    for page in reader.pages:
        w = float(page.mediabox.width)
        h = float(page.mediabox.height)

        stamp_pdf = make_logo_stamp_pdf(w, h, logo_path)
        stamp_reader = PdfReader(BytesIO(stamp_pdf))

        page.merge_page(stamp_reader.pages[0])
        writer.add_page(page)

    out = BytesIO()
    writer.write(out)
    out.seek(0)
    return out.getvalue()








def make_footer_stamp_pdf(page_w: float, page_h: float, left_text: str, right_text: str = "") -> bytes:
    """
    Create a transparent 1-page PDF with footer text.
    """
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=(page_w, page_h))

    # subtle footer (dark gray)
    c.setFillColorRGB(0.25, 0.25, 0.25)
    c.setFont("Helvetica", 9)

    y = 12  # distance from bottom
    c.drawString(18, y, left_text)

    if right_text:
        # right aligned
        c.drawRightString(page_w - 18, y, right_text)

    c.showPage()
    c.save()
    buf.seek(0)
    return buf.getvalue()


def stamp_footer_on_pdf(pdf_bytes: bytes, left_text: str) -> bytes:
    """
    Overlay footer onto every page.
    """
    reader = PdfReader(BytesIO(pdf_bytes))
    writer = PdfWriter()

    total = len(reader.pages)

    for idx, page in enumerate(reader.pages, start=1):
        w = float(page.mediabox.width)
        h = float(page.mediabox.height)

        right_text = f"Page {idx}/{total}"
        stamp_pdf = make_footer_stamp_pdf(w, h, left_text, right_text)
        stamp_reader = PdfReader(BytesIO(stamp_pdf))

        page.merge_page(stamp_reader.pages[0])
        writer.add_page(page)

    out = BytesIO()
    writer.write(out)
    out.seek(0)
    return out.getvalue()

from io import BytesIO
from datetime import datetime
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from pypdf import PdfReader, PdfWriter

def make_signature_stamp_pdf(
    page_w: float,
    page_h: float,
    inspector_name: str | None = None,
    inspector_date: str | None = None,
    manager_name: str | None = None,
    manager_date: str | None = None,
) -> bytes:
    """
    Creates a transparent overlay PDF with digital signature text.
    Positions are tuned for your QAP0600-F01 A4 portrait layout.
    """
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=(page_w, page_h))

    # Transparent-ish text if supported
    try:
        c.setFillAlpha(0.55)
    except Exception:
        pass

    c.setFont("Helvetica", 8)
    c.setFillColorRGB(0.10, 0.10, 0.10)

    # ---- Coordinates (A4 portrait typical) ----
    # Bottom signature block is near bottom; adjust if needed:
    # Inspector stamp left block
    insp_x = 50
    insp_y = 120

    # Manager stamp middle block
    mgr_x = page_w * 0.42
    mgr_y = 120

    # Inspector stamp
    if inspector_name:
        c.drawString(insp_x, insp_y + 14, f"Digitally signed by: {inspector_name}")
        if inspector_date:
            c.drawString(insp_x, insp_y, f"Date: {inspector_date}")

    # Manager stamp
    if manager_name:
        c.drawString(mgr_x, mgr_y + 14, f"Digitally approved by: {manager_name}")
        if manager_date:
            c.drawString(mgr_x, mgr_y, f"Date: {manager_date}")

    c.showPage()
    c.save()
    buf.seek(0)
    return buf.getvalue()


def make_signature_stamp_pdf_f02(
    page_w: float,
    page_h: float,
    inspector_name: str | None = None,
    inspector_dt: str | None = None,
    reviewer_name: str | None = None,
    reviewer_dt: str | None = None,
    manager_name: str | None = None,
    manager_dt: str | None = None,
    y_offset: float = 0.0,   # + moves UP, - moves DOWN
) -> bytes:
    """
    Transparent overlay PDF with digital signature text for QAP0600-F02.
    Stamps intended for the bottom signature table on the LAST page.
    """
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=(page_w, page_h))

    try:
        c.setFillAlpha(0.55)
    except Exception:
        pass

    c.setFont("Helvetica", 8)
    c.setFillColorRGB(0.10, 0.10, 0.10)

    # Base position (from bottom). Increase to move the stamp UP.
    y_base = 105 + float(y_offset)

    # Left / Middle / Right boxes
    insp_x = 55
    rev_x = page_w * 0.41
    appr_x = page_w * 0.73

    # Inspector (left)
    if inspector_name:
        c.drawString(insp_x, y_base + 14, f"Digitally signed by: {inspector_name}")
        if inspector_dt:
            c.drawString(insp_x, y_base, f"Date/Time: {inspector_dt}")

    # Reviewer (middle)
    if reviewer_name:
        c.drawString(rev_x, y_base + 14, f"Digitally reviewed by: {reviewer_name}")
        if reviewer_dt:
            c.drawString(rev_x, y_base, f"Date/Time: {reviewer_dt}")

    # Manager (right)
    if manager_name:
        c.drawString(appr_x, y_base + 14, f"Digitally approved by: {manager_name}")
        if manager_dt:
            c.drawString(appr_x, y_base, f"Date/Time: {manager_dt}")

    c.showPage()
    c.save()
    buf.seek(0)
    return buf.getvalue()

def _apply_f02_docx_signatures(
    doc: Document,
    inspector_name: str = "",
    inspector_dt: str = "",
    reviewer_name: str = "",
    reviewer_dt: str = "",
    manager_name: str = "",
    manager_dt: str = "",
) -> None:
    """
    Fill signatures INSIDE the signature boxes using bookmarks.
    Supports your bookmark names:
      BM_INSPECTED_BY
      BM_REVIEWD_BY  (your spelling)
      BM_REVIEWED_BY (correct spelling, supported too)
      BM_APPROVED_BY
    """

    # Inspector
    if inspector_name:
        _set_bookmark_signature_block(
            doc,
            "BM_INSPECTED_BY",
            f"Digitally signed by: {inspector_name}",
            f"Date: {inspector_dt}",
            font_pt=8.5,
            rgb=(70, 70, 70),
        )
    else:
        _set_bookmark_text(doc, "BM_INSPECTED_BY", "")

    # Reviewer (optional)
    if reviewer_name:
        # support both bookmark spellings
        ok = _set_bookmark_signature_block(
            doc,
            "BM_REVIEWD_BY",
            f"Digitally reviewed by: {reviewer_name}",
            f"Date: {reviewer_dt}",
            font_pt=8.5,
            rgb=(70, 70, 70),
        )
        if not ok:
            _set_bookmark_signature_block(
                doc,
                "BM_REVIEWED_BY",
                f"Digitally reviewed by: {reviewer_name}",
                f"Date: {reviewer_dt}",
                font_pt=8.5,
                rgb=(70, 70, 70),
            )
    else:
        _set_bookmark_text(doc, "BM_REVIEWD_BY", "")
        _set_bookmark_text(doc, "BM_REVIEWED_BY", "")

    # Manager approval
    if manager_name:
        _set_bookmark_signature_block(
            doc,
            "BM_APPROVED_BY",
            f"Digitally approved by: {manager_name}",
            f"Date: {manager_dt}",
            font_pt=8.5,
            rgb=(70, 70, 70),
        )
    else:
        _set_bookmark_text(doc, "BM_APPROVED_BY", "")


def stamp_signatures_on_pdf_f02(
    pdf_bytes: bytes,
    inspector_name: str | None = None,
    inspector_date: str | None = None,
    reviewer_name: str | None = None,
    reviewer_date: str | None = None,
    manager_name: str | None = None,
    manager_date: str | None = None,
) -> bytes:
    """
    Overlay F02 signatures on every page.
    """
    reader = PdfReader(BytesIO(pdf_bytes))
    writer = PdfWriter()

    for page in reader.pages:
        w = float(page.mediabox.width)
        h = float(page.mediabox.height)

        stamp_pdf = make_signature_stamp_pdf_f02(
            w, h,
            inspector_name=inspector_name,
            inspector_date=inspector_date,
            reviewer_name=reviewer_name,
            reviewer_date=reviewer_date,
            manager_name=manager_name,
            manager_date=manager_date,
        )
        stamp_reader = PdfReader(BytesIO(stamp_pdf))
        page.merge_page(stamp_reader.pages[0])
        writer.add_page(page)

    out = BytesIO()
    writer.write(out)
    out.seek(0)
    return out.getvalue()

def _as_date_str(x) -> str:
    """
    Returns YYYY-MM-DD
    Accepts datetime, ISO string, or anything convertible to str.
    """
    try:
        if isinstance(x, datetime):
            return x.strftime("%Y-%m-%d")
        s = str(x or "").strip()
        if not s:
            return ""
        try:
            dt = datetime.fromisoformat(s.replace("Z", "+00:00"))
            return dt.strftime("%Y-%m-%d")
        except Exception:
            return s[:10]
    except Exception:
        return ""


def _as_datetime_str(x) -> str:
    """
    Returns YYYY-MM-DD HH:MM (24h)
    Accepts datetime, ISO string, or anything convertible to str.
    """
    try:
        if isinstance(x, datetime):
            return x.strftime("%Y-%m-%d %H:%M")
        s = str(x or "").strip()
        if not s:
            return ""
        try:
            dt = datetime.fromisoformat(s.replace("Z", "+00:00"))
            return dt.strftime("%Y-%m-%d %H:%M")
        except Exception:
            # fallback: keep first 16 chars if it looks like datetime
            return s[:16]
    except Exception:
        return ""

def stamp_signatures_on_pdf(
    pdf_bytes: bytes,
    inspector_name: str | None = None,
    inspector_date: str | None = None,
    manager_name: str | None = None,
    manager_date: str | None = None,
) -> bytes:
    """
    Overlays signature stamp on every page.
    """
    reader = PdfReader(BytesIO(pdf_bytes))
    writer = PdfWriter()

    for page in reader.pages:
        w = float(page.mediabox.width)
        h = float(page.mediabox.height)

        stamp_pdf = make_signature_stamp_pdf(
            w, h,
            inspector_name=inspector_name,
            inspector_date=inspector_date,
            manager_name=manager_name,
            manager_date=manager_date,
        )
        stamp_reader = PdfReader(BytesIO(stamp_pdf))
        page.merge_page(stamp_reader.pages[0])
        writer.add_page(page)

    out = BytesIO()
    writer.write(out)
    out.seek(0)
    return out.getvalue()



app = FastAPI()

# Session middleware (required for request.session in require_user)
#app.add_middleware(SessionMiddleware, secret_key=SESSION_SECRET)#

BASE_DIR = os.path.dirname(__file__)

from fastapi import Depends, HTTPException, Request

def require_user(session: Session = Depends(get_session)) -> User:
    """
    Temporary auth bypass: this environment has no SessionMiddleware (itsdangerous missing).
    Returns the first user in DB if exists.
    """
    user = session.exec(select(User).order_by(User.id.asc())).first()
    if not user:
        raise HTTPException(status_code=401, detail="No users found. Create a user first.")
    return user


import io
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from pypdf import PdfReader, PdfWriter

PAGE_W, PAGE_H = A4  # 595 x 842 points

def _overlay_grid_pdf() -> PdfReader:
    """
    Creates a 1-page PDF overlay with a coordinate grid and labels.
    Useful to map where to place text on the template.
    """
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.setFont("Helvetica", 7)

    step = 50
    # vertical lines + x labels
    x = 0
    while x <= PAGE_W:
        c.line(x, 0, x, PAGE_H)
        c.drawString(x + 2, 2, f"x={int(x)}")
        x += step

    # horizontal lines + y labels
    y = 0
    while y <= PAGE_H:
        c.line(0, y, PAGE_W, y)
        c.drawString(2, y + 2, f"y={int(y)}")
        y += step

    c.save()
    buf.seek(0)
    return PdfReader(buf)

    
# =========================
# Upload storage (local FS)
# =========================
DATA_DIR = os.environ.get("DATA_DIR", "/app/data")
MRR_UPLOAD_DIR = os.path.join(DATA_DIR, "mrr_uploads")
MRR_PHOTO_DIR = os.path.join(DATA_DIR, "mrr_photos")

BURST_DIR = os.path.join(DATA_DIR, "burst")
BURST_FILES_DIR = os.path.join(BURST_DIR, "files")

os.makedirs(BURST_DIR, exist_ok=True)
os.makedirs(BURST_FILES_DIR, exist_ok=True)

os.makedirs(MRR_UPLOAD_DIR, exist_ok=True)
os.makedirs(MRR_PHOTO_DIR, exist_ok=True)



templates = Jinja2Templates(directory=os.path.join(BASE_DIR, "templates"))
app.mount("/static", StaticFiles(directory=os.path.join(BASE_DIR, "static")), name="static")


@app.exception_handler(HTTPException)
async def app_http_exception_handler(request: Request, exc: HTTPException):
    if exc.status_code == 401 and str(exc.detail) == "Not logged in":
        accept = (request.headers.get("accept") or "").lower()
        if "text/html" in accept:
            return RedirectResponse(url="/login", status_code=302)
    raise exc
# =========================
# File upload directories
# =========================
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")



IMAGE_MAP = {
    "LINER": "/static/images/liner.png",
    "REINFORCEMENT": "/static/images/reinforcement.png",
    "COVER": "/static/images/cover.png",
}

PAPER_BG_MAP = {
    "LINER": os.path.join(BASE_DIR, "static", "papers", "liner_bg.pdf"),
    "REINFORCEMENT": os.path.join(BASE_DIR, "static", "papers", "reinforcement_bg.pdf"),
    "COVER": os.path.join(BASE_DIR, "static", "papers", "cover_bg.pdf"),
}

TEMPLATE_XLSX_MAP = {
    "LINER": os.path.join(BASE_DIR, "templates", "templates_xlsx", "liner.xlsx"),
    "REINFORCEMENT": os.path.join(BASE_DIR, "templates", "templates_xlsx", "reinforcement.xlsx"),
    "COVER": os.path.join(BASE_DIR, "templates", "templates_xlsx", "cover.xlsx"),
}
def resolve_mrr_doc_path(p: str) -> str:
    """Resolve stored MRR document path to an existing file on disk.

    Backward compatible with rows that stored absolute paths, relative paths, or just filenames.
    """
    if not p:
        return ""

    p_norm = p.replace("\\", "/").lstrip("/")

    # absolute
    try:
        if os.path.isabs(p) and os.path.exists(p):
            return p
    except Exception:
        pass

    candidates = [
        p_norm,
        os.path.join(BASE_DIR, p_norm) if 'BASE_DIR' in globals() else p_norm,
        os.path.join(DATA_DIR, p_norm) if 'DATA_DIR' in globals() else p_norm,
        os.path.join(MRR_UPLOAD_DIR, p_norm),
        os.path.join(MRR_UPLOAD_DIR, os.path.basename(p_norm)),
    ]

    for c in candidates:
        if c and os.path.exists(c):
            return c

    return ""


def _resolve_template_type(lot, inspection) -> str:
    """
    Decide which template to use for:
    - UI inspection page
    - exporting the report

    Priority:
    1) inspection.template_type (if stored)
    2) lot.lot_type (if stored)
    Default: RAW
    """
    lot_type = (getattr(lot, "lot_type", None) or "").strip().upper()
    insp_type = (getattr(inspection, "template_type", None) or "").strip().upper()

    if insp_type in ["OUTSOURCED", "RAW"]:
        return insp_type

    if lot_type in ["OUTSOURCED", "RAW"]:
        return lot_type

    return "RAW"


def get_current_draft_shipment(session: Session, lot_id: int):
    """
    Returns the latest draft shipment for THIS ticket only.
    Draft = not submitted AND not manager approved.
    """
    return session.exec(
        select(MrrReceivingInspection)
        .where(MrrReceivingInspection.ticket_id == lot_id)
        .where(MrrReceivingInspection.inspector_confirmed == False)
        .where(MrrReceivingInspection.manager_approved == False)
        .order_by(MrrReceivingInspection.created_at.desc())
    ).first()


def resolve_burst_file_path(p: str) -> str:
    """
    Stored file_path is RELATIVE to BASE_DIR sometimes, or already absolute in older records.
    This mirrors MRR behavior: resolve safely.
    """
    if not p:
        return ""
    if os.path.isabs(p):
        return p
    # most of our saved paths are relative to BASE_DIR
    candidate = os.path.join(BASE_DIR, p)
    if os.path.exists(candidate):
        return candidate
    # fallback: if someone stored relative to DATA_DIR
    candidate2 = os.path.join(DATA_DIR, p)
    return candidate2


def get_run_produced_length_m(session: Session, run_id: int) -> float:
    """
    Smart produced length:
    - look at InspectionValue where param_key == 'length_m'
    - return the MAX value recorded for this run
    - if manager confirmed a larger final actual produced length, use that
    """
    run = session.get(ProductionRun, run_id)
    confirmed = 0.0
    if run:
        try:
            confirmed = float(getattr(run, "confirmed_total_length_m", 0.0) or 0.0)
        except Exception:
            confirmed = 0.0

    q = session.exec(
        select(InspectionValue.value)
        .join(InspectionEntry, InspectionEntry.id == InspectionValue.entry_id)
        .where(InspectionEntry.run_id == run_id)
        .where(InspectionValue.param_key == "length_m")
        .where(InspectionValue.value != None)  # noqa: E711
    ).all()

    vals = []
    for v in q:
        try:
            vals.append(float(v))
        except Exception:
            pass

    measured = max(vals) if vals else 0.0
    return max(measured, confirmed)

def upsert_burst_attachment(
    session: Session,
    report_id: int,
    kind: str,
    file_rel_path: str,
    caption: str,
    user,
):
    """
    Ensures ONLY ONE attachment per kind per report.
    If the kind already exists, update it (replace file).
    """
    kind = (kind or "").strip().upper()

    existing = session.exec(
        select(BurstAttachment)
        .where(BurstAttachment.report_id == report_id)
        .where(BurstAttachment.kind == kind)
    ).first()

    if existing:
        existing.file_path = file_rel_path
        existing.caption = caption
        existing.uploaded_by_user_id = user.id
        existing.uploaded_by_user_name = user.display_name
        existing.uploaded_at = datetime.utcnow()
        session.add(existing)
        session.commit()
        return existing

    att = BurstAttachment(
        report_id=report_id,
        kind=kind,
        caption=caption,
        file_path=file_rel_path,
        uploaded_by_user_id=user.id,
        uploaded_by_user_name=user.display_name,
    )
    session.add(att)
    session.commit()
    session.refresh(att)
    return att


def ensure_burst_samples(session: Session, report_id: int, desired: int = 5) -> list[BurstSample]:
 
    try:
        desired = int(desired or 5)
    except Exception:
        desired = 5
    if desired < 1:
        desired = 1

    rows = session.exec(
        select(BurstSample)
        .where(BurstSample.report_id == report_id)
        .order_by(BurstSample.id.asc())
    ).all()

    missing = desired - len(rows)
    if missing > 0:
        for _ in range(missing):
            session.add(BurstSample(report_id=report_id, sample_start_m=0.0, sample_length_m=0.0))
        session.commit()
    
    rows = session.exec(
        select(BurstSample)
        .where(BurstSample.report_id == report_id)
        .order_by(BurstSample.id.asc())
    ).all()
    
    return rows


def get_finished_cover_runs(session: Session) -> list[ProductionRun]:
    """
    Finished product = COVER run that has:
    - a batch number
    - total_length_m > 0
    - AND (liner + reinforcement exist with same batch)  [so it's a full pipe]
    """
    cover_runs = session.exec(
        select(ProductionRun)
        .where(ProductionRun.process == "COVER")
        .where(ProductionRun.dhtp_batch_no != "")
        .where(ProductionRun.total_length_m > 0)
        .order_by(ProductionRun.id.desc())
    ).all()

    finished: list[ProductionRun] = []
    for cr in cover_runs:
        batch = (cr.dhtp_batch_no or "").strip()
        liner, reinf, _cover = _find_related_runs_by_batch(session, batch)
        if liner and reinf:
            finished.append(cr)

    return finished


def _draw_signatures(c, report, y):
    w, h = A4
    if y < 45*mm:
        c.showPage()
        y = h - 32*mm

    c.setFont("Helvetica-Bold", 10)
    c.drawString(20*mm, y, "Signatures")
    y -= 12*mm

    tech_name = _txt(getattr(report, "technician_name", ""))  # typed
    insp_name = _txt(getattr(report, "created_by_user_name", ""))  # inspector = who filled

    c.setFont("Helvetica", 9)

    # Technician
    c.drawString(20*mm, y, "Technician:")
    c.line(45*mm, y-1*mm, 110*mm, y-1*mm)
    c.drawString(45*mm, y+1*mm, tech_name)
    
    # Inspector
    c.drawString(125*mm, y, "Inspector:")
    c.line(145*mm, y-1*mm, 190*mm, y-1*mm)
    c.drawString(145*mm, y+1*mm, insp_name)

    y -= 14*mm
    c.drawString(20*mm, y, "Date:")
    c.line(30*mm, y-1*mm, 80*mm, y-1*mm)
    c.drawString(30*mm, y+1*mm, _txt(getattr(report, "tested_at", "") or getattr(report, "created_at", "")))

    return y


def _create_run_with_default_params(
    session: Session,
    *,
    process: str,
    dhtp_batch_no: str,
    client_name: str,
    po_number: str,
    itp_number: str,
    pipe_specification: str,
    raw_material_spec: str,
    total_length_m: float,
) -> ProductionRun:
    run = ProductionRun(
        process=process,
        dhtp_batch_no=(dhtp_batch_no or "").strip(),
        client_name=(client_name or "").strip(),
        po_number=(po_number or "").strip(),
        itp_number=(itp_number or "").strip(),
        pipe_specification=(pipe_specification or "").strip(),
        raw_material_spec=(raw_material_spec or "").strip(),
        total_length_m=float(total_length_m or 0.0),
    )
    session.add(run)
    session.commit()
    session.refresh(run)

    for idx, (key, label, unit) in enumerate(PROCESS_PARAMS[process]):
        session.add(
            RunParameter(
                run_id=run.id,
                param_key=key,
                label=label,
                unit=unit,
                display_order=idx,
            )
        )
    session.commit()
    return run


def _next_final_phase_no(session: Session, batch_no: str) -> int:
    phases = session.exec(
        select(FinalInspectionPhase)
        .where(FinalInspectionPhase.batch_no == (batch_no or "").strip())
        .order_by(FinalInspectionPhase.phase_no.desc(), FinalInspectionPhase.id.desc())
    ).all()
    if not phases:
        return 1
    return int(phases[0].phase_no or 0) + 1


def _final_phase_total_length_m(session: Session, phase_id: int) -> float:
    reels = session.exec(
        select(FinalInspectionReel).where(FinalInspectionReel.phase_id == phase_id)
    ).all()
    total = 0.0
    for r in reels:
        try:
            total += float(r.reel_length_m or 0.0)
        except Exception:
            pass
    return round(total, 3)


def _build_final_batch_summary(session: Session, cover_run: ProductionRun) -> dict:
    batch_no = (cover_run.dhtp_batch_no or "").strip()
    produced_length_m = float(get_run_produced_length_m(session, cover_run.id) or 0.0)

    burst_ranges = _get_burst_used_ranges(session, batch_no, produced_length_m)
    burst_used_m = _range_total_length(burst_ranges)
    expected_after_burst_m = max(0.0, produced_length_m - burst_used_m)

    phases = session.exec(
        select(FinalInspectionPhase)
        .where(FinalInspectionPhase.batch_no == batch_no)
        .order_by(FinalInspectionPhase.phase_no.asc(), FinalInspectionPhase.id.asc())
    ).all()

    phase_rows = []
    approved_total_m = 0.0
    all_reels_total_m = 0.0

    for p in phases:
        reels = session.exec(
            select(FinalInspectionReel)
            .where(FinalInspectionReel.phase_id == p.id)
            .order_by(FinalInspectionReel.created_at.asc(), FinalInspectionReel.id.asc())
        ).all()

        phase_total_m = 0.0
        for r in reels:
            try:
                phase_total_m += float(r.reel_length_m or 0.0)
            except Exception:
                pass

        all_reels_total_m += phase_total_m
        if (p.status or "").upper() == "APPROVED":
            approved_total_m += phase_total_m

        phase_rows.append({
            "phase": p,
            "reels": reels,
            "reel_count": len(reels),
            "phase_total_m": round(phase_total_m, 3),
        })

    batch_note = session.exec(
        select(FinalInspectionBatchNote).where(FinalInspectionBatchNote.batch_no == batch_no)
    ).first()

    remaining_vs_confirmed_m = max(0.0, expected_after_burst_m - approved_total_m)
    pending_overrun_m = max(0.0, all_reels_total_m - expected_after_burst_m)

    return {
        "batch_no": batch_no,
        "run": cover_run,
        "client_name": getattr(cover_run, "client_name", "") or "",
        "client_po": getattr(cover_run, "po_number", "") or "",
        "pipe_specification": getattr(cover_run, "pipe_specification", "") or "",
        "produced_length_m": produced_length_m,
        "burst_used_m": burst_used_m,
        "expected_after_burst_m": expected_after_burst_m,
        "approved_total_m": round(approved_total_m, 3),
        "all_reels_total_m": round(all_reels_total_m, 3),
        "remaining_vs_confirmed_m": round(remaining_vs_confirmed_m, 3),
        "pending_overrun_m": round(pending_overrun_m, 3),
        "phases": phase_rows,
        "batch_note": batch_note,
        "burst_ranges": burst_ranges,
    }
# =========================
# BURST TESTING ROUTES
# =========================

def _find_related_runs_by_batch(session: Session, batch_no: str):
    batch_no = (batch_no or "").strip()
    if not batch_no:
        return None, None, None

    liner = session.exec(
        select(ProductionRun).where(
            (ProductionRun.dhtp_batch_no == batch_no) & (ProductionRun.process == "LINER")
        )
    ).first()

    reinf = session.exec(
        select(ProductionRun).where(
            (ProductionRun.dhtp_batch_no == batch_no) & (ProductionRun.process == "REINFORCEMENT")
        )
    ).first()

    cover = session.exec(
        select(ProductionRun).where(
            (ProductionRun.dhtp_batch_no == batch_no) & (ProductionRun.process == "COVER")
        )
    ).first()

    return liner, reinf, cover

def _draw_report_info_table(c, report, x, y):
    n = int(getattr(report, "total_no_of_specimens", 0) or 0)

    data = [
        ["Batch No", _txt(getattr(report, "batch_no", "")), "Client", _txt(getattr(report, "client_name", ""))],
        ["Client PO", _txt(getattr(report, "client_po", "")), "Pipe Spec", _txt(getattr(report, "pipe_specification", ""))],
        ["Test Medium", _txt(getattr(report, "testing_medium", "")), "Lab Temp", _txt(getattr(report, "laboratory_temperature", ""))],
        ["Standard", _txt(getattr(report, "reference_standard", "")), "Procedure", _txt(getattr(report, "reference_dhtp_procedure", ""))],
        ["System Max (MPa)", _txt(getattr(report, "system_max_pressure", "")), "Specimens", _txt(n)],
        ["Test Date", _txt(getattr(report, "tested_at", "") or getattr(report, "created_at", "")), "", ""],
    ]

    tbl = Table(data, colWidths=[34*mm, 60*mm, 34*mm, 60*mm])
    tbl.setStyle(TableStyle([
        ("FONT", (0,0), (-1,-1), "Helvetica", 9),
        ("FONT", (0,0), (0,-1), "Helvetica-Bold", 9),
        ("FONT", (2,0), (2,-1), "Helvetica-Bold", 9),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("TOPPADDING", (0,0), (-1,-1), 2),
        ("BOTTOMPADDING", (0,0), (-1,-1), 2),

        # optional: light row separators for neatness
        ("LINEBELOW", (0,0), (-1,-1), 0.25, colors.lightgrey),
    ]))

    tbl.wrapOn(c, 0, 0)
    tbl_h = len(data) * 7 * mm
    tbl.drawOn(c, x, y - tbl_h)
    return y - tbl_h


def _is_manager_or_boss(user: User) -> bool:
    return (getattr(user, "role", "") or "").strip().upper() in ["MANAGER", "BOSS"]


def _burst_report_is_complete(rep: BurstTestReport, samples: list["BurstSample"]) -> bool:
    if not (rep.reference_standard or "").strip():
        return False
    if not (rep.reference_dhtp_procedure or "").strip():
        return False
    if not (rep.testing_medium or "").strip():
        return False
    if not (rep.technician_name or "").strip():
        return False

    needed = samples[: max(1, int(rep.total_no_of_specimens or 1))]
    if not needed:
        return False

    for s in needed:
        if not (s.sample_serial_number or "").strip():
            return False
        if float(getattr(s, "sample_length_m", 0.0) or 0.0) <= 0:
            return False
        if float(getattr(s, "actual_burst_psi", 0.0) or 0.0) <= 0:
            return False
        if not (s.pressurization_time_s or "").strip():
            return False
        if not (s.test_result or "").strip():
            return False

    return True


def _burst_snapshot(rep: BurstTestReport, samples: list["BurstSample"]) -> dict:
    return {
        "report": {
            "reference_standard": rep.reference_standard,
            "reference_dhtp_procedure": rep.reference_dhtp_procedure,
            "system_max_pressure": rep.system_max_pressure,
            "laboratory_temperature": rep.laboratory_temperature,
            "testing_medium": rep.testing_medium,
            "total_no_of_specimens": rep.total_no_of_specimens,
            "effective_length_m": rep.effective_length_m,
            "liner_thickness": rep.liner_thickness,
            "reinforcement_thickness": rep.reinforcement_thickness,
            "cover_thickness": rep.cover_thickness,
            "sample_serial_number": rep.sample_serial_number,
            "actual_burst_psi": rep.actual_burst_psi,
            "pressurization_time_s": rep.pressurization_time_s,
            "test_result": rep.test_result,
            "failure_mode": rep.failure_mode,
            "notes": rep.notes,
            "qa_qc_officer_name": rep.qa_qc_officer_name,
            "technician_name": rep.technician_name,
        },
        "samples": [
            {
                "id": s.id,
                "sample_serial_number": s.sample_serial_number,
                "sample_length_m": s.sample_length_m,
                "actual_burst_psi": s.actual_burst_psi,
                "pressurization_time_s": s.pressurization_time_s,
                "test_result": s.test_result,
                "liner_material_grade": s.liner_material_grade,
                "liner_thickness_mm": s.liner_thickness_mm,
                "reinforcement_material_grade": s.reinforcement_material_grade,
                "reinforcement_thickness_mm": s.reinforcement_thickness_mm,
                "cover_material_grade": s.cover_material_grade,
                "cover_thickness_mm": s.cover_thickness_mm,
            }
            for s in samples
        ],
    }


def _apply_burst_snapshot(rep: BurstTestReport, samples: list["BurstSample"], payload: dict):
    r = payload.get("report", {}) or {}

    rep.reference_standard = r.get("reference_standard", "") or ""
    rep.reference_dhtp_procedure = r.get("reference_dhtp_procedure", "") or ""
    rep.system_max_pressure = r.get("system_max_pressure", "") or ""
    rep.laboratory_temperature = r.get("laboratory_temperature", "") or ""
    rep.testing_medium = r.get("testing_medium", "") or ""
    rep.total_no_of_specimens = int(r.get("total_no_of_specimens", rep.total_no_of_specimens or 1) or 1)
    rep.effective_length_m = r.get("effective_length_m", "") or ""
    rep.liner_thickness = r.get("liner_thickness", "") or ""
    rep.reinforcement_thickness = r.get("reinforcement_thickness", "") or ""
    rep.cover_thickness = r.get("cover_thickness", "") or ""
    rep.sample_serial_number = r.get("sample_serial_number", "") or ""
    rep.actual_burst_psi = float(r.get("actual_burst_psi", 0.0) or 0.0)
    rep.pressurization_time_s = r.get("pressurization_time_s", "") or ""
    rep.test_result = r.get("test_result", "") or ""
    rep.failure_mode = r.get("failure_mode", "") or ""
    rep.notes = r.get("notes", "") or ""
    rep.qa_qc_officer_name = r.get("qa_qc_officer_name", "") or ""
    rep.technician_name = r.get("technician_name", "") or ""

    by_id = {s.id: s for s in samples}
    for item in payload.get("samples", []) or []:
        s = by_id.get(item.get("id"))
        if not s:
            continue

        s.sample_serial_number = item.get("sample_serial_number", "") or ""
        s.sample_length_m = float(item.get("sample_length_m", 0.0) or 0.0)
        s.actual_burst_psi = float(item.get("actual_burst_psi", 0.0) or 0.0)
        s.pressurization_time_s = item.get("pressurization_time_s", "") or ""
        s.test_result = item.get("test_result", "") or ""
        s.liner_material_grade = item.get("liner_material_grade", "") or ""
        s.liner_thickness_mm = float(item.get("liner_thickness_mm", 0.0) or 0.0)
        s.reinforcement_material_grade = item.get("reinforcement_material_grade", "") or ""
        s.reinforcement_thickness_mm = float(item.get("reinforcement_thickness_mm", 0.0) or 0.0)
        s.cover_material_grade = item.get("cover_material_grade", "") or ""
        s.cover_thickness_mm = float(item.get("cover_thickness_mm", 0.0) or 0.0)
        

@app.get("/burst")
def burst_dashboard(request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)

    reports = session.exec(
        select(BurstTestReport).order_by(BurstTestReport.id.desc())
    ).all()

    run_map = {}
    burst_map = {}
    result_map = {}

    for r in reports:
        run = None
        if not r.is_unlinked and r.linked_run_id:
            run = session.get(ProductionRun, r.linked_run_id)
        run_map[r.id] = run

        sample_rows = session.exec(
            select(BurstSample)
            .where(BurstSample.report_id == r.id)
            .order_by(BurstSample.id.asc())
        ).all()

        values = [float(getattr(s, "actual_burst_psi", 0.0) or 0.0) for s in sample_rows]
        burst_map[r.id] = max(values) if values else float(getattr(r, "actual_burst_psi", 0.0) or 0.0)

        result_value = ""
        for s in sample_rows:
            tv = (getattr(s, "test_result", "") or "").strip().upper()
            if tv == "FAIL":
                result_value = "FAIL"
                break
            elif tv == "PASS" and result_value != "FAIL":
                result_value = "PASS"

        result_map[r.id] = result_value

    return templates.TemplateResponse(
        "burst_dashboard.html",
        {
            "request": request,
            "user": user,
            "reports": reports,
            "run_map": run_map,
            "burst_map": burst_map,
            "result_map": result_map,
        },
    )
    
from sqlalchemy import func  # make sure this import exists at top

from sqlalchemy import func  # make sure this import exists near your imports

@app.get("/burst/new", response_class=HTMLResponse)
def burst_new(request: Request, session: Session = Depends(get_session)):
    # Use cookie auth (your project uses get_current_user)
    user = get_current_user(request, session)

    # Show runs that have a batch number (works for your current data)
    runs = session.exec(
        select(ProductionRun)
        .where(ProductionRun.process == "COVER")
        .where(ProductionRun.dhtp_batch_no != "")
        .order_by(ProductionRun.id.desc())
    ).all()

    # IMPORTANT:
    # Some templates use "runs", others use "cover_runs".
    # Send BOTH so the dropdown never becomes empty بسبب اختلاف الاسم.
    return templates.TemplateResponse(
        "burst_new.html",
        {
            "request": request,
            "user": user,
            "runs": runs,
            "cover_runs": runs,
        },
    )
@app.post("/burst/create")
def burst_create(
    request: Request,
    mode: str = Form(...),
    linked_run_id: Optional[int] = Form(None),
    purpose: Optional[str] = Form(None),
    total_samples: Optional[int] = Form(None),
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)

    if mode == "linked" and not linked_run_id:
        raise HTTPException(status_code=400, detail="Production run is required.")

    # minimal insert only
    report = BurstTestRecord(
        batch_no="",
        linked_run_id=linked_run_id if mode == "linked" else None,
        purpose=purpose,
        total_samples=total_samples or 1,
        created_by=user.id,
        created_at=datetime.utcnow(),
    )

    session.add(report)
    session.commit()
    session.refresh(report)

    return RedirectResponse(
        url=f"/burst/{report.id}",
        status_code=303
    )


from fastapi import UploadFile, File, Form
from starlette.responses import RedirectResponse
import uuid, os, shutil

BURST_UPLOAD_DIR = "/app/uploads/burst"  # create if not exist

def _ensure_dir(p: str):
    os.makedirs(p, exist_ok=True)

def _burst_meta_path(abs_path: str) -> str:
    return f"{abs_path}.json"


def _save_burst_image_meta(abs_path: str, meta_json: str | None):
    if not abs_path:
        return

    meta = {
        "zoom": 1.0,
        "offset_x": 0.0,
        "offset_y": 0.0,
        "rotation": 0,
    }

    try:
        raw = json.loads(meta_json or "{}")
        if isinstance(raw, dict):
            meta["zoom"] = max(0.2, min(4.0, float(raw.get("zoom", 1.0) or 1.0)))
            meta["offset_x"] = max(-1.5, min(1.5, float(raw.get("offset_x", 0.0) or 0.0)))
            meta["offset_y"] = max(-1.5, min(1.5, float(raw.get("offset_y", 0.0) or 0.0)))

            rotation = int(raw.get("rotation", 0) or 0)
            rotation = rotation % 360
            if rotation not in (0, 90, 180, 270):
                rotation = 0
            meta["rotation"] = rotation
    except Exception:
        pass

    with open(_burst_meta_path(abs_path), "w", encoding="utf-8") as f:
        json.dump(meta, f)


def _load_burst_image_meta(abs_path: str) -> dict:
    meta_path = _burst_meta_path(abs_path)
    if not abs_path or not os.path.exists(meta_path):
        return {"zoom": 1.0, "offset_x": 0.0, "offset_y": 0.0, "rotation": 0}

    try:
        with open(meta_path, "r", encoding="utf-8") as f:
            raw = json.load(f)

        rotation = int(raw.get("rotation", 0) or 0) % 360
        if rotation not in (0, 90, 180, 270):
            rotation = 0

        return {
            "zoom": max(0.2, min(4.0, float(raw.get("zoom", 1.0) or 1.0))),
            "offset_x": max(-1.5, min(1.5, float(raw.get("offset_x", 0.0) or 0.0))),
            "offset_y": max(-1.5, min(1.5, float(raw.get("offset_y", 0.0) or 0.0))),
            "rotation": rotation,
        }
    except Exception:
        return {"zoom": 1.0, "offset_x": 0.0, "offset_y": 0.0, "rotation": 0}


def _render_burst_image_to_box(real_path: str, box_w: float, box_h: float, allow_rotate: bool = False, fill_box: bool = False):
    from PIL import Image, ImageOps

    img = Image.open(real_path)
    img = ImageOps.exif_transpose(img)

    meta = _load_burst_image_meta(real_path)
    rotation = int(meta.get("rotation", 0) or 0)
    zoom = float(meta.get("zoom", 1.0) or 1.0)
    offset_x = float(meta.get("offset_x", 0.0) or 0.0)
    offset_y = float(meta.get("offset_y", 0.0) or 0.0)

    if rotation in (90, 180, 270):
        img = img.rotate(-rotation, expand=True)
    elif allow_rotate and img.height > img.width:
        img = img.rotate(90, expand=True)

    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")

    iw, ih = img.size
    resample = getattr(Image, "Resampling", Image).LANCZOS

    target_px_w = max(400, int(box_w * 4))
    target_px_h = max(300, int(box_h * 4))

    canvas_img = Image.new("RGB", (target_px_w, target_px_h), "white")

    base_scale = min(target_px_w / float(iw), target_px_h / float(ih))
    final_scale = base_scale * zoom

    draw_w = max(1, int(iw * final_scale))
    draw_h = max(1, int(ih * final_scale))

    resized = img.resize((draw_w, draw_h), resample)

    max_shift_x = max(0, (draw_w - target_px_w) // 2)
    max_shift_y = max(0, (draw_h - target_px_h) // 2)

    shift_x = int(offset_x * max_shift_x)
    shift_y = int(-offset_y * max_shift_y)

    px = (target_px_w - draw_w) // 2 + shift_x
    py = (target_px_h - draw_h) // 2 + shift_y

    canvas_img.paste(resized, (px, py))

    out = BytesIO()
    canvas_img.save(out, format="JPEG", quality=92)
    out.seek(0)
    return out

    
    

@app.post("/burst/{report_id}/sample/{sample_id}/upload_attachments")
async def burst_upload_attachments(
    report_id: int,
    sample_id: int,
    request: Request,
    photo_full: UploadFile = File(None),
    photo_a: UploadFile = File(None),
    photo_b: UploadFile = File(None),
    chart: UploadFile = File(None),
    crop_photo_full: str = Form(""),
    crop_photo_a: str = Form(""),
    crop_photo_b: str = Form(""),
    crop_chart: str = Form(""),
    session: Session = Depends(get_session),
):
    rep = session.get(BurstTestReport, report_id)
    if not rep:
        raise HTTPException(404, "Burst report not found")

    sample = session.get(BurstSample, sample_id)
    if not sample or sample.report_id != report_id:
        raise HTTPException(404, "Burst sample not found")

    upload_dir = os.path.join(BURST_FILES_DIR, str(report_id), str(sample_id))
    os.makedirs(upload_dir, exist_ok=True)

    def _save_one(file: UploadFile, kind: str, crop_json: str = ""):
        att = session.exec(
            select(BurstAttachment).where(
                (BurstAttachment.report_id == report_id) &
                (BurstAttachment.sample_id == sample_id) &
                (BurstAttachment.kind == kind)
            )
        ).first()

        abs_path = ""

        if file and getattr(file, "filename", ""):
            ext = (os.path.splitext(file.filename)[1] or ".jpg").lower()
            if ext not in [".png", ".jpg", ".jpeg", ".webp"]:
                return None

            safe_name = f"{kind}{ext}"
            abs_path = os.path.join(upload_dir, safe_name)

            data = file.file.read()
            with open(abs_path, "wb") as f:
                f.write(data)

            rel_path = os.path.relpath(abs_path, BASE_DIR)

            if not att:
                att = BurstAttachment(
                    report_id=report_id,
                    sample_id=sample_id,
                    kind=kind,
                    caption=kind,
                    file_path=rel_path,
                )
            else:
                att.report_id = report_id
                att.sample_id = sample_id
                att.kind = kind
                att.caption = kind
                att.file_path = rel_path

            session.add(att)

        elif att and getattr(att, "file_path", ""):
            abs_path = resolve_burst_file_path(att.file_path or "")

        if abs_path and os.path.exists(abs_path):
            _save_burst_image_meta(abs_path, crop_json)

        return abs_path

    _save_one(photo_full, "PHOTO_FULL", crop_photo_full)
    _save_one(photo_a, "PHOTO_A", crop_photo_a)
    _save_one(photo_b, "PHOTO_B", crop_photo_b)
    _save_one(chart, "CHART", crop_chart)

    session.commit()
    return RedirectResponse(url=f"/burst/{report_id}", status_code=303)

@app.get("/burst/{report_id}")
def burst_view(report_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)

    rep = session.get(BurstTestReport, report_id)
    if not rep:
        raise HTTPException(404, "Burst report not found")

    # Always keep 5 sample rows available (UI can show 1/2/5 without extra DB clicks)
    ensure_burst_samples(session, report_id, desired=5)

    run = None
    produced_len = 0.0
    if (not rep.is_unlinked) and rep.linked_run_id:
        run = session.get(ProductionRun, rep.linked_run_id)
        if run:
            produced_len = get_run_produced_length_m(session, run.id)

    # Base values shown on page
    live = {
        "client_name": rep.client_name,
        "client_po": rep.client_po,
        "pipe_specification": rep.pipe_specification,
        "liner_material_grade": rep.liner_material_grade,
        "reinforcement_material_grade": rep.reinforcement_material_grade,
        "cover_material_grade": rep.cover_material_grade,
        "total_length_m": rep.total_length_m,
    }

    # If linked + not locked => show latest values from cover run / batch runs
    if (not rep.is_unlinked) and (not rep.is_locked) and run:
        live["client_name"] = (getattr(run, "client_name", "") or "").strip()
        live["client_po"] = (getattr(run, "po_number", "") or "").strip()
        live["pipe_specification"] = (getattr(run, "pipe_specification", "") or "").strip()
        live["cover_material_grade"] = (getattr(run, "raw_material_spec", "") or "").strip()

        run_total = float(getattr(run, "total_length_m", 0.0) or 0.0)
        if run_total > 0:
            live["total_length_m"] = run_total

        batch_no = (getattr(run, "dhtp_batch_no", "") or "").strip()
        if batch_no:
            liner_run = session.exec(
                select(ProductionRun).where(
                    (ProductionRun.dhtp_batch_no == batch_no) & (ProductionRun.process == "LINER")
                )
            ).first()
            reinf_run = session.exec(
                select(ProductionRun).where(
                    (ProductionRun.dhtp_batch_no == batch_no) & (ProductionRun.process == "REINFORCEMENT")
                )
            ).first()

            if liner_run:
                live["liner_material_grade"] = (getattr(liner_run, "raw_material_spec", "") or "").strip()
            if reinf_run:
                live["reinforcement_material_grade"] = (getattr(reinf_run, "raw_material_spec", "") or "").strip()


    # -----------------------------
    # Attachments for template
    # -----------------------------
    samples = session.exec(
        select(BurstSample).where(BurstSample.report_id == rep.id)
    ).all()

    audit = session.exec(
        select(BurstAuditLog)
        .where(BurstAuditLog.report_id == report_id)
        .order_by(BurstAuditLog.id.desc())
    ).all()

    att_rows = session.exec(
        select(BurstAttachment).where(BurstAttachment.report_id == rep.id)
    ).all()


    pending_revisions = session.exec(
        select(BurstReportRevision)
        .where(BurstReportRevision.report_id == report_id)
        .where(BurstReportRevision.status == "PENDING")
        .order_by(BurstReportRevision.id.desc())
    ).all()

    edit_direct_allowed = (
        not rep.published_at
        or not rep.edit_window_until
        or datetime.utcnow() <= rep.edit_window_until
        or _is_manager_or_boss(user)
    )

    att_status = {}
    att_map = defaultdict(dict)
    crop_meta = defaultdict(dict)

    for a in att_rows:
        sid = getattr(a, "sample_id", None)
        kind = (getattr(a, "kind", "") or "").strip().upper()
        if sid is None or not kind:
            continue



        att_status[(sid, kind)] = True
        att_map[sid][kind] = a

        real_path = resolve_burst_file_path(getattr(a, "file_path", "") or "")
        if real_path:
            crop_meta[sid][kind] = _load_burst_image_meta(real_path)
    

    return templates.TemplateResponse(
        "burst_view.html",
        {
            "request": request,
            "user": user,
            "rep": rep,
            "run": run,
            "produced_len": produced_len,
            "live": live,
            "samples": samples,
            "audit": audit,
            "att_status": att_status,
            "att_map": att_map,
            "crop_meta": crop_meta,
            "pending_revisions": pending_revisions,
            "edit_direct_allowed": edit_direct_allowed,
        },
    )


@app.post("/burst/{report_id}/update")
async def burst_update(
    report_id: int,
    request: Request,
    session: Session = Depends(get_session),
    

    # TEST DETAILS
    reference_standard: str = Form(""),
    reference_dhtp_procedure: str = Form(""),
    system_max_pressure: str = Form(""),
    laboratory_temperature: str = Form(""),
    testing_medium: str = Form("Water"),
    total_no_of_specimens: int = Form(1),

    # SPECIMENS (report-level)
    effective_length_m: str = Form(""),
    liner_thickness: str = Form(""),
    reinforcement_thickness: str = Form(""),
    cover_thickness: str = Form(""),

    # legacy single-sample fields (keep for backward compatibility)
    sample_serial_number: str = Form(""),
    actual_burst_MPa: float = Form(0.0),
    pressurization_time_s: str = Form(""),
    test_result: str = Form(""),
    failure_mode: str = Form(""),
    notes: str = Form(""),

    # signatures
    qa_qc_officer_name: str = Form(""),
    testing_operator_name: str = Form(""),

    # NEW: control behavior from UI buttons (default: save only)
    action: str = Form("save"),  # "save" or "lock"
):
    """
    Saves burst report + per-sample fields.
    - Ensures 5 BurstSample rows exist for stable UI (templates: 1/2/5)
    - Saves per-sample values for the FIRST N samples (N = total_no_of_specimens)
    - Locks only if action == "lock"
    """
    user = get_current_user(request, session)

    rep = session.get(BurstTestReport, report_id)
    if not rep:
        raise HTTPException(404, detail="Burst report not found")

    form = await request.form()

    # If already locked and user tries to save, allow saving only if you want.
    # Current behavior: locked reports can still be updated only via reopen endpoint.
    if rep.is_locked and (action or "").lower() != "lock":
        raise HTTPException(400, detail="Report is locked. Reopen it to edit.")

    try:
        n = int(total_no_of_specimens or 1)
    except Exception:
        n = 1
    if n < 1: n = 1
    if n > 50: n = 50

    # TEST DETAILS
        # TEST DETAILS
    rep.reference_standard = (form.get("reference_standard") or "").strip()
    rep.reference_dhtp_procedure = (form.get("reference_dhtp_procedure") or "").strip()
    rep.system_max_pressure = (form.get("system_max_pressure") or "").strip()
    rep.laboratory_temperature = (form.get("laboratory_temperature") or "").strip()
    rep.testing_medium = (form.get("testing_medium") or "").strip()
    rep.total_no_of_specimens = int(form.get("total_no_of_specimens") or 1)

    form_pipe_spec = (form.get("pipe_specification") or "").strip()
    form_client_po = (form.get("client_po") or "").strip()

    if form_pipe_spec:
        rep.pipe_specification = form_pipe_spec
    if form_client_po:
        rep.client_po = form_client_po

    if (not rep.is_unlinked) and rep.linked_run_id:
        linked_run = session.get(ProductionRun, rep.linked_run_id)
        if linked_run:
            if not rep.pipe_specification:
                rep.pipe_specification = (getattr(linked_run, "pipe_specification", "") or "").strip()
            if not rep.client_po:
                rep.client_po = (getattr(linked_run, "po_number", "") or "").strip()
            if not rep.client_name:
                rep.client_name = (getattr(linked_run, "client_name", "") or "").strip()

    # SPECIMENS (report-level)
    rep.effective_length_m = (form.get("effective_length_m") or "").strip()
    rep.liner_thickness = (form.get("liner_thickness") or "").strip()
    rep.reinforcement_thickness = (form.get("reinforcement_thickness") or "").strip()
    rep.cover_thickness = (form.get("cover_thickness") or "").strip()

    # legacy single-sample fields
    rep.sample_serial_number = (form.get("sample_serial_number") or "").strip()
    rep.actual_burst_psi = float(form.get("actual_burst_MPa") or form.get("actual_burst_psi") or 0.0)
    rep.pressurization_time_s = (form.get("pressurization_time_s") or "").strip()
    rep.test_result = (form.get("test_result") or "").strip()
    rep.failure_mode = (form.get("failure_mode") or "").strip()
    rep.notes = (form.get("notes") or "").strip()

    rep.qa_qc_officer_name = (form.get("qa_qc_officer_name") or "").strip()

    # ✅ get the full form FIRST (so we can read technician_name and per-sample fields)
    form = await request.form()
    
    # ✅ now this works
    rep.technician_name = (form.get("technician_name") or "").strip()

    # ---------------------------
    # Per-sample save (BurstSample rows)
    # ---------------------------
    form = await request.form()

    db_samples = ensure_burst_samples(session, report_id, desired=n)

    for s in db_samples[:n]:

        s.sample_serial_number = (form.get(f"sample_serial_number_{s.id}") or "").strip()
        s.actual_burst_psi = float(form.get(f"actual_burst_MPa_{s.id}") or form.get(f"actual_burst_psi_{s.id}") or 0.0)
        s.pressurization_time_s = (form.get(f"pressurization_time_s_{s.id}") or "").strip()
        s.test_result = (form.get(f"test_result_{s.id}") or "").strip()

        s.liner_material_grade = (form.get(f"liner_material_grade_{s.id}") or "").strip()
        s.liner_thickness_mm = float(form.get(f"liner_thickness_mm_{s.id}") or 0.0)
        
        s.reinforcement_material_grade = (form.get(f"reinforcement_material_grade_{s.id}") or "").strip()
        s.reinforcement_thickness_mm = float(form.get(f"reinforcement_thickness_mm_{s.id}") or 0.0)
        
        s.cover_material_grade = (form.get(f"cover_material_grade_{s.id}") or "").strip()
        s.cover_thickness_mm = float(form.get(f"cover_thickness_mm_{s.id}") or 0.0)

        session.add(s)

    # ---------------------------
    # Locking
    # ---------------------------
    action = (action or "save").strip().lower()
    if action == "lock":
        rep.is_locked = True
        rep.locked_at = datetime.utcnow()
        rep.locked_by_user_id = user.id
        rep.locked_by_user_name = user.display_name

        session.add(BurstAuditLog(
            report_id=report_id,
            action="LOCK",
            note="Saved and locked report",
            user_id=user.id,
            user_name=user.display_name,
        ))
    else:
        session.add(BurstAuditLog(
            report_id=report_id,
            action="EDIT",
            note="Saved report (not locked)",
            user_id=user.id,
            user_name=user.display_name,
        ))


        now_utc = datetime.utcnow()

    is_complete = _burst_report_is_complete(rep, db_samples[:n])

    if is_complete and not rep.published_at:
        rep.is_complete = True
        rep.published_at = now_utc
        rep.edit_window_until = now_utc + timedelta(hours=24)

    direct_edit_allowed = (
        not rep.published_at
        or not rep.edit_window_until
        or now_utc <= rep.edit_window_until
        or _is_manager_or_boss(user)
    )

    if not direct_edit_allowed:
        payload = _burst_snapshot(rep, db_samples[:n])

        rev = BurstReportRevision(
            report_id=report_id,
            status="PENDING",
            payload_json=json.dumps(payload),
            submitted_by_user_id=user.id,
            submitted_by_user_name=user.display_name,
        )
        session.add(rev)

        rep.current_revision_status = "PENDING_APPROVAL"

        session.add(BurstAuditLog(
            report_id=report_id,
            action="REVISION_REQUEST",
            note="Saved as pending revision after 24h lock window",
            user_id=user.id,
            user_name=user.display_name,
        ))

        session.add(rep)
        session.commit()
        return RedirectResponse(f"/burst/{report_id}?msg=pending_approval", status_code=303)

    session.add(rep)
    session.commit()

    return RedirectResponse(f"/burst/{report_id}", status_code=303)



# ------------------------------------------------------------
# Burst: AJAX endpoint to update specimen count without page refresh
# ------------------------------------------------------------
from fastapi import Request, HTTPException
from starlette.responses import RedirectResponse, JSONResponse

@app.post("/burst/{report_id}/set_specimen_count")
async def set_specimen_count(report_id: int, request: Request, session: Session = Depends(get_session)):
    form = dict(await request.form())

    try:
        n = int(form.get("total_no_of_specimens") or 1)
    except Exception:
        n = 1
    if n < 1:
        n = 1
    if n > 50:
        n = 50

    rep = session.get(BurstTestReport, report_id)
    if not rep:
        raise HTTPException(404, "Burst report not found")

    rep.total_no_of_specimens = n
    session.add(rep)
    session.commit()

    ensure_burst_samples(session, report_id, desired=n)

    accept = (request.headers.get("accept") or "").lower()
    if "application/json" in accept:
        return JSONResponse({"ok": True, "total_no_of_specimens": n})

    return RedirectResponse(url=f"/burst/{report_id}", status_code=303)

@app.get("/burst/files/{attachment_id}/view")
def burst_file_view(attachment_id: int, session: Session = Depends(get_session)):
    att = session.get(BurstAttachment, attachment_id)
    if not att:
        raise HTTPException(404, "Attachment not found")

    real_path = resolve_burst_file_path(att.file_path or "")
    if not real_path or not os.path.exists(real_path):
        raise HTTPException(404, f"Attachment file not found: {att.file_path}")

    return FileResponse(real_path)


@app.post("/burst/{report_id}/samples/add")
async def burst_add_sample(
    report_id: int,
    request: Request,
    session: Session = Depends(get_session),
    sample_start_m: float = Form(0.0),
    sample_length_m: float = Form(0.0),
):
    user = get_current_user(request, session)

    rep = session.get(BurstTestReport, report_id)
    if not rep:
        raise HTTPException(404, detail="Burst report not found")
    if rep.is_locked:
        raise HTTPException(400, detail="Report is locked. Reopen it to add samples.")

    start = float(sample_start_m or 0.0)
    length = float(sample_length_m or 0.0)
    end = start + length

    total_len = float(rep.total_length_m or 0.0)
    if total_len <= 0:
        raise HTTPException(400, detail="Total length is 0. Update run length first.")
    if start < 0 or length <= 0:
        raise HTTPException(400, detail="Invalid sample. Example: start 300m, length 3m")
    if end > total_len:
        raise HTTPException(400, detail=f"Sample ends at {end}m but total length is {total_len}m")

    s = BurstSample(report_id=report_id, sample_start_m=start, sample_length_m=length)
    session.add(s)

    session.add(BurstAuditLog(
        report_id=report_id,
        action="EDIT",
        note=f"Added sample {start}→{end}m",
        user_id=user.id,
        user_name=user.display_name,
    ))

    session.commit()
    return RedirectResponse(f"/burst/{report_id}", status_code=303)


@app.post("/burst/{report_id}/reopen")
async def burst_reopen(
    report_id: int,
    request: Request,
    session: Session = Depends(get_session),
    reason: str = Form(""),
):
    user = get_current_user(request, session)

    rep = session.get(BurstTestReport, report_id)
    if not rep:
        raise HTTPException(404, detail="Burst report not found")

    rep.is_locked = False

    session.add(BurstAuditLog(
        report_id=report_id,
        action="REOPEN",
        note=(reason or "").strip() or "Reopened for edit",
        user_id=user.id,
        user_name=user.display_name,
    ))

    session.add(rep)
    session.commit()

    return RedirectResponse(f"/burst/{report_id}", status_code=303)


PAGE_W, PAGE_H = A4  # 595 x 842


def _txt(x) -> str:
    return "" if x is None else str(x)

def _float(x):
    try:
        if x in ("", None):
            return None
        return float(x)
    except Exception:
        return None

def _logo_path() -> str:
    base_dir = os.path.dirname(__file__)
    # Put your logo here: app/static/images/logo.png
    return os.path.join(base_dir, "static", "images", "logo.png")

def _draw_header_footer(c, *, title: str, doc_control_no: str, page_num: int, page_total: int):
    w, h = A4

    logo = _logo_path()
    y_top = h - 12 * mm
    y_title = h - 26 * mm

    if os.path.exists(logo):
        try:
            img = ImageReader(logo)
            iw, ih = img.getSize()
            target_w = 55 * mm
            scale = target_w / float(iw)
            target_h = float(ih) * scale
            x = (w - target_w) / 2
            c.drawImage(img, x, y_top - target_h, width=target_w, height=target_h, mask="auto")
            y_title = (y_top - target_h) - 7 * mm
        except Exception:
            pass

    c.setFont("Helvetica-Bold", 16)
    c.drawCentredString(w / 2, y_title, title)

    line_y = y_title - 4 * mm
    c.setStrokeColor(colors.black)
    c.line(20 * mm, line_y, w - 20 * mm, line_y)

    # Footer
    c.setFont("Helvetica", 9)
    c.setFillColor(colors.grey)
    c.drawString(20 * mm, 12 * mm, doc_control_no or "")
    c.drawRightString(w - 20 * mm, 12 * mm, f"Page {page_num}/{page_total}")
    c.setFillColor(colors.black)

    return line_y - 8 * mm

def _draw_report_info_table(c, report, x, y):
    n = int(getattr(report, "total_no_of_specimens", 0) or 0)
    data = [
        ["Batch No", _txt(getattr(report, "batch_no", "")), "Client", _txt(getattr(report, "client_name", ""))],
        ["Client PO", _txt(getattr(report, "client_po", "")), "Pipe Spec", _txt(getattr(report, "pipe_specification", ""))],
        ["Test Medium", _txt(getattr(report, "testing_medium", "")), "Lab Temp", _txt(getattr(report, "laboratory_temperature", ""))],
        ["Standard", _txt(getattr(report, "reference_standard", "")), "Procedure", _txt(getattr(report, "reference_dhtp_procedure", ""))],
        ["System Max (MPa)", _txt(getattr(report, "system_max_pressure", "")), "Specimens", _txt(n)],
        ["Test Date", _txt(getattr(report, "tested_at", "") or getattr(report, "created_at", "")), "", ""],
    ]

    tbl = Table(data, colWidths=[28*mm, 62*mm, 28*mm, 62*mm])
    tbl.setStyle(TableStyle([
        ("FONT", (0,0), (-1,-1), "Helvetica", 9),
        ("FONT", (0,0), (0,-1), "Helvetica-Bold", 9),
        ("FONT", (2,0), (2,-1), "Helvetica-Bold", 9),
        ("TOPPADDING", (0,0), (-1,-1), 2),
        ("BOTTOMPADDING", (0,0), (-1,-1), 2),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
    ]))

    tbl.wrapOn(c, 0, 0)
    tbl_h = len(data) * 7 * mm
    tbl.drawOn(c, x, y - tbl_h)
    return y - tbl_h
def _draw_specimen_blocks(c, report, samples, start_y):
    w, h = A4
    y = start_y

    for idx, s in enumerate(samples, start=1):
        block_h = 38 * mm  # slightly taller to avoid overflow
        x0 = 20 * mm
        box_w = w - 40 * mm

        # Page break if needed
        if y - block_h < 55 * mm:
            c.showPage()
            y = h - 32 * mm

        c.setStrokeColor(colors.black)
        c.rect(x0, y - block_h, box_w, block_h, stroke=1, fill=0)

        c.setFont("Helvetica-Bold", 10)
        c.drawString(x0 + 3*mm, y - 6*mm, f"Specimen #{idx}")

        c.setFont("Helvetica", 9)

        # keep everything INSIDE box
        line1 = y - 14*mm
        line2 = y - 22*mm
        line3 = y - 28*mm
        line4 = y - 34*mm  # last line still inside now (because block_h increased)

        c.drawString(x0 + 3*mm,  line1, f"Serial No: {_txt(getattr(s,'sample_serial_number',''))}")
        c.drawString(x0 + 95*mm, line1, f"Length (mm): {_txt(getattr(s,'sample_length_m',''))}")

        c.drawString(x0 + 3*mm, line2,
                     f"Liner: {_txt(getattr(s,'liner_material_grade',''))} | Thk(mm): {_txt(getattr(s,'liner_thickness_mm',''))}")
        c.drawString(x0 + 3*mm, line3,
                     f"Reinf: {_txt(getattr(s,'reinforcement_material_grade',''))} | Thk(mm): {_txt(getattr(s,'reinforcement_thickness_mm',''))}")
        c.drawString(x0 + 3*mm, line4,
                     f"Cover: {_txt(getattr(s,'cover_material_grade',''))} | Thk(mm): {_txt(getattr(s,'cover_thickness_mm',''))}")

        y = y - block_h - 6*mm

    return y
    
def _draw_results_table(c, samples, y):
    w, h = A4

    data = [["#", "Serial No", "Actual Burst (MPa)", "Pressurization Time (s)", "Result"]]
    for i, s in enumerate(samples, start=1):
        data.append([
            str(i),
            _txt(getattr(s, "sample_serial_number", "")),
            _txt(getattr(s, "actual_burst_MPa", "")),
            _txt(getattr(s, "pressurization_time_s", "")),
            _txt(getattr(s, "test_result", "")),
        ])

    tbl = Table(data, colWidths=[10*mm, 45*mm, 40*mm, 50*mm, 25*mm])
    tbl.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.5, colors.black),
        ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
        ("FONT", (0,0), (-1,0), "Helvetica-Bold"),
        ("FONT", (0,1), (-1,-1), "Helvetica"),
        ("FONTSIZE", (0,0), (-1,-1), 8),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
    ]))

    # page break if needed
    approx_h = (len(data) + 1) * 8 * mm
    if y - approx_h < 25 * mm:
        c.showPage()
        y = h - 30 * mm

    tbl.wrapOn(c, w, h)
    tbl.drawOn(c, 20*mm, y - approx_h)
    return y - approx_h - 6*mm

def _make_burst_chart_png(samples) -> bytes | None:
    """
    Returns PNG bytes for chart, or None if no valid burst values.
    """
    try:
        xs = []
        ys = []
        for i, s in enumerate(samples, start=1):
            v = _float(getattr(s, "actual_burst_MPa", None))
            if v is not None:
                xs.append(i)
                ys.append(v)
        if not ys:
            return None

        fig = plt.figure()
        ax = fig.add_subplot(111)
        ax.plot(xs, ys, marker="o")
        ax.set_title("Burst Pressure by Specimen")
        ax.set_xlabel("Specimen #")
        ax.set_ylabel("Actual Burst (MPa)")
        ax.grid(True)

        buf = BytesIO()
        fig.savefig(buf, format="png", dpi=160, bbox_inches="tight")
        plt.close(fig)
        buf.seek(0)
        return buf.getvalue()
    except Exception:
        return None


@app.post("/burst/revision/{revision_id}/approve")
def burst_revision_approve(
    revision_id: int,
    request: Request,
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)
    require_manager(user)

    rev = session.get(BurstReportRevision, revision_id)
    if not rev:
        raise HTTPException(404, "Revision not found")
    if rev.status != "PENDING":
        raise HTTPException(400, "Revision already processed")

    rep = session.get(BurstTestReport, rev.report_id)
    if not rep:
        raise HTTPException(404, "Burst report not found")

    samples = ensure_burst_samples(session, rep.id, desired=rep.total_no_of_specimens or 1)

    payload = json.loads(rev.payload_json or "{}")
    _apply_burst_snapshot(rep, samples, payload)

    rep.current_revision_status = "LIVE"

    rev.status = "APPROVED"
    rev.reviewed_by_user_id = user.id
    rev.reviewed_by_user_name = user.display_name
    rev.reviewed_at = datetime.utcnow()

    session.add(rep)
    session.add(rev)
    session.add(BurstAuditLog(
        report_id=rep.id,
        action="REVISION_APPROVED",
        note=f"Approved revision #{rev.id}",
        user_id=user.id,
        user_name=user.display_name,
    ))
    session.commit()

    return RedirectResponse(f"/burst/{rep.id}?msg=revision_approved", status_code=303)


@app.post("/burst/revision/{revision_id}/reject")
async def burst_revision_reject(
    revision_id: int,
    request: Request,
    session: Session = Depends(get_session),
    review_comment: str = Form(""),
):
    user = get_current_user(request, session)
    require_manager(user)

    rev = session.get(BurstReportRevision, revision_id)
    if not rev:
        raise HTTPException(404, "Revision not found")
    if rev.status != "PENDING":
        raise HTTPException(400, "Revision already processed")

    rep = session.get(BurstTestReport, rev.report_id)
    if not rep:
        raise HTTPException(404, "Burst report not found")

    rev.status = "REJECTED"
    rev.reviewed_by_user_id = user.id
    rev.reviewed_by_user_name = user.display_name
    rev.reviewed_at = datetime.utcnow()
    rev.review_comment = (review_comment or "").strip()

    rep.current_revision_status = "LIVE"

    session.add(rep)
    session.add(rev)
    session.add(BurstAuditLog(
        report_id=rep.id,
        action="REVISION_REJECTED",
        note=f"Rejected revision #{rev.id}",
        user_id=user.id,
        user_name=user.display_name,
    ))
    session.commit()

    return RedirectResponse(f"/burst/{rep.id}?msg=revision_rejected", status_code=303)

@app.get("/burst/{report_id}/pdf")
def burst_pdf_download(report_id: int, session: Session = Depends(get_session)):
    report = session.get(BurstTestReport, report_id)
    if (not report.is_unlinked) and report.linked_run_id:
        linked_run = session.get(ProductionRun, report.linked_run_id)
        if linked_run:
            if not (report.client_name or "").strip():
                report.client_name = (getattr(linked_run, "client_name", "") or "").strip()
            if not (report.client_po or "").strip():
                report.client_po = (getattr(linked_run, "po_number", "") or "").strip()
            if not (report.pipe_specification or "").strip():
                report.pipe_specification = (getattr(linked_run, "pipe_specification", "") or "").strip()

    samples = session.exec(
        select(BurstSample)
        .where(BurstSample.report_id == report.id)
        .order_by(BurstSample.id.asc())
    ).all()

    specimen_count = int(getattr(report, "total_no_of_specimens", None) or len(samples) or 1)
    specimen_count = max(1, min(specimen_count, 50))
    samples = samples[:specimen_count] if samples else []

    atts = session.exec(
        select(BurstAttachment)
        .where(BurstAttachment.report_id == report.id)
        .order_by(BurstAttachment.sample_id.asc(), BurstAttachment.uploaded_at.asc())
    ).all()

    doc_no = _txt(getattr(report, "doc_control_no", "")) or "QAW1401-F01"
    pages: list[bytes] = []

    def new_page(title: str = "Short-Time Hydrostatic Burst Pressure Test Report"):
        buf = BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        return buf, c, title

    def _draw_boxed_image(c, path: str, x: float, y_top: float, box_w: float, box_h: float, label: str):
        c.setFont("Helvetica-Bold", 10)
        c.drawString(x, y_top, label)

        img_top = y_top - 4 * mm
        c.rect(x, img_top - box_h, box_w, box_h, stroke=1, fill=0)

        if not path or not os.path.exists(path):
            c.setFont("Helvetica", 9)
            c.drawString(x + 4 * mm, img_top - 10 * mm, "Not uploaded")
            return

        try:
            img = ImageReader(path)
            iw, ih = img.getSize()
            scale = min(box_w / float(iw), box_h / float(ih))
            dw, dh = float(iw) * scale, float(ih) * scale
            dx = x + (box_w - dw) / 2
            dy = img_top - dh - (box_h - dh) / 2
            c.drawImage(img, dx, dy, width=dw, height=dh, mask="auto")
        except Exception:
            c.setFont("Helvetica", 9)
            c.drawString(x + 4 * mm, img_top - 10 * mm, "Could not load image")

    def _stamp_page_numbers(pdf_bytes: bytes) -> bytes:
        reader = PdfReader(BytesIO(pdf_bytes))
        writer = PdfWriter()
        total_pages = len(reader.pages)

        for i, page in enumerate(reader.pages, start=1):
            w0 = float(page.mediabox.width)
            h0 = float(page.mediabox.height)

            stamp_buf = BytesIO()
            sc = canvas.Canvas(stamp_buf, pagesize=(w0, h0))

            sc.setFillColor(colors.white)
            sc.rect(0, 0, w0, 20 * mm, stroke=0, fill=1)

            sc.setFont("Helvetica", 9)
            sc.setFillColor(colors.grey)
            sc.drawString(20 * mm, 12 * mm, doc_no)
            sc.drawRightString(w0 - 20 * mm, 12 * mm, f"Page {i}/{total_pages}")
            sc.showPage()
            sc.save()
            stamp_buf.seek(0)

            stamp_reader = PdfReader(stamp_buf)
            page.merge_page(stamp_reader.pages[0])
            writer.add_page(page)

        out = BytesIO()
        writer.write(out)
        return out.getvalue()

    att_map = {}
    for a in atts:
        sample_id = getattr(a, "sample_id", None)
        kind = (getattr(a, "kind", "") or "").strip().upper()
        if sample_id is None or not kind:
            continue
        att_map[(sample_id, kind)] = a

    def _get_att_path(sample_id: int | None, *kinds: str) -> str:
        if sample_id is None:
            return ""
        for kind in kinds:
            att = att_map.get((sample_id, kind))
            if not att:
                continue
            raw_path = getattr(att, "file_path", "") or ""
            p = resolve_burst_file_path(raw_path)
            if p and os.path.exists(p):
                return p
        return ""

    # ------------------------------------------------------------
    # PAGE 1: report contents
    # ------------------------------------------------------------
    buf, c, title = new_page()
    w, h = A4
    y = _draw_header_footer(c, title=title, doc_control_no=doc_no, page_num=1, page_total=1)

    c.setFont("Helvetica-Bold", 11)
    c.drawString(20 * mm, y, "Report Information")
    y -= 6 * mm
    y = _draw_report_info_table(c, report, 20 * mm, y)
    y -= 8 * mm
    
    c.setFont("Helvetica-Bold", 11)
    c.drawString(20 * mm, y, "Specimens")
    y -= 7 * mm
    y = _draw_specimen_blocks(c, report, samples, y)



    if y < 85 * mm:
        c.showPage()
        y = _draw_header_footer(c, title=title, doc_control_no=doc_no, page_num=1, page_total=1)

    c.setFont("Helvetica-Bold", 11)
    c.drawString(20 * mm, y, "Results")
    y -= 7 * mm
    y = _draw_results_table(c, samples, y)
    _draw_signatures(c, report, y)

    c.showPage()
    c.save()
    pages.append(buf.getvalue())

    # ------------------------------------------------------------
    # ------------------------------------------------------------
    # NEXT PAGES: one attachment page per specimen
    # old-style layout
    # ------------------------------------------------------------
    for idx, s in enumerate(samples, start=1):
        buf2, c2, _ = new_page()
        w2, h2 = A4

        margin_x = 12 * mm
        content_top = _draw_header_footer(
            c2,
            title="Short-Time Hydrostatic Burst Pressure Test Report",
            doc_control_no=doc_no,
            page_num=1,
            page_total=1,
        )

        serial = _txt(getattr(s, "sample_serial_number", "")) or f"Specimen #{idx}"

        length_txt = _txt(getattr(s, "sample_length_m", "")) or "-"
        burst_txt = _txt(getattr(s, "actual_burst_MPa", "")) or "-"
        result_txt = _txt(getattr(s, "test_result", "")) or "-"

        c2.setFont("Helvetica-Bold", 12)
        c2.drawString(margin_x, content_top, f"Specimen Serial No: {serial}")

        info_y = content_top - 8 * mm
        c2.setFont("Helvetica", 9)
        c2.drawString(
            margin_x,
            info_y,
            f"Length: {length_txt}   Actual Burst: {burst_txt}   Result: {result_txt}",
        )

        gap = 4 * mm
        usable_w = w2 - (2 * margin_x)

        chart_y_top = info_y - 8 * mm
        chart_h = 70 * mm

        full_y_top = chart_y_top - chart_h - 10 * mm
        full_h = 48 * mm

        bottom_y_top = full_y_top - full_h - 10 * mm
        bottom_h = 42 * mm
        half_w = (usable_w - gap) / 2

        # ----------------------------
        # Attachment lookup
        # ----------------------------
        sid = getattr(s, "id", None)

        full_path = _get_att_path(sid, "PHOTO_FULL")
        a_path = _get_att_path(sid, "PHOTO_A")
        b_path = _get_att_path(sid, "PHOTO_B")
        chart_path = _get_att_path(sid, "CHART")

        
        def draw_labeled_image(path, x, y_top, box_w, box_h, label, allow_rotate=False, fill_box=False):
            c2.setFont("Helvetica", 9)

            # Move label a bit lower so it sits just above its own box
            label_y = y_top - 4 * mm
            c2.drawString(x, label_y, label)

            # Keep a small gap between label and box
            img_top = label_y - 3 * mm
            c2.rect(x, img_top - box_h, box_w, box_h, stroke=1, fill=0)

            if not path:
                c2.drawString(x + 4 * mm, img_top - 10 * mm, "Not uploaded")
                return

            real_path = resolve_burst_file_path(path)
            if not real_path or not os.path.exists(real_path):
                c2.drawString(x + 4 * mm, img_top - 10 * mm, "Not uploaded")
                return

            try:
                temp_buf = _render_burst_image_to_box(
                    real_path,
                    box_w,
                    box_h,
                    allow_rotate=allow_rotate,
                    fill_box=fill_box,
                )
                c2.drawImage(
                    ImageReader(temp_buf),
                    x,
                    img_top - box_h,
                    width=box_w,
                    height=box_h,
                    mask="auto",
                )
            except Exception:
                c2.drawString(x + 4 * mm, img_top - 10 * mm, "Could not load image")

        draw_labeled_image(chart_path, margin_x, chart_y_top, usable_w, chart_h, "CHART", allow_rotate=False, fill_box=False)
        draw_labeled_image(full_path, margin_x, full_y_top, usable_w, full_h, "FULL LENGTH:", allow_rotate=True, fill_box=True)
        draw_labeled_image(a_path, margin_x, bottom_y_top, half_w, bottom_h, "SIDE A:", allow_rotate=True, fill_box=True)
        draw_labeled_image(b_path, margin_x + half_w + gap, bottom_y_top, half_w, bottom_h, "SIDE B:", allow_rotate=True, fill_box=True)

        c2.showPage()
        c2.save()
        pages.append(buf2.getvalue())

    merged_writer = PdfWriter()
    for pb in pages:
        reader = PdfReader(BytesIO(pb))
        for page in reader.pages:
            merged_writer.add_page(page)

    out = BytesIO()
    merged_writer.write(out)
    pdf_bytes = out.getvalue()
    pdf_bytes = _stamp_page_numbers(pdf_bytes)

    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename=\"burst_report_{report_id}.pdf\"'},
    )

def _draw_sample_images(c, x, y, w, paths):
    # paths: dict {"A": path, "B": path, "CHART": path}
    # returns updated y

    def draw_img(path, x0, y0, max_w, max_h):
        if not path or not os.path.exists(path):
            return 0
        img = ImageReader(path)
        iw, ih = img.getSize()
        scale = min(max_w / iw, max_h / ih)
        dw, dh = iw * scale, ih * scale
        c.drawImage(img, x0, y0 - dh, width=dw, height=dh, mask="auto")
        return dh

    gap = 6 * mm
    max_w = w - 40 * mm
    col_w = (max_w - gap) / 2

    # A + B side-by-side
    top_h = 45 * mm
    used_h = 0
    hA = draw_img(paths.get("A"), x, y, col_w, top_h)
    hB = draw_img(paths.get("B"), x + col_w + gap, y, col_w, top_h)
    used_h = max(hA, hB)

    if used_h > 0:
        y = y - used_h - 6*mm

    # chart full width
    chart_h = 60 * mm
    hC = draw_img(paths.get("CHART"), x, y, max_w, chart_h)
    if hC > 0:
        y = y - hC - 6*mm

    return y


# ==========================================
# EXPORT PER-INSPECTION (SEPARATE REPORTS)
# ==========================================
@app.get("/mrr/{lot_id}/inspection/id/{inspection_id}/export/xlsx")
def mrr_export_inspection_xlsx(
    lot_id: int,
    inspection_id: int,
    request: Request,
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)

    lot = session.get(MaterialLot, lot_id)
    if not lot:
        raise HTTPException(404, "MRR Ticket not found")

    insp = session.get(MrrReceivingInspection, inspection_id)
    if not insp or insp.ticket_id != lot_id:
        raise HTTPException(404, "MRR Inspection not found")

    # Docs are ticket-level in your current DB
    docs = session.exec(
        select(MrrDocument)
        .where(MrrDocument.ticket_id == lot_id)
        .order_by(MrrDocument.created_at.asc())
    ).all()

    # Draft inspection (not submitted yet) - allow inspector to resume (currently unused)
    draft_inspection = session.exec(
        select(MrrReceivingInspection)
        .where(
            (MrrReceivingInspection.ticket_id == lot_id)
            & (MrrReceivingInspection.inspector_confirmed == False)
        )
        .order_by(MrrReceivingInspection.created_at.desc())
    ).first()

    receiving = session.exec(
        select(MrrReceiving).where(MrrReceiving.ticket_id == lot_id)
    ).first()

    # Only allow export if submitted (optional rule)
    if not (insp.inspector_confirmed or insp.manager_approved):
        raise HTTPException(400, "Inspection must be submitted before export")


    # Decide template based on inspection.template_type
    tpl = _safe_upper(getattr(insp, "template_type", "RAW"))
    if tpl != "RAW":
        raise HTTPException(
            400, f"Template type {tpl} export not wired yet (we will do F02 next)"
        )

    xlsx_bytes = fill_mrr_f01_xlsx_bytes(
        lot=lot,
        receiving=receiving,
        inspection=insp,
        docs=docs,
        photos_by_group=None,
    )

    if not xlsx_bytes:
        raise HTTPException(500, "RAW report generation failed: xlsx_bytes is empty/None.")

    filename = f"{insp.report_no or f'MRR-{lot_id}-{inspection_id}'}.xlsx"
    return Response(
        content=xlsx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _try_convert_xlsx_to_pdf_bytes(xlsx_bytes: bytes) -> bytes:
    """
    Best-effort conversion using LibreOffice.
    IMPORTANT: Do NOT zoom/shift for Excel-based templates (F01),
    because it makes the page too large and breaks the layout.
    """
    tmp_dir = "/tmp/mrr_export"
    os.makedirs(tmp_dir, exist_ok=True)

    xlsx_path = os.path.join(tmp_dir, "report.xlsx")
    with open(xlsx_path, "wb") as f:
        f.write(xlsx_bytes)

    cmd = [
        "soffice",
        "--headless",
        "--nologo",
        "--nolockcheck",
        "--nodefault",
        "--norestore",
        "--convert-to",
        "pdf",
        "--outdir",
        tmp_dir,
        xlsx_path,
    ]

    try:
        subprocess.check_call(cmd)
    except FileNotFoundError:
        raise HTTPException(
            500,
            "PDF export needs LibreOffice (soffice) installed on the server. "
            "For now use Export XLSX."
        )
    except Exception:
        raise HTTPException(500, "Failed to convert XLSX to PDF (LibreOffice error).")

    pdf_path = os.path.join(tmp_dir, "report.pdf")
    if not os.path.exists(pdf_path):
        raise HTTPException(500, "Conversion did not produce PDF output.")

    with open(pdf_path, "rb") as f:
        pdf = f.read()

    # ✅ Keep on A4 WITHOUT zoom/shift (avoid huge output)
    pdf = fit_pdf_pages_to_a4(
        pdf,
        margin_left_right=0.3,
        margin_bottom=0.3,
        header_reserved=55.0,
        zoom=1.0,
        shift_x=0.0,
        shift_y=0.0,
    )

    # ✅ Stamp logo + footer once
    base_dir = os.path.dirname(__file__)
    logo_path = os.path.join(base_dir, "static", "images", "logo.png")
    pdf = stamp_logo_on_pdf(pdf, logo_path)
    pdf = stamp_footer_on_pdf(pdf, "QAP0600-F01")

    return pdf




@app.get("/mrr/{lot_id}/inspection/id/{inspection_id}/export/pdf")
def mrr_export_inspection_pdf(
    lot_id: int,
    inspection_id: int,
    request: Request,
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)

    lot = session.get(MaterialLot, lot_id)
    if not lot:
        raise HTTPException(404, "MRR Ticket not found")

    insp = session.get(MrrReceivingInspection, inspection_id)
    if not insp or insp.ticket_id != lot_id:
        raise HTTPException(404, "MRR Inspection not found")

    # Only allow export if submitted OR manager approved
    if not (getattr(insp, "inspector_confirmed", False) or getattr(insp, "manager_approved", False)):
        raise HTTPException(400, "Inspection must be submitted before export")

    docs = session.exec(
        select(MrrDocument)
        .where(MrrDocument.ticket_id == lot_id)
        .order_by(MrrDocument.created_at.asc())
    ).all()

    receiving = session.exec(
        select(MrrReceiving).where(MrrReceiving.ticket_id == lot_id)
    ).first()

    tpl = _resolve_template_type(lot, insp)

    # Read inspection json (for timestamps / reviewer / manager)
    try:
        data = json.loads(getattr(insp, "inspection_json", None) or "{}")
        if not isinstance(data, dict):
            data = {}
    except Exception:
        data = {}

    # Inspector signature info (first submit)
    inspector_name = (getattr(insp, "inspector_name", "") or "").strip()
    inspector_date = ""
    if getattr(insp, "inspector_confirmed", False):
        ts = data.get("submitted_at_utc") or getattr(insp, "created_at", None)
        inspector_date = _as_date_str(ts) if ts else _as_date_str(datetime.utcnow())

    # Reviewer (optional)
    reviewer_name = (data.get("reviewed_by") or "").strip()
    reviewer_date = ""
    if reviewer_name:
        reviewer_date = _as_date_str(data.get("reviewed_at_utc") or "")

    # Manager approval
    manager_name = (data.get("manager_approved_by") or "").strip()
    manager_date = ""
    if getattr(insp, "manager_approved", False):
        ts2 = data.get("manager_approved_at_utc") or datetime.utcnow().isoformat()
        manager_date = _as_date_str(ts2)

    # ---------- OUTSOURCED => F02 ----------
    if tpl == "OUTSOURCED":
        docx_bytes = fill_mrr_f02_docx_bytes(
            lot=lot,
            inspection=insp,
            receiving=receiving,
            docs=docs,
        )

        pdf_bytes = docx_bytes_to_pdf_bytes(docx_bytes)

        # Footer stamp (optional)
        try:
            pdf_bytes = stamp_footer_on_pdf(pdf_bytes, "QAP0600-F02")
        except Exception:
            pass

        # Digital signatures overlay (F02)  ✅ ADD THIS
        try:
            pdf_bytes = stamp_signatures_on_pdf_f02(
                pdf_bytes,
                inspector_name=inspector_name if getattr(insp, "inspector_confirmed", False) else "",
                inspector_date=inspector_date if getattr(insp, "inspector_confirmed", False) else "",
                reviewer_name=reviewer_name or "",
                reviewer_date=reviewer_date or "",
                manager_name=manager_name if getattr(insp, "manager_approved", False) else "",
                manager_date=manager_date if getattr(insp, "manager_approved", False) else "",
            )
        except Exception:
            pass

        filename = f"{getattr(insp, 'report_no', 'MRR')}_F02.pdf"
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    # ---------- RAW => F01 ----------
    xlsx_bytes = fill_mrr_f01_xlsx_bytes(
        lot=lot,
        receiving=receiving,
        inspection=insp,
        docs=docs,
        photos_by_group=None,
    )

    pdf_bytes = _try_convert_xlsx_to_pdf_bytes(xlsx_bytes)

    # Footer stamp (optional)
    try:
        pdf_bytes = stamp_footer_on_pdf(pdf_bytes, "QAP0600-F01")
    except Exception:
        pass

    # Digital signatures (F01)
    try:
        pdf_bytes = stamp_signatures_on_pdf(
            pdf_bytes,
            inspector_name=inspector_name if getattr(insp, "inspector_confirmed", False) else "",
            inspector_date=inspector_date if getattr(insp, "inspector_confirmed", False) else "",
            manager_name=manager_name if getattr(insp, "manager_approved", False) else "",
            manager_date=manager_date if getattr(insp, "manager_approved", False) else "",
        )
    except Exception:
        pass

    return Response(pdf_bytes, media_type="application/pdf")
    
@app.get("/mrr/{lot_id}/inspection/id/{inspection_id}/export/package")
def mrr_export_inspection_package(
    lot_id: int,
    inspection_id: int,
    request: Request,
    mode: str = "zip",  # "zip" or "pdf"
    session: Session = Depends(get_session),
):
    """
    Export shipment package:
      - Standalone REPORT (PDF)
      - PO + DN + COA (+ any other docs)
      - Photos
    mode=zip -> download a bundle zip (always works)
    mode=pdf -> merge everything into ONE PDF (best-effort)
    """
    user = get_current_user(request, session)

    lot = session.get(MaterialLot, lot_id)
    if not lot:
        raise HTTPException(404, "MRR Ticket not found")

    insp = session.get(MrrReceivingInspection, inspection_id)
    if not insp or insp.ticket_id != lot_id:
        raise HTTPException(404, "MRR Inspection not found")

    # Load ticket-level receiving info (if any)
    receiving = session.exec(
        select(MrrReceiving).where(MrrReceiving.ticket_id == lot_id)
    ).first()

    # Documents: include both shipment-specific and ticket-level (PO usually ticket-level)
    all_docs = session.exec(
        select(MrrDocument)
        .where(MrrDocument.ticket_id == lot_id)
        .order_by(MrrDocument.created_at.asc())
    ).all()

   # Filter docs relevant to this shipment package:
    # 1) Always include shipment-attached docs for this inspection_id
    # 2) Always include ticket-level PO
    # 3) ALSO include ticket-level docs that match this shipment DN by doc_number
    shipment_dn = (getattr(insp, "delivery_note_no", "") or "").strip()
    
    docs_for_package = []
    for d in all_docs:
        d_insp = getattr(d, "inspection_id", None)
        dt = (getattr(d, "doc_type", "") or "").upper().strip()
        dn = (getattr(d, "doc_number", "") or "").strip()
    
        # exact shipment link
        if d_insp == inspection_id:
            docs_for_package.append(d)
            continue
    
        # ticket-level docs: include PO always
        if d_insp is None:
            if dt == "PO":
                docs_for_package.append(d)
                continue
    
            # include shipment-related docs if doc_number matches shipment DN
            if shipment_dn and dn and dn == shipment_dn:
                docs_for_package.append(d)
                continue

    # Photos for this shipment
    photos = session.exec(
        select(MrrInspectionPhoto)
        .where(
            MrrInspectionPhoto.ticket_id == lot_id,
            MrrInspectionPhoto.inspection_id == inspection_id,
        )
        .order_by(MrrInspectionPhoto.created_at.asc())
    ).all()

    # Generate the REPORT PDF
    report_pdf = _generate_shipment_report_pdf_bytes(
        lot=lot,
        insp=insp,
        receiving=receiving,
        docs=all_docs,  # report generation may need full context
    )

    report_no = getattr(insp, "report_no", "") or f"MRR_{lot_id}_{inspection_id}"
    report_no = _safe_filename(report_no)

    # Sort documents in practical order: PO -> DN -> COA -> others
    order_map = {"PO": 0, "DELIVERY_NOTE": 1, "COA": 2}
    docs_for_package.sort(key=lambda d: (order_map.get((d.doc_type or "").upper(), 99), d.created_at))

    # ------------------------
    # MODE = ONE MERGED PDF
    # ------------------------
    if (mode or "").lower() == "pdf":
        pdf_parts: list[bytes] = []
        pdf_parts.append(report_pdf)

        # Add attachments (best-effort convert to pdf)
        for d in docs_for_package:
            p = resolve_mrr_doc_path(getattr(d, "file_path", "") or "")
            if not p:
                continue
            pdf_b = _doc_path_to_pdf_bytes(p)
            if pdf_b:
                pdf_parts.append(pdf_b)

        # Add photos (each photo becomes a page in PDF)
        for ph in photos:
            p = resolve_mrr_photo_path(getattr(ph, "file_path", None))
            if not p:
                continue
            ext = os.path.splitext(p)[1].lower()
            if ext in [".png", ".jpg", ".jpeg", ".webp"]:
                try:
                    pdf_parts.append(_image_path_to_pdf_bytes(p))
                except Exception:
                    pass

        merged = _merge_pdf_bytes_in_order(pdf_parts)

        filename = f"{report_no}_PACKAGE.pdf"
        return Response(
            content=merged,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    # ------------------------
    # MODE = ZIP BUNDLE
    # ------------------------
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as z:
        # Report
        z.writestr(f"{report_no}/01_REPORT_{report_no}.pdf", report_pdf)

        # Documents
        for idx, d in enumerate(docs_for_package, start=1):
            p = resolve_mrr_doc_path(getattr(d, "file_path", "") or "")
            if not p:
                continue

            doc_type = (getattr(d, "doc_type", "") or "DOC").upper()
            doc_name = _safe_filename(getattr(d, "doc_name", "") or os.path.basename(p))
            ext = os.path.splitext(p)[1].lower() or ".bin"

            arc = f"{report_no}/02_DOCS/{idx:02d}_{doc_type}_{doc_name}{ext}"
            try:
                z.writestr(arc, _read_file_bytes(p))
            except Exception:
                pass

        # Photos
        for idx, ph in enumerate(photos, start=1):
            p = resolve_mrr_photo_path(getattr(ph, "file_path", None))
            if not p:
                continue
            group = _safe_filename(getattr(ph, "group_name", "") or "Photos")
            cap = _safe_filename(getattr(ph, "caption", "") or "")
            ext = os.path.splitext(p)[1].lower() or ".jpg"

            label = f"{idx:02d}_{group}"
            if cap:
                label += f"_{cap}"
            arc = f"{report_no}/03_PHOTOS/{label}{ext}"

            try:
                z.writestr(arc, _read_file_bytes(p))
            except Exception:
                pass

    buf.seek(0)
    filename = f"{report_no}_PACKAGE.zip"
    return Response(
        content=buf.getvalue(),
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )

@app.get("/health")
def health():
    return {"ok": True}


@app.on_event("startup")
def on_startup():
    try:
        create_db_and_tables()
        print("✅ DB ready")
    except Exception as e:
        # Do NOT kill the server if DB is temporarily unavailable
        print("⚠️ DB startup failed (server will still run). Error:", repr(e))
        ensure_default_users()


def ensure_default_users():
    from .db import engine
    with Session(engine) as session:
        u = session.exec(select(User).where(User.username == "manager")).first()
        if not u:
            session.add(User(
                username="manager",
                display_name="Manager",
                email="manager@company.local",
                role="MANAGER",
                department="QUALITY",
                password_hash=hash_password("manager123"),
            ))

        i = session.exec(select(User).where(User.username == "inspector")).first()
        if not i:
            session.add(User(
                username="inspector",
                display_name="Inspector",
                email="inspector@company.local",
                role="INSPECTOR",
                department="QUALITY",
                password_hash=hash_password("inspector"),
            ))

        b = session.exec(select(User).where(User.username == "boss")).first()
        if not b:
            session.add(User(
                username="boss",
                display_name="Boss",
                email="boss@company.local",
                role="BOSS",
                department="ADMIN",
                password_hash=hash_password("boss123"),
            ))

        session.commit()


OMAN_TZ = ZoneInfo("Asia/Muscat")

def format_oman_dt(dt_utc: datetime | None) -> str:
    if not dt_utc:
        return ""
    # dt_utc is stored as naive UTC in DB
    dt_local = dt_utc.replace(tzinfo=ZoneInfo("UTC")).astimezone(OMAN_TZ)
    return dt_local.strftime("%Y-%m-%d %H:%M")

def make_approval_stamp_pdf(page_w: float, page_h: float, text_lines: list[str]) -> bytes:
    """
    Create a 1-page transparent PDF with approval text in bottom-right corner.
    """
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=(page_w, page_h))

    # position (bottom-right)
    x = page_w - 40
    y = 60

    c.setFont("Helvetica-Bold", 10)
    for line in text_lines:
        c.drawRightString(x, y, line)
        y -= 12

    c.showPage()
    c.save()
    buf.seek(0)
    return buf.getvalue()

from io import BytesIO
from datetime import datetime
try:
    from pypdf import PdfReader, PdfWriter
except Exception:
    from pypdf import PdfReader, PdfWriter


def stamp_approval_on_pdf(
    pdf_bytes: bytes,
    approved_by: str,
    approved_at_utc: datetime | None
) -> bytes:
    """
    Pure helper: stamps APPROVED info on every page.
    Does NOT query DB, does NOT need session.
    """
    if not approved_by or not approved_at_utc:
        return pdf_bytes

    approved_local = format_oman_dt(approved_at_utc)

    reader = PdfReader(BytesIO(pdf_bytes))
    writer = PdfWriter()

    for page in reader.pages:
        w = float(page.mediabox.width)
        h = float(page.mediabox.height)

        stamp_pdf = make_approval_stamp_pdf(
            w, h,
            [
                "APPROVED",
                f"By: {approved_by}",
                f"At: {approved_local} (Oman)",
            ],
        )

        stamp_reader = PdfReader(BytesIO(stamp_pdf))
        page.merge_page(stamp_reader.pages[0])
        writer.add_page(page)

    out = BytesIO()
    writer.write(out)
    out.seek(0)
    return out.getvalue()


def get_current_user(request: Request, session: Session) -> User:
    username = request.cookies.get("user")
    if not username:
        raise HTTPException(401, "Not logged in")
    user = session.exec(select(User).where(User.username == username)).first()
    if not user:
        raise HTTPException(401, "Invalid user")
    return user

def require_manager(user: User):
    if (getattr(user, "role", "") or "").strip().upper() != "MANAGER":
        raise HTTPException(403, "Manager only")

def _can_hydro_approve(user: User) -> bool:
    return (getattr(user, "role", "") or "").strip().upper() in ["INSPECTOR", "MANAGER"]

def _can_final_approve(user: User) -> bool:
    return (getattr(user, "role", "") or "").strip().upper() in ["MANAGER"]

def _hydro_qaqc_candidates(session: Session, query: str) -> list[User]:
    q = (query or "").strip().lower()

    users = session.exec(
        select(User)
        .where(User.role.in_(["INSPECTOR", "MANAGER"]))
        .order_by(User.display_name, User.username)
    ).all()

    if not q:
        return users[:8]

    out = []
    for u in users:
        hay = f"{u.username} {u.display_name} {u.role}".lower()
        if q in hay:
            out.append(u)
        if len(out) >= 8:
            break
    return out

def forbid_boss(user: User):
    if user.role == "BOSS":
        raise HTTPException(403, "Read-only user")



pwd_context = CryptContext(schemes=["pbkdf2_sha256"], deprecated="auto")


def _verify_user_password(plain_password: str, password_hash: str) -> bool:
    try:
        if not plain_password or not password_hash:
            return False
        return pwd_context.verify(plain_password, password_hash)
    except Exception:
        return False


def _hash_user_password(password: str) -> str:
    return pwd_context.hash(password or "")


USER_DEPARTMENTS = ["QUALITY", "PRODUCTION", "ADMIN"]
USER_ROLES = ["INSPECTOR", "COORDINATOR", "RUN_CREATOR", "MANAGER", "BOSS"]

ACCESS_OPTIONS = [
    "dashboard",
    "runs_view",
    "runs_create",
    "runs_edit",
    "runs_approve",
    "mrr_view",
    "mrr_manage",
    "mrr_approve",
    "burst_view",
    "burst_manage",
    "hydro_view",
    "hydro_manage",
    "hydro_approve",
    "final_view",
    "final_manage",
    "final_approve",
    "users_manage",
]

def _default_department_for_role(role: str) -> str:
    r = (role or "").strip().upper()
    if r == "INSPECTOR":
        return "QUALITY"
    if r == "COORDINATOR":
        return "QUALITY"
    if r == "RUN_CREATOR":
        return "PRODUCTION"
    if r == "MANAGER":
        return "QUALITY"
    if r == "BOSS":
        return "ADMIN"
    return "QUALITY"

def _default_access_for_user(role: str, department: str) -> list[str]:
    r = (role or "").strip().upper()
    d = (department or "").strip().upper()

    if r == "BOSS":
        return ACCESS_OPTIONS[:]

    if r == "INSPECTOR":
        return [
            "dashboard",
            "runs_view",
            "mrr_view",
            "burst_view",
            "hydro_view",
            "final_view",
        ]

    if r == "COORDINATOR":
        if d == "PRODUCTION":
            return [
                "dashboard",
                "runs_view",
                "runs_create",
                "burst_view",
                "hydro_view",
                "final_view",
            ]
        return [
            "dashboard",
            "runs_view",
            "mrr_view",
            "burst_view",
            "hydro_view",
            "final_view",
        ]

    if r == "RUN_CREATOR":
        return [
            "dashboard",
            "runs_view",
            "runs_create",
        ]

    if r == "MANAGER":
        if d == "PRODUCTION":
            return [
                "dashboard",
                "runs_view",
                "runs_create",
                "runs_edit",
                "runs_approve",
                "burst_view",
                "burst_manage",
                "final_view",
            ]
        if d == "QUALITY":
            return [
                "dashboard",
                "runs_view",
                "mrr_view",
                "mrr_manage",
                "mrr_approve",
                "burst_view",
                "burst_manage",
                "hydro_view",
                "hydro_manage",
                "hydro_approve",
                "final_view",
                "final_manage",
                "final_approve",
                "users_manage",
            ]

    return ["dashboard"]

def _load_access_overrides(raw: str) -> list[str]:
    try:
        data = json.loads(raw or "[]")
        if not isinstance(data, list):
            return []
        out = []
        for item in data:
            key = str(item or "").strip()
            if key in ACCESS_OPTIONS and key not in out:
                out.append(key)
        return out
    except Exception:
        return []

def _save_access_overrides(values: list[str]) -> str:
    out = []
    for item in values or []:
        key = str(item or "").strip()
        if key in ACCESS_OPTIONS and key not in out:
            out.append(key)
    return json.dumps(out, ensure_ascii=False)

def _effective_access_for_user(user: User) -> list[str]:
    department = (getattr(user, "department", "") or "").strip().upper() or _default_department_for_role(user.role)
    overrides = _load_access_overrides(getattr(user, "access_overrides_json", "") or "")
    return overrides or _default_access_for_user(user.role, department)

def _access_label_map() -> dict:
    return {
        "dashboard": "Dashboard / Overview",
        "runs_view": "View In-Process Runs",
        "runs_create": "Create Production Runs",
        "runs_edit": "Edit Production Runs",
        "runs_approve": "Approve / Reopen Runs",
        "mrr_view": "View MRR",
        "mrr_manage": "Manage MRR",
        "mrr_approve": "Approve MRR",
        "burst_view": "View Burst Testing",
        "burst_manage": "Manage Burst Testing",
        "hydro_view": "View Hydro Testing",
        "hydro_manage": "Manage Hydro Testing",
        "hydro_approve": "Approve Hydro Testing",
        "final_view": "View Final Inspection",
        "final_manage": "Manage Final Inspection",
        "final_approve": "Approve Final Inspection",
        "users_manage": "User Management",
    }
def _default_department_for_role(role: str) -> str:
    r = (role or "").strip().upper()
    if r == "INSPECTOR":
        return "QUALITY"
    if r == "COORDINATOR":
        return "QUALITY"
    if r == "RUN_CREATOR":
        return "PRODUCTION"
    if r == "MANAGER":
        return "QUALITY"
    if r == "BOSS":
        return "ADMIN"
    return "QUALITY"

def _default_access_for_user(role: str, department: str) -> list[str]:
    r = (role or "").strip().upper()
    d = (department or "").strip().upper()

    if r == "BOSS":
        return [
            "dashboard",
            "runs_view",
            "mrr_view",
            "burst_view",
            "hydro_view",
            "final_view",
        ]

    if r == "INSPECTOR":
        return [
            "dashboard",
            "runs_view",
            "mrr_view",
            "burst_view",
            "hydro_view",
            "hydro_approve",
            "final_view",
        ]

    if r == "COORDINATOR":
        if d == "PRODUCTION":
            return [
                "dashboard",
                "runs_view",
                "runs_create",
                "burst_view",
                "hydro_view",
                "final_view",
            ]
        return [
            "dashboard",
            "runs_view",
            "mrr_view",
            "burst_view",
            "hydro_view",
            "final_view",
        ]

    if r == "RUN_CREATOR":
        return [
            "dashboard",
            "runs_view",
            "runs_create",
        ]

    if r == "MANAGER":
        if d == "PRODUCTION":
            return [
                "dashboard",
                "runs_view",
                "runs_create",
                "runs_edit",
                "runs_approve",
                "burst_view",
                "burst_manage",
                "final_view",
            ]
        if d == "QUALITY":
            return [
                "dashboard",
                "runs_view",
                "mrr_view",
                "mrr_manage",
                "mrr_approve",
                "burst_view",
                "burst_manage",
                "hydro_view",
                "hydro_manage",
                "hydro_approve",
                "final_view",
                "final_manage",
                "final_approve",
                "users_manage",
            ]
        return ACCESS_OPTIONS[:]

    return ["dashboard"]

def _load_access_overrides(raw: str) -> list[str]:
    try:
        data = json.loads(raw or "[]")
        if not isinstance(data, list):
            return []
        out = []
        for item in data:
            key = str(item or "").strip()
            if key in ACCESS_OPTIONS and key not in out:
                out.append(key)
        return out
    except Exception:
        return []

def _save_access_overrides(values: list[str]) -> str:
    out = []
    for item in values or []:
        key = str(item or "").strip()
        if key in ACCESS_OPTIONS and key not in out:
            out.append(key)
    return json.dumps(out, ensure_ascii=False)

def _effective_access_for_user(user: User) -> list[str]:
    department = (getattr(user, "department", "") or "").strip().upper() or _default_department_for_role(user.role)
    overrides = _load_access_overrides(getattr(user, "access_overrides_json", "") or "")
    return overrides or _default_access_for_user(user.role, department)

def _access_label_map() -> dict:
    return {
        "dashboard": "Dashboard / Overview",
        "runs_view": "View In-Process Runs",
        "runs_create": "Create Production Runs",
        "runs_edit": "Edit Production Runs",
        "runs_approve": "Approve / Reopen Runs",
        "mrr_view": "View MRR",
        "mrr_manage": "Manage MRR",
        "mrr_approve": "Approve MRR",
        "burst_view": "View Burst Testing",
        "burst_manage": "Manage Burst Testing",
        "hydro_view": "View Hydro Testing",
        "hydro_manage": "Manage Hydro Testing",
        "hydro_approve": "Approve Hydro Testing",
        "final_view": "View Final Inspection",
        "final_manage": "Manage Final Inspection",
        "final_approve": "Approve Final Inspection",
        "users_manage": "User Management",
    }

@app.post("/users/{username}/update")
async def users_update(
    username: str,
    request: Request,
    session: Session = Depends(get_session),
    display_name: str = Form(""),
    email: str = Form(""),
    role: str = Form(""),
    department: str = Form(""),
    password: str = Form(""),
):
    user = get_current_user(request, session)
    require_manager(user)

    target = session.exec(select(User).where(User.username == username)).first()
    if not target:
        raise HTTPException(404, "User not found")

    form = await request.form()

    if display_name.strip():
        target.display_name = display_name.strip()

    target.email = (email or "").strip().lower()

    if role.strip():
        r = role.strip().upper()
        if r not in USER_ROLES:
            raise HTTPException(400, "Invalid role")
        target.role = r

    new_department = (department or "").strip().upper()
    if not new_department:
        new_department = _default_department_for_role(target.role)
    if new_department not in USER_DEPARTMENTS:
        raise HTTPException(400, "Invalid department")
    target.department = new_department

    if password.strip():
        if len(password.strip()) < 4:
            raise HTTPException(400, "Password too short")
        target.password_hash = hash_password(password.strip())
        target.must_change_password = True

    access_override_values = []
    for item in form.getlist("access_overrides"):
        if item in ACCESS_OPTIONS:
            access_override_values.append(item)
    target.access_overrides_json = _save_access_overrides(access_override_values)

    session.add(target)
    session.commit()
    return RedirectResponse(f"/users/{target.username}/edit", status_code=302)


@app.post("/users/{username}/toggle-lock")
def users_toggle_lock(
    username: str,
    request: Request,
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)
    require_manager(user)

    target = session.exec(select(User).where(User.username == username)).first()
    if not target:
        raise HTTPException(404, "User not found")

    if target.username == user.username:
        raise HTTPException(400, "You cannot lock your own account.")

    target.is_locked = not bool(target.is_locked)
    session.add(target)
    session.commit()

    return RedirectResponse("/users", status_code=302)


@app.post("/users/{username}/reset-access")
def users_reset_access(
    username: str,
    request: Request,
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)
    require_manager(user)

    target = session.exec(select(User).where(User.username == username)).first()
    if not target:
        raise HTTPException(404, "User not found")

    target.access_overrides_json = ""
    session.add(target)
    session.commit()

    return RedirectResponse(f"/users/{username}/edit", status_code=302)
    

@app.post("/users/{username}/delete")
def users_delete(
    username: str,
    request: Request,
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)
    require_manager(user)

    # safety: do not delete yourself
    if user.username == username:
        raise HTTPException(400, "You cannot delete your own account")

    target = session.exec(select(User).where(User.username == username)).first()
    if not target:
        raise HTTPException(404, "User not found")

    session.delete(target)
    session.commit()
    return RedirectResponse("/users", status_code=302)


@app.get("/", response_class=HTMLResponse)
def home(request: Request, session: Session = Depends(get_session)):
    # If logged in -> dashboard, else -> login (so root stays "healthy")
    username = request.cookies.get("user")
    if username:
        u = session.exec(select(User).where(User.username == username)).first()
        if u:
            return RedirectResponse("/dashboard", status_code=302)
    return RedirectResponse("/login", status_code=302)



@app.get("/login", response_class=HTMLResponse)
def login_get(request: Request):
    return templates.TemplateResponse("login.html", {"request": request, "error": ""})


@app.post("/login")
def login_post(
    request: Request,
    username: str = Form(...),
    password: str = Form(...),
    session: Session = Depends(get_session),
):
    username = (username or "").strip().lower()
    password = password or ""

    user = session.exec(select(User).where(User.username == username)).first()

    if not user or not verify_password(password, user.password_hash):
        return templates.TemplateResponse(
            "login.html",
            {
                "request": request,
                "error": "Invalid username or password",
            },
        )

    if user.is_locked:
        return templates.TemplateResponse(
            "login.html",
            {
                "request": request,
                "error": "Your account is locked. Please contact the administrator.",
            },
        )

    response = RedirectResponse("/dashboard", status_code=302)
    response.set_cookie(
        key="username",
        value=user.username,
        httponly=True,
        secure=True,
        samesite="lax",
        path="/",
        max_age=60 * 60 * 12,
    )
    return response


@app.get("/logout")
def logout():
    resp = RedirectResponse("/login", status_code=302)
    resp.delete_cookie("user")
    return resp


def get_days_for_run(session: Session, run_id: int) -> List[date]:
    days = session.exec(
        select(InspectionEntry.actual_date)
        .where(InspectionEntry.run_id == run_id)
        .distinct()
        .order_by(InspectionEntry.actual_date)
    ).all()
    return list(days)



@app.get("/users", response_class=HTMLResponse)
def users_get(request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)
    require_manager(user)

    users = session.exec(select(User).order_by(User.created_at.desc(), User.username)).all()

    access_labels = _access_label_map()
    user_rows = []

    for u in users:
        dept = (getattr(u, "department", "") or "").strip().upper() or _default_department_for_role(u.role)
        overrides = _load_access_overrides(getattr(u, "access_overrides_json", "") or "")
        effective = _effective_access_for_user(u)

        user_rows.append({
            "user_obj": u,
            "department": dept,
            "override_access": overrides,
            "effective_access": effective,
            "status_label": "LOCKED" if getattr(u, "is_locked", False) else "ACTIVE",
        })

    return templates.TemplateResponse(
        "users.html",
        {
            "request": request,
            "user": user,
            "user_rows": user_rows,
            "departments": USER_DEPARTMENTS,
            "roles": USER_ROLES,
            "access_options": ACCESS_OPTIONS,
            "access_labels": access_labels,
            "error": "",
        },
    )

@app.post("/users")
def users_post(
    request: Request,
    session: Session = Depends(get_session),
    username: str = Form(...),
    display_name: str = Form(...),
    email: str = Form(""),
    role: str = Form(...),
    department: str = Form(""),
):
    user = get_current_user(request, session)
    require_manager(user)

    username = (username or "").strip().lower()
    display_name = (display_name or "").strip()
    email = (email or "").strip().lower()
    role = (role or "").strip().upper()
    department = (department or "").strip().upper()

    if not username:
        raise HTTPException(status_code=400, detail="Username is required.")
    if not display_name:
        raise HTTPException(status_code=400, detail="Full name is required.")
    if role not in USER_ROLES:
        raise HTTPException(status_code=400, detail="Invalid role.")

    if not department:
        department = _default_department_for_role(role)
    if department not in USER_DEPARTMENTS:
        raise HTTPException(status_code=400, detail="Invalid department.")

    existing = session.exec(select(User).where(User.username == username)).first()
    if existing:
        raise HTTPException(status_code=400, detail="Username already exists.")

    session.add(User(
        username=username,
        display_name=display_name,
        email=email,
        role=role,
        department=department,
        password_hash=hash_password("0000"),
        must_change_password=True,
        is_locked=False,
        access_overrides_json="",
    ))
    session.commit()

    return RedirectResponse("/users", status_code=302)


@app.get("/users/{username}/edit", response_class=HTMLResponse)
def user_edit_page(
    username: str,
    request: Request,
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)
    require_manager(user)

    target = session.exec(select(User).where(User.username == username)).first()
    if not target:
        raise HTTPException(404, "User not found")

    dept = (getattr(target, "department", "") or "").strip().upper() or _default_department_for_role(target.role)
    overrides = _load_access_overrides(getattr(target, "access_overrides_json", "") or "")
    effective = _effective_access_for_user(target)

    return templates.TemplateResponse(
        "user_edit.html",
        {
            "request": request,
            "user": user,
            "target": target,
            "department": dept,
            "override_access": overrides,
            "effective_access": effective,
            "departments": USER_DEPARTMENTS,
            "roles": USER_ROLES,
            "access_options": ACCESS_OPTIONS,
            "access_labels": _access_label_map(),
        },
    )

def slot_from_time_str(t: str) -> tuple[str, int]:
    """
    HARD RULE (2-hour slots):
    - HH:00 .. HH+1:30  -> HH:00
    - HH+1:31 .. HH+2:00 -> HH+2:00

    Special rule for end of day:
    - If the calculated next slot becomes 24:00, return ("00:00", 1) meaning next day.
    """
    parts = t.split(":")
    hh = int(parts[0])
    mm = int(parts[1]) if len(parts) > 1 else 0
    total_min = hh * 60 + mm

    # base even hour (00,02,04,...,22)
    base_h = (hh // 2) * 2
    base_min = base_h * 60
    delta = total_min - base_min

    # exact even hour always stays in its slot
    if delta == 0:
        slot_min = base_min
    # 0..90 minutes => same slot
    elif 0 < delta <= 90:
        slot_min = base_min
    # 91..120 => next slot
    else:
        slot_min = base_min + 120

    # If we rolled past midnight (24:00), move to next day 00:00
    if slot_min >= 24 * 60:
        return "00:00", 1

    # Normal clamp low end (shouldn't happen, but safe)
    if slot_min < 0:
        slot_min = 0

    return f"{slot_min // 60:02d}:00", 0



def get_progress_percent(session: Session, run: ProductionRun) -> int:
    if run.total_length_m <= 0:
        return 0
    max_len = 0.0
    entries = session.exec(select(InspectionEntry).where(InspectionEntry.run_id == run.id)).all()
    for e in entries:
        vals = session.exec(
            select(InspectionValue).where(
                InspectionValue.entry_id == e.id,
                InspectionValue.param_key == "length_m"
            )
        ).all()
        for v in vals:
            if v.value is not None and v.value > max_len:
                max_len = v.value
    return int(min(100.0, (max_len / run.total_length_m) * 100.0))

def get_progress_map_bulk(session: Session, runs: List[ProductionRun]) -> Dict[int, int]:
    """
    Faster bulk progress calculation for many runs at once.

    Rule stays the same as get_progress_percent():
    - look for MAX InspectionValue.value where param_key == 'length_m' per run
    - compare with run.confirmed_total_length_m
    - divide by run.total_length_m
    """
    run_ids = [r.id for r in runs if getattr(r, "id", None) is not None]
    if not run_ids:
        return {}

    # Start from confirmed lengths already stored on run
    measured_map: Dict[int, float] = {}
    total_map: Dict[int, float] = {}
    confirmed_map: Dict[int, float] = {}

    for r in runs:
        if r.id is None:
            continue
        total_map[r.id] = float(getattr(r, "total_length_m", 0.0) or 0.0)
        confirmed_map[r.id] = float(getattr(r, "confirmed_total_length_m", 0.0) or 0.0)

    # One grouped query instead of query-per-run
    rows = session.exec(
        select(
            InspectionEntry.run_id,
            func.max(InspectionValue.value),
        )
        .join(InspectionValue, InspectionValue.entry_id == InspectionEntry.id)
        .where(InspectionEntry.run_id.in_(run_ids))
        .where(InspectionValue.param_key == "length_m")
        .where(InspectionValue.value != None)  # noqa: E711
        .group_by(InspectionEntry.run_id)
    ).all()

    for run_id, max_val in rows:
        try:
            measured_map[int(run_id)] = float(max_val or 0.0)
        except Exception:
            measured_map[int(run_id)] = 0.0

    progress_map: Dict[int, int] = {}
    for run_id in run_ids:
        total_len = total_map.get(run_id, 0.0)
        if total_len <= 0:
            progress_map[run_id] = 0
            continue

        measured = measured_map.get(run_id, 0.0)
        confirmed = confirmed_map.get(run_id, 0.0)
        actual = max(measured, confirmed)

        pct = int(min(100.0, (actual / total_len) * 100.0))
        progress_map[run_id] = pct

    return progress_map
    
def get_day_latest_trace(session: Session, run_id: int, day: date) -> dict:
    entries = session.exec(
        select(InspectionEntry)
        .where(InspectionEntry.run_id == run_id, InspectionEntry.actual_date == day)
        .order_by(InspectionEntry.created_at)
    ).all()

    raw_batches = get_day_material_batches(session, run_id, day)

    tool1 = (None, None, None)
    tool2 = (None, None, None)
    for e in entries:
        if (e.tool1_name or e.tool1_serial or e.tool1_calib_due):
            tool1 = (e.tool1_name, e.tool1_serial, e.tool1_calib_due)
        if (e.tool2_name or e.tool2_serial or e.tool2_calib_due):
            tool2 = (e.tool2_name, e.tool2_serial, e.tool2_calib_due)

    tools = []
    if any(tool1):
        tools.append(tool1)
    if any(tool2):
        tools.append(tool2)

    return {"entries": entries, "raw_batches": raw_batches, "tools": tools}


def get_last_known_trace_before_day(session: Session, run_id: int, day: date) -> dict:
    # ✅ Get last known RAW batch from MaterialUseEvent (not from entries)
    raw = ""

    last_ev = session.exec(
        select(MaterialUseEvent)
        .where(
            MaterialUseEvent.run_id == run_id,
            MaterialUseEvent.day < day,
        )
        .order_by(
            MaterialUseEvent.day.desc(),
            MaterialUseEvent.slot_time.desc(),
            MaterialUseEvent.created_at.desc(),
        )
    ).first()

    if last_ev:
        lot = session.get(MaterialLot, last_ev.lot_id)
        if lot and lot.batch_no:
            raw = lot.batch_no

    # Keep existing tool carry-forward logic
    entries = session.exec(
        select(InspectionEntry)
        .where(InspectionEntry.run_id == run_id, InspectionEntry.actual_date < day)
        .order_by(InspectionEntry.actual_date, InspectionEntry.created_at)
    ).all()

    tool1 = (None, None, None)
    tool2 = (None, None, None)

    for e in entries:
        if (e.tool1_name or e.tool1_serial or e.tool1_calib_due):
            tool1 = (e.tool1_name, e.tool1_serial, e.tool1_calib_due)
        if (e.tool2_name or e.tool2_serial or e.tool2_calib_due):
            tool2 = (e.tool2_name, e.tool2_serial, e.tool2_calib_due)

    tools = []
    if any(tool1):
        tools.append(tool1)
    if any(tool2):
        tools.append(tool2)

    return {"raw": raw, "tools": tools}


def get_current_material_lot_for_slot(session: Session, run_id: int, day: date, slot_time: str):
    """
    Returns MaterialLot or None.
    Rule:
      - Find latest MaterialUseEvent for this run/day with slot_time <= current slot_time
      - If none: fallback to latest event before this day
    """
    # events for same day up to this slot
    ev = session.exec(
        select(MaterialUseEvent)
        .where(
            MaterialUseEvent.run_id == run_id,
            MaterialUseEvent.day == day,
            MaterialUseEvent.slot_time <= slot_time,
        )
        .order_by(MaterialUseEvent.slot_time.desc(), MaterialUseEvent.created_at.desc())
    ).first()

    if not ev:
        # fallback: last event before this day
        ev = session.exec(
            select(MaterialUseEvent)
            .where(
                MaterialUseEvent.run_id == run_id,
                MaterialUseEvent.day < day,
            )
            .order_by(MaterialUseEvent.day.desc(), MaterialUseEvent.slot_time.desc(), MaterialUseEvent.created_at.desc())
        ).first()

    if not ev:
        return None


def resolve_mrr_photo_path(stored_path: str | None) -> str | None:
    """Resolve stored photo path similar to documents.
    Accepts:
      - absolute paths (old data)
      - relative paths like "ticket_15/insp_10/abc.jpg"
      - filenames only
    """
    if not stored_path:
        return None

    stored_path = stored_path.strip()
    if not stored_path:
        return None

    # 1) If it is already an existing absolute path, use it
    if os.path.isabs(stored_path) and os.path.exists(stored_path):
        return stored_path

    # 2) If it's relative, try inside MRR_PHOTO_DIR
    candidate = os.path.join(MRR_PHOTO_DIR, stored_path)
    if os.path.exists(candidate):
        return candidate

    # 3) If only filename was stored, try a recursive search under photos dir (bounded)
    name = os.path.basename(stored_path)
    for root, _dirs, files in os.walk(MRR_PHOTO_DIR):
        if name in files:
            return os.path.join(root, name)

    return None

    return session.get(MaterialLot, ev.lot_id)


def get_day_material_batches(session: Session, run_id: int, day: date) -> list[str]:
    """
    Returns unique batch_no used in THIS day.
    First tries MaterialUseEvent; if none, falls back to InspectionEntry.raw_material_batch_no.
    """

    # 1) Events for this day
    events = session.exec(
        select(MaterialUseEvent)
        .where(MaterialUseEvent.run_id == run_id, MaterialUseEvent.day == day)
        .order_by(MaterialUseEvent.slot_time, MaterialUseEvent.created_at)
    ).all()

    batch_nos: list[str] = []
    seen = set()

    for ev in events:
        lot = session.get(MaterialLot, ev.lot_id)
        bn = (lot.batch_no or "").strip() if lot else ""
        if bn and bn not in seen:
            seen.add(bn)
            batch_nos.append(bn)

    # 2) Fallback: if no events found, read from entries
    if not batch_nos:
        entries = session.exec(
            select(InspectionEntry)
            .where(InspectionEntry.run_id == run_id, InspectionEntry.actual_date == day)
            .order_by(InspectionEntry.created_at)
        ).all()

        for e in entries:
            bn = (e.raw_material_batch_no or "").strip()
            if bn and bn not in seen:
                seen.add(bn)
                batch_nos.append(bn)

    return batch_nos


def format_spec_for_export(rule: str, mn: float | None, mx: float | None):
    """
    Returns (set_value, tolerance_text)
    """
    rule = (rule or "").upper()

    if rule == "RANGE" and mn is not None and mx is not None:
        set_value = (mn + mx) / 2.0
        tol = abs(mx - mn) / 2.0
        return set_value, f"±{tol:g}"

    if rule == "MAX_ONLY" and mx is not None:
        return mx, "max"   # you requested: set=35, tol=max (like temperature max 35C)

    if rule == "MIN_ONLY" and mn is not None:
        return mn, "min"

    # fallback (no spec)
    return None, ""


def _safe_float(x: Optional[str]) -> Optional[float]:
    if x is None:
        return None
    x = str(x).strip()
    if x == "":
        return None
    try:
        return float(x)
    except Exception:
        return None
# =========================
# Excel row maps (template coordinates)
# =========================
# Map param_key (from RunParameter.param_key / edit page "key: ...")
# -> Excel ROW number where that parameter's value belongs in the template.
#
# Liner + Cover template:
#   Spec column = C, Tol column = D (your apply_specs_to_template already uses this)
#
# Reinforcement template:
#   Spec column = D, Tol column = E

ROW_MAP_LINER_COVER = {
    # From your liner screenshot (sheet: In-process (Liner)):
    "length_m": 22,
    "od_mm": 23,
    "wall_thickness_mm": 24,
    "cooling_water_c": 25,
    "line_speed_m_min": 26,
    "tractor_pressure_mpa": 27,

    # If your edit page keys match these (very likely):
    "body_temp_1": 28,
    "body_temp_2": 29,
    "body_temp_3": 30,
    "body_temp_4": 31,
    "body_temp_5": 32,

    "noising_temp_1": 33,
    "noising_temp_2": 34,
    "noising_temp_3": 35,
    "noising_temp_4": 36,
    "noising_temp_5": 37,
}

ROW_MAP_REINF = {
    # From your reinforcement screenshot (sheet: In-process (Reinforcement)):
    "length_m": 21,

    # IMPORTANT:
    # These keys MUST match what your edit page shows as: key: ....
    # If your keys are different, keep the row numbers but rename the keys.
    "annular_od_1_mm": 22,
    "annular_od_2_mm": 23,
    "internal_tensile_od_mm": 24,
    "external_tensile_od_mm": 25,
    "core_mould_dia_mm": 26,
    "annular_width_1_mm": 27,
    "annular_width_2_mm": 28,

    # Next visible row in your screenshot looks like it continues at 29:
    "screw_yarn_width_mm": 29,
}



def apply_specs_to_template(ws, run: ProductionRun, session: Session):
    params = session.exec(select(RunParameter).where(RunParameter.run_id == run.id)).all()

    # pick row map per process
    row_map = ROW_MAP_REINF if run.process == "REINFORCEMENT" else ROW_MAP_LINER_COVER

    # columns per template
    if run.process in ["LINER", "COVER"]:
        SPEC_COL = "C"
        TOL_COL = "D"
    else:  # REINFORCEMENT
        SPEC_COL = "D"
        TOL_COL = "E"

    if (run.status or "").upper() == "APPROVED":
        approved_name = run.approved_by_user_name or ""
        approved_at = run.approved_at_utc

        if approved_at:
            oman_time = approved_at.astimezone(ZoneInfo("Asia/Muscat"))
            approved_at_str = oman_time.strftime("%d-%m-%Y %H:%M")
        else:
            approved_at_str = ""

        # --- Approval stamp cell differs by template/process ---
        proc = (run.process or "").strip().upper()

        # ✅ IMPORTANT FIX: proc is uppercase, so compare uppercase
        if "REINFORCEMENT" in proc:
            # Reinforcement template: merged stamp area starts at M42 (covers M42:M44)
            ws["M42"] = f"Approved by: {approved_name}"
            if approved_at_str:
                ws["M42"] = f"Approved by: {approved_name}\nApproved at: {approved_at_str}"
        else:
            # Liner + Cover templates (your original behavior)
            ws["M43"] = f"Approved by: {approved_name}"
            # If you have a separate cell for date/time in liner/cover, keep it:
            # ws["M44"] = f"Approved at: {approved_at_str}"

    for p in params:
        r = row_map.get(p.param_key)
        if not r:
            continue

        set_val, tol_txt = format_spec_for_export(p.rule, p.min_value, p.max_value)

        _set_cell_safe(ws, f"{SPEC_COL}{r}", set_val if set_val is not None else "")
        _set_cell_safe(ws, f"{TOL_COL}{r}", tol_txt)


# =========================
# HYDRO TESTING
# =========================

def _hydro_record_status_label(status: str) -> str:
    s = (status or "").upper()
    if s == "APPROVED":
        return "Approved"
    if s == "PENDING_APPROVAL":
        return "Pending Approval"
    return "Draft"

def _merge_length_ranges(ranges: list[tuple[float, float]]) -> list[tuple[float, float]]:
    cleaned = []
    for start, end in ranges:
        try:
            a = float(start or 0.0)
            b = float(end or 0.0)
        except Exception:
            continue
        if b <= a:
            continue
        cleaned.append((a, b))

    if not cleaned:
        return []

    cleaned.sort(key=lambda x: (x[0], x[1]))
    merged = [cleaned[0]]

    for start, end in cleaned[1:]:
        last_start, last_end = merged[-1]
        if start <= last_end:
            merged[-1] = (last_start, max(last_end, end))
        else:
            merged.append((start, end))

    return merged


def _subtract_length_ranges(base: list[tuple[float, float]], cuts: list[tuple[float, float]]) -> list[tuple[float, float]]:
    if not base:
        return []
    if not cuts:
        return base[:]

    cuts = _merge_length_ranges(cuts)
    result = []

    for start, end in base:
        pieces = [(start, end)]
        for cstart, cend in cuts:
            new_pieces = []
            for pstart, pend in pieces:
                if cend <= pstart or cstart >= pend:
                    new_pieces.append((pstart, pend))
                    continue

                if cstart > pstart:
                    new_pieces.append((pstart, cstart))
                if cend < pend:
                    new_pieces.append((cend, pend))

            pieces = new_pieces
            if not pieces:
                break

        result.extend(pieces)

    return _merge_length_ranges(result)


def _range_total_length(ranges: list[tuple[float, float]]) -> float:
    total = 0.0
    for start, end in ranges:
        total += max(0.0, float(end) - float(start))
    return round(total, 3)

def _ranges_overlap(a1: float, a2: float, b1: float, b2: float) -> bool:
    return max(a1, b1) < min(a2, b2)

def _clamp_ranges_to_length(ranges: list[tuple[float, float]], max_length: float) -> list[tuple[float, float]]:
    out = []
    max_length = max(0.0, float(max_length or 0.0))

    for start, end in ranges:
        a = max(0.0, min(float(start or 0.0), max_length))
        b = max(0.0, min(float(end or 0.0), max_length))
        if b > a:
            out.append((a, b))

    return _merge_length_ranges(out)


def _get_burst_used_ranges(session: Session, batch_no: str, finished_length_m: float) -> list[tuple[float, float]]:
    reports = session.exec(
        select(BurstTestReport).where(BurstTestReport.batch_no == (batch_no or "").strip())
    ).all()

    report_ids = [r.id for r in reports if r.id is not None]
    if not report_ids:
        return []

    samples = session.exec(
        select(BurstSample).where(BurstSample.report_id.in_(report_ids))
    ).all()

    ranges = []
    for s in samples:
        start = float(getattr(s, "sample_start_m", 0.0) or 0.0)
        length = float(getattr(s, "sample_length_m", 0.0) or 0.0)
        end = start + length
        if end > start:
            ranges.append((start, end))

    return _clamp_ranges_to_length(ranges, finished_length_m)


def _get_hydro_tested_ranges(
    session: Session,
    batch_no: str,
    finished_length_m: float,
    exclude_record_id: int | None = None,
) -> list[tuple[float, float]]:
    rows = session.exec(
        select(HydroTestRecord)
        .where(HydroTestRecord.batch_no == (batch_no or "").strip())
        .order_by(HydroTestRecord.start_length_m.asc())
    ).all()

    ranges = []
    for r in rows:
        if exclude_record_id and r.id == exclude_record_id:
            continue
        ranges.append((float(r.start_length_m or 0.0), float(r.end_length_m or 0.0)))

    return _clamp_ranges_to_length(ranges, finished_length_m)


def _hydro_overlap_error(start_m: float, end_m: float, used_ranges: list[tuple[float, float]]) -> bool:
    for a, b in used_ranges:
        if start_m < b and end_m > a:
            return True
    return False


def hydro_status_badge(status: str) -> str:
    s = (status or "").upper()
    if s == "COMPLETE":
        return "Complete"
    if s == "PARTIAL":
        return "Partially Tested"
    return "Not Started"


def _build_hydro_batch_summary(session: Session, cover_run: ProductionRun) -> dict:
    batch_no = (cover_run.dhtp_batch_no or "").strip()
    produced_length_m = float(get_run_produced_length_m(session, cover_run.id) or 0.0)
    total_order_m = float(getattr(cover_run, "total_length_m", 0.0) or 0.0)

    burst_ranges = _get_burst_used_ranges(session, batch_no, produced_length_m)
    hydro_rows = session.exec(
        select(HydroTestRecord)
        .where(HydroTestRecord.batch_no == batch_no)
        .order_by(HydroTestRecord.created_at.desc())
    ).all()

    hydro_ranges = _get_hydro_tested_ranges(session, batch_no, produced_length_m)
    hydro_allowed_ranges = _subtract_length_ranges([(0.0, produced_length_m)], burst_ranges)
    hydro_remaining_ranges = _subtract_length_ranges(hydro_allowed_ranges, hydro_ranges)

    burst_used_m = _range_total_length(burst_ranges)
    hydro_tested_m = _range_total_length(hydro_ranges)
    hydro_eligible_m = _range_total_length(hydro_allowed_ranges)
    hydro_remaining_m = _range_total_length(hydro_remaining_ranges)

    coverage_pct = 0
    if hydro_eligible_m > 0:
        coverage_pct = int(round(min(100.0, (hydro_tested_m / hydro_eligible_m) * 100.0)))

    if hydro_tested_m <= 0:
        status = "NOT_STARTED"
    elif hydro_remaining_m > 0.0001:
        status = "PARTIAL"
    else:
        status = "COMPLETE"

    segments = []
    if produced_length_m > 0:
        for kind, ranges in [
            ("burst", burst_ranges),
            ("tested", hydro_ranges),
            ("remaining", hydro_remaining_ranges),
        ]:
            for start, end in ranges:
                left = (start / produced_length_m) * 100.0
                width = ((end - start) / produced_length_m) * 100.0
                segments.append({
                    "kind": kind,
                    "start": start,
                    "end": end,
                    "left_pct": left,
                    "width_pct": max(width, 0.8),
                })

    return {
        "batch_no": batch_no,
        "run": cover_run,
        "client_name": getattr(cover_run, "client_name", "") or "",
        "client_po": getattr(cover_run, "po_number", "") or "",
        "pipe_specification": getattr(cover_run, "pipe_specification", "") or "",
        "total_order_m": total_order_m,
        "produced_length_m": produced_length_m,
        "burst_used_m": burst_used_m,
        "hydro_eligible_m": hydro_eligible_m,
        "hydro_tested_m": hydro_tested_m,
        "hydro_remaining_m": hydro_remaining_m,
        "coverage_pct": coverage_pct,
        "status": status,
        "burst_ranges": burst_ranges,
        "hydro_ranges": hydro_ranges,
        "hydro_remaining_ranges": hydro_remaining_ranges,
        "segments": segments,
        "records": hydro_rows,
    }

def _next_hydro_report_no(session: Session) -> str:
    rows = session.exec(
        select(HydroTestRecord).order_by(HydroTestRecord.id.desc())
    ).all()

    max_num = 0
    for r in rows:
        rn = (getattr(r, "report_no", "") or "").strip().upper()
        if rn.startswith("HT-"):
            tail = rn[3:]
            if tail.isdigit():
                max_num = max(max_num, int(tail))

    return f"HT-{max_num + 1:04d}"

def hydro_chart_upload_dir() -> str:
    path = os.path.join("app", "static", "uploads", "hydro_charts")
    os.makedirs(path, exist_ok=True)
    return path


def save_hydro_chart_file(file: UploadFile, batch_no: str) -> str:
    if not file or not file.filename:
        return ""

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in [".png", ".jpg", ".jpeg", ".webp"]:
        return ""

    safe_batch = (batch_no or "batch").replace("/", "_").replace("\\", "_").replace(" ", "_")
    filename = f"{safe_batch}_{datetime.utcnow().strftime('%Y%m%d%H%M%S%f')}{ext}"
    abs_dir = hydro_chart_upload_dir()
    abs_path = os.path.join(abs_dir, filename)

    with open(abs_path, "wb") as out:
        shutil.copyfileobj(file.file, out)

    return f"/static/uploads/hydro_charts/{filename}"

@app.get("/api/hydro/qaqc-users")
def api_hydro_qaqc_users(
    request: Request,
    q: str = "",
    session: Session = Depends(get_session),
):
    get_current_user(request, session)

    users = _hydro_qaqc_candidates(session, q)
    data = [
        {
            "id": u.id,
            "username": u.username,
            "display_name": u.display_name or u.username,
            "role": u.role,
        }
        for u in users
    ]
    return JSONResponse(data)

@app.get("/final", response_class=HTMLResponse)
def final_dashboard(request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)

    cover_runs = get_finished_cover_runs(session)
    summaries = []

    for run in cover_runs:
        summary = _build_final_batch_summary(session, run)
        if summary["produced_length_m"] > 0 or summary["phases"]:
            summaries.append(summary)

    summaries.sort(key=lambda x: x["batch_no"])

    return templates.TemplateResponse(
        "final_dashboard.html",
        {
            "request": request,
            "user": user,
            "summaries": summaries,
            "can_final_approve": _can_final_approve(user),
        },
    )


@app.get("/final/batch/{batch_no}", response_class=HTMLResponse)
def final_batch_view(batch_no: str, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)

    _liner, _reinf, cover = _find_related_runs_by_batch(session, batch_no)
    if not cover:
        raise HTTPException(404, "Cover run not found for this batch")

    summary = _build_final_batch_summary(session, cover)
    error = request.query_params.get("error") or ""

    open_phase = session.exec(
        select(FinalInspectionPhase)
        .where(FinalInspectionPhase.batch_no == (batch_no or "").strip())
        .where(FinalInspectionPhase.status == "DRAFT")
        .order_by(FinalInspectionPhase.phase_no.desc(), FinalInspectionPhase.id.desc())
    ).first()

    return templates.TemplateResponse(
        "final_view.html",
        {
            "request": request,
            "user": user,
            "summary": summary,
            "error": error,
            "open_phase": open_phase,
            "can_final_approve": _can_final_approve(user),
            "format_oman_dt": format_oman_dt,
        },
    )


@app.post("/final/batch/{batch_no}/phase/new")
def final_create_phase(batch_no: str, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)

    _liner, _reinf, cover = _find_related_runs_by_batch(session, batch_no)
    if not cover:
        raise HTTPException(404, "Cover run not found for this batch")

    existing_draft = session.exec(
        select(FinalInspectionPhase)
        .where(FinalInspectionPhase.batch_no == (batch_no or "").strip())
        .where(FinalInspectionPhase.status == "DRAFT")
    ).first()

    if existing_draft:
        return RedirectResponse(
            f"/final/batch/{batch_no}?error=There is already an open draft phase. Use it or submit it first.",
            status_code=303,
        )

    phase_no = _next_final_phase_no(session, batch_no)
    phase = FinalInspectionPhase(
        batch_no=(batch_no or "").strip(),
        linked_cover_run_id=cover.id,
        phase_no=phase_no,
        title=f"Phase {phase_no}",
        status="DRAFT",
        created_by_user_id=getattr(user, "id", None),
        created_by_user_name=getattr(user, "display_name", "") or "",
        
    )
    session.add(phase)
    session.commit()

    return RedirectResponse(f"/final/batch/{batch_no}", status_code=303)


@app.post("/final/batch/{batch_no}/phase/{phase_id}/save-reel")
def final_batch_save_reel(
    batch_no: str,
    phase_id: int,
    request: Request,
    session: Session = Depends(get_session),
    reel_no: str = Form(...),
    start_length_m: float = Form(0.0),
    end_length_m: float = Form(0.0),
    reel_length_m: float = Form(0.0),
    od_mm: float = Form(0.0),
    liner_thickness_mm: float = Form(0.0),
    reinforcement_thickness_mm: float = Form(0.0),
    cover_thickness_mm: float = Form(0.0),
    secured_ok: str = Form(""),
    condition_status: str = Form("GOOD"),
    notes: str = Form(""),
    override_burst_conflict: str = Form(""),
):
    user = get_current_user(request, session)

    phase = session.get(FinalInspectionPhase, phase_id)
    if not phase or (phase.batch_no or "").strip() != (batch_no or "").strip():
        raise HTTPException(404, "Final inspection phase not found")

    if (phase.status or "").upper() != "DRAFT":
        return RedirectResponse(
            f"/final/batch/{batch_no}?error=Only draft phases can be edited.",
            status_code=303,
        )

    _liner, _reinf, cover = _find_related_runs_by_batch(session, batch_no)
    if not cover:
        raise HTTPException(404, "Cover run not found for this batch")

    summary = _build_final_batch_summary(session, cover)

    start_m = float(start_length_m or 0.0)
    end_m = float(end_length_m or 0.0)
    reel_len = float(reel_length_m or 0.0)

    if end_m > 0 and end_m <= start_m:
        return RedirectResponse(
            f"/final/batch/{batch_no}?error=End length must be greater than start length.",
            status_code=303,
        )

    if start_m > 0 or end_m > 0:
        calculated_len = round(end_m - start_m, 3)
        if reel_len <= 0:
            reel_len = calculated_len
        elif abs(reel_len - calculated_len) > 0.05:
            return RedirectResponse(
                f"/final/batch/{batch_no}?error=Reel length does not match start/end range.",
                status_code=303,
            )

    if reel_len <= 0:
        return RedirectResponse(
            f"/final/batch/{batch_no}?error=Reel length must be greater than zero.",
            status_code=303,
        )

    burst_conflict = False
    if end_m > start_m:
        for b1, b2 in summary["burst_ranges"]:
            if _ranges_overlap(start_m, end_m, float(b1), float(b2)):
                burst_conflict = True
                break

    if burst_conflict and override_burst_conflict != "1":
        return RedirectResponse(
            f"/final/batch/{batch_no}?error=This reel range overlaps a burst-used length. Tick override only if you are sure.",
            status_code=303,
        )

    produced_confirmed = float(getattr(cover, "confirmed_total_length_m", 0.0) or 0.0)
    expected_after_burst = float(summary["expected_after_burst_m"] or 0.0)

    if produced_confirmed > 0:
        current_all = float(summary["all_reels_total_m"] or 0.0)
        if current_all + reel_len > expected_after_burst + 0.0001:
            return RedirectResponse(
                f"/final/batch/{batch_no}?error=Total inspected reels exceed confirmed available usable length ({expected_after_burst:.1f} m after burst deduction).",
                status_code=303,
            )

    od_val = float(od_mm or 0.0)
    liner_thk = float(liner_thickness_mm or 0.0)
    reinf_thk = float(reinforcement_thickness_mm or 0.0)
    cover_thk = float(cover_thickness_mm or 0.0)
    total_wall_thk = liner_thk + reinf_thk + cover_thk

    reel = FinalInspectionReel(
        phase_id=phase.id,
        batch_no=(batch_no or "").strip(),
        linked_cover_run_id=cover.id,
        reel_no=(reel_no or "").strip(),
        start_length_m=start_m,
        end_length_m=end_m,
        reel_length_m=reel_len,
        od_mm=od_val,
        wall_thickness_mm=total_wall_thk,
        liner_thickness_mm=liner_thk,
        reinforcement_thickness_mm=reinf_thk,
        cover_thickness_mm=cover_thk,
        secured_ok=(secured_ok == "1"),
        condition_status=(condition_status or "GOOD").strip().upper(),
        notes=(notes or "").strip(),
        created_by_user_id=getattr(user, "id", None),
        created_by_user_name=getattr(user, "display_name", "") or "",
    )
    session.add(reel)
    session.commit()

    return RedirectResponse(f"/final/batch/{batch_no}", status_code=303)

@app.post("/final/batch/{batch_no}/phase/{phase_id}/save-bulk")
async def final_batch_save_bulk_reels(
    batch_no: str,
    phase_id: int,
    request: Request,
    session: Session = Depends(get_session),

    row_count: int = Form(10),

    default_od_mm: float = Form(0.0),
    default_liner_thickness_mm: float = Form(0.0),
    default_reinforcement_thickness_mm: float = Form(0.0),
    default_cover_thickness_mm: float = Form(0.0),
    default_condition_status: str = Form("GOOD"),
    default_secured_ok: str = Form(""),
    default_override_burst_conflict: str = Form(""),
):
    user = get_current_user(request, session)

    phase = session.get(FinalInspectionPhase, phase_id)
    if not phase or (phase.batch_no or "").strip() != (batch_no or "").strip():
        raise HTTPException(404, "Final inspection phase not found")

    if (phase.status or "").upper() != "DRAFT":
        return RedirectResponse(
            f"/final/batch/{batch_no}?error=Only draft phases can be edited.",
            status_code=303,
        )

    _liner, _reinf, cover = _find_related_runs_by_batch(session, batch_no)
    if not cover:
        raise HTTPException(404, "Cover run not found for this batch")

    summary = _build_final_batch_summary(session, cover)
    produced_confirmed = float(getattr(cover, "confirmed_total_length_m", 0.0) or 0.0)
    expected_after_burst = float(summary["expected_after_burst_m"] or 0.0)

    try:
        row_count = max(1, min(int(row_count or 10), 100))
    except Exception:
        row_count = 10

    form = await request.form()

    def _f(name: str, default=""):
        return form.get(name, default)

    def _float(name: str, default=0.0):
        try:
            v = _f(name, "")
            if v in [None, ""]:
                return float(default)
            return float(v)
        except Exception:
            return float(default)

    def _bool(name: str, default=False):
        v = _f(name, "")
        if v in [None, ""]:
            return default
        return str(v).strip() in ["1", "true", "True", "on", "yes", "YES"]

    default_od = float(default_od_mm or 0.0)
    default_liner = float(default_liner_thickness_mm or 0.0)
    default_reinf = float(default_reinforcement_thickness_mm or 0.0)
    default_cover = float(default_cover_thickness_mm or 0.0)
    default_condition = (default_condition_status or "GOOD").strip().upper()
    default_secured = str(default_secured_ok or "").strip() in ["1", "true", "True", "on", "yes", "YES"]
    default_override = str(default_override_burst_conflict or "").strip() in ["1", "true", "True", "on", "yes", "YES"]

    current_all = float(summary["all_reels_total_m"] or 0.0)
    new_reels = []
    errors = []

    for i in range(row_count):
        reel_no = str(_f(f"reel_no_{i}", "") or "").strip()
        if not reel_no:
            continue

        start_m = _float(f"start_length_m_{i}", 0.0)
        end_m = _float(f"end_length_m_{i}", 0.0)
        reel_len = _float(f"reel_length_m_{i}", 0.0)

        od_val = _float(f"od_mm_{i}", default_od)
        liner_thk = _float(f"liner_thickness_mm_{i}", default_liner)
        reinf_thk = _float(f"reinforcement_thickness_mm_{i}", default_reinf)
        cover_thk = _float(f"cover_thickness_mm_{i}", default_cover)

        condition_val = str(_f(f"condition_status_{i}", default_condition) or default_condition).strip().upper()
        secured_val = _bool(f"secured_ok_{i}", default_secured)
        override_val = _bool(f"override_burst_conflict_{i}", default_override)
        notes_val = str(_f(f"notes_{i}", "") or "").strip()

        if end_m > 0 and end_m <= start_m:
            errors.append(f"Row {i+1}: End length must be greater than start length.")
            continue

        if start_m > 0 or end_m > 0:
            reel_len = round(end_m - start_m, 3)

        if reel_len <= 0:
            errors.append(f"Row {i+1}: Reel length must be greater than zero.")
            continue

        if reel_len <= 0:
            errors.append(f"Row {i+1}: Reel length must be greater than zero.")
            continue

        burst_conflict = False
        if end_m > start_m:
            for b1, b2 in summary["burst_ranges"]:
                if _ranges_overlap(start_m, end_m, float(b1), float(b2)):
                    burst_conflict = True
                    break

        if burst_conflict and not override_val:
            errors.append(f"Row {i+1}: Reel range overlaps burst-used length. Tick override if confirmed.")
            continue

        if produced_confirmed > 0:
            if current_all + reel_len > expected_after_burst + 0.0001:
                errors.append(
                    f"Row {i+1}: Total inspected reels exceed confirmed available usable length ({expected_after_burst:.1f} m)."
                )
                continue

        total_wall_thk = liner_thk + reinf_thk + cover_thk

        new_reels.append(
            FinalInspectionReel(
                phase_id=phase.id,
                batch_no=(batch_no or "").strip(),
                linked_cover_run_id=cover.id,
                reel_no=reel_no,
                start_length_m=start_m,
                end_length_m=end_m,
                reel_length_m=reel_len,
                od_mm=od_val,
                wall_thickness_mm=total_wall_thk,
                liner_thickness_mm=liner_thk,
                reinforcement_thickness_mm=reinf_thk,
                cover_thickness_mm=cover_thk,
                secured_ok=secured_val,
                condition_status=condition_val,
                notes=notes_val,
                created_by_user_id=getattr(user, "id", None),
                created_by_user_name=getattr(user, "display_name", "") or "",
            )
        )
        current_all += reel_len

    if errors and not new_reels:
        return RedirectResponse(
            f"/final/batch/{batch_no}?error={errors[0]}",
            status_code=303,
        )

    for reel in new_reels:
        session.add(reel)
    session.commit()

    if errors:
        return RedirectResponse(
            f"/final/batch/{batch_no}?error=Saved {len(new_reels)} reel(s). First skipped issue: {errors[0]}",
            status_code=303,
        )

    return RedirectResponse(f"/final/batch/{batch_no}", status_code=303)


@app.post("/final/batch/{batch_no}/phase/{phase_id}/delete-reel/{reel_id}")
def final_batch_delete_reel(batch_no: str, phase_id: int, reel_id: int, request: Request, session: Session = Depends(get_session)):
    get_current_user(request, session)

    phase = session.get(FinalInspectionPhase, phase_id)
    if not phase or (phase.batch_no or "").strip() != (batch_no or "").strip():
        raise HTTPException(404, "Final inspection phase not found")

    if (phase.status or "").upper() != "DRAFT":
        return RedirectResponse(
            f"/final/batch/{batch_no}?error=Only draft phases can be edited.",
            status_code=303,
        )

    reel = session.get(FinalInspectionReel, reel_id)
    if not reel or reel.phase_id != phase_id:
        raise HTTPException(404, "Reel not found")

    session.delete(reel)
    session.commit()
    return RedirectResponse(f"/final/batch/{batch_no}", status_code=303)


@app.post("/final/batch/{batch_no}/phase/{phase_id}/submit")
def final_submit_phase(batch_no: str, phase_id: int, request: Request, session: Session = Depends(get_session)):
    get_current_user(request, session)

    phase = session.get(FinalInspectionPhase, phase_id)
    if not phase or (phase.batch_no or "").strip() != (batch_no or "").strip():
        raise HTTPException(404, "Final inspection phase not found")

    if (phase.status or "").upper() != "DRAFT":
        return RedirectResponse(
            f"/final/batch/{batch_no}?error=Only draft phases can be submitted.",
            status_code=303,
        )

    reels = session.exec(
        select(FinalInspectionReel).where(FinalInspectionReel.phase_id == phase.id)
    ).all()
    if not reels:
        return RedirectResponse(
            f"/final/batch/{batch_no}?error=Add at least one reel before submit.",
            status_code=303,
        )

    phase.status = "PENDING_APPROVAL"
    phase.submitted_at = datetime.utcnow()
    session.add(phase)
    session.commit()

    return RedirectResponse(f"/final/batch/{batch_no}", status_code=303)


@app.post("/final/batch/{batch_no}/phase/{phase_id}/approve")
def final_approve_phase(batch_no: str, phase_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)
    if not _can_final_approve(user):
        raise HTTPException(403, "Manager only")

    phase = session.get(FinalInspectionPhase, phase_id)
    if not phase or (phase.batch_no or "").strip() != (batch_no or "").strip():
        raise HTTPException(404, "Final inspection phase not found")

    if (phase.status or "").upper() != "PENDING_APPROVAL":
        return RedirectResponse(
            f"/final/batch/{batch_no}?error=Only pending phases can be approved.",
            status_code=303,
        )

    _liner, _reinf, cover = _find_related_runs_by_batch(session, batch_no)
    if not cover:
        raise HTTPException(404, "Cover run not found for this batch")

    summary = _build_final_batch_summary(session, cover)
    produced_confirmed = float(getattr(cover, "confirmed_total_length_m", 0.0) or 0.0)

    if produced_confirmed > 0:
        expected_after_burst = float(summary["expected_after_burst_m"] or 0.0)

        already_approved = 0.0
        for row in summary["phases"]:
            p = row["phase"]
            if p.id != phase.id and (p.status or "").upper() == "APPROVED":
                already_approved += float(row["phase_total_m"] or 0.0)

        this_phase_total = _final_phase_total_length_m(session, phase.id)
        if already_approved + this_phase_total > expected_after_burst + 0.0001:
            return RedirectResponse(
                f"/final/batch/{batch_no}?error=Approving this phase would exceed confirmed usable length ({expected_after_burst:.1f} m).",
                status_code=303,
            )

    phase.status = "APPROVED"
    phase.approved_at = datetime.utcnow()
    phase.approved_by_user_id = getattr(user, "id", None)
    phase.approved_by_user_name = (getattr(user, "display_name", "") or getattr(user, "username", "") or "").strip()
    session.add(phase)
    session.commit()

    return RedirectResponse(f"/final/batch/{batch_no}", status_code=303)


@app.post("/final/batch/{batch_no}/phase/{phase_id}/reopen")
def final_reopen_phase(batch_no: str, phase_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)
    if not _can_final_approve(user):
        raise HTTPException(403, "Manager only")

    phase = session.get(FinalInspectionPhase, phase_id)
    if not phase or (phase.batch_no or "").strip() != (batch_no or "").strip():
        raise HTTPException(404, "Final inspection phase not found")

    phase.status = "DRAFT"
    phase.approved_at = None
    phase.approved_by_user_id = None
    phase.approved_by_user_name = ""
    session.add(phase)
    session.commit()

    return RedirectResponse(f"/final/batch/{batch_no}", status_code=303)


@app.post("/final/batch/{batch_no}/note")
def final_batch_save_note(
    batch_no: str,
    request: Request,
    session: Session = Depends(get_session),
    extra_loss_note: str = Form(""),
):
    user = get_current_user(request, session)

    _liner, _reinf, cover = _find_related_runs_by_batch(session, batch_no)
    if not cover:
        raise HTTPException(404, "Cover run not found for this batch")

    note = session.exec(
        select(FinalInspectionBatchNote).where(FinalInspectionBatchNote.batch_no == (batch_no or "").strip())
    ).first()

    if not note:
        note = FinalInspectionBatchNote(
            batch_no=(batch_no or "").strip(),
            linked_cover_run_id=cover.id,
        )

    note.extra_loss_note = (extra_loss_note or "").strip()
    note.updated_by_user_id = getattr(user, "id", None)
    note.updated_by_user_name = getattr(user, "display_name", "") or ""
    note.updated_at = datetime.utcnow()

    session.add(note)
    session.commit()

    return RedirectResponse(f"/final/batch/{batch_no}", status_code=303)


@app.get("/final/batch/{batch_no}/phase/{phase_id}/pdf")
def final_phase_pdf(batch_no: str, phase_id: int, session: Session = Depends(get_session)):
    phase = session.get(FinalInspectionPhase, phase_id)
    if not phase or (phase.batch_no or "").strip() != (batch_no or "").strip():
        raise HTTPException(404, "Final inspection phase not found")

    _liner, _reinf, cover = _find_related_runs_by_batch(session, batch_no)
    if not cover:
        raise HTTPException(404, "Cover run not found for this batch")

    reels = session.exec(
        select(FinalInspectionReel)
        .where(FinalInspectionReel.phase_id == phase.id)
        .order_by(FinalInspectionReel.created_at.asc(), FinalInspectionReel.id.asc())
    ).all()

    batch_summary = _build_final_batch_summary(session, cover)
    phase_total = sum(float(r.reel_length_m or 0.0) for r in reels)

    from io import BytesIO
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        Image,
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=14 * mm,
        rightMargin=14 * mm,
        topMargin=12 * mm,
        bottomMargin=14 * mm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "title",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=20,
        alignment=1,
        spaceAfter=6,
        textColor=colors.HexColor("#000000"),
    )
    body_style = ParagraphStyle(
        "body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10,
        leading=13,
        textColor=colors.HexColor("#111827"),
    )
    body_bold = ParagraphStyle(
        "body_bold",
        parent=body_style,
        fontName="Helvetica-Bold",
    )
    section_style = ParagraphStyle(
        "sec",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#111827"),
        spaceBefore=8,
        spaceAfter=6,
    )

    story = []

    # logo
    logo_candidates = [
        os.path.join("app", "static", "img", "logo.png"),
        os.path.join("app", "static", "logo.png"),
        os.path.join("app", "static", "images", "logo.png"),
    ]
    logo_path = next((p for p in logo_candidates if os.path.exists(p)), None)

    if logo_path:
        try:
            story.append(Image(logo_path, width=65 * mm, height=20 * mm))
            story.append(Spacer(1, 2 * mm))
        except Exception:
            pass

    story.append(Paragraph("Final Inspection Certificate", title_style))
    story.append(Spacer(1, 2 * mm))

    info_data = [
        [
            Paragraph(f"Batch: <b>{batch_no}</b>", body_style),
            Paragraph(f"Phase: <b>{phase.title or ('Phase ' + str(phase.phase_no))}</b>", body_style),
        ],
        [
            Paragraph(f"Client: {batch_summary['client_name'] or '-'}", body_style),
            Paragraph(f"No. of Reels: {len(reels)}", body_style),
        ],
        [
            Paragraph(f"PO: {batch_summary['client_po'] or '-'}", body_style),
            Paragraph(f"Phase Total Length: {phase_total:.1f} m", body_style),
        ],
        [
            Paragraph(f"Specification: {batch_summary['pipe_specification'] or '-'}", body_style),
            Paragraph("", body_style),
        ],
    ]

    info_tbl = Table(info_data, colWidths=[88 * mm, 78 * mm])
    info_tbl.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]))
    story.append(info_tbl)
    story.append(Spacer(1, 3 * mm))

    report_generated_at = datetime.utcnow()
    story.append(Paragraph(f"Phase Created: {format_oman_dt(phase.created_at)} (Oman)", body_style))
    story.append(Paragraph(f"Report Generated: {format_oman_dt(report_generated_at)} (Oman)", body_style))
    story.append(Spacer(1, 2 * mm))

    if (phase.status or "").upper() == "APPROVED" and phase.approved_by_user_name and phase.approved_at:
        story.append(Paragraph("Approval", section_style))
        story.append(Paragraph(f"Digitally signed by: <b>{phase.approved_by_user_name}</b>", body_style))
        story.append(Paragraph(f"Date/Time: {format_oman_dt(phase.approved_at)} (Oman)", body_style))
        story.append(Spacer(1, 4 * mm))

    story.append(Paragraph("Reel List", section_style))

    rows = [[
        "ID", "Reel No", "Start (m)", "End (m)", "Length (m)", "OD (mm)", "Liner Thick. (mm)", "Reinf. Thick. (mm)", "Cover Thick. (mm)", "Secured", "Condition"
    ]]

    for r in reels:
        rows.append([
            f"#{r.id}",
            r.reel_no or "-",
            f"{float(getattr(r, 'start_length_m', 0.0) or 0.0):.1f}",
            f"{float(getattr(r, 'end_length_m', 0.0) or 0.0):.1f}",
            f"{float(r.reel_length_m or 0.0):.1f}",
            f"{float(r.od_mm or 0.0):.2f}",
            f"{float(r.liner_thickness_mm or 0.0):.2f}",
            f"{float(r.reinforcement_thickness_mm or 0.0):.2f}",
            f"{float(r.cover_thickness_mm or 0.0):.2f}",
            "YES" if r.secured_ok else "NO",
            r.condition_status or "-",
        ])

    if len(rows) == 1:
        rows.append(["-", "-", "-", "-", "-", "-", "-", "-", "-", "-", "-"])

    tbl = Table(
        rows,
        colWidths=[12*mm, 20*mm, 14*mm, 14*mm, 16*mm, 14*mm, 13*mm, 13*mm, 13*mm, 16*mm, 21*mm],
        repeatRows=1
    )
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbeafe")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("BOX", (0, 0), (-1, -1), 0.75, colors.HexColor("#cbd5e1")),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(tbl)

    def _add_footer(canvas_obj, doc_obj):
        canvas_obj.saveState()
        canvas_obj.setFont("Helvetica", 9)
        canvas_obj.setFillColor(colors.grey)
        canvas_obj.drawString(14 * mm, 8 * mm, "QAP0700-F03")
        canvas_obj.drawRightString(A4[0] - 14 * mm, 8 * mm, f"Page {doc_obj.page}")
        canvas_obj.restoreState()

    doc.build(story, onFirstPage=_add_footer, onLaterPages=_add_footer)

    pdf_bytes = buf.getvalue()
    filename = f"final_phase_{batch_no}_p{phase.phase_no}.pdf"
    return StreamingResponse(
        BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'inline; filename=\"{filename}\"'},
    )


@app.get("/final/batch/{batch_no}/pdf")
def final_batch_pdf(batch_no: str, session: Session = Depends(get_session)):
    _liner, _reinf, cover = _find_related_runs_by_batch(session, batch_no)
    if not cover:
        raise HTTPException(404, "Cover run not found for this batch")

    summary = _build_final_batch_summary(session, cover)

    from io import BytesIO
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=12*mm, rightMargin=12*mm, topMargin=12*mm, bottomMargin=12*mm)
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle("title", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=16, leading=20, spaceAfter=8)
    body_style = ParagraphStyle("body", parent=styles["BodyText"], fontName="Helvetica", fontSize=9, leading=12)
    section_style = ParagraphStyle("sec", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=10, leading=12, spaceBefore=8, spaceAfter=6)

    story = []
    story.append(Paragraph("Final Inspection Full Batch Summary", title_style))
    story.append(Paragraph(f"Batch: <b>{summary['batch_no']}</b>", body_style))
    story.append(Paragraph(f"Client: {summary['client_name'] or '-'}", body_style))
    story.append(Paragraph(f"PO: {summary['client_po'] or '-'}", body_style))
    report_generated_at = datetime.utcnow()

    story.append(Paragraph(f"Specification: {summary['pipe_specification'] or '-'}", body_style))
    story.append(Paragraph(f"Report Generated: {format_oman_dt(report_generated_at)} (Oman)", body_style))
    story.append(Spacer(1, 4 * mm))


    batch_rows = [
        ["Produced Length", f"{float(summary['produced_length_m'] or 0.0):.1f} m"],
        ["Burst Used", f"{float(summary['burst_used_m'] or 0.0):.1f} m"],
        ["Expected After Burst", f"{float(summary['expected_after_burst_m'] or 0.0):.1f} m"],
        ["Approved Total", f"{float(summary['approved_total_m'] or 0.0):.1f} m"],
        ["All Phase Reel Total", f"{float(summary['all_reels_total_m'] or 0.0):.1f} m"],
        ["Remaining vs Confirmed", f"{float(summary['remaining_vs_confirmed_m'] or 0.0):.1f} m"],
    ]
    batch_tbl = Table(batch_rows, colWidths=[65*mm, 45*mm])
    batch_tbl.setStyle(TableStyle([
        ("BOX", (0, 0), (-1, -1), 0.75, colors.HexColor("#cbd5e1")),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(batch_tbl)
    story.append(Spacer(1, 5 * mm))

    story.append(Paragraph("Phase Summary", section_style))
    phase_rows = [["Phase", "Status", "Reels", "Length (m)", "Approved By", "Approved At"]]
    for row in summary["phases"]:
        p = row["phase"]
        phase_rows.append([
            p.title or f"Phase {p.phase_no}",
            p.status or "-",
            str(row["reel_count"]),
            f"{float(row['phase_total_m'] or 0.0):.1f}",
            p.approved_by_user_name or "-",
            format_oman_dt(p.approved_at) if p.approved_at else "-",
        ])

    if len(phase_rows) == 1:
        phase_rows.append(["-", "-", "-", "-", "-", "-"])

    phase_tbl = Table(phase_rows, colWidths=[28*mm, 28*mm, 18*mm, 22*mm, 45*mm, 35*mm], repeatRows=1)
    phase_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e2e8f0")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("BOX", (0, 0), (-1, -1), 0.75, colors.HexColor("#cbd5e1")),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(phase_tbl)
    story.append(Spacer(1, 5 * mm))

    story.append(Paragraph("Loss / Justification Note", section_style))
    story.append(Paragraph(
        f"Auto Reason: Burst test consumption: {float(summary['burst_used_m'] or 0.0):.1f} m",
        body_style
    ))
    story.append(Spacer(1, 2 * mm))
    story.append(Paragraph(
        summary["batch_note"].extra_loss_note if summary["batch_note"] and summary["batch_note"].extra_loss_note else "No additional loss note provided.",
        body_style
    ))

    doc.build(story)

    pdf_bytes = buf.getvalue()
    filename = f"final_batch_{batch_no}.pdf"
    return StreamingResponse(
        BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'inline; filename=\"{filename}\"'},
    )


@app.get("/hydro", response_class=HTMLResponse)
def hydro_dashboard(
    request: Request,
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)

    cover_runs = get_finished_cover_runs(session)
    summaries = []

    for run in cover_runs:
        summary = _build_hydro_batch_summary(session, run)
        if summary["produced_length_m"] > 0 or summary["records"]:
            summaries.append(summary)

    summaries.sort(key=lambda x: x["batch_no"])

    can_hydro_approve = _can_hydro_approve(user)

    my_pending_approvals = []
    if can_hydro_approve and user:
        my_pending_approvals = session.exec(
            select(HydroTestRecord)
            .where(HydroTestRecord.approval_status == "PENDING_APPROVAL")
            .where(HydroTestRecord.assigned_qaqc_user_id == user.id)
            .order_by(HydroTestRecord.submitted_at.desc())
        ).all()

    return templates.TemplateResponse(
        "hydro_dashboard.html",
        {
            "request": request,
            "user": user,
            "summaries": summaries,
            "can_hydro_approve": can_hydro_approve,
            "my_pending_approvals": my_pending_approvals,
            "hydro_status_badge": hydro_status_badge,
        },
    )
    
@app.get("/hydro/batch/{batch_no}", response_class=HTMLResponse)
def hydro_batch_view(batch_no: str, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)
    _liner, _reinf, cover = _find_related_runs_by_batch(session, batch_no)
    if not cover:
        raise HTTPException(404, "Cover run not found for this batch")

    summary = _build_hydro_batch_summary(session, cover)

    error = request.query_params.get("error") or ""
    edit_id_raw = request.query_params.get("edit_id") or ""
    edit_record = None

    if edit_id_raw.isdigit():
        edit_record = session.get(HydroTestRecord, int(edit_id_raw))
        if edit_record and (edit_record.batch_no or "").strip() != (batch_no or "").strip():
            edit_record = None

    return templates.TemplateResponse(
        "hydro_view.html",
        {
            "request": request,
            "user": user,
            "summary": summary,
            "error": error,
            "edit_record": edit_record,
            "hydro_status_badge": hydro_status_badge,
            "hydro_record_status_label": _hydro_record_status_label,
            "format_oman_dt": format_oman_dt,
            "can_hydro_approve": _can_hydro_approve(user),
        },
    )

@app.post("/hydro/batch/{batch_no}/save")
def hydro_batch_save_record(
    batch_no: str,
    request: Request,
    session: Session = Depends(get_session),
    record_id: str = Form(""),
    start_length_m: float = Form(...),
    end_length_m: float = Form(...),
    hydrotest_pressure_mpa: float = Form(0.0),
    hold_time_s: str = Form(""),
    test_medium: str = Form("Water"),
    laboratory_temperature: str = Form(""),
    test_result: str = Form("PASS"),
    technician_name: str = Form(""),
    notes: str = Form(""),
    reference_standard: str = Form("API 15S"),
    reference_dhtp_procedure: str = Form("QAW2000"),
    machine_model: str = Form(""),
    calibration_status: str = Form(""),
    highest_pressure_recorded_mpa: float = Form(0.0),
    lowest_pressure_recorded_mpa: float = Form(0.0),
    pressure_holding_time_min: str = Form(""),
    qaqc_name: str = Form(""),
    assigned_qaqc_user_id: str = Form(""),
    testing_operator_name: str = Form(""),
    chart_image: UploadFile | None = File(None),
):
    user = get_current_user(request, session)
    _liner, _reinf, cover = _find_related_runs_by_batch(session, batch_no)
    if not cover:
        raise HTTPException(404, "Cover run not found for this batch")

    summary = _build_hydro_batch_summary(session, cover)
    produced = float(summary["produced_length_m"] or 0.0)
    start_m = float(start_length_m or 0.0)
    end_m = float(end_length_m or 0.0)

    current_id = int(record_id) if str(record_id).isdigit() else None

    def _redir(msg: str):
        if current_id:
            return RedirectResponse(f"/hydro/batch/{batch_no}?error={msg}&edit_id={current_id}", status_code=303)
        return RedirectResponse(f"/hydro/batch/{batch_no}?error={msg}", status_code=303)

    if produced <= 0:
        return _redir("This batch has no finished produced length yet.")
    if end_m <= start_m:
        return _redir("End length must be greater than start length.")
    if start_m < 0 or end_m > produced:
        return _redir(f"Range must stay inside finished produced length 0.0 to {produced:.1f} m.")
    if _hydro_overlap_error(start_m, end_m, summary["burst_ranges"]):
        return _redir("This range overlaps material already consumed by burst testing.")

    existing_hydro_ranges = _get_hydro_tested_ranges(
        session=session,
        batch_no=batch_no,
        finished_length_m=produced,
        exclude_record_id=current_id,
    )
    if _hydro_overlap_error(start_m, end_m, existing_hydro_ranges):
        return _redir("This range overlaps an already hydrotested range.")

    if current_id:
        record = session.get(HydroTestRecord, current_id)
        if not record or (record.batch_no or "").strip() != (batch_no or "").strip():
            raise HTTPException(404, "Hydrotest record not found")
    else:
        record = HydroTestRecord(
            report_no=_next_hydro_report_no(session),
            batch_no=(batch_no or "").strip(),
            linked_run_id=cover.id,
            created_by_user_id=getattr(user, "id", None),
            created_by_user_name=getattr(user, "display_name", "") or "",
        )

    if not (record.report_no or "").strip():
        record.report_no = _next_hydro_report_no(session)

    record.batch_no = (batch_no or "").strip()
    record.linked_run_id = cover.id
    record.client_name = getattr(cover, "client_name", "") or ""
    record.client_po = getattr(cover, "po_number", "") or ""
    record.pipe_specification = getattr(cover, "pipe_specification", "") or ""
    record.finished_length_m = produced
    record.start_length_m = start_m
    record.end_length_m = end_m
    record.tested_length_m = round(end_m - start_m, 3)

    record.hydrotest_pressure_mpa = float(hydrotest_pressure_mpa or 0.0)
    record.hold_time_s = (hold_time_s or "").strip()
    record.test_medium = (test_medium or "Water").strip()
    record.laboratory_temperature = (laboratory_temperature or "").strip()
    record.test_result = (test_result or "PASS").strip().upper()
    record.technician_name = (technician_name or "").strip()
    record.notes = (notes or "").strip()

    record.reference_standard = (reference_standard or "API 15S").strip()
    record.reference_dhtp_procedure = (reference_dhtp_procedure or "QAW2000").strip()
    record.machine_model = (machine_model or "").strip()
    record.calibration_status = (calibration_status or "").strip()
    record.highest_pressure_recorded_mpa = float(highest_pressure_recorded_mpa or 0.0)
    record.lowest_pressure_recorded_mpa = float(lowest_pressure_recorded_mpa or 0.0)
    record.pressure_holding_time_min = (pressure_holding_time_min or "").strip()
    record.testing_operator_name = (testing_operator_name or "").strip()

    selected_qaqc = None
    if str(assigned_qaqc_user_id).isdigit():
        selected_qaqc = session.get(User, int(assigned_qaqc_user_id))
        if not selected_qaqc:
            return _redir("Selected QA/QC approver was not found.")
        if (selected_qaqc.role or "").strip().upper() not in ["INSPECTOR", "MANAGER"]:
            return _redir("Assigned QA/QC approver must be INSPECTOR or MANAGER.")

    if selected_qaqc:
        record.assigned_qaqc_user_id = selected_qaqc.id
        record.assigned_qaqc_username = selected_qaqc.username or ""
        record.assigned_qaqc_display_name = selected_qaqc.display_name or selected_qaqc.username or ""
        record.qaqc_name = record.assigned_qaqc_display_name
    else:
        manual_qaqc = (qaqc_name or "").strip()
        if manual_qaqc.startswith("@"):
            manual_qaqc = manual_qaqc[1:].strip()

        record.assigned_qaqc_user_id = None
        record.assigned_qaqc_username = ""
        record.assigned_qaqc_display_name = ""
        record.qaqc_name = manual_qaqc

    if chart_image and getattr(chart_image, "filename", ""):
        saved = save_hydro_chart_file(chart_image, batch_no)
        if saved:
            record.chart_image_path = saved

    record.approval_status = "DRAFT"
    record.submitted_at = None
    record.approved_at = None
    record.approved_by_user_id = None
    record.approved_by_user_name = ""

    session.add(record)
    session.commit()

    return RedirectResponse(f"/hydro/batch/{batch_no}", status_code=303)

    return RedirectResponse(f"/hydro/batch/{batch_no}", status_code=303)

@app.post("/hydro/batch/{batch_no}/delete/{record_id}")
def hydro_batch_delete_record(batch_no: str, record_id: int, request: Request, session: Session = Depends(get_session)):
    get_current_user(request, session)

    record = session.get(HydroTestRecord, record_id)
    if not record or (record.batch_no or "").strip() != (batch_no or "").strip():
        raise HTTPException(404, "Hydrotest record not found")

    session.delete(record)
    session.commit()

    return RedirectResponse(f"/hydro/batch/{batch_no}", status_code=303)


@app.post("/hydro/batch/{batch_no}/submit/{record_id}")
def hydro_batch_submit_record(batch_no: str, record_id: int, request: Request, session: Session = Depends(get_session)):
    get_current_user(request, session)

    record = session.get(HydroTestRecord, record_id)
    if not record or (record.batch_no or "").strip() != (batch_no or "").strip():
        raise HTTPException(404, "Hydrotest record not found")

    if not record.assigned_qaqc_user_id:
        return RedirectResponse(
            f"/hydro/batch/{batch_no}?error=Please assign a QA/QC approver before submit.",
            status_code=303,
        )

    record.approval_status = "PENDING_APPROVAL"
    record.submitted_at = datetime.utcnow()
    session.add(record)
    session.commit()

    return RedirectResponse(f"/hydro/batch/{batch_no}", status_code=303)

@app.post("/hydro/batch/{batch_no}/approve/{record_id}")
def hydro_batch_approve_record(batch_no: str, record_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)

    if not _can_hydro_approve(user):
        raise HTTPException(403, "Only INSPECTOR or MANAGER can approve hydro reports.")

    record = session.get(HydroTestRecord, record_id)
    if not record or (record.batch_no or "").strip() != (batch_no or "").strip():
        raise HTTPException(404, "Hydrotest record not found")

    approver_name = (getattr(user, "display_name", "") or getattr(user, "username", "") or "").strip()

    record.approval_status = "APPROVED"
    record.approved_at = datetime.utcnow()
    record.approved_by_user_id = getattr(user, "id", None)
    record.approved_by_user_name = approver_name

    # after approval, QA/QC shown on report becomes the real approver
    record.qaqc_name = approver_name

    session.add(record)
    session.commit()

    return RedirectResponse(f"/hydro/batch/{batch_no}", status_code=303)


@app.get("/hydro/record/{record_id}/pdf")
def hydro_record_pdf(record_id: int, session: Session = Depends(get_session)):
    record = session.get(HydroTestRecord, record_id)
    if not record:
        raise HTTPException(404, "Hydrotest record not found")

    docx_bytes = fill_hydro_template_docx_bytes(record)
    pdf_bytes = docx_bytes_to_pdf_bytes(docx_bytes)

    filename = f"hydro_record_{record_id}.pdf"
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'inline; filename=\"{filename}\"'},
    )

@app.get("/hydro/batch/{batch_no}/pdf")
def hydro_batch_pdf(batch_no: str, session: Session = Depends(get_session)):
    _liner, _reinf, cover = _find_related_runs_by_batch(session, batch_no)
    if not cover:
        raise HTTPException(404, "Cover run not found for this batch")

    summary = _build_hydro_batch_summary(session, cover)

    from io import BytesIO
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    w, h = A4

    def draw_header(title: str):
        y_top = h - 18 * mm
        c.setFont("Helvetica-Bold", 16)
        c.drawCentredString(w / 2, y_top, title)
        c.setStrokeColor(colors.black)
        c.setLineWidth(1)
        c.line(10 * mm, y_top - 5 * mm, w - 10 * mm, y_top - 5 * mm)

    def draw_pair_row(y, left_label, left_value, right_label="", right_value=""):
        c.setFont("Helvetica-Bold", 10)
        c.setFillColor(colors.black)
        c.drawString(12 * mm, y, left_label)
        c.setFont("Helvetica", 10)
        c.drawString(42 * mm, y, str(left_value or ""))

        if right_label:
            c.setFont("Helvetica-Bold", 10)
            c.drawString(95 * mm, y, right_label)
            c.setFont("Helvetica", 10)
            c.drawString(150 * mm, y, str(right_value or ""))

        return y - 8 * mm

    def draw_legend_dot(x, y, fill_color, label):
        c.setFillColor(fill_color)
        c.circle(x, y, 2.4 * mm, stroke=0, fill=1)
        c.setFillColor(colors.black)
        c.setFont("Helvetica-Bold", 9)
        c.drawString(x + 5 * mm, y - 1.5 * mm, label)

    def report_no_text(r):
        return (getattr(r, "report_no", "") or f"HT-{r.id:04d}")

    draw_header(f"Hydro Testing Batch Summary - {batch_no}")
    y = h - 34 * mm

    y = draw_pair_row(y, "Batch No", summary["batch_no"], "Finished Produced", f"{summary['produced_length_m']:.1f} m")
    y = draw_pair_row(y, "Client", summary["client_name"], "Hydrotested", f"{summary['hydro_tested_m']:.1f} m")
    y = draw_pair_row(y, "Client PO", summary["client_po"], "Burst Excluded", f"{summary['burst_used_m']:.1f} m")
    y = draw_pair_row(y, "Pipe Spec", summary["pipe_specification"], "Remaining", f"{summary['hydro_remaining_m']:.1f} m")

    y -= 2 * mm
    c.setFont("Helvetica-Bold", 11)
    c.setFillColor(colors.black)
    c.drawString(12 * mm, y, "Coverage Visual")
    c.setFont("Helvetica-Bold", 11)
    c.drawRightString(w - 12 * mm, y, f"Coverage  {summary['coverage_pct']}%")
    y -= 10 * mm

    # Legend row
    legend_y = y
    draw_legend_dot(20 * mm, y, colors.HexColor("#22c55e"), "Hydrotested")
    draw_legend_dot(72 * mm, y, colors.HexColor("#f59e0b"), "Used in burst")
    draw_legend_dot(128 * mm, y, colors.HexColor("#cbd5e1"), "Remaining")
    
    # leave clear gap under legend before drawing the bar
    y =legend_y - 16 * mm

    # Main coverage bar
    track_x = 12 * mm
    track_y = y
    track_w = w - 24 * mm
    track_h = 12 * mm

    c.setFillColor(colors.HexColor("#e5e7eb"))
    c.roundRect(track_x, track_y, track_w, track_h, 3, stroke=0, fill=1)

    # Draw remaining first, then hydro, then burst on top so burst is very visible
    if summary["produced_length_m"] > 0:
        ordered_kinds = ["remaining", "tested", "burst"]
        for kind in ordered_kinds:
            for seg in summary["segments"]:
                if seg["kind"] != kind:
                    continue

                if kind == "tested":
                    c.setFillColor(colors.HexColor("#22c55e"))
                elif kind == "burst":
                    c.setFillColor(colors.HexColor("#f59e0b"))
                else:
                    c.setFillColor(colors.HexColor("#cbd5e1"))

                sx = track_x + (seg["left_pct"] / 100.0) * track_w
                sw = (seg["width_pct"] / 100.0) * track_w

                # Make burst areas clearly visible, even if tiny
                if kind == "burst":
                    sw = max(sw, 8)  # thicker minimum width for burst zones
                else:
                    sw = max(sw, 2)

                # keep inside track
                if sx + sw > track_x + track_w:
                    sw = (track_x + track_w) - sx

                c.roundRect(sx, track_y, sw, track_h, 2.5, stroke=0, fill=1)

    y -= 16 * mm

    # Records section
    c.setFont("Helvetica-Bold", 11)
    c.setFillColor(colors.black)
    c.drawString(12 * mm, y, "Hydrotest Records")
    y -= 8 * mm

    headers = ["Report No", "Range", "Length", "Pressure", "Holding Period", "Result"]
    xs = [12 * mm, 38 * mm, 78 * mm, 110 * mm, 136 * mm, 176 * mm]

    c.setFont("Helvetica-Bold", 9)
    c.setFillColor(colors.black)
    for i, head in enumerate(headers):
        c.drawString(xs[i], y, head)
    y -= 5 * mm
    c.setLineWidth(0.8)
    c.line(12 * mm, y, w - 12 * mm, y)
    y -= 7 * mm

    for r in summary["records"]:
        if y < 20 * mm:
            c.showPage()
            draw_header(f"Hydro Testing Batch Summary - {batch_no}")
            y = h - 28 * mm

            c.setFont("Helvetica-Bold", 11)
            c.setFillColor(colors.black)
            c.drawString(12 * mm, y, "Hydrotest Records")
            y -= 8 * mm

            c.setFont("Helvetica-Bold", 9)
            for i, head in enumerate(headers):
                c.drawString(xs[i], y, head)
            y -= 5 * mm
            c.line(12 * mm, y, w - 12 * mm, y)
            y -= 7 * mm

        # Keep records text black
        c.setFont("Helvetica", 8.5)
        c.setFillColor(colors.black)
        c.drawString(xs[0], y, report_no_text(r))
        c.drawString(xs[1], y, f"{r.start_length_m:.1f}-{r.end_length_m:.1f} m")
        c.drawString(xs[2], y, f"{r.tested_length_m:.1f} m")
        c.drawString(xs[3], y, f"{float(r.hydrotest_pressure_mpa or 0):.2f}")
        c.drawString(xs[4], y, (getattr(r, "pressure_holding_time_min", "") or getattr(r, "hold_time_s", "") or "-"))
        c.drawString(xs[5], y, (r.test_result or "-"))
        y -= 6.5 * mm

    c.showPage()
    c.save()
    buf.seek(0)

    filename = f"hydro_batch_{batch_no}.pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'inline; filename="{filename}"'},
    )
    
@app.get("/dashboard", response_class=HTMLResponse)
def dashboard(request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)

    # Optional filter (?process=LINER)
    q_process = (request.query_params.get("process") or "").strip().upper()
    if q_process and q_process not in ["LINER", "REINFORCEMENT", "COVER"]:
        q_process = ""

    # Production runs (optionally filtered)
    q = select(ProductionRun).order_by(ProductionRun.created_at.desc())
    if q_process:
        q = q.where(ProductionRun.process == q_process)

    runs = session.exec(q).all()

    grouped: Dict[str, List[ProductionRun]] = {}
    for r in runs:
        grouped.setdefault(r.dhtp_batch_no, []).append(r)

    # ✅ bulk progress map instead of query-per-run
    progress_map = get_progress_map_bulk(session, runs)

    # Build "batch cards" for the dashboard UI
    batch_cards = []
    for batch_no, batch_runs in grouped.items():
        if not batch_no:
            continue

        batch_runs_sorted = sorted(
            batch_runs,
            key=lambda r: (r.created_at or datetime.min),
            reverse=True,
        )
        rep = batch_runs_sorted[0]

        pcts = [progress_map.get(r.id, 0) for r in batch_runs if r.id is not None]
        avg_progress = int(sum(pcts) / max(1, len(pcts)))

        any_open = any((r.status or "").upper() == "OPEN" for r in batch_runs)
        priority = "HIGH PRIORITY" if any_open else ""

        processes = sorted({(r.process or "").strip().upper() for r in batch_runs if r.process})

        batch_cards.append(
            {
                "batch_no": batch_no,
                "client_name": rep.client_name or "",
                "po_number": rep.po_number or "",
                "itp_number": rep.itp_number or "",
                "created_at": rep.created_at,
                "avg_progress": avg_progress,
                "priority": priority,
                "processes": processes,
                "run_count": len(batch_runs),
            }
        )

    batch_cards.sort(key=lambda x: x.get("created_at") or datetime.min, reverse=True)

    # ---------- Process tiles stats ----------
    # Reuse the already-loaded runs when possible
    if q_process:
        all_runs = session.exec(select(ProductionRun).order_by(ProductionRun.created_at.desc())).all()
        all_progress_map = get_progress_map_bulk(session, all_runs)
    else:
        all_runs = runs
        all_progress_map = progress_map

    process_stats = {}
    for proc in ["LINER", "REINFORCEMENT", "COVER"]:
        proc_runs = [r for r in all_runs if (r.process or "").upper() == proc]
        if not proc_runs:
            continue

        open_cnt = sum(1 for r in proc_runs if (r.status or "").upper() == "OPEN")
        closed_cnt = sum(1 for r in proc_runs if (r.status or "").upper() == "CLOSED")
        approved_cnt = sum(1 for r in proc_runs if (r.status or "").upper() == "APPROVED")

        proc_pcts = [all_progress_map.get(r.id, 0) for r in proc_runs if r.id is not None]
        avg_progress = int(sum(proc_pcts) / max(1, len(proc_pcts)))

        process_stats[proc] = {
            "open": open_cnt,
            "closed": closed_cnt,
            "approved": approved_cnt,
            "avg_progress": avg_progress,
            "icon": IMAGE_MAP.get(proc, ""),
        }

    # ---------- MRR status summary ----------
    lots = session.exec(
        select(MaterialLot)
        .where(MaterialLot.status != MRR_CANCELED_STATUS)
        .order_by(MaterialLot.created_at.desc())
    ).all()

    lot_ids = [l.id for l in lots if l and l.id is not None]

    receiving_map = {}
    inspection_map = {}

    if lot_ids:
        receivings = session.exec(
            select(MrrReceiving).where(MrrReceiving.ticket_id.in_(lot_ids))
        ).all()
        receiving_map = {r.ticket_id: r for r in receivings}

        inspections = session.exec(
            select(MrrReceivingInspection).where(MrrReceivingInspection.ticket_id.in_(lot_ids))
        ).all()
        inspection_map = {i.ticket_id: i for i in inspections}

    mrr_pending_docs = 0
    mrr_docs_cleared = 0
    mrr_insp_submitted = 0
    mrr_final_approved = 0

    for lot in lots:
        rec = receiving_map.get(lot.id)
        insp = inspection_map.get(lot.id)

        docs_ok = bool(rec and (rec.inspector_confirmed_po or rec.manager_confirmed_po))
        insp_submitted = bool(insp and insp.inspector_confirmed)
        insp_ok = bool(insp and insp.manager_approved)

        if (lot.status or "").upper() == "APPROVED":
            mrr_final_approved += 1
        else:
            if not docs_ok:
                mrr_pending_docs += 1
            else:
                mrr_docs_cleared += 1

            if insp_submitted and not insp_ok:
                mrr_insp_submitted += 1

    mrr_stats = {
        "pending_docs": mrr_pending_docs,
        "docs_cleared": mrr_docs_cleared,
        "insp_submitted": mrr_insp_submitted,
        "final_approved": mrr_final_approved,
    }

    return templates.TemplateResponse(
        "dashboard.html",
        {
            "request": request,
            "user": user,
            "grouped": grouped,
            "progress_map": progress_map,
            "mrr_stats": mrr_stats,
            "process_stats": process_stats,
            "selected_process": q_process,
            "batch_cards": batch_cards,
        },
    )

@app.get("/search")
def global_search(
    q: str = "",
    session: Session = Depends(get_session),
):
    q = (q or "").strip()
    q_lower = q.lower()

    results = []
    if len(q_lower) < 2:
        return {"results": results}

    # Production batches / runs
    runs = session.exec(
        select(ProductionRun).order_by(ProductionRun.created_at.desc())
    ).all()

    seen_batches = set()

    for r in runs:
        batch_no = (r.dhtp_batch_no or "").strip()
        hay = " ".join([
            batch_no,
            r.client_name or "",
            r.po_number or "",
            r.itp_number or "",
            r.process or "",
            r.pipe_specification or "",
        ]).lower()

        if q_lower in hay and batch_no and batch_no not in seen_batches:
            results.append({
                "type": "batch",
                "title": f"Batch {batch_no}",
                "subtitle": f"{r.client_name or '-'} · PO: {r.po_number or '-'} · ITP: {r.itp_number or '-'}",
                "url": f"/batches/{batch_no}",
            })
            seen_batches.add(batch_no)

    # MRR
    lots = session.exec(
        select(MaterialLot).order_by(MaterialLot.created_at.desc())
    ).all()
    for lot in lots:
        hay = " ".join([
            lot.lot_no or "",
            lot.batch_no or "",
            lot.material_name or "",
            lot.vendor_name or "",
        ]).lower()
        if q_lower in hay:
            results.append({
                "type": "mrr",
                "title": f"MRR {lot.lot_no or '-'}",
                "subtitle": f"Batch: {lot.batch_no or '-'} · {lot.material_name or '-'}",
                "url": f"/mrr",
            })

    # Hydro
    hydro_records = session.exec(
        select(HydroTestRecord).order_by(HydroTestRecord.created_at.desc())
    ).all()
    for rec in hydro_records:
        hay = " ".join([
            rec.batch_no or "",
            rec.report_no or "",
            rec.client_name or "",
        ]).lower()
        if q_lower in hay:
            results.append({
                "type": "hydro",
                "title": f"Hydro {rec.batch_no or '-'}",
                "subtitle": f"Report: {rec.report_no or '-'}",
                "url": f"/hydro/batch/{rec.batch_no}",
            })

    # Burst
    burst_reports = session.exec(
        select(BurstTestReport).order_by(BurstTestReport.created_at.desc())
    ).all()
    for rec in burst_reports:
        hay = " ".join([
            rec.batch_no or "",
            rec.report_no or "",
            rec.client_name or "",
        ]).lower()
        if q_lower in hay:
            results.append({
                "type": "burst",
                "title": f"Burst {rec.batch_no or '-'}",
                "subtitle": f"Report: {rec.report_no or '-'}",
                "url": f"/burst/{rec.id}",
            })

    # Final inspection
    final_phases = session.exec(
        select(FinalInspectionPhase).order_by(FinalInspectionPhase.created_at.desc())
    ).all()
    for phase in final_phases:
        hay = " ".join([
            phase.batch_no or "",
            phase.title or "",
            phase.notes or "",
        ]).lower()
        if q_lower in hay:
            results.append({
                "type": "final",
                "title": f"Final Inspection {phase.batch_no or '-'}",
                "subtitle": f"{phase.title or '-'} · {phase.status or '-'}",
                "url": f"/final/batch/{phase.batch_no}",
            })

    # RFI
    rfi_records = session.exec(
        select(RfiRecord).order_by(RfiRecord.created_at.desc())
    ).all()
    for rfi in rfi_records:
        hay = " ".join([
            rfi.batch_no or "",
            rfi.rfi_no or "",
            rfi.client_name or "",
            rfi.inspection_stage or "",
            rfi.status or "",
        ]).lower()
        if q_lower in hay:
            results.append({
                "type": "rfi",
                "title": f"RFI {rfi.rfi_no or '-'}",
                "subtitle": f"Batch: {rfi.batch_no or '-'} · {rfi.inspection_stage or '-'}",
                "url": f"/rfi/{rfi.id}",
            })

    return {"results": results[:20]}

@app.get("/me", response_class=HTMLResponse)
def my_profile(
    request: Request,
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)

    return templates.TemplateResponse(
        "me.html",
        {
            "request": request,
            "user": user,
            "overtime_summary": {
                "approved_hours": 0,
                "pending_hours": 0,
                "approved_amount": 0,
                "pending_amount": 0,
            },
        },
    )


@app.post("/me/update")
def my_profile_update(
    request: Request,
    display_name: str = Form(""),
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)

    clean_name = (display_name or "").strip()
    if not clean_name:
        raise HTTPException(status_code=400, detail="Display name is required.")

    user.display_name = clean_name
    session.add(user)
    session.commit()

    return RedirectResponse(url="/me", status_code=303)

@app.post("/me/change-password")
def my_profile_change_password(
    request: Request,
    current_password: str = Form(""),
    new_password: str = Form(""),
    confirm_password: str = Form(""),
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)

    if not current_password or not new_password or not confirm_password:
        raise HTTPException(status_code=400, detail="All password fields are required.")

    if not _verify_user_password(current_password, user.password_hash):
        raise HTTPException(status_code=400, detail="Current password is not correct.")

    if len(new_password) < 6:
        raise HTTPException(status_code=400, detail="New password must be at least 6 characters.")

    if new_password != confirm_password:
        raise HTTPException(status_code=400, detail="New password and confirmation do not match.")

    user.password_hash = _hash_user_password(new_password)
    user.must_change_password = False
    session.add(user)
    session.commit()

    return RedirectResponse(url="/me", status_code=303)

@app.get("/batches/{batch_no}", response_class=HTMLResponse)
def batch_detail(batch_no: str, request: Request, session: Session = Depends(get_session)):
    """Batch detail page: shows all production runs (processes) under the same batch."""
    user = get_current_user(request, session)

    batch_no = (batch_no or "").strip()
    if not batch_no:
        return RedirectResponse("/dashboard", status_code=302)

    runs = session.exec(
        select(ProductionRun)
        .where(ProductionRun.dhtp_batch_no == batch_no)
        .order_by(ProductionRun.created_at.desc())
    ).all()

    if not runs:
        return RedirectResponse("/dashboard", status_code=302)

    progress_map = {r.id: get_progress_percent(session, r) for r in runs}

    rep = runs[0]
    pcts = [progress_map.get(r.id, 0) for r in runs if r.id is not None]
    avg_progress = int(sum(pcts) / max(1, len(pcts)))

    return templates.TemplateResponse(
        "batch_detail.html",
        {
            "request": request,
            "user": user,
            "batch_no": batch_no,
            "runs": runs,
            "progress_map": progress_map,
            "avg_progress": avg_progress,
            "client_name": rep.client_name or "",
            "po_number": rep.po_number or "",
            "itp_number": rep.itp_number or "",
        },
    )


from sqlalchemy import or_, cast
from sqlalchemy.types import String as SqlString


@app.get("/mrr", response_class=HTMLResponse)
def mrr_reports_list(request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)

    q = (request.query_params.get("q") or "").strip().lower()

    # show only submitted or approved inspections (real reports)
    reports = session.exec(
        select(MrrReceivingInspection)
        .where(
            (MrrReceivingInspection.inspector_confirmed == True) |
            (MrrReceivingInspection.manager_approved == True)
        )
        .order_by(MrrReceivingInspection.created_at.desc())
    ).all()

    # load their tickets to show Material/Supplier/PO in the list
    ticket_ids = sorted({r.ticket_id for r in reports if r.ticket_id})
    lots = []
    if ticket_ids:
        lots = session.exec(select(MaterialLot).where(MaterialLot.id.in_(ticket_ids))).all()
    lot_map = {l.id: l for l in lots}

    # simple search across report + ticket + ticket fields
    if q:
        filtered = []
        for r in reports:
            lot = lot_map.get(r.ticket_id)
            fields = [
                str(r.report_no or ""),
                str(r.delivery_note_no or ""),
                str(r.ticket_id or ""),
                str(lot.material_name if lot else ""),
                str(lot.supplier_name if lot else ""),
                str(lot.po_number if lot else ""),
                str(lot.batch_no if lot else ""),
            ]
            if any(q in (f.lower()) for f in fields if f):
                filtered.append(r)
        reports = filtered

    return templates.TemplateResponse(
        "mrr_reports_list.html",
        {
            "request": request,
            "user": user,
            "q": request.query_params.get("q") or "",
            "reports": reports,
            "lot_map": lot_map,
        },
    )

@app.get("/mrr/tickets", response_class=HTMLResponse)
def mrr_tickets_list(request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)

    q = (request.query_params.get("q") or "").strip()
    like = f"%{q}%"

    base_stmt = select(MaterialLot).where(MaterialLot.lot_type.in_(["RAW", "OUTSOURCED"]))

    # ---- SEARCH FILTERING ----
    if not q:
        lots = session.exec(base_stmt.order_by(MaterialLot.id.desc())).all()
    else:
        # Search on lots directly
        lot_stmt = base_stmt.where(
            or_(
                cast(MaterialLot.id, SqlString).ilike(like),
                (MaterialLot.material_name or "").ilike(like),
                (MaterialLot.supplier_name or "").ilike(like),
                (MaterialLot.po_number or "").ilike(like),
                (MaterialLot.batch_no or "").ilike(like),
                (MaterialLot.status or "").ilike(like),
            )
        )
        lots_direct = session.exec(lot_stmt.order_by(MaterialLot.id.desc())).all()
        direct_ids = {l.id for l in lots_direct}

        # Search on inspection side (DN, report no, JSON: batches/grade/etc)
        insp_hits = session.exec(
            select(MrrReceivingInspection.ticket_id).where(
                or_(
                    (MrrReceivingInspection.delivery_note_no or "").ilike(like),
                    (MrrReceivingInspection.report_no or "").ilike(like),
                    (MrrReceivingInspection.inspection_json or "").ilike(like),
                )
            )
        ).all()
        insp_ticket_ids = {x for x in insp_hits if x is not None}

        all_ids = list(direct_ids.union(insp_ticket_ids))

        if not all_ids:
            lots = []
        else:
            lots = session.exec(
                select(MaterialLot).where(MaterialLot.id.in_(all_ids)).order_by(MaterialLot.id.desc())
            ).all()

    # ---- BUILD MAPS REQUIRED BY TEMPLATE ----
    lot_ids = [l.id for l in lots]

    receiving_map = {}
    inspection_map = {}
    docs_status_map = {}

    if lot_ids:
        receivings = session.exec(
            select(MrrReceiving).where(MrrReceiving.ticket_id.in_(lot_ids))
        ).all()
        receiving_map = {r.ticket_id: r for r in receivings}

        inspections = session.exec(
            select(MrrReceivingInspection)
            .where(MrrReceivingInspection.ticket_id.in_(lot_ids))
            .order_by(MrrReceivingInspection.created_at.desc())
        ).all()

        # latest inspection per ticket
        for ins in inspections:
            if ins.ticket_id not in inspection_map:
                inspection_map[ins.ticket_id] = ins

        docs = session.exec(
            select(MrrDocument).where(MrrDocument.ticket_id.in_(lot_ids))
        ).all()

        # very lightweight docs status: Pending/Done
        # (Your UI probably expects this)
        docs_by_ticket = {}
        for d in docs:
            docs_by_ticket.setdefault(d.ticket_id, []).append(d)

        for tid in lot_ids:
            docs_list = docs_by_ticket.get(tid, [])
            # Consider docs done if at least one PO exists (since you require PO)
            has_po = any((x.doc_type or "").upper() == "PO" for x in docs_list)
            docs_status_map[tid] = "Done" if has_po else "Pending"

    return templates.TemplateResponse(
        "mrr_list.html",
        {
            "request": request,
            "user": user,
            "lots": lots,
            "q": q,
            "receiving_map": receiving_map,
            "inspection_map": inspection_map,
            "docs_status_map": docs_status_map,
        },
    )




@app.get("/mrr/canceled", response_class=HTMLResponse)
def mrr_list_canceled(request: Request, session: Session = Depends(get_session)):
    """Show canceled (soft-deleted) MRR tickets for audit/review."""
    user = get_current_user(request, session)
    require_manager(user)


    lots = session.exec(
        select(MaterialLot)
        .where(MaterialLot.status == MRR_CANCELED_STATUS)
        .order_by(MaterialLot.id.desc())
    ).all()

    lot_ids = [l.id for l in lots if l and l.id is not None]

    receiving_map = {}
    inspection_map = {}

    if lot_ids:
        receivings = session.exec(
            select(MrrReceiving).where(MrrReceiving.ticket_id.in_(lot_ids))
        ).all()
        receiving_map = {r.ticket_id: r for r in receivings}

        inspections = session.exec(
            select(MrrReceivingInspection).where(MrrReceivingInspection.ticket_id.in_(lot_ids))
        ).all()
        inspection_map = {i.ticket_id: i for i in inspections}

    return templates.TemplateResponse(
        "mrr_list.html",
        {
            "request": request,
            "user": user,
            "lots": lots,
            "receiving_map": receiving_map,
            "inspection_map": inspection_map,
            "error": "",
            "showing_canceled": True,
        },
    )

@app.get("/mrr/new", response_class=HTMLResponse)
def mrr_new_get(request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)
    require_manager(user)

    return templates.TemplateResponse(
        "mrr_new.html",
        {
            "request": request,
            "user": user,
            "error": "",
        },
    )

@app.post("/mrr/new")
async def mrr_new(request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)
    require_manager(user)

    form = await request.form()

    material_name = str(form.get("material_name", "")).strip()
    supplier_name = str(form.get("supplier_name", "")).strip()

    lot_type = str(form.get("lot_type", "RAW")).strip().upper()
    if lot_type not in ["RAW", "OUTSOURCED"]:
        lot_type = "RAW"

    po_number = str(form.get("po_number", "")).strip()

    qty_raw = str(form.get("quantity", "")).strip()
    if qty_raw == "":
        raise HTTPException(400, "PO Quantity is required")
    try:
        quantity = float(qty_raw)
    except Exception:
        raise HTTPException(400, "Invalid PO Quantity")

    quantity_unit = str(form.get("quantity_unit", "KG")).strip().upper()
    if quantity_unit not in ["KG", "T", "PCS"]:
        quantity_unit = "KG"

    lot = MaterialLot(
        lot_type=lot_type,
        batch_no="",  # ✅ DO NOT auto-generate batch here
        material_name=material_name,
        supplier_name=supplier_name,
        po_number=po_number,
        quantity=quantity,
        quantity_unit=quantity_unit,
        received_total=0.0,
        status="PENDING",
    )

    session.add(lot)
    session.commit()

    return RedirectResponse("/mrr", status_code=303)


@app.post("/mrr/{lot_id}/approve")
def mrr_approve(lot_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)
    require_manager(user)

    lot = session.get(MaterialLot, lot_id)
    if not lot:
        raise HTTPException(404, "MRR Ticket not found")

    # ✅ Do NOT approve the lot for production here anymore.
    # This button now only confirms the ticket is valid to proceed (still not production-approved).
    # We'll mark it as PENDING and rely on Receiving Inspection approval to set APPROVED.
    lot.status = "PENDING"
    session.add(lot)
    session.commit()

    return RedirectResponse(f"/mrr/{lot_id}", status_code=303)



@app.post("/mrr/{lot_id}/reject")
def mrr_reject(lot_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)
    require_manager(user)

    lot = session.get(MaterialLot, lot_id)
    if not lot:
        raise HTTPException(404, "MRR Ticket not found")

    lot.status = "REJECTED"
    session.add(lot)
    session.commit()

    return RedirectResponse(f"/mrr/{lot_id}", status_code=303)



@app.post("/mrr/{lot_id}/cancel")
def mrr_cancel(lot_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)
    require_manager(user)

    lot = session.get(MaterialLot, lot_id)
    if not lot:
        raise HTTPException(404, "MRR Ticket not found")

    lot.status = MRR_CANCELED_STATUS
    session.add(lot)
    session.commit()

    return RedirectResponse("/mrr", status_code=303)





    
@app.get("/mrr/{lot_id}", response_class=HTMLResponse)
def mrr_view(lot_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)

    lot = session.get(MaterialLot, lot_id)
    if not lot:
        raise HTTPException(404, "MRR Ticket not found")

    readonly = is_mrr_canceled(lot)

    docs = session.exec(
        select(MrrDocument)
        .where(MrrDocument.ticket_id == lot_id)
        .order_by(MrrDocument.created_at.desc())
    ).all()

    receiving = session.exec(
        select(MrrReceiving).where(MrrReceiving.ticket_id == lot_id)
    ).first()

    # Latest inspection (could be draft OR submitted)
    inspection = session.exec(
        select(MrrReceivingInspection)
        .where(MrrReceivingInspection.ticket_id == lot_id)
        .order_by(MrrReceivingInspection.created_at.desc())
    ).first()

    # Submitted shipments (for showing DN list + photos)
    submitted_shipments = session.exec(
        select(MrrReceivingInspection)
        .where(
            (MrrReceivingInspection.ticket_id == lot_id) &
            (MrrReceivingInspection.inspector_confirmed == True)
        )
        .order_by(MrrReceivingInspection.created_at.asc())
    ).all()

    # Draft inspection (saved but not submitted) - allow resume
    draft_inspection = session.exec(
        select(MrrReceivingInspection)
        .where(
            (MrrReceivingInspection.ticket_id == lot_id) &
            (MrrReceivingInspection.inspector_confirmed == False)
        )
        .order_by(MrrReceivingInspection.created_at.desc())
    ).first()

    # ✅ Build a map: inspection_id -> batch_numbers list (safe for SQLModel)
    batch_numbers_map = {}
    for s in submitted_shipments:
        data = safe_json_loads(getattr(s, "inspection_json", None))
        batch_numbers_map[int(s.id)] = data.get("batch_numbers", []) or []

    

    used_dns = [
        (s.delivery_note_no or "").strip()
        for s in submitted_shipments
        if (s.delivery_note_no or "").strip()
    ]

    # ✅ Needs-approval inspection (latest submitted not approved)
    inspection_to_approve = session.exec(
        select(MrrReceivingInspection)
        .where(
            (MrrReceivingInspection.ticket_id == lot_id) &
            (MrrReceivingInspection.inspector_confirmed == True) &
            (MrrReceivingInspection.manager_approved == False)
        )
        .order_by(MrrReceivingInspection.created_at.desc())
    ).first()

    # ✅ Latest approved inspection (for Unapprove button)
    latest_approved_inspection = session.exec(
        select(MrrReceivingInspection)
        .where(
            (MrrReceivingInspection.ticket_id == lot_id) &
            (MrrReceivingInspection.inspector_confirmed == True) &
            (MrrReceivingInspection.manager_approved == True)
        )
        .order_by(MrrReceivingInspection.created_at.desc())
    ).first()

    # Photos grouped by inspection
    all_photos = session.exec(
        select(MrrInspectionPhoto)
        .where(MrrInspectionPhoto.ticket_id == lot_id)
        .order_by(MrrInspectionPhoto.created_at.asc())
    ).all()

    photos_by_inspection: Dict[int, Dict[str, List[MrrInspectionPhoto]]] = {}
    for p in all_photos:
        photos_by_inspection.setdefault(int(p.inspection_id), {})
        g = (p.group_name or "General").strip() or "General"
        photos_by_inspection[int(p.inspection_id)].setdefault(g, []).append(p)

    docs_ok = bool(
        receiving and (
            getattr(receiving, "inspector_confirmed_po", False) or
            getattr(receiving, "manager_confirmed_po", False)
        )
    )
    insp_submitted = bool(inspection and getattr(inspection, "inspector_confirmed", False))
    insp_ok = bool(inspection and getattr(inspection, "manager_approved", False))

    return templates.TemplateResponse(
        "mrr_view.html",
        {
            "request": request,
            "user": user,
            "lot": lot,
            "readonly": readonly,
            "error": request.query_params.get("error", ""),
            "docs": docs,
            "receiving": receiving,
            "docs_ok": docs_ok,
            "inspection": inspection,
            "insp_submitted": insp_submitted,
            "insp_ok": insp_ok,
            "submitted_shipments": submitted_shipments,
            "used_dns": used_dns,
            "photos_by_inspection": photos_by_inspection,
            "inspection_to_approve": inspection_to_approve,
            "latest_approved_inspection": latest_approved_inspection,
            "batch_numbers_map": batch_numbers_map,
            "draft_inspection": draft_inspection,

        },
    )

@app.get("/mrr/{lot_id}/docs")
def mrr_docs_page(
    lot_id: int,
    request: Request,
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)

    lot = session.get(MaterialLot, lot_id)
    if not lot:
        raise HTTPException(404, "MRR Ticket not found")

    receiving = session.exec(
        select(MrrReceiving).where(MrrReceiving.ticket_id == lot_id)
    ).first()

    docs = session.exec(
        select(MrrDocument).where(
            (MrrDocument.ticket_id == lot_id) &
            (MrrDocument.is_deleted == False)
        ).order_by(MrrDocument.created_at.desc())
    ).all()
    
    deleted_docs = session.exec(
        select(MrrDocument).where(
            (MrrDocument.ticket_id == lot_id) &
            (MrrDocument.is_deleted == True)
        ).order_by(MrrDocument.deleted_at_utc.desc())
    ).all()

    readonly = is_mrr_canceled(lot)

    # Show ONLY submitted/approved shipments in the dropdown
    inspections = session.exec(
        select(MrrReceivingInspection)
        .where(MrrReceivingInspection.ticket_id == lot_id)
        .where(
            (MrrReceivingInspection.inspector_confirmed == True) |
            (MrrReceivingInspection.manager_approved == True)
        )
        .order_by(MrrReceivingInspection.created_at.asc())
    ).all()

    # Defaults
    default_insp_id = None
    draft = None
    try:
        draft = get_latest_draft_shipment(session, lot_id)
    except Exception:
        draft = None

    if draft and getattr(draft, "id", None) is not None:
        default_insp_id = int(draft.id)
    elif inspections:
        default_insp_id = int(inspections[-1].id)

    # NEW: allow preselect from query params
    qp_attach_to = (request.query_params.get("attach_to") or "").strip().upper()
    qp_insp = (request.query_params.get("attach_inspection_id") or "").strip()

    initial_attach_to = qp_attach_to if qp_attach_to in ("AUTO", "TICKET", "SHIPMENT") else "AUTO"
    if qp_insp:
        try:
            picked = int(qp_insp)
            # validate belongs to this ticket
            found = any(int(s.id) == picked for s in inspections)
            if found:
                default_insp_id = picked
                initial_attach_to = "SHIPMENT"
        except Exception:
            pass

    return templates.TemplateResponse(
        "mrr_doc_upload.html",
        {
            "request": request,
            "user": user,
            "lot": lot,
            "receiving": receiving,
            "docs": docs,
            "readonly": readonly,
            "inspections": inspections,
            "default_insp_id": default_insp_id,
            "initial_attach_to": initial_attach_to,   # NEW
        },
    )
@app.post("/mrr/{lot_id}/docs/upload")
async def mrr_doc_upload(
    lot_id: int,
    request: Request,
    session: Session = Depends(get_session),
    doc_type: str = Form(...),
    doc_title: str = Form(""),
    doc_number: str = Form(""),
    attach_to: str = Form("AUTO"),                 # NEW: AUTO / TICKET / SHIPMENT
    attach_inspection_id: str = Form(""),          # NEW: chosen inspection id
    file: UploadFile = File(...),
):
    user = get_current_user(request, session)

    lot = session.get(MaterialLot, lot_id)
    if not lot:
        raise HTTPException(404, "MRR Ticket not found")
    block_if_mrr_canceled(lot)

    safe_original = os.path.basename(file.filename or "upload.bin")
    filename = f"{lot_id}_{int(datetime.utcnow().timestamp())}_{safe_original}"
    abs_path = os.path.join(MRR_UPLOAD_DIR, filename)

    # write file
    with open(abs_path, "wb") as f:
        f.write(await file.read())

    dt = (doc_type or "").strip().upper()

    # For PO docs, doc number is the ticket PO number (auto) and must be ticket-level
    if dt == "PO":
        doc_number = (lot.po_number or "").strip()
        attach_to = "TICKET"
        attach_inspection_id = ""

    # Auto doc name unless RELATED
    title = (doc_title or "").strip()
    if dt != "RELATED":
        title = {
            "PO": "PO Copy",
            "DELIVERY_NOTE": "Delivery Note",
            "COA": "COA / Lab Test",
            "GENERAL": "General Document",
        }.get(dt, dt)

    if dt == "RELATED" and not title:
        raise HTTPException(400, "Document Name is required when type is RELATED")

    # store RELATIVE path (portable)
    rel_path = os.path.relpath(abs_path, BASE_DIR)

    if dt != "PO" and not (doc_number or "").strip():
        # allow GENERAL without number
        if dt not in ("GENERAL", "RELATED"):
            raise HTTPException(400, "Document Number is required")

    # Decide target inspection id based on attach_to
    target_insp_id = None

    mode = (attach_to or "AUTO").strip().upper()

    if mode == "TICKET":
        target_insp_id = None

    elif mode == "SHIPMENT":
        try:
            chosen = int((attach_inspection_id or "").strip())
        except Exception:
            chosen = 0
        if chosen <= 0:
            raise HTTPException(400, "Please select a shipment to attach this document to.")
        # validate shipment belongs to this ticket
        insp = session.get(MrrReceivingInspection, chosen)
        if not insp or insp.ticket_id != lot_id:
            raise HTTPException(400, "Invalid shipment selected.")
        
        # Only allow attaching to SUBMITTED/APPROVED shipments
        if not (insp.inspector_confirmed or insp.manager_approved):
            raise HTTPException(400, "You can only attach to a submitted shipment.")
        
        target_insp_id = chosen

    else:
        # AUTO behavior (your current flow): attach to latest draft shipment if exists
        if dt != "PO":
            draft = get_latest_draft_shipment(session, lot_id)
            if draft and getattr(draft, "id", None) is not None:
                target_insp_id = int(draft.id)
            else:
                # no draft exists yet -> keep ticket-level (will not mis-attach)
                target_insp_id = None

    doc = MrrDocument(
        ticket_id=lot_id,
        inspection_id=target_insp_id,
        doc_type=dt,
        doc_name=title,
        doc_number=(doc_number or "").strip(),
        file_path=rel_path,
        uploaded_by_user_id=user.id,
        uploaded_by_user_name=user.display_name,
    )

    session.add(doc)
    session.commit()

    return RedirectResponse(f"/mrr/{lot_id}/docs", status_code=303)

from fastapi import Form
from datetime import datetime
from fastapi.responses import RedirectResponse

@app.post("/mrr/docs/{doc_id}/trash", response_class=RedirectResponse)
def trash_mrr_doc(
    doc_id: int,
    request: Request,
    confirm_doc_number: str = Form(""),
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)

    doc = session.get(MrrDocument, doc_id)
    if not doc:
        raise HTTPException(404, "Document not found")

    expected = (doc.doc_number or doc.doc_name or "").strip()
    if (confirm_doc_number or "").strip() != expected:
        raise HTTPException(400, f"Type exactly: {expected}")

    if not doc.is_deleted:
        doc.is_deleted = True
        doc.deleted_at_utc = datetime.utcnow()
        doc.deleted_by_user_id = getattr(user, "id", None)
        doc.deleted_by_user_name = (getattr(user, "display_name", None) or getattr(user, "username", None) or "").strip()
        session.add(doc)
        session.commit()

    return RedirectResponse(url=f"/mrr/{doc.ticket_id}/docs", status_code=303)


@app.post("/mrr/docs/{doc_id}/restore", response_class=RedirectResponse)
def restore_mrr_doc(
    doc_id: int,
    request: Request,
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)

    doc = session.get(MrrDocument, doc_id)
    if not doc:
        raise HTTPException(404, "Document not found")

    doc.is_deleted = False
    doc.deleted_at_utc = None
    doc.deleted_by_user_id = None
    doc.deleted_by_user_name = None
    session.add(doc)
    session.commit()

    return RedirectResponse(url=f"/mrr/{doc.ticket_id}/docs", status_code=303)

@app.get("/mrr/docs/{doc_id}/download")
def mrr_doc_download(doc_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)

    doc = session.get(MrrDocument, doc_id)
    if not doc:
        raise HTTPException(404, "File not found")

    real_path = resolve_mrr_doc_path(doc.file_path)
    if not real_path:
        raise HTTPException(404, "File not found")

    return FileResponse(
        real_path,
        filename=os.path.basename(real_path),
    )

@app.get("/mrr/docs/{doc_id}/inline")
def mrr_doc_inline(doc_id: int, request: Request, session: Session = Depends(get_session)):
    """
    Open document in browser (inline). Used by the "View" button in templates.

    Important:
    - Always resolve the stored file_path via resolve_mrr_doc_path(), because old records might store
      absolute paths from a different machine, or relative paths.
    """
    user = get_current_user(request, session)
    forbid_none(user)

    doc = session.get(MrrDocument, doc_id)
    if not doc:
        raise HTTPException(404, "Document not found")

    resolved = resolve_mrr_doc_path(doc.file_path)
    if not resolved or not os.path.exists(resolved):
        raise HTTPException(404, "File not found")

    media_type, _ = mimetypes.guess_type(resolved)
    media_type = media_type or "application/octet-stream"
    return FileResponse(
        resolved,
        media_type=media_type,
        filename=os.path.basename(resolved),
        headers={"Content-Disposition": f'inline; filename="{os.path.basename(resolved)}"'},
    )


@app.get("/mrr/docs/{doc_id}/view")
def mrr_doc_view(
    doc_id: int,
    request: Request,
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)

    d = session.get(MrrDocument, doc_id)
    if not d:
        raise HTTPException(404, "Document not found")

    resolved = resolve_mrr_doc_path(getattr(d, "file_path", "") or "")
    if not resolved or not os.path.exists(resolved):
        raise HTTPException(404, "File missing on server")

    # Detect content type
    ctype, _ = mimetypes.guess_type(resolved)
    if not ctype:
        ctype = "application/octet-stream"

    # Decide inline vs download:
    # - PDFs & images open inline
    # - everything else downloads
    ext = os.path.splitext(resolved)[1].lower()
    inline = (ctype == "application/pdf") or (ctype.startswith("image/")) or (ext in [".pdf", ".png", ".jpg", ".jpeg", ".webp"])

    disposition_type = "inline" if inline else "attachment"

    # Filename handling: sanitize hidden RTL/LTR marks + use RFC5987 for UTF-8
    raw_name = getattr(d, "doc_name", "") or os.path.basename(resolved)

    # Remove invisible direction marks that break latin-1 header encoding
    raw_name = (
        raw_name.replace("\u200e", "")
                .replace("\u200f", "")
                .replace("\u202a", "")
                .replace("\u202b", "")
                .replace("\u202c", "")
                .strip()
    )

    # ASCII-safe fallback (for filename="")
    safe_ascii = _safe_filename(raw_name)
    if not safe_ascii:
        safe_ascii = "document"

    # RFC5987 filename*=UTF-8''...
    from urllib.parse import quote
    utf8_name = quote(raw_name.encode("utf-8"))

    headers = {
        "Content-Disposition": f'{disposition_type}; filename="{safe_ascii}"; filename*=UTF-8\'\'{utf8_name}'
    }

    return FileResponse(resolved, media_type=ctype, headers=headers)

@app.post("/mrr/{lot_id}/docs/submit")
def mrr_docs_submit(
    lot_id: int,
    request: Request,
    session: Session = Depends(get_session),
    inspector_po_number: str = Form(...),
):
    user = get_current_user(request, session)

    lot = session.get(MaterialLot, lot_id)
    if not lot:
        raise HTTPException(404, "MRR Ticket not found")
    block_if_mrr_canceled(lot)

    # create or load receiving record (this is ONLY for PO verification now)
    receiving = session.exec(
        select(MrrReceiving).where(MrrReceiving.ticket_id == lot_id)
    ).first()

    if not receiving:
        receiving = MrrReceiving(ticket_id=lot_id)

    receiving.inspector_po_number = inspector_po_number.strip()

    # PO match check
    receiving.po_match = (receiving.inspector_po_number == (lot.po_number or "").strip())

    session.add(receiving)
    session.commit()

    return RedirectResponse(f"/mrr/{lot_id}", status_code=303)


@app.post("/mrr/{lot_id}/docs/confirm")
def mrr_docs_confirm(lot_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)
    forbid_boss(user)

    lot = session.get(MaterialLot, lot_id)
    if not lot:
        raise HTTPException(404, "MRR Ticket not found")
    block_if_mrr_canceled(lot)

    rec = session.exec(select(MrrReceiving).where(MrrReceiving.ticket_id == lot_id)).first()
    docs = session.exec(select(MrrDocument).where(MrrDocument.ticket_id == lot_id)).all()

    if not rec:
        raise HTTPException(400, "Documentation not saved")

    has_po_doc = any(((d.doc_type or "").strip().upper() == "PO") for d in docs)
    po_match = rec.inspector_po_number.strip() == (lot.po_number or "").strip()

    rec.po_match = po_match

    if has_po_doc and po_match:
        rec.inspector_confirmed_po = True
    else:
        rec.inspector_confirmed_po = False
        rec.manager_confirmed_po = False

    session.add(rec)
    session.commit()

    return RedirectResponse(f"/mrr/{lot_id}", status_code=303)

@app.post("/mrr/{lot_id}/docs/submit_and_confirm")
def mrr_docs_submit_and_confirm(
    lot_id: int,
    request: Request,
    session: Session = Depends(get_session),
    inspector_po_number: str = Form(...),
):
    user = get_current_user(request, session)
    forbid_boss(user)

    lot = session.get(MaterialLot, lot_id)
    if not lot:
        raise HTTPException(404, "MRR Ticket not found")
    block_if_mrr_canceled(lot)

    # Load or create receiving record
    rec = session.exec(select(MrrReceiving).where(MrrReceiving.ticket_id == lot_id)).first()
    if not rec:
        rec = MrrReceiving(ticket_id=lot_id)

    # Save inspector PO
    rec.inspector_po_number = (inspector_po_number or "").strip()

    # PO match
    ticket_po = (lot.po_number or "").strip()
    po_match = bool(rec.inspector_po_number and ticket_po and rec.inspector_po_number == ticket_po)
    rec.po_match = po_match

    # Check if PO document uploaded
    docs = session.exec(select(MrrDocument).where(MrrDocument.ticket_id == lot_id)).all()
    has_po_doc = any(((d.doc_type or "").strip().upper() == "PO") for d in docs)

    # Confirm rules
    if has_po_doc and po_match:
        rec.inspector_confirmed_po = True
    else:
        rec.inspector_confirmed_po = False
        rec.manager_confirmed_po = False

    session.add(rec)
    session.commit()

    return RedirectResponse(f"/mrr/{lot_id}", status_code=303)

@app.post("/mrr/{lot_id}/docs/approve")
def mrr_docs_manager_approve(lot_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)
    require_manager(user)

    rec = session.exec(select(MrrReceiving).where(MrrReceiving.ticket_id == lot_id)).first()
    if not rec:
        raise HTTPException(404, "Receiving record not found")

    rec.manager_confirmed_po = True
    rec.inspector_confirmed_po = True

    session.add(rec)
    session.commit()

    return RedirectResponse(f"/mrr/{lot_id}", status_code=303)

def mrr_docs_are_cleared(session: Session, lot_id: int) -> bool:
    rec = session.exec(
        select(MrrReceiving).where(MrrReceiving.ticket_id == lot_id)
    ).first()

    if not rec:
        return False


# =========================
# MRR Cancel helpers (soft delete)
# =========================
MRR_CANCELED_STATUS = "CANCELED"

def is_mrr_canceled(lot):
    return bool(lot and (getattr(lot, "status", "") or "").upper() == MRR_CANCELED_STATUS)

def block_if_mrr_canceled(lot):
    """Guard helper: prevent actions on canceled MRR tickets."""
    if is_mrr_canceled(lot):
        raise HTTPException(status_code=403, detail="This MRR Ticket is canceled")




def mrr_inspection_approve(lot_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)
    require_manager(user)

    insp = session.exec(
        select(MrrReceivingInspection)
        .where(MrrReceivingInspection.ticket_id == lot_id)
    ).first()

    if not insp:
        raise HTTPException(404, "Receiving Inspection not found")

    # ✅ Manager approves the receiving inspection
    insp.manager_approved = True
    session.add(insp)

    # ✅ THIS is the ONLY place the batch becomes usable in production
    lot = session.get(MaterialLot, lot_id)
    if not lot:
        raise HTTPException(404, "MRR Ticket not found")
    block_if_mrr_canceled(lot)

    lot.status = "APPROVED"   # <-- makes it appear in production dropdown
    session.add(lot)

    session.commit()
    return RedirectResponse(f"/mrr/{lot_id}", status_code=303)

@app.get("/mrr-pending-approvals", response_class=HTMLResponse)
def mrr_pending(request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)
    require_manager(user)

    # Find inspections submitted by inspector but not yet manager-approved
    pending = session.exec(
        select(MrrReceivingInspection)
        .where(MrrReceivingInspection.inspector_confirmed == True)
        .where(MrrReceivingInspection.manager_approved == False)
        .order_by(MrrReceivingInspection.created_at.desc())
    ).all()

    # Load lots to show ticket info
    lot_ids = [p.ticket_id for p in pending]
    lots = []
    lot_map = {}
    if lot_ids:
        lots = session.exec(select(MaterialLot).where(MaterialLot.id.in_(lot_ids))).all()
        lot_map = {l.id: l for l in lots}

    return templates.TemplateResponse(
        "mrr_pending.html",
        {
            "request": request,
            "user": user,
            "pending": pending,
            "lot_map": lot_map,
        },
    )

@app.get("/mrr/{lot_id}/inspection/new", response_class=HTMLResponse)
def new_shipment_inspection_page(
    lot_id: int,
    request: Request,
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)
    forbid_boss(user)

    lot = session.get(MaterialLot, lot_id)
    if not lot:
        raise HTTPException(404, "MRR Ticket not found")

    if (lot.status or "").upper() == "CANCELED":
        return RedirectResponse(f"/mrr/{lot_id}?error=Ticket%20is%20canceled", status_code=303)

    # ✅ LOCK: if ticket is approved (fully received) no new shipment inspection allowed
    if (lot.status or "").upper() == "APPROVED":
        return RedirectResponse(
            f"/mrr/{lot_id}?error=Ticket%20is%20APPROVED%20(receiving%20closed).%20Manager%20must%20unapprove%20to%20reopen.",
            status_code=303,
        )

    # Documentation prerequisites
    receiving = session.exec(select(MrrReceiving).where(MrrReceiving.ticket_id == lot_id)).first()

    has_po_doc = session.exec(
        select(MrrDocument).where((MrrDocument.ticket_id == lot_id) & (MrrDocument.doc_type == "PO"))
    ).first() is not None

    if not has_po_doc:
        return RedirectResponse(
            f"/mrr/{lot_id}?error=Upload%20PO%20document%20(Type:%20PO)%20before%20starting%20Receiving%20Inspection",
            status_code=303,
        )

        # ✅ Require at least ONE quality document before starting inspection
        if not quality_doc_exists(session, lot_id):
            return RedirectResponse(
                f"/mrr/{lot_id}?error=Upload%20one%20Quality%20Document%20(COA%20or%20MTC%20or%20Inspection%20Report)%20before%20starting%20Receiving%20Inspection",
                status_code=303,
            )


    if not receiving:
        return RedirectResponse(
            f"/mrr/{lot_id}?error=Documentation%20is%20not%20saved.%20Please%20fill%20Documentation%20and%20click%20Save%20first",
            status_code=303,
        )

    docs_ok = bool(getattr(receiving, "inspector_confirmed_po", False) or getattr(receiving, "manager_confirmed_po", False))
    if not docs_ok:
        return RedirectResponse(
            f"/mrr/{lot_id}?error=Documentation%20is%20not%20cleared.%20Click%20Confirm%20Documentation%20Complete%20first",
            status_code=303,
        )

    # If everything ok -> show your existing "new shipment" page
    return templates.TemplateResponse(
        "mrr_new_shipment.html",
        {
            "request": request,
            "user": user,
            "lot": lot,
            "error": request.query_params.get("error", ""),
        },
    )

@app.post("/mrr/{lot_id}/inspection/{inspection_id}/submit")
async def mrr_inspection_submit(
    lot_id: int,
    inspection_id: int,
    request: Request,
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)
    forbid_boss(user)

    lot = session.get(MaterialLot, lot_id)
    if not lot:
        raise HTTPException(404, "MRR Ticket not found")

    insp = session.get(MrrReceivingInspection, inspection_id)
    if not insp or insp.ticket_id != lot_id:
        raise HTTPException(404, "MRR Inspection not found")

    # LOCK after manager approval or ticket approved/closed
    if insp.manager_approved or (lot.status in ["APPROVED", "CLOSED"]):
        return RedirectResponse(
            url=f"/mrr/{lot_id}/inspection/id/{inspection_id}?error=Inspection%20is%20locked%20(manager%20approved%20/%20ticket%20approved).",
            status_code=303
        )

    form = await request.form()
    tpl = _resolve_template_type(lot, insp)
    action = (form.get("action") or "submit").strip().lower()

    batch_numbers = form.getlist("batch_numbers") if hasattr(form, "getlist") else []
    batch_numbers = [str(x).strip() for x in (batch_numbers or []) if str(x).strip()]

    vis_heat = form.getlist("vis_batch[]") if hasattr(form, "getlist") else []
    vis_heat = [str(x).strip() for x in (vis_heat or []) if str(x).strip()]

    # Validation on submit
    if action == "submit":
        if tpl == "OUTSOURCED":
            if not vis_heat:
                return RedirectResponse(
                    url=f"/mrr/{lot_id}/inspection/id/{inspection_id}?error=Heat%20number%20is%20required%20to%20submit.",
                    status_code=303,
                )
        else:
            if not batch_numbers:
                return RedirectResponse(
                    url=f"/mrr/{lot_id}/inspection/id/{inspection_id}?error=Batch%20number%20is%20required%20to%20submit.",
                    status_code=303,
                )

    # Load existing json
    try:
        existing = json.loads(insp.inspection_json or "{}")
        if not isinstance(existing, dict):
            existing = {}
    except Exception:
        existing = {}

    data = dict(existing)

    # store scalar fields safely
    for k, v in form.items():
        if k in ("action",):
            continue
        if k.endswith("[]"):
            continue
        if k != "batch_numbers":
            data[k] = (str(v).strip() if v is not None else "")

    data["batch_numbers"] = batch_numbers

    # Preserve list fields (keep row alignment)
    list_keys = [
        "items_item[]", "items_desc[]", "items_size[]", "items_type[]",
        "items_pressure[]", "items_qty[]", "items_mtc[]",
        "vis_batch[]", "vis_flange[]", "vis_surface[]", "vis_damage[]",
        "vis_package[]", "vis_marking[]", "vis_result[]",
    ]
    for key in list_keys:
        if hasattr(form, "getlist"):
            data[key] = [str(x).strip() for x in form.getlist(key)]

    def _trim_table(prefix_keys):
        cols = [data.get(k, []) for k in prefix_keys]
        if not cols or not all(isinstance(c, list) for c in cols):
            return
        max_len = max((len(c) for c in cols), default=0)
        last_keep = max_len
        for i in range(max_len - 1, -1, -1):
            row_empty = True
            for c in cols:
                v = c[i] if i < len(c) else ""
                if str(v).strip() != "":
                    row_empty = False
                    break
            if row_empty:
                last_keep = i
            else:
                break
        for k in prefix_keys:
            data[k] = data.get(k, [])[:last_keep]

    _trim_table([
        "items_item[]","items_desc[]","items_size[]",
        "items_type[]","items_pressure[]","items_qty[]","items_mtc[]"
    ])
    _trim_table([
        "vis_batch[]","vis_flange[]","vis_surface[]",
        "vis_damage[]","vis_package[]","vis_marking[]","vis_result[]"
    ])

    # Save model fields needed for header
    dn = (form.get("delivery_note_no") or "").strip()
    qty_arrived = form.get("qty_arrived")
    qty_unit = (form.get("qty_unit") or "").strip()

    if dn:
        insp.delivery_note_no = dn
    if qty_arrived is not None and str(qty_arrived).strip() != "":
        try:
            insp.qty_arrived = float(qty_arrived)
        except Exception:
            pass
    if qty_unit:
        insp.qty_unit = qty_unit

    # Always keep report_no from model
    data["report_no"] = getattr(insp, "report_no", "") or data.get("report_no", "")

    # --- Signature logic ---
    now_utc = datetime.utcnow().isoformat()
    current_name = (getattr(user, "display_name", "") or "").strip()
    current_id = getattr(user, "id", None)

    # If draft save: do not sign anything
    if action == "draft":
        insp.inspection_json = json.dumps(data, ensure_ascii=False)
        insp.inspector_confirmed = False
        session.add(insp)
        session.commit()
        return RedirectResponse(f"/mrr/{lot_id}/inspection/id/{inspection_id}?saved=draft", status_code=303)

    # SUBMIT behavior:
    # 1) If not yet submitted: current user becomes INSPECTOR signature
    # 2) If already submitted AND current user is different: optional REVIEWER signature (one time)
    # 3) If already submitted AND same user: do nothing (no overwrite)

    # Ensure inspector fields exist when first submitted
    if not getattr(insp, "inspector_confirmed", False):
        # First-time submission -> inspector signs
        insp.inspector_id = current_id
        insp.inspector_name = current_name
        insp.inspector_confirmed = True

        data["inspected_by"] = current_name
        data["submitted_by"] = current_name
        data["submitted_at_utc"] = now_utc

    else:
        # Already submitted: reviewer is optional if different user and not already reviewed
        already_reviewed = bool(data.get("reviewed_by"))
        same_user = (current_id is not None and current_id == getattr(insp, "inspector_id", None)) or (
            current_name and (current_name == (getattr(insp, "inspector_name", "") or "").strip())
        )

        if (not same_user) and (not already_reviewed):
            data["reviewed_by"] = current_name
            data["reviewed_at_utc"] = now_utc

        # Do NOT overwrite submitted_by/submitted_at_utc

    insp.inspection_json = json.dumps(data, ensure_ascii=False)
    session.add(insp)
    session.commit()

    return RedirectResponse(f"/mrr/{lot_id}", status_code=303)




@app.post("/mrr/{lot_id}/inspection/new")
async def create_shipment_inspection(
    lot_id: int,
    request: Request,
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)
    forbid_boss(user)

    lot = session.get(MaterialLot, lot_id)
    if not lot:
        raise HTTPException(404, "MRR Ticket not found")

    if (lot.status or "").upper() == "CANCELED":
        return RedirectResponse(f"/mrr/{lot_id}?error=Ticket%20is%20canceled", status_code=303)

    # LOCK: approved tickets should not accept new inspections
    if (lot.status or "").upper() == "APPROVED":
        return RedirectResponse(
            f"/mrr/{lot_id}?error=Ticket%20is%20APPROVED%20(receiving%20closed).%20Manager%20must%20unapprove%20to%20reopen.",
            status_code=303,
        )

    # Require at least one quality doc (COA/MTC/INSPECTION_REPORT) before inspection
    has_quality_doc = (
        session.exec(
            select(MrrDocument).where(
                (MrrDocument.ticket_id == lot_id)
                & (MrrDocument.doc_type.in_(["COA", "MTC", "INSPECTION_REPORT"]))
            )
        ).first()
        is not None
    )
    if not has_quality_doc:
        return RedirectResponse(
            f"/mrr/{lot_id}?error=Upload%20Quality%20Document%20(COA%20or%20MTC%20or%20INSPECTION%20REPORT)%20before%20starting%20Receiving%20Inspection",
            status_code=303,
        )

    # Require PO document
    has_po_doc = (
        session.exec(
            select(MrrDocument).where(
                (MrrDocument.ticket_id == lot_id) & (MrrDocument.doc_type == "PO")
            )
        ).first()
        is not None
    )
    if not has_po_doc:
        return RedirectResponse(
            f"/mrr/{lot_id}?error=Upload%20PO%20document%20(Type:%20PO)%20before%20starting%20Receiving%20Inspection",
            status_code=303,
        )

    # Documentation prerequisites: receiving must exist and be cleared
    receiving = session.exec(
        select(MrrReceiving).where(MrrReceiving.ticket_id == lot_id)
    ).first()
    if not receiving:
        return RedirectResponse(
            f"/mrr/{lot_id}?error=Documentation%20is%20not%20saved.%20Please%20fill%20Documentation%20and%20click%20Save%20first",
            status_code=303,
        )

    docs_ok = bool(
        getattr(receiving, "inspector_confirmed_po", False)
        or getattr(receiving, "manager_confirmed_po", False)
    )
    if not docs_ok:
        return RedirectResponse(
            f"/mrr/{lot_id}?error=Documentation%20is%20not%20cleared.%20Click%20Confirm%20Documentation%20Complete%20first",
            status_code=303,
        )

    form = await request.form()

    dn = (form.get("delivery_note_no") or "").strip()
    qty_arrived = form.get("qty_arrived")
    qty_unit = (form.get("qty_unit") or "KG").strip().upper()

    if not dn:
        return RedirectResponse(
            f"/mrr/{lot_id}/inspection/new?error=Delivery%20Note%20is%20required",
            status_code=303,
        )

    # ✅ If a DRAFT shipment already exists for this DN, reuse it (do NOT create duplicates)
    existing_draft = get_draft_shipment_by_dn(session, lot_id, dn)
    if existing_draft and getattr(existing_draft, "id", None) is not None:
        return RedirectResponse(f"/mrr/{lot_id}/inspection/id/{existing_draft.id}", status_code=303)

    # Must have DN document uploaded with matching doc_number
    dn_doc = session.exec(
        select(MrrDocument).where(
            (MrrDocument.ticket_id == lot_id)
            & (MrrDocument.doc_type == "DELIVERY_NOTE")
            & (MrrDocument.doc_number == dn)
        )
    ).first()
    if not dn_doc:
        return RedirectResponse(
            f"/mrr/{lot_id}/inspection/new?error=Upload%20Delivery%20Note%20document%20(Type:%20DELIVERY_NOTE)%20and%20set%20Document%20Number%20exactly%20=%20{dn}",
            status_code=303,
        )

    try:
        qty_arrived_val = float(qty_arrived or 0.0)
    except Exception:
        return RedirectResponse(
            f"/mrr/{lot_id}/inspection/new?error=Invalid%20quantity",
            status_code=303,
        )

    if qty_arrived_val <= 0:
        return RedirectResponse(
            f"/mrr/{lot_id}/inspection/new?error=Quantity%20must%20be%20greater%20than%200",
            status_code=303,
        )

    # DN must not be reused in another SUBMITTED/APPROVED shipment
    dn_used = session.exec(
        select(MrrReceivingInspection).where(
            (MrrReceivingInspection.ticket_id == lot_id)
            & (MrrReceivingInspection.delivery_note_no == dn)
            & (
                (MrrReceivingInspection.inspector_confirmed == True)
                | (MrrReceivingInspection.manager_approved == True)
            )
        )
    ).first()
    if dn_used:
        return RedirectResponse(
            f"/mrr/{lot_id}/inspection/new?error=This%20Delivery%20Note%20was%20already%20used%20in%20a%20submitted%20shipment",
            status_code=303,
        )

    # Resolve template type for this inspection (RAW/F01 vs OUTSOURCED/F02)
    tpl = _resolve_template_type(lot, None)

    # Create inspection record
    report_no = generate_report_no(lot_id, 1)

    insp = MrrReceivingInspection(
        ticket_id=lot_id,
        inspector_id=user.id,
        inspector_name=user.display_name,
        delivery_note_no=dn,
        qty_arrived=qty_arrived_val,
        qty_unit=qty_unit,
        report_no=report_no,
        template_type=tpl,
        inspection_json="{}",
        inspector_confirmed=False,
        manager_approved=False,
    )
    session.add(insp)
    session.commit()

    return RedirectResponse(f"/mrr/{lot_id}/inspection/id/{insp.id}", status_code=303)

@app.get("/runs", response_class=HTMLResponse)
def runs_list(
    request: Request,
    q: str = "",
    view: str = "open",   # open | approved
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)

    # ✅ Only 2 views:
    # - open: everything except APPROVED (includes CLOSED)
    # - approved: only APPROVED
    view = (view or "open").lower().strip()
    if view not in ["open", "approved"]:
        view = "open"

    stmt = select(ProductionRun)

    # ✅ If user is searching, show ALL (open + approved) regardless of tab
    # ✅ If no search, keep the split logic
    q_clean = (q or "").strip()
    if not q_clean:
        if view == "approved":
            stmt = stmt.where(ProductionRun.status == "APPROVED")
        else:
            stmt = stmt.where(ProductionRun.status != "APPROVED")


    # ✅ Search across multiple fields
    q_clean = (q or "").strip()
    if q_clean:
        like = f"%{q_clean}%"
        stmt = stmt.where(
            or_(
                ProductionRun.dhtp_batch_no.ilike(like),
                ProductionRun.process.ilike(like),
                ProductionRun.client_name.ilike(like),
                ProductionRun.po_number.ilike(like),
                ProductionRun.itp_number.ilike(like),
                ProductionRun.pipe_specification.ilike(like),
                ProductionRun.raw_material_spec.ilike(like),
            )
        )

    stmt = stmt.order_by(ProductionRun.created_at.desc())
    runs = session.exec(stmt).all()

    # ✅ Group by batch number
    grouped = defaultdict(list)
    for r in runs:
        grouped[(r.dhtp_batch_no or "-").strip()].append(r)

    # ✅ Build stable client colors (soft background)
    def client_color(name: str) -> str:
        n = (name or "").strip().lower()
        if not n:
            return "#f7f7f7"
        h = hashlib.md5(n.encode("utf-8")).hexdigest()
        hue = int(h[:2], 16) * 360 // 255
        # very light HSL-like pastel using a fixed palette trick
        # (keeps consistent color per client)
        return f"hsl({hue}, 70%, 95%)"

    client_bg = {}
    for r in runs:
        cn = (r.client_name or "").strip()
        if cn and cn not in client_bg:
            client_bg[cn] = client_color(cn)

    # ✅ Convert dict to list (for template)
    batch_cards = []
    for batch_no, items in grouped.items():
        # Sort processes inside each batch in a nice order
        order = {"LINER": 1, "REINFORCEMENT": 2, "COVER": 3}
        items_sorted = sorted(items, key=lambda x: order.get((x.process or "").upper(), 99))

        # take summary info from first row
        client_name = items_sorted[0].client_name if items_sorted else ""
        po_number = items_sorted[0].po_number if items_sorted else ""

        batch_cards.append({
            "batch_no": batch_no,
            "client_name": client_name,
            "po_number": po_number,
            "runs": items_sorted,
        })

    # Sort batches newest-first by newest run created_at
    batch_cards.sort(
        key=lambda b: b["runs"][0].created_at if b["runs"] else datetime.min,
        reverse=True,
    )

    return templates.TemplateResponse(
        "run_list.html",
        {
            "request": request,
            "user": user,
            "q": q_clean,
            "view": view,
            "batch_cards": batch_cards,
            "client_bg": client_bg,
        },
    )




@app.get("/runs/new", response_class=HTMLResponse)
def run_new_get(request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)
    if (user.role or "").upper() not in ["MANAGER", "RUN_CREATOR"]:
        raise HTTPException(403, "Manager only")

    return templates.TemplateResponse(
        "run_new.html",
        {
            "request": request,
            "user": user,
            "error": "",
            "default_processes": ["LINER", "REINFORCEMENT", "COVER"],
        },
    )



@app.get("/mrr/{lot_id}/inspection")
def old_inspection_redirect(lot_id: int):
    return RedirectResponse(f"/mrr/{lot_id}/inspection/new", status_code=302)

@app.post("/runs/new")
def run_new_post(
    request: Request,
    session: Session = Depends(get_session),
    process: str = Form(""),
    selected_processes: List[str] = Form([]),
    dhtp_batch_no: str = Form(...),
    client_name: str = Form(...),
    po_number: str = Form(...),
    itp_number: str = Form(...),
    pipe_specification: str = Form(...),
    raw_material_spec: str = Form(""),
    total_length_m: float = Form(0.0),
    allow_duplicate: str = Form(""),
):
    user = get_current_user(request, session)
    if (user.role or "").upper() not in ["MANAGER", "RUN_CREATOR"]:
        raise HTTPException(403, "Manager only")

    chosen = [p.upper().strip() for p in selected_processes if (p or "").strip()]
    if not chosen and (process or "").strip():
        chosen = [(process or "").upper().strip()]

    chosen = [p for p in chosen if p in PROCESS_PARAMS]
    # keep order predictable
    ordered = [p for p in ["LINER", "REINFORCEMENT", "COVER"] if p in chosen]

    if not ordered:
        return templates.TemplateResponse(
            "run_new.html",
            {
                "request": request,
                "user": user,
                "error": "Select at least one layer/process.",
                "default_processes": ["LINER", "REINFORCEMENT", "COVER"],
            },
        )

    batch_no = (dhtp_batch_no or "").strip()
    raw_material_spec = (raw_material_spec or "").strip()
    if len(ordered) > 1:
        raw_material_spec = ""

    for proc in ordered:
        existing_open = session.exec(
            select(ProductionRun).where(
                ProductionRun.dhtp_batch_no == batch_no,
                ProductionRun.process == proc,
                ProductionRun.status == "OPEN",
            )
        ).first()

        if existing_open and allow_duplicate != "1":
            return templates.TemplateResponse(
                "run_new.html",
                {
                    "request": request,
                    "user": user,
                    "error": f"There is already an OPEN {proc} run for Batch {batch_no}. Tick 'Allow duplicate line' only if you really need another line.",
                    "default_processes": ordered,
                },
            )

    created_runs = []
    for proc in ordered:
        created_runs.append(
            _create_run_with_default_params(
                session,
                process=proc,
                dhtp_batch_no=batch_no,
                client_name=client_name,
                po_number=po_number,
                itp_number=itp_number,
                pipe_specification=pipe_specification,
                raw_material_spec=raw_material_spec,
                total_length_m=total_length_m or 0.0,
            )
        )

    if len(created_runs) == 1:
        return RedirectResponse(f"/runs/{created_runs[0].id}", status_code=303)

    return RedirectResponse("/runs", status_code=303)
    
@app.get("/mrr/{lot_id}/inspection/id/{inspection_id}", response_class=HTMLResponse)
def shipment_inspection_form(
    lot_id: int,
    inspection_id: int,
    request: Request,
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)

    lot = session.get(MaterialLot, lot_id)
    if not lot:
        raise HTTPException(404, "Ticket not found")

    inspection = session.get(MrrReceivingInspection, inspection_id)
    if not inspection or inspection.ticket_id != lot_id:
        raise HTTPException(404, "Inspection not found")

    # Load stored JSON safely
    data = safe_json_loads(getattr(inspection, "inspection_json", None)) or {}
    if not isinstance(data, dict):
        data = {}

    # ✅ IMPORTANT: read query messages so reviewer understands what happened
    error = request.query_params.get("error", "")
    success = request.query_params.get("success", "")

    # Template selection (same as your current logic)
    template_name = "mrr_inspection_outsourced.html" if (
        (getattr(lot, "inspection_type", "") or "").upper() == "OUTSOURCED"
    ) else "mrr_inspection.html"

    return templates.TemplateResponse(
        template_name,
        {
            "request": request,
            "user": user,
            "ticket": lot,
            "inspection": inspection,

            # ✅ IMPORTANT: give template the name it actually uses
            "inspection_data": data,

            # Keep compatibility with any older fields in template(s)
            "data": data,
            "form_data": data,
            
            # ✅ provide both names (templates in your repo use both)
            "lot": lot,
            "ticket": lot,

            # ✅ show messages instead of “refresh only”
            "error": error,
            "success": success,
        },
    )
@app.get("/runs/{run_id}", response_class=HTMLResponse)
def run_view(run_id: int, request: Request, session: Session = Depends(get_session)):
    try:
        user = get_current_user(request, session)

        run = session.get(ProductionRun, run_id)
        if not run:
            raise HTTPException(404, "Run not found")

        machines = session.exec(select(RunMachine).where(RunMachine.run_id == run_id)).all()
        params = session.exec(
            select(RunParameter).where(RunParameter.run_id == run_id).order_by(RunParameter.display_order)
        ).all()

        days = get_days_for_run(session, run_id)

        selected_day = None
        day_q = request.query_params.get("day")
        if day_q:
            try:
                selected_day = date.fromisoformat(day_q)
            except Exception:
                selected_day = None

        if selected_day is None:
            selected_day = days[-1] if days else date.today()

        entries = session.exec(
            select(InspectionEntry)
            .where(InspectionEntry.run_id == run_id, InspectionEntry.actual_date == selected_day)
            .order_by(InspectionEntry.created_at)
        ).all()
        
        # ✅ NEW: for the UI to know which slots have an InspectionEntry
        slot_entry_ids: Dict[str, int] = {}
        for e in entries:
            if not e.slot_time or e.slot_time not in SLOTS:
                continue
            slot_entry_ids[e.slot_time] = e.id



        users = session.exec(select(User)).all()
        user_map = {u.id: u for u in users}

        slot_inspectors: Dict[str, str] = {s: "" for s in SLOTS}
        slot_actual_times: Dict[str, str] = {s: "" for s in SLOTS}
        grid: Dict[str, Dict[str, dict]] = {s: {} for s in SLOTS}

        for e in entries:
            if not e.slot_time or e.slot_time not in SLOTS:
                continue

            slot_actual_times[e.slot_time] = e.actual_time or ""

            if e.inspector_id and e.inspector_id in user_map:
                slot_inspectors[e.slot_time] = user_map[e.inspector_id].display_name or ""
            else:
                slot_inspectors[e.slot_time] = ""


            vals = session.exec(select(InspectionValue).where(InspectionValue.entry_id == e.id)).all()
            for v in vals:
                grid[e.slot_time][v.param_key] = {
                    "value_id": v.id,
                    "value": v.value,
                    "out": bool(v.is_out_of_spec),
                    "note": v.spec_note or "",
                    "pending_value": v.pending_value,
                    "pending_status": v.pending_status or "",
                }

        trace_today = get_day_latest_trace(session, run_id, selected_day)
        carry = get_last_known_trace_before_day(session, run_id, selected_day)

        raw_batches = trace_today["raw_batches"] or ([carry["raw"]] if carry["raw"] else [])
        tools = trace_today["tools"] or carry["tools"]

        progress = get_progress_percent(session, run)

        return templates.TemplateResponse(
            "run_view.html",
            {
                "request": request,
                "user": user,
                "run": run,
                "machines": machines,
                "params": params,
                "days": days,
                "selected_day": selected_day,
                "grid": grid,
                "slot_inspectors": slot_inspectors,
                "slot_entry_ids": slot_entry_ids, 
                "image_url": IMAGE_MAP.get(run.process, ""),
                "progress": progress,
                "raw_batches": raw_batches,
                "tools": tools,
                "slot_actual_times": slot_actual_times,
               
            },
        )

    except Exception:
        # TEMP DEBUG: show the real error on the page
        return HTMLResponse(
            "<pre style='white-space:pre-wrap;font-size:14px'>"
            + traceback.format_exc()
            + "</pre>",
            status_code=500,
        )

@app.post("/runs/{run_id}/close")
def run_close(
    run_id: int,
    request: Request,
    session: Session = Depends(get_session),
    confirmed_total_length_m: float = Form(0.0),
    confirmed_length_note: str = Form(""),
):
    user = get_current_user(request, session)
    require_manager(user)

    run = session.get(ProductionRun, run_id)
    if not run:
        raise HTTPException(404, "Run not found")

    measured = float(get_run_produced_length_m(session, run_id) or 0.0)
    confirmed = float(confirmed_total_length_m or 0.0)

    if confirmed > 0:
        if confirmed < measured:
            raise HTTPException(400, f"Confirmed total length cannot be less than measured length ({measured:.1f} m).")
        run.confirmed_total_length_m = confirmed
        run.confirmed_length_note = (confirmed_length_note or "").strip()

    run.status = "CLOSED"
    session.add(run)
    session.commit()

    return RedirectResponse(f"/runs/{run_id}", status_code=302)
@app.post("/runs/{run_id}/approve")
def run_approve(run_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)
    require_manager(user)  # only managers can approve

    run = session.get(ProductionRun, run_id)
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")

    # Require CLOSED before approving (keep or remove depending on your rule)
    if (run.status or "").upper() != "CLOSED":
        raise HTTPException(status_code=400, detail="Run must be CLOSED before approving")

    run.status = "APPROVED"
    run.approved_by_user_id = user.id
    run.approved_by_user_name = user.display_name or ""
    run.approved_at_utc = datetime.utcnow()

    session.add(run)
    session.commit()

    return RedirectResponse(url=f"/runs/{run_id}", status_code=302)


@app.post("/runs/{run_id}/reopen")
def run_reopen(run_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)
    require_manager(user)  # only manager can reopen

    run = session.get(ProductionRun, run_id)
    if not run:
        raise HTTPException(404, "Run not found")

    # Reopen allowed from CLOSED or APPROVED
    run.status = "OPEN"

    # If it was approved before, clear approval fields so it behaves like a normal open run again
    run.approved_by_user_id = None
    run.approved_by_user_name = ""
    run.approved_at_utc = None

    session.add(run)
    session.commit()

    return RedirectResponse(f"/runs/{run_id}", status_code=302)


@app.get("/runs/{run_id}/edit", response_class=HTMLResponse)
def run_edit_get(run_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)
    require_manager(user)
        

    run = session.get(ProductionRun, run_id)
    if not run:
        raise HTTPException(404, "Run not found")

    machines = session.exec(select(RunMachine).where(RunMachine.run_id == run_id)).all()
    params = session.exec(
        select(RunParameter).where(RunParameter.run_id == run_id).order_by(RunParameter.display_order)
    ).all()

    return templates.TemplateResponse(
        "run_edit.html",
        {"request": request, "user": user, "run": run, "machines": machines, "params": params, "error": ""},
    )


@app.post("/runs/{run_id}/edit")
async def run_edit_post(run_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)
    require_manager(user)
        

    run = session.get(ProductionRun, run_id)
    if not run:
        raise HTTPException(404, "Run not found")

    form = await request.form()

    run.client_name = str(form.get("client_name", "")).strip()
    run.po_number = str(form.get("po_number", "")).strip()
    run.itp_number = str(form.get("itp_number", "")).strip()
    run.pipe_specification = str(form.get("pipe_specification", "")).strip()
    run.raw_material_spec = str(form.get("raw_material_spec", "")).strip()
    try:
        run.total_length_m = float(form.get("total_length_m") or 0.0)
    except Exception:
        run.total_length_m = 0.0

    session.add(run)
    session.commit()

    existing = session.exec(select(RunMachine).where(RunMachine.run_id == run_id)).all()
    for m in existing:
        session.delete(m)
    session.commit()

    def _m(name_key: str, tag_key: str):
        name = str(form.get(name_key, "")).strip()
        tag = str(form.get(tag_key, "")).strip()
        if name:
            session.add(RunMachine(run_id=run_id, machine_name=name, machine_tag=tag))

    _m("machine1_name", "machine1_tag")
    _m("machine2_name", "machine2_tag")
    _m("machine3_name", "machine3_tag")
    _m("machine4_name", "machine4_tag")
    _m("machine5_name", "machine5_tag")
    session.commit()

    params = session.exec(select(RunParameter).where(RunParameter.run_id == run_id)).all()
    for p in params:
        new_label = str(form.get(f"label_{p.param_key}", "")).strip()
        if new_label:
            p.label = new_label

        rule = str(form.get(f"rule_{p.param_key}", "")).strip().upper()
        if rule not in ["", "RANGE", "MIN_ONLY", "MAX_ONLY"]:
            rule = ""
        p.rule = rule

        set_raw = form.get(f"set_{p.param_key}", "")
        tol_raw = form.get(f"tol_{p.param_key}", "")
        
        set_v = _safe_float(set_raw)
        tol_v = _safe_float(tol_raw)
        
        # convert Set/Tolerance to Min/Max (keeps old system working)
        if p.rule == "RANGE":
            if set_v is None or tol_v is None:
                p.min_value = None
                p.max_value = None
            else:
                t = abs(tol_v)
                p.min_value = set_v - t
                p.max_value = set_v + t
        
        elif p.rule == "MAX_ONLY":
            # your example: temperature max 35 -> set=35, tol can be empty
            p.min_value = None
            p.max_value = set_v
        
        elif p.rule == "MIN_ONLY":
            p.min_value = set_v
            p.max_value = None
        
        else:
            p.min_value = None
            p.max_value = None


        session.add(p)

    session.commit()
    return RedirectResponse(f"/runs/{run.id}", status_code=302)


@app.get("/runs/{run_id}/entry/new", response_class=HTMLResponse)
def entry_new_get(run_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)

    run = session.get(ProductionRun, run_id)
    if not run:
        raise HTTPException(404, "Run not found")

    params = session.exec(
        select(RunParameter).where(RunParameter.run_id == run_id).order_by(RunParameter.display_order)
    ).all()

    has_any = session.exec(select(InspectionEntry.id).where(InspectionEntry.run_id == run_id)).first() is not None
    error = request.query_params.get("error", "")

    approved_lots = session.exec(
        select(MaterialLot)
        .where(MaterialLot.status == "APPROVED", MaterialLot.lot_type == "RAW")
        .order_by(MaterialLot.batch_no)
    ).all()

    # ✅ TRUE check: does the run already have any batch event?
    has_any_event = session.exec(
        select(MaterialUseEvent.id).where(MaterialUseEvent.run_id == run_id).limit(1)
    ).first() is not None

    # ✅ Show current batch as the latest batch in the run (not today 00:00)
    today_lot = get_latest_material_lot_for_run(session, run_id)

    return templates.TemplateResponse(
        "entry_new.html",
        {
            "request": request,
            "user": user,
            "run": run,
            "params": params,
            "has_any": has_any,
            "error": error,
            "approved_lots": approved_lots,
            "has_any_event": has_any_event,
            "current_lot_preview": today_lot,
        },
    )


def get_latest_material_lot_for_run(session: Session, run_id: int):
    last_ev = session.exec(
        select(MaterialUseEvent)
        .where(MaterialUseEvent.run_id == run_id)
        .order_by(
            MaterialUseEvent.day.desc(),
            MaterialUseEvent.slot_time.desc(),
            MaterialUseEvent.created_at.desc(),
        )
    ).first()

    if not last_ev:
        return None
    return session.get(MaterialLot, last_ev.lot_id)


def apply_spec_check(param, v):
    """
    Returns (is_oos, note)
    - is_oos: True if value is out of spec
    - note: human-readable explanation
    """

    # Handle empty / missing values safely
    if v is None or (isinstance(v, str) and v.strip() == ""):
        return False, ""

    # Convert value to float if possible (common for measurements)
    try:
        val = float(v)
    except (TypeError, ValueError):
        # If it's not numeric, we can't spec-check it here
        return False, ""

    # Extract limits from param in a flexible way (dict or object)
    def get_attr(obj, *names):
        if obj is None:
            return None
        # dict style
        if isinstance(obj, dict):
            for n in names:
                if n in obj and obj[n] is not None:
                    return obj[n]
        # object style
        for n in names:
            if hasattr(obj, n):
                x = getattr(obj, n)
                if x is not None:
                    return x
        return None

    min_v = get_attr(param, "min_value", "min", "lower", "low_limit")
    max_v = get_attr(param, "max_value", "max", "upper", "high_limit")

    # Try to convert min/max if they exist
    try:
        min_v = float(min_v) if min_v is not None else None
    except (TypeError, ValueError):
        min_v = None
    try:
        max_v = float(max_v) if max_v is not None else None
    except (TypeError, ValueError):
        max_v = None

    # Spec checks
    if min_v is not None and val < min_v:
        return True, f"Below min ({val} < {min_v})"
    if max_v is not None and val > max_v:
        return True, f"Above max ({val} > {max_v})"

    return False, ""
   
@app.post("/runs/{run_id}/entry/new")
async def entry_new_post(
    run_id: int,
    request: Request,
    session: Session = Depends(get_session),
    actual_date: str = Form(...),
    actual_time: str = Form(...),
    operator_1: str = Form(""),
    operator_2: str = Form(""),
    operator_annular_12: str = Form(""),
    operator_int_ext_34: str = Form(""),
    remarks: str = Form(""),
    tool1_name: str = Form(""),
    tool1_serial: str = Form(""),
    tool1_calib_due: str = Form(""),
    tool2_name: str = Form(""),
    tool2_serial: str = Form(""),
    tool2_calib_due: str = Form(""),
    start_lot_id: str = Form(""),   # ✅ for first ever entry
):
    user = get_current_user(request, session)
    forbid_boss(user)

    run = session.get(ProductionRun, run_id)
    if not run:
        raise HTTPException(404, "Run not found")

    if run.status in ["CLOSED", "APPROVED"] and user.role != "MANAGER":
        raise HTTPException(403, "Run is not open")

    slot_time, day_add = slot_from_time_str(actual_time)
    day_obj = date.fromisoformat(actual_date) + timedelta(days=day_add)

    # prevent duplicate slot entry
    existing_for_slot = session.exec(
        select(InspectionEntry).where(
            InspectionEntry.run_id == run_id,
            InspectionEntry.actual_date == day_obj,
            InspectionEntry.slot_time == slot_time
        )
    ).first()
    if existing_for_slot:
        msg = "This timing slot is already inspected. Please confirm the time, or use Edit to change the existing record."
        return RedirectResponse(f"/runs/{run_id}/entry/new?error={msg}", status_code=302)

    form = await request.form()

    # batch-change checkbox + selected lot
    batch_changed = str(form.get("batch_changed", "")).strip() == "1"
    new_lot_id_raw = str(form.get("new_lot_id", "")).strip()
    
    # ✅ If user selected a new lot from dropdown, treat it as batch changed (even if checkbox not ticked)
    if new_lot_id_raw.isdigit():
        batch_changed = True


    # check if run has ANY material event yet
    has_any_event = session.exec(
        select(MaterialUseEvent.id).where(MaterialUseEvent.run_id == run_id).limit(1)
    ).first() is not None

    # If this is the FIRST entry EVER (no event exists), require start_lot_id
    if not has_any_event:
        if not str(start_lot_id).isdigit():
            msg = "Please select the STARTING approved RAW batch (first entry only)."
            return RedirectResponse(f"/runs/{run_id}/entry/new?error={msg}", status_code=302)

        lot = session.get(MaterialLot, int(start_lot_id))
        if (not lot) or (lot.status != "APPROVED") or (getattr(lot, "lot_type", "RAW") != "RAW"):
            msg = "Selected starting batch is not an APPROVED RAW batch."
            return RedirectResponse(f"/runs/{run_id}/entry/new?error={msg}", status_code=302)
            

        # Create the inspection entry first (NO batch yet)
    entry = InspectionEntry(
        run_id=run_id,
        actual_date=day_obj,
        actual_time=actual_time,
        slot_time=slot_time,
        inspector_id=user.id,
        operator_1=operator_1,
        operator_2=operator_2,
        operator_annular_12=operator_annular_12,
        operator_int_ext_34=operator_int_ext_34,
        remarks=remarks,
        tool1_name=tool1_name,
        tool1_serial=tool1_serial,
        tool1_calib_due=tool1_calib_due,
        tool2_name=tool2_name,
        tool2_serial=tool2_serial,
        tool2_calib_due=tool2_calib_due,
    )
    session.add(entry)
    session.commit()
    session.refresh(entry)

    # ✅ FIRST entry: create starting MaterialUseEvent at this slot
    if not has_any_event:
        session.add(MaterialUseEvent(
            run_id=run_id,
            day=day_obj,
            slot_time=slot_time,
            lot_id=int(start_lot_id),
            created_by_user_id=user.id,
        ))
        session.commit()

    # ✅ Batch changed: create new event
    if batch_changed:
        if not new_lot_id_raw.isdigit():
            msg = "Please select the NEW approved RAW batch when 'batch changed' is checked."
            return RedirectResponse(f"/runs/{run_id}/entry/new?error={msg}", status_code=302)

        new_lot = session.get(MaterialLot, int(new_lot_id_raw))
        if (not new_lot) or (new_lot.status != "APPROVED") or (getattr(new_lot, "lot_type", "RAW") != "RAW"):
            msg = "Selected NEW batch is not an APPROVED RAW batch."
            return RedirectResponse(f"/runs/{run_id}/entry/new?error={msg}", status_code=302)

        session.add(MaterialUseEvent(
            run_id=run_id,
            day=day_obj,
            slot_time=slot_time,
            lot_id=int(new_lot_id_raw),
            created_by_user_id=user.id,
        ))
        session.commit()

    # ✅ NOW we can read the current batch (because the events exist) and save it into the entry
    current_lot = get_current_material_lot_for_slot(session, run_id, day_obj, slot_time)
    entry.raw_material_batch_no = (current_lot.batch_no or "").strip() if current_lot else ""
    session.add(entry)
    session.commit()


    # save values (unchanged logic)
    params = session.exec(select(RunParameter).where(RunParameter.run_id == run_id)).all()
    by_key = {p.param_key: p for p in params}

    for key, param in by_key.items():
        raw = form.get(f"v_{key}")
        if raw in [None, ""]:
            continue
        v = _safe_float(raw)
        if v is None:
            continue
        is_oos, note = apply_spec_check(param, v)
        session.add(InspectionValue(
            entry_id=entry.id,
            param_key=key,
            value=v,
            is_out_of_spec=is_oos,
            spec_note=note,
        ))

    session.commit()
    return RedirectResponse(f"/runs/{run_id}?day={entry.actual_date.isoformat()}", status_code=302)


@app.get("/runs/{run_id}/entry/{slot_time}/fill/{param_key}", response_class=HTMLResponse)
def fill_missing_value_get(
    run_id: int,
    slot_time: str,
    param_key: str,
    request: Request,
    day: str,
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)
    forbid_boss(user)

    run = session.get(ProductionRun, run_id)
    if not run:
        raise HTTPException(404, "Run not found")

    # permissions
    if run.status in ["CLOSED", "APPROVED"] and (user.role or "").upper() != "MANAGER":
        raise HTTPException(403, "Run is closed")

    day_obj = date.fromisoformat(day)

    entry = session.exec(
        select(InspectionEntry).where(
            InspectionEntry.run_id == run_id,
            InspectionEntry.actual_date == day_obj,
            InspectionEntry.slot_time == slot_time,
        )
    ).first()
    if not entry:
        raise HTTPException(404, "No inspection entry for this slot/day")

    # param
    param = session.exec(
        select(RunParameter).where(RunParameter.run_id == run_id, RunParameter.param_key == param_key)
    ).first()
    if not param:
        raise HTTPException(404, "Parameter not found")

    # already exists? send to normal edit
    existing = session.exec(
        select(InspectionValue).where(InspectionValue.entry_id == entry.id, InspectionValue.param_key == param_key)
    ).first()
    if existing:
        return RedirectResponse(f"/values/{existing.id}/edit", status_code=302)

    return templates.TemplateResponse(
        "value_fill.html",
        {"request": request, "user": user, "run": run, "entry": entry, "param": param, "error": ""},
    )

@app.post("/runs/{run_id}/entry/{slot_time}/fill/{param_key}")
async def fill_missing_value_post(
    run_id: int,
    slot_time: str,
    param_key: str,
    request: Request,
    day: str = Form(...),
    new_value: str = Form(...),
    note: str = Form(""),
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)
    forbid_boss(user)

    run = session.get(ProductionRun, run_id)
    if not run:
        raise HTTPException(404, "Run not found")

    # permissions
    if run.status in ["CLOSED", "APPROVED"] and (user.role or "").upper() != "MANAGER":
        raise HTTPException(403, "Run is closed")

    day_obj = date.fromisoformat(day)

    entry = session.exec(
        select(InspectionEntry).where(
            InspectionEntry.run_id == run_id,
            InspectionEntry.actual_date == day_obj,
            InspectionEntry.slot_time == slot_time,
        )
    ).first()
    if not entry:
        raise HTTPException(404, "No inspection entry for this slot/day")

    param = session.exec(
        select(RunParameter).where(RunParameter.run_id == run_id, RunParameter.param_key == param_key)
    ).first()
    if not param:
        raise HTTPException(404, "Parameter not found")

    # prevent duplicates
    existing = session.exec(
        select(InspectionValue).where(InspectionValue.entry_id == entry.id, InspectionValue.param_key == param_key)
    ).first()
    if existing:
        return RedirectResponse(f"/runs/{run_id}?day={day_obj.isoformat()}", status_code=302)

    v = _safe_float(new_value)
    if v is None:
        return RedirectResponse(f"/runs/{run_id}?day={day_obj.isoformat()}", status_code=302)

    is_oos, spec_note = apply_spec_check(param, v)

    new_iv = InspectionValue(
        entry_id=entry.id,
        param_key=param_key,
        value=v,
        is_out_of_spec=is_oos,
        spec_note=spec_note,
    )
    session.add(new_iv)
    session.commit()
    session.refresh(new_iv)

    # optional: store audit (if your audit table allows it)
    session.add(InspectionValueAudit(
        inspection_value_id=new_iv.id,
        action="CREATED",
        old_value=None,
        new_value=v,
        by_user_id=user.id,
        by_user_name=user.display_name,
        note=note or "",
    ))
    session.commit()

    return RedirectResponse(f"/runs/{run_id}?day={day_obj.isoformat()}", status_code=302)


# ===== VALUE EDIT + APPROVAL (kept as your working logic) =====
@app.get("/values/{value_id}/edit", response_class=HTMLResponse)
def value_edit_get(value_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)

    v = session.get(InspectionValue, value_id)
    if not v:
        raise HTTPException(404, "Value not found")
    entry = session.get(InspectionEntry, v.entry_id)
    run = session.get(ProductionRun, entry.run_id) if entry else None
    if not entry or not run:
        raise HTTPException(404, "Run/Entry not found")

    if run.status in ["CLOSED", "APPROVED"] and user.role != "MANAGER":
        raise HTTPException(403, "Run is not open")

    param = session.exec(
        select(RunParameter).where(RunParameter.run_id == run.id, RunParameter.param_key == v.param_key)
    ).first()

    return templates.TemplateResponse(
        "value_edit.html",
        {"request": request, "user": user, "run": run, "entry": entry, "param": param, "v": v, "error": ""},
    )


@app.post("/values/{value_id}/edit")
async def value_edit_post(
    value_id: int,
    request: Request,
    session: Session = Depends(get_session),
    new_value: str = Form(...),
    note: str = Form(""),
):
    user = get_current_user(request, session)
    forbid_boss(user)

    v = session.get(InspectionValue, value_id)
    if not v:
        raise HTTPException(404, "Value not found")
    entry = session.get(InspectionEntry, v.entry_id)
    run = session.get(ProductionRun, entry.run_id) if entry else None
    if not entry or not run:
        raise HTTPException(404, "Run/Entry not found")

    if run.status in ["CLOSED", "APPROVED"] and user.role != "MANAGER":
        raise HTTPException(403, "Run is not open")

    nv = _safe_float(new_value)
    if nv is None:
        return RedirectResponse(f"/values/{value_id}/edit", status_code=302)

    v.pending_value = nv
    v.pending_status = "PENDING"
    v.pending_by_user_id = user.id
    v.pending_at = datetime.utcnow()
    session.add(v)

    session.add(InspectionValueAudit(
        inspection_value_id=v.id,
        action="PROPOSED",
        old_value=v.value,
        new_value=nv,
        by_user_id=user.id,
        by_user_name=user.display_name,
        note=note or "",
    ))

    session.commit()
    return RedirectResponse(f"/runs/{run.id}?day={entry.actual_date.isoformat()}", status_code=302)


@app.get("/runs/{run_id}/pending", response_class=HTMLResponse)
def pending_list(run_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)
    require_manager(user)
        

    run = session.get(ProductionRun, run_id)
    if not run:
        raise HTTPException(404, "Run not found")

    pending_items = []
    entries = session.exec(select(InspectionEntry).where(InspectionEntry.run_id == run_id)).all()
    entry_ids = [e.id for e in entries]
    if entry_ids:
        vals = session.exec(
            select(InspectionValue).where(
                InspectionValue.entry_id.in_(entry_ids),
                InspectionValue.pending_status == "PENDING"
            )
        ).all()

        params = session.exec(select(RunParameter).where(RunParameter.run_id == run_id)).all()
        pmap = {p.param_key: p for p in params}
        emap = {e.id: e for e in entries}

        for v in vals:
            pending_items.append({"value": v, "entry": emap.get(v.entry_id), "param": pmap.get(v.param_key)})

    return templates.TemplateResponse(
        "pending_list.html",
        {"request": request, "user": user, "run": run, "items": pending_items},
    )


@app.post("/values/{value_id}/approve")
def value_approve(value_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)
    require_manager(user)
        

    v = session.get(InspectionValue, value_id)
    if not v:
        raise HTTPException(404, "Value not found")
    entry = session.get(InspectionEntry, v.entry_id)
    run = session.get(ProductionRun, entry.run_id) if entry else None
    if not entry or not run:
        raise HTTPException(404, "Run/Entry not found")

    if v.pending_status != "PENDING" or v.pending_value is None:
        return RedirectResponse(f"/runs/{run.id}/pending", status_code=302)

    old = v.value
    new = v.pending_value

    v.value = new
    v.pending_status = "APPROVED"
    session.add(v)

    session.add(InspectionValueAudit(
        inspection_value_id=v.id,
        action="APPROVED",
        old_value=old,
        new_value=new,
        by_user_id=user.id,
        by_user_name=user.display_name,
        note="",
    ))

    v.pending_value = None
    v.pending_status = ""
    v.pending_by_user_id = None
    v.pending_at = None
    session.add(v)

    session.commit()
    return RedirectResponse(f"/runs/{run.id}/pending", status_code=302)


@app.post("/values/{value_id}/reject")
def value_reject(value_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)
    require_manager(user)
       

    v = session.get(InspectionValue, value_id)
    if not v:
        raise HTTPException(404, "Value not found")
    entry = session.get(InspectionEntry, v.entry_id)
    run = session.get(ProductionRun, entry.run_id) if entry else None
    if not entry or not run:
        raise HTTPException(404, "Run/Entry not found")

    if v.pending_status != "PENDING":
        return RedirectResponse(f"/runs/{run.id}/pending", status_code=302)

    session.add(InspectionValueAudit(
        inspection_value_id=v.id,
        action="REJECTED",
        old_value=v.value,
        new_value=v.pending_value,
        by_user_id=user.id,
        by_user_name=user.display_name,
        note="",
    ))

    v.pending_value = None
    v.pending_status = ""
    v.pending_by_user_id = None
    v.pending_at = None
    session.add(v)

    session.commit()
    return RedirectResponse(f"/runs/{run.id}/pending", status_code=302)
from openpyxl.cell.cell import MergedCell
import subprocess
import tempfile
from pathlib import Path
from fastapi.responses import Response


def _set_cell_safe(ws, addr: str, value, number_format: str | None = None):
    """
    Write value to Excel, even if the target address is part of a merged range.
    Also optionally apply number_format to the actual written cell.
    """
    target_addr = addr
    for rng in ws.merged_cells.ranges:
        if addr in rng:
            target_addr = rng.coord.split(":")[0]  # top-left
            break

    c = ws[target_addr]
    c.value = value
    if number_format:
        c.number_format = number_format



def _clone_sheet_no_drawings(wb, src_ws, title: str):
    """
    Clone a worksheet WITHOUT drawings/images (prevents crash).
    Preserves:
      - values
      - styles
      - merged cells
      - row/col dimensions
    """
    dst = wb.create_sheet(title=title)

    for col, dim in src_ws.column_dimensions.items():
        dst.column_dimensions[col].width = dim.width
    for row, dim in src_ws.row_dimensions.items():
        dst.row_dimensions[row].height = dim.height

    for merged_range in list(src_ws.merged_cells.ranges):
        dst.merge_cells(str(merged_range))

    for row in src_ws.iter_rows():
        for cell in row:
            new_cell = dst.cell(row=cell.row, column=cell.col_idx, value=cell.value)
            if cell.has_style:
                new_cell._style = cell._style
                new_cell.number_format = cell.number_format
                new_cell.font = cell.font
                new_cell.border = cell.border
                new_cell.fill = cell.fill
                new_cell.alignment = cell.alignment
                new_cell.protection = cell.protection

    return dst


def convert_xlsx_bytes_to_pdf_bytes(xlsx_bytes: bytes) -> bytes:
    """
    Uses LibreOffice headless to convert XLSX -> PDF.
    Requires 'soffice' available in the runtime image.
    """
    with tempfile.TemporaryDirectory() as tmp:
        tmpdir = Path(tmp)
        xlsx_path = tmpdir / "report.xlsx"
        out_dir = tmpdir / "out"
        out_dir.mkdir(parents=True, exist_ok=True)

        xlsx_path.write_bytes(xlsx_bytes)

        cmd = [
            "soffice",
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--invisible",
            "--convert-to", "pdf",
            "--outdir", str(out_dir),
            str(xlsx_path),
        ]
        # capture output for debugging
        # run conversion and capture output
        try:
            r = subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        except FileNotFoundError:
            raise RuntimeError("PDF export failed: LibreOffice 'soffice' is not installed / not found on this server.")
        except subprocess.CalledProcessError as e:
            out = (e.stdout or b"").decode("utf-8", errors="ignore")
            err = (e.stderr or b"").decode("utf-8", errors="ignore")
            raise RuntimeError(
                "PDF export failed: LibreOffice conversion error.\n\n"
                f"CMD: {' '.join(cmd)}\n\nSTDOUT:\n{out}\n\nSTDERR:\n{err}\n"
            )

        # find produced pdf
        pdfs = list(out_dir.glob("*.pdf"))
        if not pdfs:
            out = (r.stdout or b"").decode("utf-8", errors="ignore")
            err = (r.stderr or b"").decode("utf-8", errors="ignore")
            raise RuntimeError(
                "PDF conversion failed: LibreOffice produced no PDF output.\n\n"
                f"CMD: {' '.join(cmd)}\n\nSTDOUT:\n{out}\n\nSTDERR:\n{err}\n"
            )

        pdf_path = pdfs[0]
        pdf_bytes = pdf_path.read_bytes()

        # ✅ critical validation
        if not pdf_bytes or len(pdf_bytes) < 10:
            out = (r.stdout or b"").decode("utf-8", errors="ignore")
            err = (r.stderr or b"").decode("utf-8", errors="ignore")
            raise RuntimeError(
                "PDF conversion failed: LibreOffice created an EMPTY PDF (0 bytes).\n\n"
                f"PDF path: {pdf_path}\n"
                f"PDF size: {len(pdf_bytes)} bytes\n\n"
                f"CMD: {' '.join(cmd)}\n\nSTDOUT:\n{out}\n\nSTDERR:\n{err}\n"
            )

        # optional but very helpful: check pdf signature
        if not pdf_bytes.startswith(b"%PDF"):
            out = (r.stdout or b"").decode("utf-8", errors="ignore")
            err = (r.stderr or b"").decode("utf-8", errors="ignore")
            head = pdf_bytes[:80]
            raise RuntimeError(
                "PDF conversion failed: output is not a valid PDF.\n\n"
                f"PDF path: {pdf_path}\n"
                f"PDF size: {len(pdf_bytes)} bytes\n"
                f"First bytes: {head!r}\n\n"
                f"CMD: {' '.join(cmd)}\n\nSTDOUT:\n{out}\n\nSTDERR:\n{err}\n"
            )

        return pdf_bytes


def build_one_day_workbook_bytes(run_id: int, day: date, session: Session) -> bytes:
    run = session.get(ProductionRun, run_id)
    if not run:
        raise HTTPException(404, "Run not found")

    template_path = TEMPLATE_XLSX_MAP.get(run.process)
    if not template_path or not os.path.exists(template_path):
        raise HTTPException(404, f"Template not found: {template_path}")

    # Load the template fresh (images/logos stay)
    wb = openpyxl.load_workbook(template_path)
    ws = wb.worksheets[0]

    # ✅ Print setup (keep 1-page)
    apply_pdf_page_setup(ws)
    apply_specs_to_template(ws, run, session)
    

    # ✅ Per-process coordinates (THIS was missing and causing crashes)
    if run.process in ["LINER", "COVER"]:
        col_start = 5  # E
        date_row = 20
        time_row = 21
        inspector_row = 38
        op1_row = 39
        op2_row = 40
        row_map = ROW_MAP_LINER_COVER

        # Header cells (LINER/COVER)
        _set_cell_safe(ws, "E5", run.dhtp_batch_no)
        _set_cell_safe(ws, "I5", run.client_name)
        _set_cell_safe(ws, "I6", run.po_number)
        _set_cell_safe(ws, "E6", run.pipe_specification)
        _set_cell_safe(ws, "E7", run.raw_material_spec)
        _set_cell_safe(ws, "E9", run.itp_number)

    else:  # REINFORCEMENT
        col_start = 6   # F
        date_row = 19
        time_row = 20
        inspector_row = 35
        op1_row = 36
        op2_row = 37
        row_map = ROW_MAP_REINF

        # ✅ IMPORTANT:
        # Put ONLY the reinforcement header cells here (do NOT write E5/I5 again)
        # Use the real cells from your reinforcement.xlsx template
        _set_cell_safe(ws, "D4", run.dhtp_batch_no)
        _set_cell_safe(ws, "I4", run.client_name)
        _set_cell_safe(ws, "I5", run.po_number)
        _set_cell_safe(ws, "D5", run.pipe_specification)
        _set_cell_safe(ws, "D6", run.raw_material_spec)
        _set_cell_safe(ws, "D8", run.itp_number)

    # Machines Used
    machines = session.exec(select(RunMachine).where(RunMachine.run_id == run_id)).all()
    for idx in range(5):
        r = 4 + idx
        name = machines[idx].machine_name if idx < len(machines) else ""
        tag = machines[idx].machine_tag if (idx < len(machines) and machines[idx].machine_tag) else ""
        _set_cell_safe(ws, f"M{r}", name)
        _set_cell_safe(ws, f"P{r}", tag)
    
    # ✅ Date row (IMPORTANT: format as DATE so it doesn't show 12:00 AM)
    for slot_idx, slot in enumerate(SLOTS):
        col = openpyxl.utils.get_column_letter(col_start + slot_idx)
        _set_cell_safe(ws, f"{col}{date_row}", day, number_format="m/d/yyyy")
    
    # ✅ Time row: leave blank by default (we will fill actual times from entries)
    for slot_idx, _slot in enumerate(SLOTS):
        col = openpyxl.utils.get_column_letter(col_start + slot_idx)
        _set_cell_safe(ws, f"{col}{time_row}", "", number_format="h:mm")
    


    # Trace for THIS day: raw batch + tools
    trace_today = get_day_latest_trace(session, run_id, day)
    carry = get_last_known_trace_before_day(session, run_id, day)

    raw_batches = trace_today["raw_batches"] or ([carry["raw"]] if carry["raw"] else [])
    raw_str = ", ".join([x for x in raw_batches if x])
    if raw_str:
        if run.process in ["LINER", "COVER"]:
            _set_cell_safe(ws, "E8", raw_str)   # LINER/COVER: Raw Material Batch No.
        else:
            _set_cell_safe(ws, "D7", raw_str)   # REINFORCEMENT (keep as-is)


    tools = trace_today["tools"] or carry["tools"]
    for t_idx in range(2):
        r = 8 + t_idx
        if t_idx < len(tools):
            name, serial, calib = tools[t_idx]
            _set_cell_safe(ws, f"G{r}", name or "")
            _set_cell_safe(ws, f"I{r}", serial or "")
            _set_cell_safe(ws, f"K{r}", calib or "")

    apply_specs_to_template(ws, run, session)

    # Fill inspector/operators per slot + values
    day_entries = session.exec(
        select(InspectionEntry)
        .where(InspectionEntry.run_id == run_id, InspectionEntry.actual_date == day)
        .order_by(InspectionEntry.created_at)
    ).all()

    user_map = {u.id: u for u in session.exec(select(User)).all()}

    for e in day_entries:
        if e.slot_time not in SLOTS:
            continue
        slot_idx = SLOTS.index(e.slot_time)
        col = openpyxl.utils.get_column_letter(col_start + slot_idx)
        # ✅ write ACTUAL time under the slot header (export uses actual time)
        try:
            hh, mm = (e.actual_time or "00:00").split(":")
            _set_cell_safe(ws, f"{col}{time_row}", dtime(int(hh), int(mm)), number_format="h:mm")
        except Exception:
            # if actual_time is weird, keep it blank instead of crashing export
            pass


        inspector_name = user_map.get(e.inspector_id).display_name if e.inspector_id in user_map else ""
        _set_cell_safe(ws, f"{col}{inspector_row}", inspector_name)

        if run.process in ["LINER", "COVER"]:
            _set_cell_safe(ws, f"{col}{op1_row}", e.operator_1 or "")
            _set_cell_safe(ws, f"{col}{op2_row}", e.operator_2 or "")
        else:
            _set_cell_safe(ws, f"{col}{op1_row}", e.operator_annular_12 or "")
            _set_cell_safe(ws, f"{col}{op2_row}", e.operator_int_ext_34 or "")

        vals = session.exec(select(InspectionValue).where(InspectionValue.entry_id == e.id)).all()
        for v in vals:
            r = row_map.get(v.param_key)
            if not r:
                continue
            _set_cell_safe(ws, f"{col}{r}", v.value)

    out = BytesIO()
    wb.save(out)
    out.seek(0)
    return out.getvalue()


from pypdf import PdfReader, PdfWriter
from pypdf import Transformation
from io import BytesIO
import os






def build_export_xlsx_bytes(run_id: int, request: Request, session: Session) -> tuple[bytes, str]:
    """
    Build the XLSX export in memory and return (bytes, filename_base).
    This is used by BOTH /export/xlsx and /export/pdf.
    """
    user = get_current_user(request, session)

    run = session.get(ProductionRun, run_id)
    if not run:
        raise HTTPException(404, "Run not found")

    template_path = TEMPLATE_XLSX_MAP.get(run.process)
    if not template_path or not os.path.exists(template_path):
        raise HTTPException(404, f"Template not found: {template_path}")

    days = get_days_for_run(session, run_id)
    if not days:
        raise HTTPException(400, "No entries to export")

    base_wb = openpyxl.load_workbook(template_path)
    base_ws = base_wb.worksheets[0]

    # per process coordinates
    if run.process in ["LINER", "COVER"]:
        col_start = 5  # E
        date_row = 20
        inspector_row = 38
        op1_row = 39
        op2_row = 40
        row_map = ROW_MAP_LINER_COVER
    else:
        col_start = 6  # F
        date_row = 20
        inspector_row = 36
        op1_row = 37
        op2_row = 38
        row_map = ROW_MAP_REINF

    machines = session.exec(select(RunMachine).where(RunMachine.run_id == run_id)).all()
    user_map = {u.id: u for u in session.exec(select(User)).all()}

    for i, day in enumerate(days):
        title = f"Day {i+1} ({day.isoformat()})"

        if i == 0:
            ws = base_ws
            ws.title = title
        else:
            ws = _clone_sheet_no_drawings(base_wb, base_ws, title)

        # ----- HEADER (fixed cells) -----
        _set_cell_safe(ws, "E5", run.dhtp_batch_no)      # Batch
        _set_cell_safe(ws, "I5", run.client_name)        # Client
        _set_cell_safe(ws, "I6", run.po_number)          # PO
        _set_cell_safe(ws, "E6", run.pipe_specification) # Pipe Spec
        _set_cell_safe(ws, "E7", run.raw_material_spec)  # Raw Spec
        _set_cell_safe(ws, "E9", run.itp_number)         # ITP

        # ----- Machines (M4:P8) -----
        for idx in range(5):
            r = 4 + idx
            name = machines[idx].machine_name if idx < len(machines) else ""
            tag = machines[idx].machine_tag if (idx < len(machines) and machines[idx].machine_tag) else ""
            _set_cell_safe(ws, f"M{r}", name)
            _set_cell_safe(ws, f"P{r}", tag)

        # ----- Day trace (raw batch + tools) -----
        trace_today = get_day_latest_trace(session, run_id, day)
        carry = get_last_known_trace_before_day(session, run_id, day)

        raw_batches = trace_today["raw_batches"] or ([carry["raw"]] if carry["raw"] else [])
        raw_str = ", ".join([x for x in raw_batches if x])
        if raw_str:
            if run.process in ["LINER", "COVER"]:
                _set_cell_safe(ws, "E8", raw_str)
            else:
                _set_cell_safe(ws, "D7", raw_str)


        tools = trace_today["tools"] or carry["tools"]
        for t_idx in range(2):
            r = 8 + t_idx
            if t_idx < len(tools):
                name, serial, calib = tools[t_idx]
                _set_cell_safe(ws, f"G{r}", name or "")
                _set_cell_safe(ws, f"I{r}", serial or "")
                _set_cell_safe(ws, f"K{r}", calib or "")

        # ----- Date row for each time slot -----
        for slot_idx, slot in enumerate(SLOTS):
            col = openpyxl.utils.get_column_letter(col_start + slot_idx)
            _set_cell_safe(ws, f"{col}{date_row}", day)
        # ----- Time header row: blank by default (we will write actual times if entry exists) -----
        time_row = 21
        for slot_idx, _slot in enumerate(SLOTS):
            col = openpyxl.utils.get_column_letter(col_start + slot_idx)
            _set_cell_safe(ws, f"{col}{time_row}", "", number_format="h:mm")

        # ✅ print setup (important for PDF export)
        apply_pdf_page_setup(ws)
                

        
        # ----- Fill per-slot inspector/operators + values -----
        day_entries = session.exec(
            select(InspectionEntry)
            .where(InspectionEntry.run_id == run_id, InspectionEntry.actual_date == day)
            .order_by(InspectionEntry.created_at)
        ).all()

        for e in day_entries:
            if e.slot_time not in SLOTS:
                continue

            slot_idx = SLOTS.index(e.slot_time)
            col = openpyxl.utils.get_column_letter(col_start + slot_idx)

            inspector_name = user_map.get(e.inspector_id).display_name if e.inspector_id in user_map else ""
            _set_cell_safe(ws, f"{col}{inspector_row}", inspector_name)

            if run.process in ["LINER", "COVER"]:
                _set_cell_safe(ws, f"{col}{op1_row}", e.operator_1 or "")
                _set_cell_safe(ws, f"{col}{op2_row}", e.operator_2 or "")
            else:
                _set_cell_safe(ws, f"{col}{op1_row}", e.operator_annular_12 or "")
                _set_cell_safe(ws, f"{col}{op2_row}", e.operator_int_ext_34 or "")

            vals = session.exec(select(InspectionValue).where(InspectionValue.entry_id == e.id)).all()
            for v in vals:
                r = row_map.get(v.param_key)
                if not r:
                    continue
                _set_cell_safe(ws, f"{col}{r}", v.value)

            # ✅ write ACTUAL time for this slot
            try:
                hh, mm = (e.actual_time or "00:00").split(":")
                _set_cell_safe(ws, f"{col}{time_row}", dtime(int(hh), int(mm)), number_format="h:mm")
            except Exception:
                pass


    out = BytesIO()
    base_wb.save(out)
    out.seek(0)

    filename_base = f"{run.process}_{run.dhtp_batch_no}_ALL_DAYS"
    return out.getvalue(), filename_base
    

# =========================
# MRR EXPORT (per MaterialLot)
# =========================

def build_mrr_xlsx_bytes(lot_id: int, session: Session) -> bytes:
    """
    Build an MRR XLSX directly in code (no template file needed).
    This avoids missing-template issues and makes export stable.
    """
    lot = session.get(MaterialLot, lot_id)
    if not lot:
        raise HTTPException(404, "MRR Ticket not found")

    docs = session.exec(
        select(MrrDocument)
        .where(MrrDocument.ticket_id == lot_id)
        .order_by(MrrDocument.created_at.asc())
    ).all()

    receiving = session.exec(
        select(MrrReceiving).where(MrrReceiving.ticket_id == lot_id)
    ).first()

    shipments = session.exec(
        select(MrrReceivingInspection)
        .where(
            (MrrReceivingInspection.ticket_id == lot_id) &
            (MrrReceivingInspection.inspector_confirmed == True)
        )
        .order_by(MrrReceivingInspection.created_at.asc())
    ).all()

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "MRR"

    # Basic formatting
    ws["A1"] = f"MRR Ticket #{lot.id}"
    ws["A1"].font = openpyxl.styles.Font(bold=True, size=16)

    row = 3
    def put(label, value):
        nonlocal row
        ws[f"A{row}"] = label
        ws[f"B{row}"] = value
        ws[f"A{row}"].font = openpyxl.styles.Font(bold=True)
        row += 1

    put("Type", lot.lot_type or "")
    put("Batch No", lot.batch_no or "")
    put("Status", lot.status or "")
    put("Material", lot.material_name or "")
    put("Supplier", lot.supplier_name or "")
    put("PO Number", lot.po_number or "")
    put("PO Quantity", f"{lot.quantity or 0} {lot.quantity_unit or ''}".strip())

    # received_total is stored normalized in KG for weight units
    unit = (lot.quantity_unit or "KG").upper().strip()
    if unit in ["PC", "PCS"]:
        put("Received Total", f"{lot.received_total or 0} {unit}")
    else:
        put("Received Total", f"{lot.received_total or 0} KG (normalized)")

    row += 1

    # Documentation status
    ws[f"A{row}"] = "Documentation"
    ws[f"A{row}"].font = openpyxl.styles.Font(bold=True, size=12)
    row += 1

    if receiving:
        put("Saved", "YES")
        cleared = bool(getattr(receiving, "inspector_confirmed_po", False) or getattr(receiving, "manager_confirmed_po", False))
        put("Cleared", "YES" if cleared else "NO")

        put("Inspector PO No.", getattr(receiving, "inspector_po_number", "") or "")

        # Receiving doc qty fields (these DO exist in your model)
        doc_qty = getattr(receiving, "qty_arrived", None)
        doc_unit = getattr(receiving, "qty_unit", "KG") or "KG"
        put("Arrived Qty (Doc)", f"{doc_qty if doc_qty is not None else ''} {doc_unit}".strip())

        is_partial = bool(getattr(receiving, "is_partial_delivery", False))
        put("Partial Delivery", "YES" if is_partial else "NO")
        put("Mismatch/Partial Reason", getattr(receiving, "qty_mismatch_reason", "") or "")

        put("Received Date", str(getattr(receiving, "received_date", "") or ""))
        put("Remarks (Doc)", getattr(receiving, "remarks", "") or "")
    else:
        put("Saved", "NO")
        put("Cleared", "NO")

    row += 1

    # Documents table
    ws[f"A{row}"] = "Documents"
    ws[f"A{row}"].font = openpyxl.styles.Font(bold=True, size=12)
    row += 1

    headers = ["Type", "Name", "Number", "Uploaded By", "Uploaded At"]
    for col, h in enumerate(headers, start=1):
        cell = ws.cell(row=row, column=col, value=h)
        cell.font = openpyxl.styles.Font(bold=True)
    row += 1

    for d in docs:
        ws.cell(row=row, column=1, value=d.doc_type or "")
        ws.cell(row=row, column=2, value=d.doc_name or "")
        ws.cell(row=row, column=3, value=d.doc_number or "")
        ws.cell(row=row, column=4, value=d.uploaded_by_user_name or "")
        ws.cell(row=row, column=5, value=str(getattr(d, "created_at", "") or ""))
        row += 1

    row += 1

    # Shipments table
    ws[f"A{row}"] = "Submitted Shipments"
    ws[f"A{row}"].font = openpyxl.styles.Font(bold=True, size=12)
    row += 1

    ship_headers = ["#", "DN", "Arrived Qty", "Unit", "Report No", "Submitted", "Manager Approved"]
    for col, h in enumerate(ship_headers, start=1):
        cell = ws.cell(row=row, column=col, value=h)
        cell.font = openpyxl.styles.Font(bold=True)
    row += 1

    for idx, s in enumerate(shipments, start=1):
        ws.cell(row=row, column=1, value=idx)
        ws.cell(row=row, column=2, value=s.delivery_note_no or "")
        ws.cell(row=row, column=3, value=float(s.qty_arrived or 0))
        ws.cell(row=row, column=4, value=s.qty_unit or "")
        ws.cell(row=row, column=5, value=s.report_no or "")
        ws.cell(row=row, column=6, value="YES" if s.inspector_confirmed else "NO")
        ws.cell(row=row, column=7, value="YES" if s.manager_approved else "NO")
        row += 1

    # Column widths
    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 40
    ws.column_dimensions["C"].width = 18
    ws.column_dimensions["D"].width = 20
    ws.column_dimensions["E"].width = 22
    ws.column_dimensions["F"].width = 18
    ws.column_dimensions["G"].width = 18

    apply_pdf_page_setup(ws)

    out = BytesIO()
    wb.save(out)
    out.seek(0)
    return out.getvalue()


def build_mrr_photo_appendix_pdf_bytes(lot_id: int, session: Session) -> bytes:
    """
    Build a PDF appendix containing all photo evidence (grouped) for submitted shipments.
    Appended to the base MRR PDF.
    """
    submitted_inspections = session.exec(
        select(MrrReceivingInspection)
        .where(
            (MrrReceivingInspection.ticket_id == lot_id) &
            (MrrReceivingInspection.inspector_confirmed == True)
        )
        .order_by(MrrReceivingInspection.created_at.asc())
    ).all()

    insp_ids = [i.id for i in submitted_inspections if i.id is not None]
    if not insp_ids:
        return b""

    photos = session.exec(
        select(MrrInspectionPhoto)
        .where(
            (MrrInspectionPhoto.ticket_id == lot_id) &
            (MrrInspectionPhoto.inspection_id.in_(insp_ids))
        )
        .order_by(MrrInspectionPhoto.created_at.asc())
    ).all()

    if not photos:
        return b""

    insp_by_id = {i.id: i for i in submitted_inspections if i.id is not None}

    grouped: Dict[int, Dict[str, List[MrrInspectionPhoto]]] = {}
    for p in photos:
        grouped.setdefault(int(p.inspection_id), {})
        g = (p.group_name or "General").strip() or "General"
        grouped[int(p.inspection_id)].setdefault(g, []).append(p)

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    page_w, page_h = A4
    margin = 46  # ~0.65 inch

    def header(title: str, subtitle: str = ""):
        y = page_h - margin
        c.setFont("Helvetica-Bold", 16)
        c.drawString(margin, y, title)
        y -= 18
        if subtitle:
            c.setFont("Helvetica", 10)
            c.drawString(margin, y, subtitle)
            y -= 14
        c.setLineWidth(0.8)
        c.line(margin, y, page_w - margin, y)
        return y - 16

    y = header("MRR Photo Evidence Appendix", f"Ticket #{lot_id}")
    c.setFont("Helvetica", 10)
    c.drawString(margin, y, "This appendix contains photo evidence attached to submitted shipment inspections.")
    c.showPage()

    max_w = page_w - 2 * margin
    max_h = 320  # keep room for captions

    for iid in insp_ids:
        insp = insp_by_id.get(iid)
        if not insp:
            continue

        dn = (insp.delivery_note_no or "").strip() or "-"
        rep = (insp.report_no or "").strip() or "-"

        y = header("Shipment Inspection Photos", f"DN: {dn} | Report: {rep} | Inspection ID: {iid}")

        for gname, items in grouped.get(iid, {}).items():
            if y < margin + 160:
                c.showPage()
                y = header("Shipment Inspection Photos", f"DN: {dn} | Report: {rep} | Inspection ID: {iid}")

            c.setFont("Helvetica-Bold", 12)
            c.drawString(margin, y, f"Group: {gname}")
            y -= 16

            for p in items:
                path = p.file_path or ""
                if not path or not os.path.exists(path):
                    continue

                caption = (p.caption or "").strip()

                needed = max_h + (40 if caption else 20)
                if y < margin + needed:
                    c.showPage()
                    y = header("Shipment Inspection Photos", f"DN: {dn} | Report: {rep} | Inspection ID: {iid}")
                    c.setFont("Helvetica-Bold", 12)
                    c.drawString(margin, y, f"Group: {gname}")
                    y -= 16

                try:
                    img = ImageReader(path)
                    iw, ih = img.getSize()
                    if iw <= 0 or ih <= 0:
                        continue
                    scale = min(max_w / iw, max_h / ih)
                    dw = iw * scale
                    dh = ih * scale

                    c.drawImage(img, margin, y - dh, width=dw, height=dh, preserveAspectRatio=True, mask="auto")
                    y -= dh + 8

                    if caption:
                        c.setFont("Helvetica", 9)
                        c.drawString(margin, y, caption[:180])
                        y -= 14
                    else:
                        y -= 8

                    c.setLineWidth(0.3)
                    c.line(margin, y, page_w - margin, y)
                    y -= 12
                except Exception:
                    continue

        c.showPage()

    c.save()
    out = buf.getvalue()
    buf.close()
    return out


def merge_pdf_bytes(base_pdf: bytes, appendix_pdf: bytes) -> bytes:
    if not appendix_pdf:
        return base_pdf

    base_reader = PdfReader(BytesIO(base_pdf))
    app_reader = PdfReader(BytesIO(appendix_pdf))

    w = PdfWriter()
    for p in base_reader.pages:
        w.add_page(p)
    for p in app_reader.pages:
        w.add_page(p)

    out = BytesIO()
    w.write(out)
    out.seek(0)
    return out.getvalue()


@app.get("/mrr/{lot_id}/export/xlsx")
def mrr_export_xlsx(lot_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)  # allow any logged-in reviewer

    xlsx_bytes = build_mrr_xlsx_bytes(lot_id, session)
    filename = f"MRR_{lot_id}.xlsx"

    return Response(
        content=xlsx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/mrr/{lot_id}/export/pdf")
def mrr_export_pdf(lot_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)  # allow any logged-in reviewer

    # 1) Build XLSX
    xlsx_bytes = build_mrr_xlsx_bytes(lot_id, session)

    # 2) Convert to PDF (existing pipeline)
    pdf_bytes = convert_xlsx_bytes_to_pdf_bytes(xlsx_bytes)

    # 3) Append photos
    appendix = build_mrr_photo_appendix_pdf_bytes(lot_id, session)
    final_pdf = merge_pdf_bytes(pdf_bytes, appendix)

    filename = f"MRR_{lot_id}.pdf"
    return Response(
        content=final_pdf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )

# =========================
# MRR MANAGER APPROVAL
# =========================

@app.post("/mrr/{lot_id}/inspection/id/{inspection_id}/approve")
def mrr_approve_receiving_inspection(
    lot_id: int,
    inspection_id: int,
    request: Request,
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)
    require_manager(user)

    lot = session.get(MaterialLot, lot_id)
    if not lot:
        raise HTTPException(404, "MRR Ticket not found")

    if (lot.status or "").upper() == "CANCELED":
        return RedirectResponse(f"/mrr/{lot_id}?error=Ticket%20is%20canceled", status_code=303)

    insp = session.get(MrrReceivingInspection, inspection_id)
    if not insp or insp.ticket_id != lot_id:
        raise HTTPException(404, "MRR Inspection not found")

    # Must be submitted first
    if not insp.inspector_confirmed:
        return RedirectResponse(
            f"/mrr/{lot_id}?error=Inspection%20must%20be%20submitted%20before%20approval",
            status_code=303,
        )

    # Approve this inspection
    insp.manager_approved = True
    session.add(insp)

        # ✅ Save manager identity + approval timestamp into JSON (no guessing)
    try:
        data = json.loads(getattr(insp, "inspection_json", None) or "{}")
    except Exception:
        data = {}

    data["manager_approved_by"] = (getattr(user, "display_name", "") or "").strip()
    data["manager_approved_at_utc"] = datetime.utcnow().isoformat()

    insp.inspection_json = json.dumps(data, ensure_ascii=False)


    # =========================
    # Recompute totals from APPROVED shipments (source of truth)
    # =========================
    approved_shipments = session.exec(
        select(MrrReceivingInspection).where(
            (MrrReceivingInspection.ticket_id == lot_id)
            & (MrrReceivingInspection.inspector_confirmed == True)
            & (MrrReceivingInspection.manager_approved == True)
        )
    ).all()

    po_unit = (lot.quantity_unit or "KG").upper().strip()
    try:
        po_qty = float(lot.quantity or 0.0)
    except Exception:
        po_qty = 0.0

    approved_total = 0.0

    # If PO is PCS/PC => sum as-is in PCS
    if po_unit in ["PC", "PCS"]:
        for s in approved_shipments:
            try:
                approved_total += float(s.qty_arrived or 0.0)
            except Exception:
                pass

        remaining = po_qty - approved_total
        if remaining < 0:
            remaining = 0.0

        lot.received_total = approved_total  # PCS total
        lot.status = "APPROVED" if remaining <= 0 else ("PARTIAL" if approved_total > 0 else "PENDING")

    else:
        # Weight-based => normalize everything into KG
        po_kg = normalize_qty_to_kg(po_qty, po_unit)

        for s in approved_shipments:
            try:
                approved_total += float(normalize_qty_to_kg(float(s.qty_arrived or 0.0), s.qty_unit or "KG"))
            except Exception:
                pass

        remaining_kg = po_kg - approved_total
        if remaining_kg < 0:
            remaining_kg = 0.0

        lot.received_total = approved_total  # KG normalized total
        lot.status = "APPROVED" if remaining_kg <= 0 else ("PARTIAL" if approved_total > 0 else "PENDING")

    session.add(lot)
    session.commit()

    return RedirectResponse(f"/mrr/{lot_id}", status_code=303)

@app.post("/mrr/{lot_id}/inspection/id/{inspection_id}/unapprove")
def mrr_unapprove_receiving_inspection(
    lot_id: int,
    inspection_id: int,
    request: Request,
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)
    require_manager(user)

    lot = session.get(MaterialLot, lot_id)
    if not lot:
        raise HTTPException(404, "MRR Ticket not found")

    if (lot.status or "").upper() == "CANCELED":
        return RedirectResponse(f"/mrr/{lot_id}?error=Ticket%20is%20canceled", status_code=303)

    insp = session.get(MrrReceivingInspection, inspection_id)
    if not insp or insp.ticket_id != lot_id:
        raise HTTPException(404, "MRR Inspection not found")

    # Only unapprove if it was approved already
    if not insp.manager_approved:
        return RedirectResponse(
            f"/mrr/{lot_id}?error=This%20inspection%20is%20not%20approved",
            status_code=303,
        )

    # Unapprove
    insp.manager_approved = False
    session.add(insp)

    # =========================
    # Recompute totals after unapprove
    # =========================
    approved_shipments = session.exec(
        select(MrrReceivingInspection).where(
            (MrrReceivingInspection.ticket_id == lot_id)
            & (MrrReceivingInspection.inspector_confirmed == True)
            & (MrrReceivingInspection.manager_approved == True)
        )
    ).all()

    po_unit = (lot.quantity_unit or "KG").upper().strip()
    try:
        po_qty = float(lot.quantity or 0.0)
    except Exception:
        po_qty = 0.0

    approved_total = 0.0

    if po_unit in ["PC", "PCS"]:
        for s in approved_shipments:
            try:
                approved_total += float(s.qty_arrived or 0.0)
            except Exception:
                pass

        remaining = po_qty - approved_total
        if remaining < 0:
            remaining = 0.0

        lot.received_total = approved_total
        # If nothing approved anymore -> go back to PENDING. Else PARTIAL.
        lot.status = "PARTIAL" if approved_total > 0 else "PENDING"

    else:
        po_kg = normalize_qty_to_kg(po_qty, po_unit)

        for s in approved_shipments:
            try:
                approved_total += float(normalize_qty_to_kg(float(s.qty_arrived or 0.0), s.qty_unit or "KG"))
            except Exception:
                pass

        remaining_kg = po_kg - approved_total
        if remaining_kg < 0:
            remaining_kg = 0.0

        lot.received_total = approved_total  # KG normalized
        lot.status = "PARTIAL" if approved_total > 0 else "PENDING"

    session.add(lot)
    session.commit()

    return RedirectResponse(f"/mrr/{lot_id}", status_code=303)


@app.get("/runs/{run_id}/export/xlsx")
def export_xlsx(run_id: int, request: Request, session: Session = Depends(get_session)):
    xlsx_bytes, filename_base = build_export_xlsx_bytes(run_id, request, session)

    return Response(
        content=xlsx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename_base}.xlsx"'},
    )




from fastapi import HTTPException
from starlette.responses import Response

@app.get("/runs/{run_id}/export/pdf")
def export_pdf(run_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)

    run = session.get(ProductionRun, run_id)
    if not run:
        raise HTTPException(404, "Run not found")

    days = get_days_for_run(session, run_id)
    if not days:
        raise HTTPException(400, "No entries to export")

    writer = PdfWriter()

    for day in days:
        # 1) Build 1-day excel
        xlsx_bytes = build_one_day_workbook_bytes(run_id, day, session)

        # 2) Convert to PDF
        pdf_bytes = convert_xlsx_bytes_to_pdf_bytes(xlsx_bytes)

        # 3) Stamp approval if approved
        if (run.status or "").upper() == "APPROVED":
            pdf_bytes = stamp_approval_on_pdf(
                pdf_bytes,
                approved_by=getattr(run, "approved_by_user_name", "") or "",
                approved_at_utc=getattr(run, "approved_at_utc", None),
            )

        # 4) Merge pages
        reader = PdfReader(BytesIO(pdf_bytes))
        for page in reader.pages:
            writer.add_page(page)

    # 5) Write final output ONCE
    out = BytesIO()
    writer.write(out)
    out.seek(0)

    filename = f"{run.process}_{run.dhtp_batch_no}_ALL_DAYS.pdf"
    return Response(
        content=out.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )



def apply_pdf_page_setup(ws):
    # A4 Portrait
    ws.page_setup.orientation = ws.ORIENTATION_PORTRAIT
    ws.page_setup.paperSize = ws.PAPERSIZE_A4

    # ✅ Back to "everything in 1 page"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 1
    ws.sheet_properties.pageSetUpPr.fitToPage = True

    # Margins (keep as before)
    ws.page_margins.left = 0.25
    ws.page_margins.right = 0.25
    ws.page_margins.top = 0.35
    ws.page_margins.bottom = 0.70


# =========================
# MRR PHOTO EVIDENCE
# =========================

def _safe_ext(filename: str) -> str:
    name = (filename or "").lower().strip()
    if "." not in name:
        return ".jpg"
    ext = "." + name.split(".")[-1]
    if ext not in [".jpg", ".jpeg", ".png", ".webp"]:
        return ".jpg"
    return ext


@app.get("/mrr/photos/{photo_id}/view")
def mrr_photo_view(photo_id: int, request: Request, session: Session = Depends(get_session)):
    user = get_current_user(request, session)

    p = session.get(MrrInspectionPhoto, photo_id)
    if not p:
        raise HTTPException(404, "Photo not found")

    # (optional) basic permission gate: any logged-in user can view
    # You can tighten later by checking ticket access rules if needed.

    resolved = resolve_mrr_photo_path(p.file_path)
    if not resolved or not os.path.exists(resolved):
        raise HTTPException(404, "Photo file missing")

    mt, _ = mimetypes.guess_type(resolved)
    return FileResponse(resolved, media_type=mt or "image/jpeg")


@app.post("/mrr/{lot_id}/inspection/id/{inspection_id}/photos/upload")
async def mrr_photo_upload(
    lot_id: int,
    inspection_id: int,
    request: Request,
    session: Session = Depends(get_session),
    group_name: str = Form(...),
    caption: str = Form(""),
    photos: List[UploadFile] = File(...),
):
    user = get_current_user(request, session)
    forbid_boss(user)

    lot = session.get(MaterialLot, lot_id)
    if not lot:
        raise HTTPException(404, "MRR Ticket not found")

    insp = session.get(MrrReceivingInspection, inspection_id)
    if not insp or insp.ticket_id != lot_id:
        raise HTTPException(404, "Shipment inspection not found")

    # Prevent edits after approval (manager lock)
    if insp.manager_approved:
        return RedirectResponse(
            f"/mrr/{lot_id}/inspection/id/{inspection_id}?photo_error=Inspection%20is%20approved.%20Photos%20cannot%20be%20changed.",
            status_code=303,
        )

    g = (group_name or "General").strip() or "General"
    cap = (caption or "").strip()

    if not photos or len(photos) == 0:
        return RedirectResponse(
            f"/mrr/{lot_id}/inspection/id/{inspection_id}?photo_error=Please%20select%20at%20least%20one%20photo",
            status_code=303,
        )

    # Store under: DATA_DIR/mrr_photos/<ticket>/<inspection>/
    base = os.path.join(MRR_PHOTO_DIR, f"ticket_{lot_id}", f"insp_{inspection_id}")
    os.makedirs(base, exist_ok=True)

    for f in photos:
        # basic image validation
        ct = (f.content_type or "").lower()
        if not ct.startswith("image/"):
            continue

        ext = _safe_ext(f.filename)
        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S_%f")
        out_path = os.path.join(base, f"{ts}{ext}")

        content = await f.read()
        with open(out_path, "wb") as w:
            w.write(content)

        rec = MrrInspectionPhoto(
            ticket_id=lot_id,
            inspection_id=inspection_id,
            group_name=g,
            caption=cap,
            file_path=os.path.relpath(out_path, MRR_PHOTO_DIR),
            uploaded_by_user_id=user.id,
            uploaded_by_user_name=user.display_name,
        )
        session.add(rec)

    session.commit()

    return RedirectResponse(f"/mrr/{lot_id}/inspection/id/{inspection_id}", status_code=303)


@app.post("/mrr/{lot_id}/inspection/id/{inspection_id}/photos/{photo_id}/delete")
def mrr_photo_delete(
    lot_id: int,
    inspection_id: int,
    photo_id: int,
    request: Request,
    session: Session = Depends(get_session),
):
    user = get_current_user(request, session)
    forbid_boss(user)

    lot = session.get(MaterialLot, lot_id)
    if not lot:
        raise HTTPException(404, "MRR Ticket not found")

    insp = session.get(MrrReceivingInspection, inspection_id)
    if not insp or insp.ticket_id != lot_id:
        raise HTTPException(404, "Shipment inspection not found")

    # Prevent edits after approval (manager lock)
    if insp.manager_approved:
        return RedirectResponse(f"/mrr/{lot_id}/inspection/id/{inspection_id}", status_code=303)

    p = session.get(MrrInspectionPhoto, photo_id)
    if not p or p.ticket_id != lot_id or p.inspection_id != inspection_id:
        raise HTTPException(404, "Photo not found")

    # Delete file if exists
    try:
        if p.file_path and os.path.exists(p.file_path):
            os.remove(p.file_path)
    except Exception:
        pass

    session.delete(p)
    session.commit()

    return RedirectResponse(f"/mrr/{lot_id}/inspection/id/{inspection_id}", status_code=303)




