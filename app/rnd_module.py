from __future__ import annotations

import math
import os
import json
from datetime import datetime, date
from typing import Optional, List

from fastapi import APIRouter, Depends, Form, HTTPException, Request, UploadFile, File
from fastapi.responses import RedirectResponse, FileResponse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fastapi.templating import Jinja2Templates
from sqlmodel import SQLModel, Field, Session, select
from pathlib import Path
import shutil
import uuid
import mimetypes

from .db import get_session
from .models import (
    User,
    RndQualificationProgram,
    RndQualificationTest,
    RndQualificationSpecimen,
    RndMaterialQualification,
    RndMaterialTestRecord,
    RndAttachmentRegister,
)

router = APIRouter(prefix="/rnd", tags=["R&D Qualification"])
TEMPLATES = Jinja2Templates(directory=os.path.join(os.path.dirname(__file__), "templates"))
RCRT_HOURS = 175000.0
CYCLIC_BASIS_CYCLES = 1_000_000.0
DESIGN_FACTOR_NONMETALLIC = 0.67

BASE_DIR = Path(__file__).resolve().parent
RND_UPLOAD_DIR = BASE_DIR / "uploaded_rnd_files"
RND_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

ALLOWED_RND_FILE_EXTENSIONS = {
    ".pdf", ".doc", ".docx", ".xls", ".xlsx", ".csv",
    ".jpg", ".jpeg", ".png", ".webp",
    ".txt", ".zip"
}

def _safe_filename(filename: str) -> str:
    raw = (filename or "").strip()
    if not raw:
        raw = "file"
    safe = "".join(c if c.isalnum() or c in {".", "-", "_"} else "_" for c in raw)
    return safe[:200]

def _program_upload_dir(program_id: int) -> Path:
    folder = RND_UPLOAD_DIR / f"program_{program_id}"
    folder.mkdir(parents=True, exist_ok=True)
    return folder

def _save_rnd_upload(program_id: int, uploaded_file: UploadFile) -> dict:
    if uploaded_file is None:
        raise HTTPException(status_code=400, detail="No file uploaded.")

    original_name = _safe_filename(uploaded_file.filename or "file")
    suffix = Path(original_name).suffix.lower()

    if suffix not in ALLOWED_RND_FILE_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"File type not allowed: {suffix or 'unknown'}"
        )

    target_dir = _program_upload_dir(program_id)
    stored_name = f"{uuid.uuid4().hex}{suffix}"
    target_path = target_dir / stored_name

    size_bytes = 0
    with target_path.open("wb") as buffer:
        while True:
            chunk = uploaded_file.file.read(1024 * 1024)
            if not chunk:
                break
            size_bytes += len(chunk)
            buffer.write(chunk)

    content_type = uploaded_file.content_type or mimetypes.guess_type(original_name)[0] or "application/octet-stream"

    return {
        "original_filename": original_name,
        "stored_filename": stored_name,
        "file_path": str(target_path),
        "content_type": content_type,
        "file_size_bytes": size_bytes,
    }
def _fmt_date(value) -> str:
    if not value:
        return "-"
    try:
        return value.strftime("%Y-%m-%d")
    except Exception:
        return str(value)

def _fmt_datetime(value) -> str:
    if not value:
        return "-"
    try:
        return value.strftime("%Y-%m-%d %H:%M UTC")
    except Exception:
        return str(value)

def _fmt_number(value, digits: int = 2) -> str:
    if value is None or value == "":
        return "-"
    try:
        return f"{float(value):,.{digits}f}"
    except Exception:
        return str(value)

def _fmt_bool(value) -> str:
    return "Yes" if bool(value) else "No"

def _generated_report_options(test_code: str) -> list[dict]:
    code = (test_code or "").strip().upper()

    common = [
        {"key": "TECHNICAL_REPORT", "label": "Technical Report"},
        {"key": "RESULT_SHEET", "label": "Result Sheet"},
    ]

    if code in {"MPR_REG", "CYCLIC_REG"}:
        return [
            {"key": "TECHNICAL_REPORT", "label": "Technical Report"},
            {"key": "REGRESSION_REPORT", "label": "Regression Report"},
            {"key": "REGRESSION_GRAPH", "label": "Regression Graph"},
            {"key": "PRESSURE_LOG", "label": "Pressure Log"},
            {"key": "RESULT_SHEET", "label": "Result Sheet"},
        ]

    if code in {"PV_1000H", "TEMP_ELEV", "TEMP_CYCLE", "RAPID_DECOMP", "OPERATING_MBR", "AXIAL_LOAD", "CRUSH", "LAOT", "IMPACT"}:
        return [
            {"key": "TECHNICAL_REPORT", "label": "Technical Report"},
            {"key": "RESULT_SHEET", "label": "Result Sheet"},
            {"key": "PRESSURE_LOG", "label": "Pressure Log"},
        ]

    if code in {"TEC", "GROWTH"}:
        return [
            {"key": "TECHNICAL_REPORT", "label": "Technical Report"},
            {"key": "RESULT_SHEET", "label": "Result Sheet"},
            {"key": "DATA_SHEET", "label": "Data Sheet"},
        ]

    return common

def _preferred_generated_document_type(test_code: str) -> str:
    options = _generated_report_options(test_code)
    return options[0]["key"] if options else "TECHNICAL_REPORT"

def _auto_report_reference(program: RndQualificationProgram, test: RndQualificationTest) -> str:
    stamp = datetime.utcnow().strftime("%Y%m%d-%H%M%S")
    code = (program.program_code or f"PROGRAM-{program.id}").replace(" ", "-")
    test_code = (test.code or "TEST").replace(" ", "-")
    return f"RND-{code}-{test_code}-{stamp}"

def _auto_report_title(test: RndQualificationTest, document_type: str) -> str:
    base = (test.title or test.code or "Qualification Test").strip()
    doc_type = (document_type or "TECHNICAL_REPORT").strip().upper()

    mapping = {
        "TECHNICAL_REPORT": f"{base} Technical Report",
        "REGRESSION_REPORT": f"{base} Regression Report",
        "REGRESSION_GRAPH": f"{base} Regression Graph",
        "RESULT_SHEET": f"{base} Result Sheet",
        "PRESSURE_LOG": f"{base} Pressure Log",
        "DATA_SHEET": f"{base} Data Sheet",
    }
    return mapping.get(doc_type, f"{base} Report")

DEFAULT_SPECIMEN_RULE = {
    "min_specimens": "As defined by the applicable qualification test requirement.",
    "minimum_cut_length": "Specimens shall be cut on the basis of total cut length and shall not be cut to active or effective test length only.",
    "effective_length": "The effective test length shall be the active section defined by the applicable test setup, fixture arrangement, support span, or pressurized free length.",
    "od_basis_notice": "Before any diameter-based preparation rule is applied, the actual qualified pipe outside diameter (OD) shall be confirmed from the controlled product definition, dimensional record, or approved size specification. Diameter-based rules shall not be calculated from nominal size designation alone unless the nominal size has already been mapped to a controlled OD value in the qualification basis.",
    "calculation_rule_notice": "Any preparation rule expressed as D, OD, 1D, 3D, 5D, or 1.5 x OD shall be calculated using the confirmed pipe outside diameter and then converted into millimetres for specimen release and cutting control.",
    "recorded_dimensions_notice": "Before cutting, the specimen register shall record the confirmed pipe OD, the rule basis used, the calculated effective length, the end allowances, the trimming margin, and the total cut length in millimetres.",
    "end_allowance_each_side": "Unless a stricter test-specific requirement applies, each specimen shall include an end allowance on both sides sufficient for gripping, sealing, end fittings, support, trimming, and handling. The baseline preparation allowance shall be not less than 1.0 x outside diameter on each side.",
    "total_length_formula": "Total cut length = effective test length + left end allowance + right end allowance + trimming margin.",
    "marking_requirements": [
        "Each specimen shall be assigned a unique identification number before cutting.",
        "Each specimen shall be marked with pipe size, batch or lot reference, and intended test code.",
        "Where applicable, the centerline or active test section shall be marked prior to preparation.",
        "Identification markings should be placed outside the critical observation or expected failure zone whenever practical."
    ],
    "visual_acceptance": [
        "Specimens shall be free from visible cuts, gouges, cracks, crushed areas, and other preparation damage.",
        "Specimens shall be free from obvious ovality or deformation caused by handling.",
        "Cut ends shall be square and suitable for the required end preparation.",
        "Traceability markings shall remain legible after preparation."
    ],
    "preconditioning": {
        "required": "depends",
        "when_required": "Preconditioning shall be applied whenever required by the qualification basis, applicable procedure, test condition, or environmental exposure requirement.",
        "medium": "Ambient air unless otherwise required by the applicable test method or internal procedure.",
        "target_temperature": "As required by the selected test condition.",
        "minimum_process": [
            "Prepare, cut, and identify the specimen.",
            "Verify dimensions, cut quality, and end preparation before conditioning.",
            "Place the specimen in the required conditioning environment at the specified target temperature.",
            "Allow the specimen to stabilize before commencement of testing.",
            "Record conditioning start time, target temperature, observed temperature, and responsible operator.",
            "Testing shall not begin until specimen stabilization has been confirmed."
        ],
        "records_required": [
            "Conditioning start time",
            "Conditioning end time or release time",
            "Target temperature",
            "Observed temperature",
            "Conditioning medium or environment",
            "Operator or approver"
        ],
    },
    "technician_tips": [
        "Do not cut specimens to effective test length only.",
        "Record both total cut length and effective test length before release for test.",
        "Where fixture or fitting engagement length is uncertain, confirm the requirement before cutting.",
        "Reject any specimen damaged during cutting, machining, or handling."
    ],
    "release_checks": [
        "Specimen identification is assigned and traceable to production batch.",
        "Total cut length is measured and recorded.",
        "Effective test length is identified.",
        "End preparation is complete.",
        "No visible damage is present after preparation.",
        "Preconditioning is complete where required.",
        "Specimen is released for testing by the responsible person."
    ],
}

SPECIMEN_PREP_RULES = {
    "mpr_reg": {
        "min_specimens": 18,
        "minimum_cut_length": "Specimens for long-term hydrostatic regression shall be cut to provide a consistent free pressurized section for the regression setup together with sufficient additional length for end terminations, sealing, and trimming.",
        "effective_length": "The effective length shall be the free pressurized section between end terminations used for long-term hydrostatic exposure.",
        "end_allowance_each_side": "A preparation allowance of not less than 1.5 x outside diameter shall be provided on each side as a baseline for long-term end termination and sealing requirements. Greater allowance shall be used where the termination system requires additional engagement length.",
        "total_length_formula": "Total cut length = effective regression section + 2 x termination allowance + trimming margin.",
        "preconditioning": {
            "required": True,
            "when_required": "Required before pressurization at the selected regression test temperature.",
            "medium": "Controlled temperature environment matching the regression condition",
            "target_temperature": "Selected regression temperature",
        },
    },
    "pv_1000h": {
        "min_specimens": 2,
        "minimum_cut_length": "Specimens for the 1000-hour pressure confirmation test shall be cut to provide the required active pressurized section together with full end-sealing allowance.",
        "effective_length": "The effective length shall be the active pressurized section between end fittings for the PV confirmation specimen.",
        "end_allowance_each_side": "A baseline sealing allowance of not less than 1.0 x outside diameter shall be provided on each side.",
        "preconditioning": {
            "required": True,
            "when_required": "Condition before the 1000-hour hold at the selected test temperature.",
            "medium": "Controlled environment",
            "target_temperature": "Selected PV confirmation temperature"
        },
    },
    "temp_elev": {
        "min_specimens": 1,
        "minimum_cut_length": "Specimens shall be cut to provide the active elevated-temperature test section together with sufficient end engagement and trimming allowance.",
        "effective_length": "The effective length shall be the section exposed to the elevated temperature and pressure condition.",
        "preconditioning": {
            "required": True,
            "when_required": "Elevated-temperature stabilization is required before test start.",
            "medium": "Controlled temperature environment",
            "target_temperature": "Selected elevated test temperature"
        },
    },
    "temp_cycle": {
        "min_specimens": 2,
        "minimum_cut_length": "Specimens shall be cut to provide the active thermal cycling section together with secure end engagement for the cycling setup.",
        "effective_length": "The effective length shall be the section subjected to thermal cycling.",
        "preconditioning": {
            "required": True,
            "when_required": "Stabilization at the starting temperature is required before thermal cycling begins.",
            "medium": "Controlled temperature environment",
            "target_temperature": "Cycle start temperature"
        },
    },
    "rapid_decomp": {
        "min_specimens": 1,
        "minimum_cut_length": "Specimens shall be cut to provide the decompression exposure section together with safe end engagement and sealing allowance.",
        "effective_length": "The effective length shall be the pressurized section exposed to gas decompression.",
        "preconditioning": {
            "required": True,
            "when_required": "Conditioning shall be applied as required by the gas charging and temperature setup.",
            "medium": "Gas exposure environment",
            "target_temperature": "Selected rapid decompression temperature"
        },
    },
    "operating_mbr": {
        "min_specimens": 2,
        "minimum_cut_length": "Specimens shall be prepared using the full length required to achieve the specified bending radius and safe handling arrangement.",
        "effective_length": "The effective length shall be the full section involved in bending or respooling exposure.",
        "end_allowance_each_side": "Additional length shall be included as required for gripping, handling, and fixture engagement.",
    },
    "axial_load": {
        "min_specimens": 2,
        "minimum_cut_length": "Specimens shall be cut to provide the required gauge section together with sufficient end gripping allowance.",
        "effective_length": "The effective length shall be the section over which axial response is evaluated.",
    },
    "crush": {
        "min_specimens": 3,
        "minimum_cut_length": "Specimens shall be cut to a length sufficient to suit the crush test arrangement with proper support and alignment.",
        "effective_length": "The effective length shall be the section subjected to radial external loading.",
    },
    "laot": {
        "min_specimens": 2,
        "minimum_cut_length": "Specimens shall be cut to provide the low-temperature qualification section together with the required end engagement.",
        "effective_length": "The effective length shall be the section exposed to the lowest allowable operating temperature qualification condition.",
        "preconditioning": {
            "required": True,
            "when_required": "Condition at the selected low temperature before test start.",
            "medium": "Low-temperature environment",
            "target_temperature": "Selected LAOT"
        },
    },
    "impact": {
        "min_specimens": 2,
        "minimum_cut_length": "Specimens shall be cut to provide the required support span or impact location together with trimming margin.",
        "effective_length": "The effective length shall be the supported test section containing the impact location.",
        "end_allowance_each_side": "Preparation allowance shall be based on the support arrangement rather than pressure sealing requirements.",
        "preconditioning": {
            "required": True,
            "when_required": "Required when impact testing is performed at a controlled test temperature.",
            "medium": "Conditioning environment at test temperature",
            "target_temperature": "Selected impact test temperature"
        },
    },
    "tec": {
        "min_specimens": 2,
        "minimum_cut_length": "Specimens shall be cut to provide sufficient length for axial response measurement over the defined gauge or effective section.",
        "effective_length": "The effective length shall be the section used for thermal expansion coefficient measurement.",
        "preconditioning": {
            "required": True,
            "when_required": "Thermal stabilization is required before thermal expansion measurement.",
            "medium": "Controlled temperature environment",
            "target_temperature": "Selected TEC temperature"
        },
    },
    "growth": {
        "min_specimens": 2,
        "minimum_cut_length": "Specimens shall be cut to provide sufficient length for measurement of growth or shrinkage over the intended gauge section.",
        "effective_length": "The effective length shall be the section over which dimensional response is recorded.",
    },
    "cyclic_reg": {
        "min_specimens": 18,
        "minimum_cut_length": "Specimens for cyclic pressure regression shall be cut to provide the active cycling section together with full termination or fitting allowance.",
        "effective_length": "The effective length shall be the active section subjected to repeated pressure cycling.",
        "end_allowance_each_side": "Sufficient length shall be provided for termination and sealing without compromising the active cycling section.",
        "preconditioning": {
            "required": True,
            "when_required": "Required when the cyclic test is conducted at controlled temperature or conditioned service state.",
            "medium": "Controlled environment as required by the test condition",
            "target_temperature": "Selected cyclic qualification temperature"
        },
    },
}

def _save_generated_test_report_docx(
    *,
    program: RndQualificationProgram,
    test: RndQualificationTest,
    guidance: dict,
    specimens: list,
    attachments: list,
    evidence: dict,
    materials: list,
    document_type: str,
) -> dict:
    target_dir = _program_upload_dir(program.id)
    stored_name = f"{uuid.uuid4().hex}.docx"
    target_path = target_dir / stored_name

    safe_document_type = (document_type or "TECHNICAL_REPORT").strip().upper()
    pretty_type = safe_document_type.replace("_", " ").title()

    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(9)

    title = doc.add_heading(_auto_report_title(test, safe_document_type), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Program: {program.program_code or '-'} | Test Code: {test.code or '-'}").bold = True
    subtitle.add_run(f"\nDocument Type: {pretty_type}")
    subtitle.add_run(f"\nGenerated: {_fmt_datetime(datetime.utcnow())}")

    doc.add_heading("1. Qualification summary", level=1)
    summary = doc.add_table(rows=0, cols=2)
    summary.style = "Table Grid"
    for label, value in [
        ("Program code", program.program_code or "-"),
        ("Program title", program.title or "-"),
        ("Qualification standard", program.qualification_standard or "-"),
        ("Route", program.pfr_or_pv or "-"),
        ("Nominal size (in)", _fmt_number(program.nominal_size_in)),
        ("NPR (MPa)", _fmt_number(program.npr_mpa)),
        ("MAOT (°C)", _fmt_number(program.maot_c)),
        ("LAOT (°C)", _fmt_number(program.laot_c)),
        ("Service medium", program.service_medium or "-"),
        ("Test code", test.code or "-"),
        ("Test title", test.title or "-"),
        ("Clause reference", test.clause_ref or "-"),
        ("Status", test.status or "-"),
        ("Result summary", test.result_summary or "-"),
    ]:
        cells = summary.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)

    if safe_document_type == "REGRESSION_GRAPH":
        doc.add_heading("2. Regression graph summary", level=1)
        p = doc.add_paragraph()
        p.add_run("Purpose: ").bold = True
        p.add_run("This file is intended to hold the regression graph issue record for this test.")

        graph_table = doc.add_table(rows=1, cols=5)
        graph_table.style = "Table Grid"
        headers = ["Specimen ID", "Failure Hours", "Failure Pressure (MPa)", "Failure Mode", "Include in Regression"]
        for idx, h in enumerate(headers):
            graph_table.rows[0].cells[idx].text = h

        if specimens:
            for s in specimens:
                row = graph_table.add_row().cells
                row[0].text = s.specimen_id or "-"
                row[1].text = _fmt_number(s.failure_hours)
                row[2].text = _fmt_number(s.actual_pressure_at_failure_mpa)
                row[3].text = s.failure_mode or "-"
                row[4].text = _fmt_bool(s.include_in_regression)
        else:
            row = graph_table.add_row().cells
            row[0].text = "No regression specimens recorded"
            for idx in range(1, 5):
                row[idx].text = "-"

        doc.add_paragraph("Note: chart image can be inserted later after external plotting / signed issue, then uploaded back as signed evidence.")
    elif safe_document_type == "PRESSURE_LOG":
        doc.add_heading("2. Pressure log", level=1)
        pressure_table = doc.add_table(rows=1, cols=8)
        pressure_table.style = "Table Grid"
        headers = ["Specimen ID", "Date", "Planned Pressure", "Actual Pressure", "Hold Pressure", "Failure Time", "Result", "Notes"]
        for idx, h in enumerate(headers):
            pressure_table.rows[0].cells[idx].text = h

        if specimens:
            for s in specimens:
                row = pressure_table.add_row().cells
                row[0].text = s.specimen_id or "-"
                row[1].text = _fmt_date(s.sample_date)
                row[2].text = _fmt_number(s.planned_pressure_mpa)
                row[3].text = _fmt_number(s.actual_pressure_at_failure_mpa)
                row[4].text = _fmt_number(s.pressure_at_hold_mpa)
                row[5].text = _fmt_number(s.failure_time_sec)
                row[6].text = s.result_status or "-"
                row[7].text = s.notes or "-"
        else:
            row = pressure_table.add_row().cells
            row[0].text = "No pressure records recorded"
            for idx in range(1, 8):
                row[idx].text = "-"
    elif safe_document_type in {"RESULT_SHEET", "DATA_SHEET"}:
        doc.add_heading("2. Result sheet", level=1)
        result_table = doc.add_table(rows=1, cols=10)
        result_table.style = "Table Grid"
        headers = [
            "Specimen ID",
            "Material Ref",
            "Sample Date",
            "Target / Planned",
            "Actual",
            "Failure Mode",
            "Failure Location",
            "Result",
            "QA Review",
            "Remarks",
        ]
        for idx, h in enumerate(headers):
            result_table.rows[0].cells[idx].text = h

        if specimens:
            for s in specimens:
                row = result_table.add_row().cells
                row[0].text = s.specimen_id or "-"
                row[1].text = s.material_ref or "-"
                row[2].text = _fmt_date(s.sample_date)
                row[3].text = _fmt_number(s.planned_pressure_mpa)
                row[4].text = _fmt_number(s.actual_pressure_at_failure_mpa)
                row[5].text = s.failure_mode or "-"
                row[6].text = s.failure_location or "-"
                row[7].text = s.result_status or "-"
                row[8].text = s.qa_review_status or "-"
                row[9].text = s.notes or "-"
        else:
            row = result_table.add_row().cells
            row[0].text = "No specimen results recorded"
            for idx in range(1, 10):
                row[idx].text = "-"
    else:
        doc.add_heading("2. Test guidance and acceptance basis", level=1)
        for label, value in [
            ("When required", guidance.get("when_required", "")),
            ("Specimen count", guidance.get("specimen_count", "")),
            ("API clause", guidance.get("api_clause", "")),
            ("External standard", guidance.get("external_standard", "")),
            ("Conditioning required", guidance.get("conditioning_required", "")),
            ("Retest logic", guidance.get("retest_logic", "")),
        ]:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(value or "-")

        for heading, items in [
            ("Conditioning steps", guidance.get("conditioning_steps", [])),
            ("Core process", guidance.get("core_process", [])),
            ("Acceptance criteria", guidance.get("acceptance", [])),
            ("Practical notes", guidance.get("practical_notes", [])),
        ]:
            doc.add_paragraph(heading, style="List Bullet")
            if items:
                for item in items:
                    doc.add_paragraph(str(item), style="List Bullet 2")
            else:
                doc.add_paragraph("-", style="List Bullet 2")

        doc.add_heading("3. Material register", level=1)
        material_table = doc.add_table(rows=1, cols=6)
        material_table.style = "Table Grid"
        headers = ["Component", "Material", "Manufacturer", "Grade", "Batch / Lot", "Status"]
        for idx, h in enumerate(headers):
            material_table.rows[0].cells[idx].text = h

        if materials:
            for m in materials:
                row = material_table.add_row().cells
                row[0].text = m.component or "-"
                row[1].text = m.material_name or "-"
                row[2].text = m.manufacturer_name or "-"
                row[3].text = m.grade_name or "-"
                row[4].text = " / ".join([x for x in [m.batch_ref, m.lot_ref] if x]) or "-"
                row[5].text = m.status or "-"
        else:
            row = material_table.add_row().cells
            row[0].text = "No materials recorded"
            for idx in range(1, 6):
                row[idx].text = "-"

        doc.add_heading("4. Specimen execution record", level=1)
        specimen_table = doc.add_table(rows=1, cols=10)
        specimen_table.style = "Table Grid"
        specimen_headers = [
            "Specimen ID",
            "Date",
            "Material ref",
            "Planned P (MPa)",
            "Actual P (MPa)",
            "Hours",
            "Cycles",
            "Result",
            "QA review",
            "Failure mode",
        ]
        for idx, h in enumerate(specimen_headers):
            specimen_table.rows[0].cells[idx].text = h

        if specimens:
            for s in specimens:
                row = specimen_table.add_row().cells
                row[0].text = s.specimen_id or "-"
                row[1].text = _fmt_date(s.sample_date)
                row[2].text = s.material_ref or "-"
                row[3].text = _fmt_number(s.planned_pressure_mpa)
                row[4].text = _fmt_number(s.actual_pressure_at_failure_mpa)
                row[5].text = _fmt_number(s.failure_hours)
                row[6].text = _fmt_number(s.failure_cycles)
                row[7].text = s.result_status or "-"
                row[8].text = s.qa_review_status or "-"
                row[9].text = s.failure_mode or "-"
        else:
            row = specimen_table.add_row().cells
            row[0].text = "No specimens recorded"
            for idx in range(1, 10):
                row[idx].text = "-"

        if specimens:
            doc.add_heading("5. Detailed specimen observations", level=1)
            for idx, s in enumerate(specimens, start=1):
                p = doc.add_paragraph()
                p.add_run(f"{idx}. {s.specimen_id or f'Specimen {idx}'}").bold = True

                details = [
                    f"Batch ref: {s.batch_ref or '-'}",
                    f"Source pipe ref: {s.source_pipe_ref or '-'}",
                    f"Preparation rule: {s.preparation_rule_basis or '-'}",
                    f"Conditioning complete: {_fmt_bool(s.conditioning_complete)}",
                    f"Pre-test visual OK: {_fmt_bool(s.pretest_visual_ok)}",
                    f"Released for test: {_fmt_bool(s.released_for_test)}",
                    f"Failure location: {s.failure_location or '-'}",
                    f"Failure description: {s.failure_description or '-'}",
                    f"Leak observation: {s.leak_observation or '-'}",
                    f"Notes: {s.notes or '-'}",
                ]
                for item in details:
                    doc.add_paragraph(item, style="List Bullet 2")

        doc.add_heading("6. Evidence package status", level=1)
        evidence_table = doc.add_table(rows=1, cols=3)
        evidence_table.style = "Table Grid"
        for idx, h in enumerate(["Document type", "Required", "Present"]):
            evidence_table.rows[0].cells[idx].text = h

        for row_data in evidence.get("rows", []):
            row = evidence_table.add_row().cells
            row[0].text = row_data.get("document_type", "-")
            row[1].text = "Yes" if row_data.get("document_type") in evidence.get("required", []) else "No"
            row[2].text = "Yes" if row_data.get("present") else "No"

        missing = evidence.get("missing", [])
        p = doc.add_paragraph()
        p.add_run("Missing evidence items: ").bold = True
        p.add_run(", ".join(missing) if missing else "None")

        doc.add_heading("7. Attachment register", level=1)
        attachment_table = doc.add_table(rows=1, cols=6)
        attachment_table.style = "Table Grid"
        for idx, h in enumerate(["Title", "Type", "Source", "Status", "Signed", "File name"]):
            attachment_table.rows[0].cells[idx].text = h

        if attachments:
            for a in attachments:
                row = attachment_table.add_row().cells
                row[0].text = a.title or "-"
                row[1].text = a.document_type or "-"
                row[2].text = a.source_mode or "-"
                row[3].text = a.approval_status or "-"
                row[4].text = _fmt_bool(a.is_signed_copy)
                row[5].text = a.original_filename or "-"
        else:
            row = attachment_table.add_row().cells
            row[0].text = "No attachments recorded"
            for idx in range(1, 6):
                row[idx].text = "-"

        doc.add_heading("8. Approval / close-out note", level=1)
        close_note = doc.add_paragraph()
        close_note.add_run("Readiness for closure: ").bold = True
        if specimens and not missing and (test.result_summary or "").strip():
            close_note.add_run("Ready for technical review and issue.")
        else:
            close_note.add_run("Not yet ready for closure. Complete missing evidence, specimen execution, or result summary first.")

    doc.save(target_path)

    return {
        "original_filename": _safe_filename(f"{(program.program_code or 'program')}_{(test.code or 'test')}_{safe_document_type.lower()}.docx"),
        "stored_filename": stored_name,
        "file_path": str(target_path),
        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "file_size_bytes": target_path.stat().st_size,
    }
    
def get_specimen_prep(test_code: str):
    key = (test_code or '').strip().lower()
    base = {
        **DEFAULT_SPECIMEN_RULE,
        "marking_requirements": list(DEFAULT_SPECIMEN_RULE["marking_requirements"]),
        "visual_acceptance": list(DEFAULT_SPECIMEN_RULE["visual_acceptance"]),
        "technician_tips": list(DEFAULT_SPECIMEN_RULE["technician_tips"]),
        "release_checks": list(DEFAULT_SPECIMEN_RULE["release_checks"]),
        "preconditioning": {
            **DEFAULT_SPECIMEN_RULE["preconditioning"],
            "minimum_process": list(DEFAULT_SPECIMEN_RULE["preconditioning"]["minimum_process"]),
            "records_required": list(DEFAULT_SPECIMEN_RULE["preconditioning"]["records_required"]),
        },
    }
    override = SPECIMEN_PREP_RULES.get(key, {})
    for k, v in override.items():
        if isinstance(v, dict) and k in base and isinstance(base[k], dict):
            merged = dict(base[k]); merged.update(v); base[k] = merged
        else:
            base[k] = v
    return base


def get_test_guidance(test_code: str, test_row: Optional["RndQualificationTest"] = None) -> dict:
    code = (test_code or "").strip().upper()
    base = TEST_GUIDANCE.get(code, {
        "when_required": "Refer to the approved qualification basis.",
        "specimen_count": "As required by the approved route.",
        "api_clause": "",
        "external_standard": "",
        "conditioning_required": "Refer to applicable procedure",
        "conditioning_steps": ["Refer to the approved procedure and setup instructions."],
        "core_process": ["Perform the test in accordance with the approved method."],
        "acceptance": ["Apply the approved acceptance basis."],
        "retest_logic": "Follow the approved retest logic.",
        "practical_notes": [],
    })

    enriched = {
        "when_required": base.get("when_required", ""),
        "specimen_count": base.get("specimen_count", ""),
        "api_clause": base.get("api_clause", ""),
        "external_standard": base.get("external_standard", ""),
        "conditioning_required": base.get("conditioning_required", ""),
        "conditioning_steps": list(base.get("conditioning_steps", [])),
        "core_process": list(base.get("core_process", [])),
        "acceptance": list(base.get("acceptance", [])),
        "retest_logic": base.get("retest_logic", ""),
        "practical_notes": list(base.get("practical_notes", [])),
    }

    if test_row:
        if (test_row.guidance_when_required_override or "").strip():
            enriched["when_required"] = test_row.guidance_when_required_override.strip()

        if (test_row.guidance_specimen_count_override or "").strip():
            enriched["specimen_count"] = test_row.guidance_specimen_count_override.strip()
        elif test_row.specimen_count is not None and test_row.specimen_count > 0:
            unit = "specimen" if test_row.specimen_count == 1 else "specimens"
            enriched["specimen_count"] = f"{test_row.specimen_count} {unit}"
        elif (test_row.specimen_requirement or "").strip():
            enriched["specimen_count"] = test_row.specimen_requirement.strip()

        if (test_row.guidance_api_clause_override or "").strip():
            enriched["api_clause"] = test_row.guidance_api_clause_override.strip()

        if (test_row.guidance_external_standard_override or "").strip():
            enriched["external_standard"] = test_row.guidance_external_standard_override.strip()

        if (test_row.guidance_conditioning_required_override or "").strip():
            enriched["conditioning_required"] = test_row.guidance_conditioning_required_override.strip()

        if (test_row.guidance_conditioning_steps_override or "").strip():
            enriched["conditioning_steps"] = [
                line.strip() for line in test_row.guidance_conditioning_steps_override.splitlines() if line.strip()
            ]

        if (test_row.guidance_core_process_override or "").strip():
            enriched["core_process"] = [
                line.strip() for line in test_row.guidance_core_process_override.splitlines() if line.strip()
            ]

        if (test_row.guidance_acceptance_override or "").strip():
            enriched["acceptance"] = [
                line.strip() for line in test_row.guidance_acceptance_override.splitlines() if line.strip()
            ]

        if (test_row.guidance_retest_logic_override or "").strip():
            enriched["retest_logic"] = test_row.guidance_retest_logic_override.strip()

        if (test_row.guidance_practical_notes_override or "").strip():
            enriched["practical_notes"] = [
                line.strip() for line in test_row.guidance_practical_notes_override.splitlines() if line.strip()
            ]

    test_procedure = []
    for item in enriched.get("conditioning_steps", []):
        test_procedure.append(item)
    for item in enriched.get("core_process", []):
        test_procedure.append(item)

    operator_checks = [
        "Verify specimen identity, batch traceability, and test assignment before setup.",
        "Confirm the applicable fittings, fixtures, and calibrated instruments are available.",
        "Confirm conditioning has been completed where required.",
        "Record setup condition before applying load, pressure, temperature, or cycling.",
        "Capture any abnormal observation immediately and do not rely on memory after the test.",
    ]

    records_to_capture = [
        "Specimen ID",
        "Operator / witness",
        "Date and time",
        "Applied setup / fixture basis",
        "Pressure / temperature / time / cycles as applicable",
        "Observed failure mode or survival condition",
        "Acceptance decision",
    ]

    if code == "MPR_REG":
        records_to_capture.extend([
            "Failure hours",
            "Pressure at failure",
            "Failure location",
            "Permissible / excluded failure decision",
        ])
    elif code == "PV_1000H":
        records_to_capture.extend([
            "Hold pressure",
            "Elapsed hours",
            "Leak / survival result",
        ])
    elif code == "TEMP_CYCLE":
        records_to_capture.extend([
            "Cycle range",
            "Cycle count",
            "Leak / post-cycle condition",
        ])
    elif code == "RAPID_DECOMP":
        records_to_capture.extend([
            "Soak pressure",
            "Soak duration",
            "Decompression result",
            "Damage / blister / disbondment observation",
        ])
    elif code == "IMPACT":
        records_to_capture.extend([
            "Impact setup / energy",
            "Post-impact condition",
            "Follow-up proof result",
        ])
    elif code in {"AXIAL_LOAD", "AXIAL"}:
        records_to_capture.extend([
            "Applied axial load",
            "Hold duration",
            "Post-load proof result",
        ])

    enriched["test_procedure"] = test_procedure
    enriched["operator_checks"] = operator_checks
    enriched["records_to_capture"] = records_to_capture
    return enriched

def _fmt_date(value) -> str:
    if not value:
        return "-"
    try:
        return value.strftime("%Y-%m-%d")
    except Exception:
        return str(value)

def _fmt_datetime(value) -> str:
    if not value:
        return "-"
    try:
        return value.strftime("%Y-%m-%d %H:%M UTC")
    except Exception:
        return str(value)

def _fmt_number(value, digits: int = 2) -> str:
    if value is None or value == "":
        return "-"
    try:
        return f"{float(value):,.{digits}f}"
    except Exception:
        return str(value)

def _fmt_bool(value) -> str:
    return "Yes" if bool(value) else "No"

def _generated_report_options(test_code: str) -> list[dict]:
    code = (test_code or "").strip().upper()

    common = [
        {"key": "TECHNICAL_REPORT", "label": "Technical Report"},
        {"key": "RESULT_SHEET", "label": "Result Sheet"},
    ]

    if code in {"MPR_REG", "CYCLIC_REG"}:
        return [
            {"key": "TECHNICAL_REPORT", "label": "Technical Report"},
            {"key": "REGRESSION_REPORT", "label": "Regression Report"},
            {"key": "REGRESSION_GRAPH", "label": "Regression Graph"},
            {"key": "PRESSURE_LOG", "label": "Pressure Log"},
            {"key": "RESULT_SHEET", "label": "Result Sheet"},
        ]

    if code in {"PV_1000H", "TEMP_ELEV", "TEMP_CYCLE", "RAPID_DECOMP", "OPERATING_MBR", "AXIAL_LOAD", "CRUSH", "LAOT", "IMPACT"}:
        return [
            {"key": "TECHNICAL_REPORT", "label": "Technical Report"},
            {"key": "RESULT_SHEET", "label": "Result Sheet"},
            {"key": "PRESSURE_LOG", "label": "Pressure Log"},
        ]

    if code in {"TEC", "GROWTH"}:
        return [
            {"key": "TECHNICAL_REPORT", "label": "Technical Report"},
            {"key": "RESULT_SHEET", "label": "Result Sheet"},
            {"key": "DATA_SHEET", "label": "Data Sheet"},
        ]

    return common

def _preferred_generated_document_type(test_code: str) -> str:
    options = _generated_report_options(test_code)
    return options[0]["key"] if options else "TECHNICAL_REPORT"

def _auto_report_reference(program: RndQualificationProgram, test: RndQualificationTest) -> str:
    stamp = datetime.utcnow().strftime("%Y%m%d-%H%M%S")
    code = (program.program_code or f"PROGRAM-{program.id}").replace(" ", "-")
    test_code = (test.code or "TEST").replace(" ", "-")
    return f"RND-{code}-{test_code}-{stamp}"

def _auto_report_title(test: RndQualificationTest, document_type: str) -> str:
    base = (test.title or test.code or "Qualification Test").strip()
    doc_type = (document_type or "TECHNICAL_REPORT").strip().upper()

    mapping = {
        "TECHNICAL_REPORT": f"{base} Technical Report",
        "REGRESSION_REPORT": f"{base} Regression Report",
        "REGRESSION_GRAPH": f"{base} Regression Graph",
        "RESULT_SHEET": f"{base} Result Sheet",
        "PRESSURE_LOG": f"{base} Pressure Log",
        "DATA_SHEET": f"{base} Data Sheet",
    }
    return mapping.get(doc_type, f"{base} Report")

def _save_generated_test_report_docx(
    *,
    program: RndQualificationProgram,
    test: RndQualificationTest,
    guidance: dict,
    specimens: list,
    attachments: list,
    evidence: dict,
    materials: list,
    document_type: str,
) -> dict:
    target_dir = _program_upload_dir(program.id)
    stored_name = f"{uuid.uuid4().hex}.docx"
    target_path = target_dir / stored_name

    safe_document_type = (document_type or "TECHNICAL_REPORT").strip().upper()
    pretty_type = safe_document_type.replace("_", " ").title()

    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(9)

    title = doc.add_heading(_auto_report_title(test, safe_document_type), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Program: {program.program_code or '-'} | Test Code: {test.code or '-'}").bold = True
    subtitle.add_run(f"\nDocument Type: {pretty_type}")
    subtitle.add_run(f"\nGenerated: {_fmt_datetime(datetime.utcnow())}")

    doc.add_heading("1. Qualification summary", level=1)
    summary = doc.add_table(rows=0, cols=2)
    summary.style = "Table Grid"
    for label, value in [
        ("Program code", program.program_code or "-"),
        ("Program title", program.title or "-"),
        ("Qualification standard", program.qualification_standard or "-"),
        ("Route", program.pfr_or_pv or "-"),
        ("Nominal size (in)", _fmt_number(program.nominal_size_in)),
        ("NPR (MPa)", _fmt_number(program.npr_mpa)),
        ("MAOT (°C)", _fmt_number(program.maot_c)),
        ("LAOT (°C)", _fmt_number(program.laot_c)),
        ("Service medium", program.service_medium or "-"),
        ("Test code", test.code or "-"),
        ("Test title", test.title or "-"),
        ("Clause reference", test.clause_ref or "-"),
        ("Status", test.status or "-"),
        ("Result summary", test.result_summary or "-"),
    ]:
        cells = summary.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)

    if safe_document_type == "REGRESSION_GRAPH":
        doc.add_heading("2. Regression graph summary", level=1)
        p = doc.add_paragraph()
        p.add_run("Purpose: ").bold = True
        p.add_run("This file is intended to hold the regression graph issue record for this test.")

        graph_table = doc.add_table(rows=1, cols=5)
        graph_table.style = "Table Grid"
        headers = ["Specimen ID", "Failure Hours", "Failure Pressure (MPa)", "Failure Mode", "Include in Regression"]
        for idx, h in enumerate(headers):
            graph_table.rows[0].cells[idx].text = h

        if specimens:
            for s in specimens:
                row = graph_table.add_row().cells
                row[0].text = s.specimen_id or "-"
                row[1].text = _fmt_number(s.failure_hours)
                row[2].text = _fmt_number(s.actual_pressure_at_failure_mpa)
                row[3].text = s.failure_mode or "-"
                row[4].text = _fmt_bool(s.include_in_regression)
        else:
            row = graph_table.add_row().cells
            row[0].text = "No regression specimens recorded"
            for idx in range(1, 5):
                row[idx].text = "-"

        doc.add_paragraph("Note: chart image can be inserted later after external plotting / signed issue, then uploaded back as signed evidence.")
    elif safe_document_type == "PRESSURE_LOG":
        doc.add_heading("2. Pressure log", level=1)
        pressure_table = doc.add_table(rows=1, cols=8)
        pressure_table.style = "Table Grid"
        headers = ["Specimen ID", "Date", "Planned Pressure", "Actual Pressure", "Hold Pressure", "Failure Time", "Result", "Notes"]
        for idx, h in enumerate(headers):
            pressure_table.rows[0].cells[idx].text = h

        if specimens:
            for s in specimens:
                row = pressure_table.add_row().cells
                row[0].text = s.specimen_id or "-"
                row[1].text = _fmt_date(s.sample_date)
                row[2].text = _fmt_number(s.planned_pressure_mpa)
                row[3].text = _fmt_number(s.actual_pressure_at_failure_mpa)
                row[4].text = _fmt_number(s.pressure_at_hold_mpa)
                row[5].text = _fmt_number(s.failure_time_sec)
                row[6].text = s.result_status or "-"
                row[7].text = s.notes or "-"
        else:
            row = pressure_table.add_row().cells
            row[0].text = "No pressure records recorded"
            for idx in range(1, 8):
                row[idx].text = "-"
    elif safe_document_type in {"RESULT_SHEET", "DATA_SHEET"}:
        doc.add_heading("2. Result sheet", level=1)
        result_table = doc.add_table(rows=1, cols=10)
        result_table.style = "Table Grid"
        headers = [
            "Specimen ID",
            "Material Ref",
            "Sample Date",
            "Target / Planned",
            "Actual",
            "Failure Mode",
            "Failure Location",
            "Result",
            "QA Review",
            "Remarks",
        ]
        for idx, h in enumerate(headers):
            result_table.rows[0].cells[idx].text = h

        if specimens:
            for s in specimens:
                row = result_table.add_row().cells
                row[0].text = s.specimen_id or "-"
                row[1].text = s.material_ref or "-"
                row[2].text = _fmt_date(s.sample_date)
                row[3].text = _fmt_number(s.planned_pressure_mpa)
                row[4].text = _fmt_number(s.actual_pressure_at_failure_mpa)
                row[5].text = s.failure_mode or "-"
                row[6].text = s.failure_location or "-"
                row[7].text = s.result_status or "-"
                row[8].text = s.qa_review_status or "-"
                row[9].text = s.notes or "-"
        else:
            row = result_table.add_row().cells
            row[0].text = "No specimen results recorded"
            for idx in range(1, 10):
                row[idx].text = "-"
    else:
        doc.add_heading("2. Test guidance and acceptance basis", level=1)
        for label, value in [
            ("When required", guidance.get("when_required", "")),
            ("Specimen count", guidance.get("specimen_count", "")),
            ("API clause", guidance.get("api_clause", "")),
            ("External standard", guidance.get("external_standard", "")),
            ("Conditioning required", guidance.get("conditioning_required", "")),
            ("Retest logic", guidance.get("retest_logic", "")),
        ]:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(value or "-")

        for heading, items in [
            ("Conditioning steps", guidance.get("conditioning_steps", [])),
            ("Core process", guidance.get("core_process", [])),
            ("Acceptance criteria", guidance.get("acceptance", [])),
            ("Practical notes", guidance.get("practical_notes", [])),
        ]:
            doc.add_paragraph(heading, style="List Bullet")
            if items:
                for item in items:
                    doc.add_paragraph(str(item), style="List Bullet 2")
            else:
                doc.add_paragraph("-", style="List Bullet 2")

        doc.add_heading("3. Material register", level=1)
        material_table = doc.add_table(rows=1, cols=6)
        material_table.style = "Table Grid"
        headers = ["Component", "Material", "Manufacturer", "Grade", "Batch / Lot", "Status"]
        for idx, h in enumerate(headers):
            material_table.rows[0].cells[idx].text = h

        if materials:
            for m in materials:
                row = material_table.add_row().cells
                row[0].text = m.component or "-"
                row[1].text = m.material_name or "-"
                row[2].text = m.manufacturer_name or "-"
                row[3].text = m.grade_name or "-"
                row[4].text = " / ".join([x for x in [m.batch_ref, m.lot_ref] if x]) or "-"
                row[5].text = m.status or "-"
        else:
            row = material_table.add_row().cells
            row[0].text = "No materials recorded"
            for idx in range(1, 6):
                row[idx].text = "-"

        doc.add_heading("4. Specimen execution record", level=1)
        specimen_table = doc.add_table(rows=1, cols=10)
        specimen_table.style = "Table Grid"
        specimen_headers = [
            "Specimen ID",
            "Date",
            "Material ref",
            "Planned P (MPa)",
            "Actual P (MPa)",
            "Hours",
            "Cycles",
            "Result",
            "QA review",
            "Failure mode",
        ]
        for idx, h in enumerate(specimen_headers):
            specimen_table.rows[0].cells[idx].text = h

        if specimens:
            for s in specimens:
                row = specimen_table.add_row().cells
                row[0].text = s.specimen_id or "-"
                row[1].text = _fmt_date(s.sample_date)
                row[2].text = s.material_ref or "-"
                row[3].text = _fmt_number(s.planned_pressure_mpa)
                row[4].text = _fmt_number(s.actual_pressure_at_failure_mpa)
                row[5].text = _fmt_number(s.failure_hours)
                row[6].text = _fmt_number(s.failure_cycles)
                row[7].text = s.result_status or "-"
                row[8].text = s.qa_review_status or "-"
                row[9].text = s.failure_mode or "-"
        else:
            row = specimen_table.add_row().cells
            row[0].text = "No specimens recorded"
            for idx in range(1, 10):
                row[idx].text = "-"

        if specimens:
            doc.add_heading("5. Detailed specimen observations", level=1)
            for idx, s in enumerate(specimens, start=1):
                p = doc.add_paragraph()
                p.add_run(f"{idx}. {s.specimen_id or f'Specimen {idx}'}").bold = True

                details = [
                    f"Batch ref: {s.batch_ref or '-'}",
                    f"Source pipe ref: {s.source_pipe_ref or '-'}",
                    f"Preparation rule: {s.preparation_rule_basis or '-'}",
                    f"Conditioning complete: {_fmt_bool(s.conditioning_complete)}",
                    f"Pre-test visual OK: {_fmt_bool(s.pretest_visual_ok)}",
                    f"Released for test: {_fmt_bool(s.released_for_test)}",
                    f"Failure location: {s.failure_location or '-'}",
                    f"Failure description: {s.failure_description or '-'}",
                    f"Leak observation: {s.leak_observation or '-'}",
                    f"Notes: {s.notes or '-'}",
                ]
                for item in details:
                    doc.add_paragraph(item, style="List Bullet 2")

        doc.add_heading("6. Evidence package status", level=1)
        evidence_table = doc.add_table(rows=1, cols=3)
        evidence_table.style = "Table Grid"
        for idx, h in enumerate(["Document type", "Required", "Present"]):
            evidence_table.rows[0].cells[idx].text = h

        for row_data in evidence.get("rows", []):
            row = evidence_table.add_row().cells
            row[0].text = row_data.get("document_type", "-")
            row[1].text = "Yes" if row_data.get("document_type") in evidence.get("required", []) else "No"
            row[2].text = "Yes" if row_data.get("present") else "No"

        missing = evidence.get("missing", [])
        p = doc.add_paragraph()
        p.add_run("Missing evidence items: ").bold = True
        p.add_run(", ".join(missing) if missing else "None")

        doc.add_heading("7. Attachment register", level=1)
        attachment_table = doc.add_table(rows=1, cols=6)
        attachment_table.style = "Table Grid"
        for idx, h in enumerate(["Title", "Type", "Source", "Status", "Signed", "File name"]):
            attachment_table.rows[0].cells[idx].text = h

        if attachments:
            for a in attachments:
                row = attachment_table.add_row().cells
                row[0].text = a.title or "-"
                row[1].text = a.document_type or "-"
                row[2].text = a.source_mode or "-"
                row[3].text = a.approval_status or "-"
                row[4].text = _fmt_bool(a.is_signed_copy)
                row[5].text = a.original_filename or "-"
        else:
            row = attachment_table.add_row().cells
            row[0].text = "No attachments recorded"
            for idx in range(1, 6):
                row[idx].text = "-"

        doc.add_heading("8. Approval / close-out note", level=1)
        close_note = doc.add_paragraph()
        close_note.add_run("Readiness for closure: ").bold = True
        if specimens and not missing and (test.result_summary or "").strip():
            close_note.add_run("Ready for technical review and issue.")
        else:
            close_note.add_run("Not yet ready for closure. Complete missing evidence, specimen execution, or result summary first.")

    doc.save(target_path)

    return {
        "original_filename": _safe_filename(f"{(program.program_code or 'program')}_{(test.code or 'test')}_{safe_document_type.lower()}.docx"),
        "stored_filename": stored_name,
        "file_path": str(target_path),
        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "file_size_bytes": target_path.stat().st_size,
    }

def _specimen_lifecycle_summary(specimens: list[RndQualificationSpecimen]) -> dict:
    total = len(specimens)
    prepared = sum(1 for s in specimens if s.pretest_visual_ok)
    released = sum(1 for s in specimens if s.released_for_test)
    executed = sum(1 for s in specimens if s.result_status and s.result_status != 'PENDING')
    qa_ready = sum(1 for s in specimens if (s.qa_review_status or 'PENDING') in {'ACCEPTED', 'APPROVED'})
    failed_logged = sum(1 for s in specimens if s.actual_pressure_at_failure_mpa is not None or s.failure_mode or s.failure_location)
    return {
        'total': total,
        'prepared': prepared,
        'released': released,
        'executed': executed,
        'qa_ready': qa_ready,
        'failed_logged': failed_logged,
    }


def _test_progress_snapshot(test: RndQualificationTest, specimens: list[RndQualificationSpecimen], attachments: list[RndAttachmentRegister]) -> dict:
    lifecycle = _specimen_lifecycle_summary(specimens)
    evidence = _evidence_status(test.code, attachments)
    has_result = bool((test.result_summary or '').strip())
    ready_for_close = lifecycle['executed'] > 0 and evidence['complete'] and has_result
    return {
        'lifecycle': lifecycle,
        'evidence': evidence,
        'has_result': has_result,
        'ready_for_close': ready_for_close,
    }


def _phase_cards(program: RndQualificationProgram, tests: list[RndQualificationTest], materials: list[RndMaterialQualification], specimens: list[RndQualificationSpecimen], attachments: list[RndAttachmentRegister]) -> list[dict]:
    materials_ok = all((m.status or '').upper() in {'APPROVED', 'ACCEPTED', 'COMPLETE'} for m in materials) if materials else False
    tests_with_plan = sum(1 for t in tests if (t.specimen_requirement or '').strip())
    executed_tests = 0
    evidence_complete_tests = 0
    closed_tests = 0
    for t in tests:
        ts = [s for s in specimens if s.test_id == t.id]
        ta = [a for a in attachments if a.test_id == t.id]
        snap = _test_progress_snapshot(t, ts, ta)
        if snap['lifecycle']['executed']:
            executed_tests += 1
        if snap['evidence']['complete']:
            evidence_complete_tests += 1
        if (t.status or '').upper() in {'COMPLETE', 'CLOSED', 'ACCEPTED'}:
            closed_tests += 1
    total_tests = max(len(tests), 1)
    return [
        {'name': 'Definition', 'status': 'Complete' if (program.title and program.program_code) else 'Open', 'detail': f"{program.qualification_standard or 'Qualification basis'} / {program.pfr_or_pv}"},
        {'name': 'Material Control', 'status': 'Complete' if materials_ok else 'In Progress', 'detail': f"{sum(1 for m in materials if (m.status or '').upper() in {'APPROVED', 'ACCEPTED', 'COMPLETE'})}/{len(materials)} approved" if materials else 'No materials confirmed yet'},
        {'name': 'Test Planning', 'status': 'Complete' if tests_with_plan == len(tests) and tests else 'In Progress', 'detail': f"{tests_with_plan}/{len(tests)} tests with defined specimen requirement" if tests else 'No tests listed yet'},
        {'name': 'Execution', 'status': 'Complete' if executed_tests == len(tests) and tests else 'In Progress', 'detail': f"{executed_tests}/{len(tests)} tests with execution records" if tests else 'Waiting for specimens'},
        {'name': 'Evidence Review', 'status': 'Complete' if evidence_complete_tests == len(tests) and tests else 'In Progress', 'detail': f"{evidence_complete_tests}/{len(tests)} tests with complete evidence checklist" if tests else 'Waiting for evidence'},
        {'name': 'Closure / Monogram', 'status': 'Ready' if closed_tests == len(tests) and len(tests) > 0 else 'Pending', 'detail': f"{closed_tests}/{len(tests)} tests closed with result summary" if tests else 'No closeout yet'},
    ]


def _require_user(session: Session = Depends(get_session)) -> User:
    user = session.exec(select(User).order_by(User.id.asc())).first()
    if not user:
        raise HTTPException(status_code=401, detail="No users found.")
    return user


def _touch_program(program: RndQualificationProgram) -> None:
    program.updated_at = datetime.utcnow()


def _touch_row(row) -> None:
    row.updated_at = datetime.utcnow()




def _material_required_test_types(material: 'RndMaterialQualification', program: Optional[RndQualificationProgram] = None) -> list[str]:
    component = (material.component or "").strip().upper()
    reinforcement_type = (material.reinforcement_type or "").strip().upper()
    material_family = (material.material_family or "").strip().upper()
    service_text = " ".join([
        (program.intended_service if program else "") or "",
        material.service_fluid_basis or "",
        material.service_notes or "",
    ]).lower()

    required = []

    if component == "LINER":
        required.extend(["FLUID_COMPATIBILITY", "AGING"])
        if "gas" in service_text or "multiphase" in service_text:
            required.append("BLISTER_RESISTANCE")
        required.append("PERMEABILITY")

        if material_family in {"PE", "PE100", "PE-RT", "POLYETHYLENE"}:
            required.append("PE_CLASSIFICATION_BASIS")

    elif component == "REINFORCEMENT":
        required.extend(["LOAD_CAPABILITY", "FLUID_COMPATIBILITY", "AGING"])

        if reinforcement_type == "STEEL":
            required.append("CORROSION_REVIEW")
            if "cathod" in service_text or "seawater" in service_text or "subsea" in service_text:
                required.append("CATHODIC_CHARGING")
            if "sour" in service_text or "h2s" in service_text:
                required.extend(["SSC", "HIC"])

        elif reinforcement_type in {"GLASS", "ARAMID", "POLYESTER", "NONMETALLIC"}:
            required.append("FIBER_MECHANICAL_BASIS")
            required.append("HYDROLYSIS_PH_REVIEW")

    elif component == "MATRIX":
        required.extend(["FLUID_COMPATIBILITY", "AGING"])
        if (material.matrix_resin_type or "").strip().upper() in {"THERMOSET", "EPOXY"}:
            required.append("TG_OR_CURE_BASIS")

    elif component == "COVER":
        required.append("WEATHERING_UV")
        required.append("LOW_TEMP_DUCTILITY")

    elif component in {"FITTING", "COUPLING"}:
        required.append("MATERIAL_SPEC_BASIS")
        if "sour" in service_text or "h2s" in service_text:
            required.append("SOUR_SERVICE_REVIEW")

    # unique, keep order
    seen = set()
    ordered = []
    for item in required:
        if item not in seen:
            seen.add(item)
            ordered.append(item)
    return ordered


def _material_test_map(material_tests: List['RndMaterialTestRecord']) -> dict:
    out = {}
    for row in material_tests:
        key = (row.test_type or "").strip().upper()
        if not key:
            continue
        out.setdefault(key, []).append(row)
    return out


def _material_review(material: 'RndMaterialQualification', material_tests: List['RndMaterialTestRecord'], program: Optional[RndQualificationProgram] = None) -> dict:
    required = _material_required_test_types(material, program)
    test_map = _material_test_map(material_tests)

    missing = []
    failed = []
    pending = []

    for test_type in required:
        rows = test_map.get(test_type, [])
        if not rows:
            missing.append(test_type)
            continue

        decisions = {(r.decision or "PENDING").strip().upper() for r in rows}
        if "FAIL" in decisions or "REJECTED" in decisions or "NOT_ACCEPTABLE" in decisions:
            failed.append(test_type)
        elif decisions <= {"PENDING", "UNDER_REVIEW"}:
            pending.append(test_type)

    component = (material.component or "").strip().upper()
    reinforcement_type = (material.reinforcement_type or "").strip().upper()

    clarifications = []
    if component == "REINFORCEMENT":
        if not reinforcement_type:
            clarifications.append("Reinforcement type is not defined.")
        if material.reinforcement_layer_count is None:
            clarifications.append("Reinforcement layer count is not defined.")
        if not (material.reinforcement_form or "").strip():
            clarifications.append("Reinforcement form is not defined.")
    if component == "LINER":
        if not (material.standard_ref or "").strip():
            clarifications.append("Liner classification/standard basis is not defined.")
    if component == "COVER":
        if not (material.material_family or "").strip():
            clarifications.append("Cover material family is not defined.")

    if failed:
        outcome = "NOT_ACCEPTABLE"
    elif missing:
        outcome = "ADDITIONAL_TESTING_REQUIRED" if component in {"LINER", "REINFORCEMENT", "MATRIX", "COVER"} else "MORE_DATA_REQUIRED"
    elif pending or clarifications:
        outcome = "MATCH_WITH_CLARIFICATION"
    else:
        outcome = "MATCH"

    compatibility_status = "SUPPORTED"
    if failed:
        compatibility_status = "NOT_SUPPORTED"
    elif missing or pending:
        compatibility_status = "PARTIAL"

    evidence_status = "COMPLETE"
    if failed:
        evidence_status = "FAILED"
    elif missing:
        evidence_status = "MISSING"
    elif pending:
        evidence_status = "PENDING_REVIEW"

    summary_parts = []
    if outcome == "MATCH":
        summary_parts.append("Material basis and available evidence align with the current qualification input.")
    if missing:
        summary_parts.append("Missing required test/evidence: " + ", ".join(missing))
    if failed:
        summary_parts.append("Failed or unacceptable evidence in: " + ", ".join(failed))
    if pending:
        summary_parts.append("Pending review for: " + ", ".join(pending))
    if clarifications:
        summary_parts.append("Clarification needed: " + " ".join(clarifications))

    additional_tests = ", ".join(missing) if missing else ""
    clarification_text = " ".join(clarifications).strip()

    return {
        "required_tests": required,
        "missing_tests": missing,
        "failed_tests": failed,
        "pending_tests": pending,
        "compatibility_status": compatibility_status,
        "evidence_status": evidence_status,
        "review_outcome": outcome,
        "review_summary": " ".join(summary_parts).strip(),
        "clarification_required": clarification_text,
        "additional_tests_required": additional_tests,
        "change_requalification_flag": bool(
            component in {"LINER", "REINFORCEMENT", "MATRIX", "COVER"} and (
                missing or failed or clarification_text
            )
        ),
    }


def _refresh_material_review(session: Session, material: 'RndMaterialQualification', program: Optional[RndQualificationProgram] = None) -> 'RndMaterialQualification':
    tests = session.exec(
        select(RndMaterialTestRecord)
        .where(RndMaterialTestRecord.material_id == material.id)
        .order_by(RndMaterialTestRecord.test_date.desc(), RndMaterialTestRecord.id.desc())
    ).all()

    review = _material_review(material, tests, program)

    material.compatibility_status = review["compatibility_status"]
    material.evidence_status = review["evidence_status"]
    material.review_outcome = review["review_outcome"]
    material.review_summary = review["review_summary"]
    material.clarification_required = review["clarification_required"]
    material.additional_tests_required = review["additional_tests_required"]
    material.change_requalification_flag = review["change_requalification_flag"]
    _touch_row(material)
    session.add(material)
    return material

def _material_reference_options(materials: List['RndMaterialQualification']) -> list[dict]:
    options: list[dict] = []
    for row in materials:
        label_parts = []

        component = (row.component or '').strip().title() if (row.component or '').strip() else 'Material'
        label_parts.append(component)

        if (row.material_family or '').strip():
            label_parts.append((row.material_family or '').strip())

        if (row.material_name or '').strip():
            label_parts.append((row.material_name or '').strip())

        if (row.grade_name or '').strip():
            label_parts.append((row.grade_name or '').strip())

        if (row.reinforcement_type or '').strip():
            label_parts.append((row.reinforcement_type or '').strip())

        if row.reinforcement_layer_count is not None:
            label_parts.append(f"{row.reinforcement_layer_count}L")

        if (row.batch_ref or '').strip():
            label_parts.append(f"Batch {row.batch_ref.strip()}")

        value = " | ".join(part for part in label_parts if part)
        if value:
            options.append({
                'value': value,
                'label': value,
                'component': row.component or '',
                'review_outcome': row.review_outcome or 'MORE_DATA_REQUIRED',
            })
    return options

def _coalesce_material_ref(material_ref: str, batch_ref: str, source_pipe_ref: str, materials: List['RndMaterialQualification']) -> str:
    raw = (material_ref or '').strip()
    if raw:
        return raw
    options = _material_reference_options(materials)
    if len(options) == 1:
        return options[0]['value']
    batch_ref = (batch_ref or '').strip()
    source_pipe_ref = (source_pipe_ref or '').strip()
    if batch_ref and source_pipe_ref:
        return f"Batch {batch_ref} | Pipe {source_pipe_ref}"
    if batch_ref:
        return f"Batch {batch_ref}"
    if source_pipe_ref:
        return f"Pipe {source_pipe_ref}"
    return 'UNASSIGNED'




def _conditioning_required_flag(value: str | None) -> bool:
    text = (value or '').strip().lower()
    if not text:
        return False
    no_tokens = ['no ', 'not required', 'no specific', 'optional', 'conditional', 'refer to applicable procedure']
    if any(token in text for token in no_tokens):
        return False
    return 'yes' in text or 'conditioning' in text or 'temperature' in text

def _material_test_rows_by_material(material_tests: List['RndMaterialTestRecord']) -> dict[int, list['RndMaterialTestRecord']]:
    grouped: dict[int, list[RndMaterialTestRecord]] = {}
    for row in material_tests:
        grouped.setdefault(row.material_id, []).append(row)
    return grouped


def _material_dashboard_rows(
    materials: List['RndMaterialQualification'],
    material_tests: List['RndMaterialTestRecord'],
    program: Optional[RndQualificationProgram] = None,
) -> list[dict]:
    grouped = _material_test_rows_by_material(material_tests)
    rows = []

    for material in materials:
        tests = grouped.get(material.id or 0, [])
        required = _material_required_test_types(material, program)
        review = {
            "required_tests": required,
            "test_count": len(tests),
            "review_outcome": material.review_outcome or "MORE_DATA_REQUIRED",
            "review_summary": material.review_summary or "",
            "compatibility_status": material.compatibility_status or "UNKNOWN",
            "evidence_status": material.evidence_status or "MISSING",
            "clarification_required": material.clarification_required or "",
            "additional_tests_required": material.additional_tests_required or "",
            "change_requalification_flag": bool(material.change_requalification_flag),
        }
        rows.append({
            "material": material,
            "tests": tests,
            "review": review,
        })

    return rows


def _required_attachment_types(test_code: str) -> list[str]:
    common = [
        "TEST_PROCEDURE",
        "SPECIMEN_PHOTO",
        "DIMENSION_RECORD",
        "CALIBRATION_CERTIFICATE",
        "RESULT_SHEET",
    ]
    specific = {
        "MPR_REG": ["PRESSURE_LOG", "TEMPERATURE_LOG", "REGRESSION_REPORT"],
        "PV_1000H": ["PRESSURE_LOG", "TEMPERATURE_LOG"],
        "TEMP_ELEV": ["TEMPERATURE_LOG", "SETUP_PHOTO"],
        "TEMP_CYCLE": ["TEMPERATURE_LOG", "CYCLE_LOG", "SETUP_PHOTO"],
        "RAPID_DECOMP": ["GAS_CHARGE_RECORD", "TEMPERATURE_LOG", "SETUP_PHOTO"],
        "OPERATING_MBR": ["SETUP_PHOTO", "BEND_LAYOUT", "RESULT_SHEET"],
        "AXIAL_LOAD": ["SETUP_PHOTO", "LOAD_RECORD"],
        "CRUSH": ["SETUP_PHOTO", "LOAD_RECORD"],
        "LAOT": ["TEMPERATURE_LOG", "SETUP_PHOTO"],
        "IMPACT": ["SETUP_PHOTO", "IMPACT_RECORD"],
        "TEC": ["TEMPERATURE_LOG", "MEASUREMENT_RECORD"],
        "GROWTH": ["PRESSURE_LOG", "MEASUREMENT_RECORD"],
        "CYCLIC_REG": ["PRESSURE_LOG", "CYCLE_LOG", "REGRESSION_REPORT"],
    }
    return common + specific.get((test_code or "").upper(), [])


def _execution_requirements(test_code: str) -> dict:
    code = (test_code or "").upper()
    base = {
        "equipment": [
            "Calibrated test rig suitable for the test type",
            "Calibrated pressure, temperature, and dimensional measuring devices as applicable",
            "Controlled fixtures, grips, supports, or end terminations suited to the specimen",
        ],
        "records": [
            "Operator name",
            "Test date and time",
            "Specimen identification",
            "Applied test conditions",
            "Observed result",
            "Acceptance decision",
        ],
        "hold_points": [
            "Specimen preparation completed",
            "Pre-test visual verification completed",
            "Conditioning completed where required",
            "Calibration validity verified before test",
        ],
    }

    specific = {
        "MPR_REG": {
            "equipment": [
                "Long-term hydrostatic pressure system",
                "Controlled temperature environment",
                "Qualified end terminations for long-duration exposure",
            ],
            "records": [
                "Pressure level",
                "Temperature",
                "Failure time in hours",
                "Failure mode",
                "Runout status",
            ],
        },
        "PV_1000H": {
            "equipment": [
                "Constant pressure hold system",
                "Controlled temperature environment",
            ],
            "records": [
                "Pressure hold value",
                "Temperature",
                "Exposure duration",
                "Result after 1000 hours",
            ],
        },
        "TEMP_CYCLE": {
            "records": [
                "Cycle start temperature",
                "Cycle end temperature",
                "Cycle count",
                "Leakage or damage observation",
            ],
        },
        "IMPACT": {
            "records": [
                "Impact location",
                "Impact energy or setup basis",
                "Post-impact condition",
                "Follow-up acceptance result",
            ],
        },
    }

    merged = dict(base)
    extra = specific.get(code, {})
    if "equipment" in extra:
        merged["equipment"] = extra["equipment"]
    if "records" in extra:
        merged["records"] = extra["records"]
    return merged


def _acceptance_criteria(test_code: str) -> list[str]:
    code = (test_code or "").upper()
    criteria = {
        "MPR_REG": [
            "Only valid and permissible failures shall be included in the regression data set.",
            "Excluded points shall be documented with technical justification.",
            "Regression output shall satisfy the qualification basis and required confidence treatment.",
        ],
        "PV_1000H": [
            "Required specimens shall complete the 1000-hour confirmation without disqualifying failure.",
        ],
        "TEMP_ELEV": [
            "Specimen shall maintain integrity under the elevated-temperature qualification condition.",
        ],
        "TEMP_CYCLE": [
            "Specimen shall complete the required thermal cycling without disqualifying damage or leakage.",
        ],
        "RAPID_DECOMP": [
            "Specimen shall meet the acceptance basis defined for decompression resistance.",
        ],
        "OPERATING_MBR": [
            "Pipe shall demonstrate acceptable performance at the qualified operating bending radius.",
        ],
        "AXIAL_LOAD": [
            "Specimen shall satisfy the required axial load capability without disqualifying failure.",
        ],
        "CRUSH": [
            "Specimen shall satisfy the external load or crush acceptance basis.",
        ],
        "LAOT": [
            "Specimen shall satisfy qualification at the lowest allowable operating temperature.",
        ],
        "IMPACT": [
            "Specimen shall satisfy impact acceptance and any required follow-up confirmation.",
        ],
        "TEC": [
            "Measured thermal expansion values shall be recorded and accepted against the qualification basis.",
        ],
        "GROWTH": [
            "Measured dimensional growth or shrinkage shall remain within the acceptance basis.",
        ],
        "CYCLIC_REG": [
            "Cyclic regression data shall satisfy the qualification basis and required confidence treatment.",
        ],
    }
    return criteria.get(code, ["Test shall comply with the approved qualification basis and internal acceptance procedure."])


def _evidence_status(test_code: str, attachments: list) -> dict:
    required = _required_attachment_types(test_code)
    uploaded = {(a.document_type or "").upper() for a in attachments}
    rows = []
    missing = []
    for item in required:
        ok = item in uploaded
        rows.append({"document_type": item, "present": ok})
        if not ok:
            missing.append(item)
    return {
        "required": required,
        "rows": rows,
        "missing": missing,
        "complete": len(missing) == 0,
    }


def _specimen_readiness(specimens: list, prep: dict | None = None) -> dict:
    total = len(specimens)
    released = 0
    conditioning_pending = 0
    visual_pending = 0

    preconditioning_required = None
    if prep:
        preconditioning_required = prep.get("preconditioning", {}).get("required")

    for s in specimens:
        if s.released_for_test:
            released += 1
        if preconditioning_required is True and not s.conditioning_complete:
            conditioning_pending += 1
        if not s.pretest_visual_ok:
            visual_pending += 1

    return {
        "total": total,
        "released": released,
        "conditioning_pending": conditioning_pending,
        "visual_pending": visual_pending,
        "complete": total > 0 and conditioning_pending == 0 and visual_pending == 0,
    }


def _default_test_matrix(pfr_or_pv: str) -> list[dict]:
    route = (pfr_or_pv or '').strip().upper()

    base = [
        {"code": "MPR_REG", "title": "Long-term hydrostatic regression", "description": "Primary regression basis for qualification. Use ASTM D2992 Procedure B logic, exclude points below 10 h, calculate mean line, LCL, LPL, and LCL at RCRT.", "specimen_requirement": "18+ target", "clause_ref": "API 15S 5.3.2.3 / Annex E / Annex G", "applicability": "CORE", "scope_tag": "PFR", "source_standard": "API_15S"},
        {"code": "PV_1000H", "title": "1000-hour constant pressure confirmation", "description": "1000-hour proof / confirmation exposure used as a core verification step and for PV relationship confirmation.", "specimen_requirement": "2", "clause_ref": "API 15S 5.3.4.2", "applicability": "CORE", "scope_tag": "PV", "source_standard": "API_15S"},
        {"code": "TEMP_ELEV", "title": "Elevated temperature test", "description": "Seal and polymer creep or relaxation confirmation above MAOT.", "specimen_requirement": "1", "clause_ref": "API 15S 5.3.5", "applicability": "CORE", "scope_tag": "BOTH", "source_standard": "API_15S"},
        {"code": "TEMP_CYCLE", "title": "Temperature cycling", "description": "Thermal cycling confirmation for qualified size and rating combinations.", "specimen_requirement": "2", "clause_ref": "API 15S 5.3.6", "applicability": "CORE", "scope_tag": "BOTH", "source_standard": "API_15S"},
        {"code": "RAPID_DECOMP", "title": "Rapid decompression", "description": "Required for gas or multiphase service.", "specimen_requirement": "1", "clause_ref": "API 15S 5.3.7 / Annex B", "applicability": "SERVICE_DEP", "scope_tag": "BOTH", "source_standard": "API_15S"},
        {"code": "OPERATING_MBR", "title": "Operating MBR", "description": "Confirm operating bending performance.", "specimen_requirement": "2", "clause_ref": "API 15S 5.3.8.1", "applicability": "CORE", "scope_tag": "BOTH", "source_standard": "API_15S"},
        {"code": "HANDLING_MBR", "title": "Handling MBR preconditioning", "description": "Applies when handling MBR is smaller than operating MBR.", "specimen_requirement": "2", "clause_ref": "API 15S 5.3.8.2", "applicability": "RANGE_DEP", "scope_tag": "BOTH", "source_standard": "API_15S"},
        {"code": "HANDLING_AND_SPOOLING", "title": "Handling and spooling durability", "description": "Handling / spooling preconditioning followed by proof confirmation.", "specimen_requirement": "2", "clause_ref": "API 15S 5.3.8.3", "applicability": "RANGE_DEP", "scope_tag": "PFR", "source_standard": "API_15S"},
        {"code": "RESPOOLING", "title": "Respooling qualification", "description": "Used where respooling is claimed or allowed.", "specimen_requirement": "2", "clause_ref": "API 15S 5.3.8.4", "applicability": "RANGE_DEP", "scope_tag": "PFR", "source_standard": "API_15S"},
        {"code": "AXIAL_LOAD", "title": "Axial load capability", "description": "Maximum allowable axial load followed by proof confirmation.", "specimen_requirement": "2", "clause_ref": "API 15S 5.3.9", "applicability": "CORE", "scope_tag": "BOTH", "source_standard": "API_15S"},
        {"code": "CRUSH", "title": "External load / crush", "description": "Radial crush / external load characterization.", "specimen_requirement": "3", "clause_ref": "API 15S 5.3.10", "applicability": "RANGE_DEP", "scope_tag": "BOTH", "source_standard": "API_15S"},
        {"code": "LAOT", "title": "Lowest allowable operating temperature", "description": "Minimum operating temperature qualification.", "specimen_requirement": "2", "clause_ref": "API 15S 5.3.11", "applicability": "CORE", "scope_tag": "BOTH", "source_standard": "API_15S"},
        {"code": "IMPACT", "title": "Impact resistance", "description": "Impact followed by proof confirmation.", "specimen_requirement": "2", "clause_ref": "API 15S 5.3.12", "applicability": "CORE", "scope_tag": "BOTH", "source_standard": "API_15S"},
        {"code": "TEC", "title": "Thermal expansion coefficient", "description": "Axial TEC measurement and hoop TEC where clearance is critical.", "specimen_requirement": "2", "clause_ref": "API 15S 5.3.13", "applicability": "CORE", "scope_tag": "BOTH", "source_standard": "API_15S"},
        {"code": "GROWTH", "title": "Growth / shrinkage under pressure", "description": "Pressure elongation and dimensional response.", "specimen_requirement": "2", "clause_ref": "API 15S 5.3.14", "applicability": "CORE", "scope_tag": "BOTH", "source_standard": "API_15S"},
        {"code": "CYCLIC_REG", "title": "Cyclic pressure regression", "description": "For cyclic service. Use cyclic regression and lower confidence basis.", "specimen_requirement": "18+ target", "clause_ref": "API 15S 5.3.16 / Annex D", "applicability": "SERVICE_DEP", "scope_tag": "BOTH", "source_standard": "API_15S"},
    ]

    items = []
    for row in base:
        entry = dict(row)

        if route == 'PFR':
            if row["code"] == "MPR_REG":
                entry["route_note"] = "Primary PFR qualification regression"
            elif row["code"] == "PV_1000H":
                entry["route_note"] = "PV-only requirement kept visible for family planning"
            else:
                entry["route_note"] = "PFR qualification matrix item"
        elif route == 'PV':
            if row["code"] == "PV_1000H":
                entry["route_note"] = "Primary PV confirmation test"
            elif row["code"] == "MPR_REG":
                entry["route_note"] = "Reference PFR regression basis from the same qualification family"
            else:
                entry["route_note"] = "PV verification / inherited qualification matrix item"
        else:
            entry["route_note"] = "Qualification matrix item"

        items.append(entry)

    return items

def _seed_test_matrix(session: Session, program: RndQualificationProgram) -> None:
    existing = session.exec(
        select(RndQualificationTest).where(RndQualificationTest.program_id == program.id)
    ).all()
    if existing:
        return

    if not session.exec(
        select(RndMaterialQualification).where(RndMaterialQualification.program_id == program.id)
    ).first():
        for component, material in [
            ('LINER', program.liner_material),
            ('REINFORCEMENT', program.reinforcement_material),
            ('COVER', program.cover_material),
        ]:
            row = RndMaterialQualification(
                program_id=program.id,
                component=component,
                material_name=material,
                material_family=(
                    'POLYMER' if component in {'LINER', 'COVER'}
                    else 'REINFORCEMENT'
                ),
                reinforcement_type=(
                    'NONMETALLIC'
                    if component == 'REINFORCEMENT' and 'steel' not in (material or '').lower()
                    else ('STEEL' if component == 'REINFORCEMENT' else '')
                ),
                reinforcement_layer_count=(2 if component == 'REINFORCEMENT' else None),
                status='PLANNED',
                review_outcome='MORE_DATA_REQUIRED',
                evidence_status='MISSING',
                compatibility_status='UNKNOWN',
            )
            session.add(row)

    for idx, item in enumerate(_default_test_matrix(program.pfr_or_pv), start=1):
        session.add(
            RndQualificationTest(
                program_id=program.id,
                sort_order=idx,
                clause_ref=item['clause_ref'],
                code=item['code'],
                title=item['title'],
                description=item['description'],
                specimen_requirement=item['specimen_requirement'],
                applicability=item['applicability'],
                scope_tag=item.get('scope_tag', 'BOTH'),
                source_standard=item.get('source_standard', 'API_15S'),
            )
        )

    session.commit()

def _ensure_complete_test_matrix(session: Session, program: RndQualificationProgram) -> None:
    existing = session.exec(
        select(RndQualificationTest)
        .where(RndQualificationTest.program_id == program.id)
        .order_by(RndQualificationTest.sort_order.asc(), RndQualificationTest.id.asc())
    ).all()

    existing_by_code = {(row.code or '').strip().upper(): row for row in existing}
    matrix = _default_test_matrix(program.pfr_or_pv)

    next_order = max([row.sort_order for row in existing], default=0)

    changed = False
    for item in matrix:
        code = (item['code'] or '').strip().upper()
        row = existing_by_code.get(code)

        if row is None:
            next_order += 1
            session.add(
                RndQualificationTest(
                    program_id=program.id,
                    sort_order=next_order,
                    clause_ref=item['clause_ref'],
                    code=code,
                    title=item['title'],
                    description=item['description'],
                    specimen_requirement=item['specimen_requirement'],
                    applicability=item['applicability'],
                    scope_tag=item.get('scope_tag', 'BOTH'),
                    source_standard=item.get('source_standard', 'API_15S'),
                    status='PLANNED',
                )
            )
            changed = True
        else:
            row.clause_ref = item['clause_ref']
            row.title = item['title']
            row.description = item['description']
            row.specimen_requirement = item['specimen_requirement']
            row.applicability = item['applicability']
            row.scope_tag = item.get('scope_tag', row.scope_tag or 'BOTH')
            row.source_standard = item.get('source_standard', row.source_standard or 'API_15S')
            _touch_row(row)
            session.add(row)
            changed = True

    if changed:
        _touch_program(program)
        session.add(program)
        session.commit()


def _pv_extension_matrix() -> list[dict]:
    return [
        {
            "code": "PV_1000H",
            "title": "1000-hour constant pressure confirmation",
            "description": "Primary PV confirmation test for the added PV scope.",
            "specimen_requirement": "2",
            "clause_ref": "API 15S 5.3.4.2",
            "applicability": "CORE",
            "scope_tag": "PV",
            "source_standard": "API_15S",
        },
        {
            "code": "TEMP_ELEV",
            "title": "Elevated temperature test",
            "description": "Seal and polymer creep or relaxation confirmation above MAOT for the PV scope.",
            "specimen_requirement": "1",
            "clause_ref": "API 15S 5.3.5",
            "applicability": "CORE",
            "scope_tag": "BOTH",
            "source_standard": "API_15S",
        },
        {
            "code": "TEMP_CYCLE",
            "title": "Temperature cycling",
            "description": "Thermal cycling confirmation for the PV scope.",
            "specimen_requirement": "2",
            "clause_ref": "API 15S 5.3.6",
            "applicability": "CORE",
            "scope_tag": "BOTH",
            "source_standard": "API_15S",
        },
        {
            "code": "RAPID_DECOMP",
            "title": "Rapid decompression",
            "description": "Required for gas or multiphase service where the PV scope is applicable.",
            "specimen_requirement": "1",
            "clause_ref": "API 15S 5.3.7 / Annex B",
            "applicability": "SERVICE_DEP",
            "scope_tag": "BOTH",
            "source_standard": "API_15S",
        },
        {
            "code": "OPERATING_MBR",
            "title": "Operating MBR",
            "description": "Confirm operating bending performance for the PV scope.",
            "specimen_requirement": "2",
            "clause_ref": "API 15S 5.3.8.1",
            "applicability": "CORE",
            "scope_tag": "BOTH",
            "source_standard": "API_15S",
        },
        {
            "code": "AXIAL_LOAD",
            "title": "Axial load capability",
            "description": "Maximum allowable axial load followed by proof confirmation for the PV scope.",
            "specimen_requirement": "2",
            "clause_ref": "API 15S 5.3.9",
            "applicability": "CORE",
            "scope_tag": "BOTH",
            "source_standard": "API_15S",
        },
        {
            "code": "LAOT",
            "title": "Lowest allowable operating temperature",
            "description": "Minimum operating temperature qualification for the PV scope.",
            "specimen_requirement": "2",
            "clause_ref": "API 15S 5.3.11",
            "applicability": "CORE",
            "scope_tag": "BOTH",
            "source_standard": "API_15S",
        },
        {
            "code": "IMPACT",
            "title": "Impact resistance",
            "description": "Impact followed by proof confirmation for the PV scope.",
            "specimen_requirement": "2",
            "clause_ref": "API 15S 5.3.12",
            "applicability": "CORE",
            "scope_tag": "BOTH",
            "source_standard": "API_15S",
        },
        {
            "code": "TEC",
            "title": "Thermal expansion coefficient",
            "description": "Axial TEC measurement and hoop TEC where clearance is critical for the PV scope.",
            "specimen_requirement": "2",
            "clause_ref": "API 15S 5.3.13",
            "applicability": "CORE",
            "scope_tag": "BOTH",
            "source_standard": "API_15S",
        },
        {
            "code": "GROWTH",
            "title": "Growth / shrinkage under pressure",
            "description": "Pressure elongation and dimensional response for the PV scope.",
            "specimen_requirement": "2",
            "clause_ref": "API 15S 5.3.14",
            "applicability": "CORE",
            "scope_tag": "BOTH",
            "source_standard": "API_15S",
        },
        {
            "code": "CYCLIC_REG",
            "title": "Cyclic pressure regression",
            "description": "For cyclic service when the PV scope is intended for cyclic duty.",
            "specimen_requirement": "18+ target",
            "clause_ref": "API 15S 5.3.16 / Annex D",
            "applicability": "SERVICE_DEP",
            "scope_tag": "BOTH",
            "source_standard": "API_15S",
        },
    ]


def _add_pv_extension_to_program(session: Session, program: RndQualificationProgram) -> dict:
    existing_tests = session.exec(
        select(RndQualificationTest)
        .where(RndQualificationTest.program_id == program.id)
        .order_by(RndQualificationTest.sort_order.asc(), RndQualificationTest.id.asc())
    ).all()

    existing_by_code = {(row.code or '').strip().upper(): row for row in existing_tests}
    next_order = max([row.sort_order for row in existing_tests], default=0)

    created = []
    updated = []

    for item in _pv_extension_matrix():
        code = (item["code"] or "").strip().upper()
        row = existing_by_code.get(code)

        if row is None:
            next_order += 1
            row = RndQualificationTest(
                program_id=program.id,
                sort_order=next_order,
                clause_ref=item["clause_ref"],
                code=code,
                title=item["title"],
                description=item["description"],
                specimen_requirement=item["specimen_requirement"],
                applicability=item["applicability"],
                scope_tag=item.get("scope_tag", "PV"),
                source_standard=item.get("source_standard", "API_15S"),
                status="PLANNED",
            )
            session.add(row)
            created.append(code)
        else:
            # Do not destroy existing row content; only widen scope where needed
            current_scope = (row.scope_tag or "BOTH").strip().upper()
            wanted_scope = (item.get("scope_tag") or "PV").strip().upper()

            if current_scope != wanted_scope:
                if {current_scope, wanted_scope} & {"PFR", "PV"}:
                    row.scope_tag = "BOTH"
                elif current_scope == "CUSTOM":
                    row.scope_tag = "BOTH"
                else:
                    row.scope_tag = wanted_scope

            if not (row.clause_ref or "").strip():
                row.clause_ref = item["clause_ref"]
            if not (row.description or "").strip():
                row.description = item["description"]
            if not (row.specimen_requirement or "").strip():
                row.specimen_requirement = item["specimen_requirement"]
            if not (row.applicability or "").strip():
                row.applicability = item["applicability"]
            if not (row.source_standard or "").strip():
                row.source_standard = item.get("source_standard", "API_15S")

            _touch_row(row)
            session.add(row)
            updated.append(code)

    # mark program as now supporting PV in the same qualification family
    if (program.pfr_or_pv or "").strip().upper() == "PFR":
        program.pfr_or_pv = "PFR"
    _touch_program(program)
    session.add(program)

    session.commit()

    return {
        "created_codes": created,
        "updated_codes": updated,
    }
    

def _t_critical_975(df: int) -> float:
    table = {1: 12.706, 2: 4.303, 3: 3.182, 4: 2.776, 5: 2.571, 6: 2.447, 7: 2.365, 8: 2.306, 9: 2.262, 10: 2.228, 11: 2.201, 12: 2.179, 13: 2.160, 14: 2.145, 15: 2.131, 16: 2.120, 17: 2.110, 18: 2.101, 19: 2.093, 20: 2.086, 21: 2.080, 22: 2.074, 23: 2.069, 24: 2.064, 25: 2.060, 26: 2.056, 27: 2.052, 28: 2.048, 29: 2.045, 30: 2.042, 40: 2.021, 60: 2.000, 120: 1.980}
    if df <= 1:
        return table[1]
    keys = sorted(table.keys())
    if df in table:
        return table[df]
    if df > keys[-1]:
        return 1.960
    lower = max(k for k in keys if k < df)
    upper = min(k for k in keys if k > df)
    low_v = table[lower]
    up_v = table[upper]
    ratio = (df - lower) / (upper - lower)
    return low_v + (up_v - low_v) * ratio


def _regression_from_specimens(specimens: List[RndQualificationSpecimen], mode: str = 'STATIC_REGRESSION', target_npr_mpa: float = 0.0, design_factor: float | None = None) -> dict:
    filtered = []
    excluded = []

    for s in specimens:
        specimen_type = (s.test_type or '').strip().upper()
        mode_key = (mode or '').strip().upper()

        static_aliases = {'STATIC_REGRESSION', 'MPR_REG'}
        cyclic_aliases = {'CYCLIC_REGRESSION', 'CYCLIC_REG'}

        if mode_key in static_aliases:
            if specimen_type not in static_aliases:
                continue
        elif mode_key in cyclic_aliases:
            if specimen_type not in cyclic_aliases:
                continue
        else:
            if specimen_type != mode_key:
                continue

        if not s.include_in_regression or not s.permissible_failure:
            excluded.append(s)
            continue

        x_raw = s.failure_hours if mode_key in static_aliases else s.failure_cycles
        y_raw = s.pressure_mpa

        if x_raw is None or y_raw is None or x_raw <= 0 or y_raw <= 0:
            continue

        if mode_key in static_aliases and x_raw < 10:
            excluded.append(s)
            continue

        filtered.append(s)

    n = len(filtered)
    required_minimum = 18 if mode.upper() in {'STATIC_REGRESSION', 'CYCLIC_REGRESSION', 'MPR_REG', 'CYCLIC_REG'} else 2

    result = {
        'count': n,
        'required_minimum': required_minimum,
        'points': [],
        'excluded_count': len(excluded),
        'excluded_ids': [s.specimen_id for s in excluded],
        'warning': '',
    }

    if n < 2:
        result['warning'] = 'Need at least 2 valid points to calculate a regression line.'
        return result

    pts, xs, ys = [], [], []
    for s in filtered:
        x_raw = s.failure_hours if mode_key in static_aliases else s.failure_cycles
        y_raw = s.pressure_mpa
        x = math.log10(float(x_raw))
        y = math.log10(float(y_raw))
        xs.append(x)
        ys.append(y)
        pts.append({
            'specimen_id': s.specimen_id,
            'x_raw': x_raw,
            'y_raw': y_raw,
            'x': x,
            'y': y,
            'temperature_c': s.temperature_c,
            'failure_mode': s.failure_mode
        })

    x_bar = sum(xs) / n
    y_bar = sum(ys) / n
    sxx = sum((x - x_bar) ** 2 for x in xs)

    if sxx == 0:
        result.update({'points': pts, 'warning': 'All time values are identical; regression cannot be calculated.'})
        return result

    sxy = sum((xs[i] - x_bar) * (ys[i] - y_bar) for i in range(n))
    slope = sxy / sxx
    intercept = y_bar - slope * x_bar

    residuals = [ys[i] - (intercept + slope * xs[i]) for i in range(n)]
    df = max(1, n - 2)
    syx = math.sqrt(sum(r * r for r in residuals) / df)
    tcrit = _t_critical_975(df)

    def _predict_components(x_val: float) -> dict:
        mean_y = intercept + slope * x_val
        mean_se = syx * math.sqrt((1 / n) + ((x_val - x_bar) ** 2 / sxx))
        pred_se = syx * math.sqrt(1 + (1 / n) + ((x_val - x_bar) ** 2 / sxx))
        lcl_y = mean_y - tcrit * mean_se
        lpl_y = mean_y - tcrit * pred_se
        return {
            'x_val': x_val,
            'mean_y': mean_y,
            'mean_se': mean_se,
            'pred_se': pred_se,
            'lcl_y': lcl_y,
            'lpl_y': lpl_y,
            'mean_pressure': 10 ** mean_y,
            'lcl_pressure': 10 ** lcl_y,
            'lpl_pressure': 10 ** lpl_y,
        }

    basis_x = math.log10(RCRT_HOURS if mode_key in static_aliases else CYCLIC_BASIS_CYCLES)
    basis_calc = _predict_components(basis_x)

    mean_basis_mpa = basis_calc['mean_pressure']
    lcl_basis_mpa = basis_calc['lcl_pressure']
    lpl_basis_mpa = basis_calc['lpl_pressure']

    effective_design_factor = (
        float(design_factor) if (design_factor is not None and mode_key in static_aliases)
        else (DESIGN_FACTOR_NONMETALLIC if mode_key in static_aliases else 1.0)
    )
    mpr_mpa = lcl_basis_mpa * effective_design_factor if mode_key in static_aliases else lcl_basis_mpa
    margin_mpa = mpr_mpa - target_npr_mpa if target_npr_mpa else None
    pass_status = None if not target_npr_mpa else ('PASS' if mpr_mpa >= target_npr_mpa else 'FAIL')

    chart_points = []
    x_min = min(xs)
    x_max = max(max(xs), basis_x)
    steps = 24
    for i in range(steps + 1):
        x_val = x_min + (x_max - x_min) * i / steps
        comp = _predict_components(x_val)
        chart_points.append({
            'x': x_val,
            'time_or_cycles': round(10 ** x_val, 3),
            'mean_pressure': round(comp['mean_pressure'], 4),
            'lcl_pressure': round(comp['lcl_pressure'], 4),
            'lpl_pressure': round(comp['lpl_pressure'], 4)
        })

    result.update({
        'points': pts,
        'slope': slope,
        'intercept': intercept,
        'syx': syx,
        'tcrit': tcrit,
        'x_bar': x_bar,
        'y_bar': y_bar,
        'sxx': sxx,
        'n': n,
        'df': df,
        'x_basis': basis_x,
        'rcrt_hours': RCRT_HOURS,
        'cyclic_basis_cycles': CYCLIC_BASIS_CYCLES,
        'mean_rcrt_mpa': mean_basis_mpa,
        'lcl_rcrt_mpa': lcl_basis_mpa,
        'lpl_rcrt_mpa': lpl_basis_mpa,
        'chart_points': chart_points,
        'design_factor': effective_design_factor,
        'mpr_mpa': mpr_mpa,
        'target_npr_mpa': target_npr_mpa,
        'margin_mpa': margin_mpa,
        'pass_status': pass_status,
        'formula_text': 'log10(P) = intercept + slope * log10(time)',
        'basis_calc': basis_calc,
    })

    if n < required_minimum:
        result['warning'] = 'Regression is calculated, but you are below the readiness target for a full qualification set.'

    return result


def _matrix_counts(tests: List[RndQualificationTest]) -> dict:
    counts = {'PLANNED': 0, 'IN_PROGRESS': 0, 'PASSED': 0, 'FAILED': 0, 'WAIVED': 0}
    for t in tests:
        key = (t.status or 'PLANNED').upper()
        counts[key] = counts.get(key, 0) + 1
    return counts


def _status_pct(counts: dict, total: int) -> int:
    if total <= 0:
        return 0
    done = counts.get('PASSED', 0) + counts.get('WAIVED', 0)
    return int(round((done / total) * 100))


def _qualification_guide(program: Optional[RndQualificationProgram] = None) -> dict:
    size = f"{program.nominal_size_in:g} in" if program else '4 in'
    npr = f"{program.npr_mpa:g} MPa" if program else '10 MPa'
    maot = f"{program.maot_c:g} °C" if program else '65 °C'
    return {
        'summary': (f'This workspace organizes API 15S qualification for LLRTP with PE-RT liner, polyester fiber reinforcement, and PE100 cover. It guides the user through product definition, test matrix, specimen tracking, and regression review for {size} / {npr} / {maot}.' if (program and not (program.qualification_standard or '').strip().upper().startswith('OTHER')) else f'This workspace organizes a custom qualification program with checklist tracking, evidence collection, test guidance, and result review for {size} / {npr} / {maot}.'),
        'steps': [
            {'title': '1. Define the qualification basis', 'text': 'Create the program as PFR or PV, set size, NPR, MAOT, service, and material stack. Use the most demanding representative as the PFR when possible.'},
            {'title': '2. Lock the material system', 'text': 'Record liner, reinforcement, and cover grade, supplier, batch, and certificate references before test execution.'},
            {'title': '3. Build the specimen plan', 'text': 'Prepare static regression specimens, cyclic specimens if cyclic service applies, and the rest of the API 15S matrix such as temperature cycling, MBR, impact, axial load, and decompression when applicable.'},
            {'title': '4. Run regression correctly', 'text': 'For static regression, record pressure, temperature, time to failure, and failure mode. Exclude invalid failures and any point below 10 h from the regression dataset.'},
            {'title': '5. Review the lower confidence basis', 'text': 'Use the lower confidence result at 175,000 h for nonmetallic reinforcement, then apply the design factor to compare against the target NPR.'},
            {'title': '6. Close the program only with full evidence', 'text': 'A program is ready to close only when the matrix is complete, materials are traceable, exclusions are justified, and the final qualification package is signed off.'},
        ],
        'observe': [
            'Use the correct end fittings so fitting failures do not corrupt the pipe qualification dataset.',
            'Keep test temperature stable and recorded for every specimen.',
            'Keep a clear reason whenever a point is excluded from regression.',
            'Do not use average pressure alone for acceptance; review LCL and MPR basis.',
        ],
        'avoid': [
            'Do not mix different designs or reinforcement constructions in one regression set.',
            'Do not include points below 10 h in static regression.',
            'Do not treat the software as a substitute for engineering review or third-party witness requirements.',
            'Do not close a qualification with missing raw records, certificates, or failure descriptions.',
        ],
        'formula_examples': [
            {'label': 'Regression line', 'expr': 'log10(P) = intercept + slope * log10(time)'},
            {'label': 'Lower confidence at design life', 'expr': 'LCL_175000h = lower confidence pressure at 175,000 h'},
            {'label': 'Nonmetallic MPR', 'expr': 'MPR = LCL_175000h x 0.67'},
            {'label': 'PV helper', 'expr': 'PPV1000 = PPFR1000 x (NPR_PV / NPR_PFR)'},
        ],
    }

def _program_answers(program: RndQualificationProgram) -> dict:
    raw = (program.notes or '').strip()
    if raw.startswith('__RNDJSON__'):
        try:
            return json.loads(raw[len('__RNDJSON__'):])
        except Exception:
            return {}
    return {}


def _save_program_answers(program: RndQualificationProgram, answers: dict) -> None:
    program.notes = '__RNDJSON__' + json.dumps(answers, ensure_ascii=False)
    _touch_program(program)


def _wizard_state(program: RndQualificationProgram) -> dict:
    answers = _program_answers(program)
    launch_size = answers.get('launch_size_in') or f"{program.nominal_size_in:g}"
    sister_size = answers.get('sister_size_in') or '6'
    service_route = answers.get('service_route') or (
        'gas_multiphase' if 'gas' in (program.intended_service or '').lower() else 'static_liquid'
    )
    cyclic_service = answers.get('cyclic_service', 'no')

    decision = {
        'launch_size': launch_size,
        'sister_size': sister_size,
        'service_route': service_route,
        'cyclic_service': cyclic_service,
        'family_decision': f'Use {launch_size} in as the main qualified size and handle {sister_size} in as a PV only if materials, pressure class, and construction remain matched.',
        'service_decision': 'Rapid decompression is required.' if service_route == 'gas_multiphase' else 'Rapid decompression is not required for static liquid service.',
        'cyclic_decision': 'Cyclic regression route is required.' if cyclic_service == 'yes' else 'Cyclic route is not required unless the field duty exceeds the API cyclic trigger.',
        'temperature_decision': f'Qualification temperature should be at least the claimed MAOT of {program.maot_c:g} °C. Higher claims need their own basis.',
        'wizard_complete': True,
    }
    return {'answers': answers, 'decision': decision}


def _material_screening_state(materials: List[RndMaterialQualification]) -> dict:
    rows = []
    all_ready = True

    for m in materials:
        missing = []

        if not (m.component or '').strip():
            missing.append('component')
        if not (m.material_name or '').strip():
            missing.append('material name')
        if not (m.supplier_name or '').strip():
            missing.append('supplier')
        if not (m.grade_name or '').strip():
            missing.append('grade')
        if not (m.certificate_ref or '').strip():
            missing.append('certificate ref')
        if not (m.batch_ref or '').strip():
            missing.append('batch ref')
        if not (m.standard_ref or '').strip():
            missing.append('standard basis')

        if (m.component or '').strip().upper() == 'REINFORCEMENT':
            if not (m.reinforcement_type or '').strip():
                missing.append('reinforcement type')
            if m.reinforcement_layer_count is None:
                missing.append('reinforcement layer count')
            if not (m.reinforcement_form or '').strip():
                missing.append('reinforcement form')

        ready = len(missing) == 0 and (m.review_outcome or '') in {'MATCH', 'MATCH_WITH_CLARIFICATION'}
        all_ready = all_ready and ready

        rows.append({
            'row': m,
            'missing': missing,
            'ready': ready,
            'review_outcome': m.review_outcome or 'MORE_DATA_REQUIRED',
            'review_summary': m.review_summary or '',
        })

    return {
        'rows': rows,
        'complete': all_ready and len(rows) >= 3,
        'status_label': 'Accepted' if all_ready and len(rows) >= 3 else 'More data required',
        'headline': 'Record traceable material identity, standards basis, qualification evidence, and reinforcement construction before structural testing starts.',
    }


def _burst_threshold(program: RndQualificationProgram) -> float:
    if program.npr_mpa <= 0:
        return 0.0
    return round(program.npr_mpa / DESIGN_FACTOR_NONMETALLIC, 3)


def _burst_state(program: RndQualificationProgram, specimens: List[RndQualificationSpecimen]) -> dict:
    burst_rows = [s for s in specimens if (s.test_type or '').upper() == 'BURST_QUALIFICATION']
    threshold = _burst_threshold(program)
    evaluated = []
    accepted = 0
    review_needed = 0

    for s in burst_rows:
        flags = []
        if s.pressure_mpa <= 0:
            flags.append('Burst pressure missing.')
        if s.temperature_c and abs(float(s.temperature_c) - float(program.maot_c)) > 5.0:
            flags.append(f'Test temperature {s.temperature_c:g} °C is outside the ±5 °C window around the qualification basis.')
        mode = (s.failure_mode or '').strip().lower()
        if mode and mode not in {'burst', 'rupture'}:
            flags.append('Failure mode is not a clear burst/rupture and needs engineering review.')
        if s.pressure_mpa and s.pressure_mpa < threshold:
            flags.append(f'Burst pressure is below the minimum screen threshold of {threshold:.3f} MPa.')

        status = 'ACCEPTED' if not flags else 'REVIEW'
        if status == 'ACCEPTED':
            accepted += 1
        else:
            review_needed += 1

        evaluated.append({'specimen': s, 'flags': flags, 'status': status})

    require_count = 5
    complete = accepted >= require_count and review_needed == 0

    return {
        'threshold_mpa': threshold,
        'required_count': require_count,
        'accepted_count': accepted,
        'review_count': review_needed,
        'rows': evaluated,
        'complete': complete,
        'headline': 'Run burst testing first as a design screen. The system will only unlock the next step once five acceptable specimens are recorded against the minimum burst threshold.',
    }


def _active_stage(program: RndQualificationProgram, materials: List[RndMaterialQualification], specimens: List[RndQualificationSpecimen]) -> dict:
    wizard = _wizard_state(program)
    material_state = _material_screening_state(materials)
    burst_state = _burst_state(program, specimens)
    static_reg = _regression_from_specimens(specimens, 'STATIC_REGRESSION', program.npr_mpa, program.service_factor)

    if not material_state['complete']:
        current = 'materials'
        percent = 33
    elif not burst_state['complete']:
        current = 'burst'
        percent = 58
    elif static_reg['count'] < static_reg['required_minimum']:
        current = 'regression'
        percent = 78
    else:
        current = 'review'
        percent = 100

    return {
        'wizard': wizard,
        'materials': material_state,
        'burst': burst_state,
        'static_reg': static_reg,
        'current': current,
        'progress_pct': percent,
    }


@router.get('')
def rnd_home() -> RedirectResponse:
    return RedirectResponse(url='/rnd/qualifications', status_code=303)


@router.get('/qualifications')
def rnd_dashboard(request: Request, session: Session = Depends(get_session), user: User = Depends(_require_user)):
    from sqlmodel import text

    rows = session.exec(
        text("""
            SELECT
                id,
                program_code,
                title,
                nominal_size_in,
                npr_mpa,
                maot_c,
                pfr_or_pv,
                program_type,
                qualification_standard,
                service_medium,
                intended_service,
                status,
                COALESCE(is_archived, FALSE) AS is_archived,
                updated_at
            FROM rndqualificationprogram
            ORDER BY updated_at DESC, id DESC
        """)
    ).all()

    class DashboardProgram:
        def __init__(self, row):
            self.id = row.id
            self.program_code = row.program_code
            self.title = row.title
            self.nominal_size_in = row.nominal_size_in
            self.npr_mpa = row.npr_mpa
            self.maot_c = row.maot_c
            self.pfr_or_pv = row.pfr_or_pv
            self.program_type = row.program_type
            self.qualification_standard = row.qualification_standard
            self.service_medium = row.service_medium
            self.intended_service = row.intended_service
            self.status = row.status
            self.is_archived = bool(row.is_archived)

    dashboard = []
    active_count = 0
    archived_count = 0

    for row in rows:
        program = DashboardProgram(row)

        if program.is_archived:
            archived_count += 1
        else:
            active_count += 1

        flow_name = 'review'
        status_key = (program.status or 'DRAFT').strip().upper()
        if status_key in {'DRAFT', 'PLANNED'}:
            flow_name = 'materials'
        elif status_key in {'IN_PROGRESS', 'ACTIVE'}:
            flow_name = 'regression'
        elif status_key in {'COMPLETE', 'CLOSED', 'APPROVED'}:
            flow_name = 'review'

        dashboard.append({
            "program": program,
            "flow": {
                "current": flow_name,
                "progress_pct": 35 if flow_name == 'materials' else (75 if flow_name == 'regression' else 100),
            },
        })

    return TEMPLATES.TemplateResponse(
        request=request,
        name='rnd_dashboard.html',
        context={
            'request': request,
            'user': user,
            'dashboard': dashboard,
            'guide': _qualification_guide(),
            'design_factor_nonmetallic': DESIGN_FACTOR_NONMETALLIC,
            'rcrt_hours': RCRT_HOURS,
            'view': 'all',
            'active_count': active_count,
            'archived_count': archived_count,
        },
    )


@router.get('/qualifications/new')
def rnd_new_program_form(request: Request, user: User = Depends(_require_user)):
    return TEMPLATES.TemplateResponse(request,'rnd_program_form.html', {'request': request, 'user': user})


@router.post('/qualifications/new')
def rnd_create_program(
    session: Session = Depends(get_session),
    user: User = Depends(_require_user),

    title: str = Form(...),
    program_code: str = Form(...),

    program_type: str = Form('API_15S'),
    qualification_standard: str = Form('API 15S R3'),

    nominal_size_in: float = Form(4.0),
    npr_mpa: float = Form(10.0),
    maot_c: float = Form(65.0),
    laot_c: float = Form(0.0),

    pfr_or_pv: str = Form('PFR'),
    parent_program_id: Optional[int] = Form(None),

    service_medium: str = Form('WATER'),
    service_factor: float = Form(1.0),
    intended_service: str = Form('Static water service'),

    product_family: str = Form('LLRTP-PE-RT-PET-PE100'),
    reinforcement_type: str = Form('NONMETALLIC'),
    liner_material: str = Form('PE-RT'),
    reinforcement_material: str = Form('Polyester Fiber'),
    cover_material: str = Form('PE100'),

    custom_requirements: str = Form(''),
    custom_acceptance_criteria: str = Form(''),
    custom_tests: str = Form(''),

    selected_test_title: List[str] = Form(default=[], alias='selected_test_title[]'),
    selected_test_code: List[str] = Form(default=[], alias='selected_test_code[]'),
    selected_test_clause_ref: List[str] = Form(default=[], alias='selected_test_clause_ref[]'),
    selected_test_specimen_count: List[str] = Form(default=[], alias='selected_test_specimen_count[]'),
    selected_test_scope_tag: List[str] = Form(default=[], alias='selected_test_scope_tag[]'),
    selected_test_source_standard: List[str] = Form(default=[], alias='selected_test_source_standard[]'),
    selected_test_description: List[str] = Form(default=[], alias='selected_test_description[]'),

    notes: str = Form(''),
):
    safe_program_type = (program_type or 'API_15S').strip().upper()
    if safe_program_type not in {'API_15S', 'OTHER'}:
        safe_program_type = 'API_15S'

    safe_service_medium = (service_medium or 'WATER').strip().upper()
    if safe_service_medium not in {'WATER', 'GAS', 'HYDROCARBON', 'LIQUIDS'}:
        safe_service_medium = 'WATER'

    safe_pfr_or_pv = (pfr_or_pv or 'PFR').strip().upper()
    if safe_pfr_or_pv not in {'PFR', 'PV'}:
        safe_pfr_or_pv = 'PFR'

    if safe_program_type == 'OTHER':
        safe_standard = (qualification_standard or 'OTHER QUALIFICATION').strip()
        safe_pfr_or_pv = 'PFR'
        parent_program_id = None
    else:
        safe_standard = (qualification_standard or 'API 15S R3').strip()

    def _safe_list_value(values: List[str], idx: int, default: str = '') -> str:
        if idx < len(values):
            return (values[idx] or '').strip()
        return default

    selected_rows = []
    selected_max = max(
        len(selected_test_title),
        len(selected_test_code),
        len(selected_test_clause_ref),
        len(selected_test_specimen_count),
        len(selected_test_scope_tag),
        len(selected_test_source_standard),
        len(selected_test_description),
        0,
    )

    for idx in range(selected_max):
        title_part = _safe_list_value(selected_test_title, idx)
        if not title_part:
            continue

        code_part = _safe_list_value(selected_test_code, idx, f'CUSTOM_{idx + 1}')
        code_part = code_part.upper().replace(' ', '_').replace('-', '_')
        if not code_part:
            code_part = f'CUSTOM_{idx + 1}'

        clause_part = _safe_list_value(selected_test_clause_ref, idx, 'CUSTOM')
        specimen_count_raw = _safe_list_value(selected_test_specimen_count, idx, '')
        specimen_count_value = None
        if specimen_count_raw:
            try:
                specimen_count_value = int(float(specimen_count_raw))
            except Exception:
                specimen_count_value = None

        if specimen_count_value is not None and specimen_count_value > 0:
            unit = 'specimen' if specimen_count_value == 1 else 'specimens'
            specimens_part = f'{specimen_count_value} {unit}'
        else:
            specimens_part = 'As required'
        scope_part = _safe_list_value(selected_test_scope_tag, idx, 'CUSTOM').upper()
        source_part = _safe_list_value(selected_test_source_standard, idx, 'CUSTOM').upper()
        desc_part = _safe_list_value(selected_test_description, idx, 'Custom qualification requirement.')

        if scope_part not in {'BOTH', 'PFR', 'PV', 'CUSTOM'}:
            scope_part = 'CUSTOM'

        if source_part not in {'API_15S', 'CUSTOM', 'CLIENT', 'INTERNAL'}:
            source_part = 'CUSTOM'

        selected_rows.append({
            'title': title_part,
            'code': code_part,
            'clause_ref': clause_part,
            'specimen_requirement': specimens_part,
            'specimen_count': specimen_count_value,
            'scope_tag': scope_part,
            'source_standard': source_part,
            'description': desc_part,
        })

    if not selected_rows and (custom_tests or '').strip():
        legacy_rows = []
        for raw_line in (custom_tests or '').splitlines():
            line = (raw_line or '').strip()
            if not line:
                continue

            parts = [p.strip() for p in line.split('|')]
            title_part = parts[0] if len(parts) > 0 else 'Custom Test'
            specimens_part = parts[1] if len(parts) > 1 else 'As required'
            clause_part = parts[2] if len(parts) > 2 else 'CUSTOM'
            desc_part = parts[3] if len(parts) > 3 else 'Custom qualification requirement.'

            legacy_rows.append({
                'title': title_part,
                'code': f'OTHER_{len(legacy_rows) + 1}',
                'specimen_requirement': specimens_part,
                'clause_ref': clause_part,
                'description': desc_part,
                'scope_tag': 'CUSTOM',
                'source_standard': 'CUSTOM',
            })

        selected_rows = legacy_rows

    program = RndQualificationProgram(
        program_code=(program_code or '').strip().upper(),
        title=(title or '').strip(),

        program_type=safe_program_type,
        qualification_standard=safe_standard,

        nominal_size_in=nominal_size_in,
        npr_mpa=npr_mpa,
        maot_c=maot_c,
        laot_c=laot_c,

        pfr_or_pv=safe_pfr_or_pv,
        parent_program_id=parent_program_id,

        service_medium=safe_service_medium,
        service_factor=service_factor,
        intended_service=(intended_service or '').strip(),

        product_family=(product_family or '').strip(),
        reinforcement_type=(reinforcement_type or 'NONMETALLIC').strip().upper(),
        liner_material=(liner_material or '').strip(),
        reinforcement_material=(reinforcement_material or '').strip(),
        cover_material=(cover_material or '').strip(),

        custom_requirements=(custom_requirements or '').strip(),
        custom_acceptance_criteria=(custom_acceptance_criteria or '').strip(),

        notes=(notes or '').strip(),
        created_by_name=(getattr(user, 'display_name', '') or getattr(user, 'username', '') or ''),
    )

    session.add(program)
    session.commit()
    session.refresh(program)

    if program.parent_program_id:
        parent = session.get(RndQualificationProgram, program.parent_program_id)
        if parent:
            program.pfr_reference_code = parent.program_code
            _touch_program(program)
            session.add(program)
            session.commit()

    if program.program_type == 'API_15S':
        _seed_test_matrix(session, program)
    else:
        if not session.exec(
            select(RndMaterialQualification).where(RndMaterialQualification.program_id == program.id)
        ).first():
            for component, material in [
                ('LINER', program.liner_material or 'CUSTOM'),
                ('REINFORCEMENT', program.reinforcement_material or 'CUSTOM'),
                ('COVER', program.cover_material or 'CUSTOM'),
            ]:
                row = RndMaterialQualification(
                    program_id=program.id,
                    component=component,
                    material_name=material,
                    material_family=('POLYMER' if component in {'LINER', 'COVER'} else 'REINFORCEMENT'),
                    reinforcement_type=(
                        'NONMETALLIC'
                        if component == 'REINFORCEMENT' and 'steel' not in (material or '').lower()
                        else ('STEEL' if component == 'REINFORCEMENT' else '')
                    ),
                    reinforcement_layer_count=(2 if component == 'REINFORCEMENT' else None),
                    status='PLANNED',
                    review_outcome='MORE_DATA_REQUIRED',
                    evidence_status='MISSING',
                    compatibility_status='UNKNOWN',
                )
                session.add(row)
            session.commit()

    existing_tests = session.exec(
        select(RndQualificationTest)
        .where(RndQualificationTest.program_id == program.id)
        .order_by(RndQualificationTest.sort_order.asc(), RndQualificationTest.id.asc())
    ).all()

    existing_codes = {((t.code or '').strip().upper()) for t in existing_tests if (t.code or '').strip()}
    next_order = max([t.sort_order for t in existing_tests], default=0)

    for offset, item in enumerate(selected_rows, start=1):
        base_code = (item.get('code') or f'CUSTOM_{offset}').strip().upper()
        final_code = base_code
        suffix = 2
        while final_code in existing_codes:
            final_code = f'{base_code}_{suffix}'
            suffix += 1
        existing_codes.add(final_code)

        session.add(
            RndQualificationTest(
                program_id=program.id,
                sort_order=next_order + offset,
                clause_ref=(item.get('clause_ref') or 'CUSTOM').strip(),
                code=final_code,
                title=(item.get('title') or 'Custom Test').strip(),
                description=(item.get('description') or '').strip(),
                specimen_requirement=(item.get('specimen_requirement') or 'As required').strip(),
                specimen_count=item.get('specimen_count'),
                applicability='CUSTOM',
                scope_tag=(item.get('scope_tag') or 'CUSTOM').strip().upper(),
                source_standard=(item.get('source_standard') or 'CUSTOM').strip().upper(),
                status='PLANNED',
            )
        )

    session.commit()

    return RedirectResponse(url=f'/rnd/qualifications/{program.id}', status_code=303)


@router.get('/qualifications/{program_id}')
def rnd_program_view(program_id: int, request: Request, session: Session = Depends(get_session), user: User = Depends(_require_user)):
    from sqlmodel import text

    program_row = session.exec(
        text("""
            SELECT
                id,
                program_code,
                title,
                program_type,
                qualification_standard,
                nominal_size_in,
                npr_mpa,
                maot_c,
                laot_c,
                pfr_or_pv,
                parent_program_id,
                pfr_reference_code,
                service_medium,
                service_factor,
                intended_service,
                product_family,
                reinforcement_type,
                liner_material,
                reinforcement_material,
                cover_material,
                custom_requirements,
                custom_acceptance_criteria,
                status,
                notes,
                created_by_name,
                created_at,
                updated_at,
                COALESCE(is_archived, FALSE) AS is_archived
            FROM rndqualificationprogram
            WHERE id = :program_id
        """).bindparams(program_id=program_id)
    ).first()

    if not program_row:
        raise HTTPException(404, 'Program not found')

    class RowObj:
        def __init__(self, r):
            self.__dict__.update(dict(r._mapping))

    program = RowObj(program_row)

    tests = session.exec(
        select(RndQualificationTest)
        .where(RndQualificationTest.program_id == program_id)
        .order_by(RndQualificationTest.sort_order.asc(), RndQualificationTest.id.asc())
    ).all()

    specimens = session.exec(
        select(RndQualificationSpecimen)
        .where(RndQualificationSpecimen.program_id == program_id)
        .order_by(RndQualificationSpecimen.created_at.desc())
    ).all()

    materials = session.exec(
        select(RndMaterialQualification)
        .where(RndMaterialQualification.program_id == program_id)
        .order_by(RndMaterialQualification.id.asc())
    ).all()

    attachments = session.exec(
        select(RndAttachmentRegister)
        .where(RndAttachmentRegister.program_id == program_id)
        .order_by(RndAttachmentRegister.created_at.desc())
    ).all()

    material_tests = session.exec(
        select(RndMaterialTestRecord)
        .where(RndMaterialTestRecord.program_id == program_id)
        .order_by(RndMaterialTestRecord.test_date.desc(), RndMaterialTestRecord.id.desc())
    ).all()

    material_dashboard = _material_dashboard_rows(materials, material_tests, program)
    phase_cards = _phase_cards(program, tests, materials, specimens, attachments)

    static_reg = _regression_from_specimens(
        specimens,
        'STATIC_REGRESSION',
        getattr(program, 'npr_mpa', 0.0),
        getattr(program, 'service_factor', 1.0),
    )

    cyclic_reg = _regression_from_specimens(
        specimens,
        'CYCLIC_REGRESSION',
        getattr(program, 'npr_mpa', 0.0),
    )

    counts = {'PLANNED': 0, 'IN_PROGRESS': 0, 'PASSED': 0, 'FAILED': 0, 'WAIVED': 0, 'COMPLETE': 0}
    for t in tests:
        key = (getattr(t, 'status', 'PLANNED') or 'PLANNED').upper()
        counts[key] = counts.get(key, 0) + 1

    total_tests = len(tests)
    done_tests = counts.get('PASSED', 0) + counts.get('WAIVED', 0) + counts.get('COMPLETE', 0)
    progress_pct = int(round((done_tests / total_tests) * 100)) if total_tests else 0

    return TEMPLATES.TemplateResponse(
        request,
        'rnd_program_view.html',
        {
            'request': request,
            'user': user,
            'program': program,
            'tests': tests,
            'specimens': specimens,
            'materials': materials,
            'material_tests': material_tests,
            'material_dashboard': material_dashboard,
            'attachments': attachments,
            'static_reg': static_reg,
            'cyclic_reg': cyclic_reg,
            'counts': counts,
            'progress_pct': progress_pct,
            'guide': _qualification_guide(program),
            'phase_cards': phase_cards,
            'design_factor_nonmetallic': DESIGN_FACTOR_NONMETALLIC,
            'rcrt_hours': RCRT_HOURS,
            'qual_is_other': ((getattr(program, 'program_type', 'API_15S') or 'API_15S').strip().upper() == 'OTHER'),
        }
    )
    
@router.post('/qualifications/{program_id}/status')
def rnd_update_program_status(program_id: int, status: str = Form(...), session: Session = Depends(get_session), user: User = Depends(_require_user)):
    program = session.get(RndQualificationProgram, program_id)
    if not program:
        raise HTTPException(404, 'Program not found')
    program.status = (status or 'DRAFT').strip().upper(); _touch_program(program); session.add(program); session.commit()
    return RedirectResponse(url=f'/rnd/qualifications/{program_id}', status_code=303)


@router.post('/qualifications/{program_id}/archive')
def rnd_archive_program(
    program_id: int,
    session: Session = Depends(get_session),
    user: User = Depends(_require_user),
):
    program = session.get(RndQualificationProgram, program_id)
    if not program:
        raise HTTPException(404, 'Program not found')

    program.is_archived = True
    program.archived_at = datetime.utcnow()
    program.archived_by_name = (getattr(user, 'display_name', '') or getattr(user, 'username', '') or '')
    _touch_program(program)
    session.add(program)
    session.commit()

    return RedirectResponse(url='/rnd/qualifications?view=active', status_code=303)


@router.post('/qualifications/{program_id}/restore')
def rnd_restore_program(
    program_id: int,
    session: Session = Depends(get_session),
    user: User = Depends(_require_user),
):
    program = session.get(RndQualificationProgram, program_id)
    if not program:
        raise HTTPException(404, 'Program not found')

    program.is_archived = False
    program.archived_at = None
    program.archived_by_name = ''
    _touch_program(program)
    session.add(program)
    session.commit()

    return RedirectResponse(url='/rnd/qualifications?view=archived', status_code=303)
    

@router.get('/qualifications/{program_id}/edit')
def rnd_edit_program_form(
    program_id: int,
    request: Request,
    session: Session = Depends(get_session),
    user: User = Depends(_require_user),
):
    program = session.get(RndQualificationProgram, program_id)
    if not program:
        raise HTTPException(404, 'Program not found')

    return TEMPLATES.TemplateResponse(
        request,
        'rnd_program_edit.html',
        {
            'request': request,
            'user': user,
            'program': program,
        }
    )


@router.post('/qualifications/{program_id}/edit')
def rnd_edit_program(
    program_id: int,
    session: Session = Depends(get_session),
    user: User = Depends(_require_user),

    title: str = Form(...),
    program_code: str = Form(...),

    program_type: str = Form('API_15S'),
    qualification_standard: str = Form('API 15S R3'),

    nominal_size_in: float = Form(4.0),
    npr_mpa: float = Form(10.0),
    maot_c: float = Form(65.0),
    laot_c: float = Form(0.0),

    pfr_or_pv: str = Form('PFR'),
    parent_program_id: Optional[int] = Form(None),

    service_medium: str = Form('WATER'),
    service_factor: float = Form(1.0),
    intended_service: str = Form(''),

    product_family: str = Form(''),
    reinforcement_type: str = Form('NONMETALLIC'),
    liner_material: str = Form(''),
    reinforcement_material: str = Form(''),
    cover_material: str = Form(''),

    custom_requirements: str = Form(''),
    custom_acceptance_criteria: str = Form(''),
    notes: str = Form(''),
):
    program = session.get(RndQualificationProgram, program_id)
    if not program:
        raise HTTPException(404, 'Program not found')

    old_program_type = (program.program_type or 'API_15S').strip().upper()

    safe_program_type = (program_type or 'API_15S').strip().upper()
    if safe_program_type not in {'API_15S', 'OTHER'}:
        safe_program_type = 'API_15S'

    safe_service_medium = (service_medium or 'WATER').strip().upper()
    if safe_service_medium not in {'WATER', 'GAS', 'HYDROCARBON', 'LIQUIDS'}:
        safe_service_medium = 'WATER'

    safe_pfr_or_pv = (pfr_or_pv or 'PFR').strip().upper()
    if safe_pfr_or_pv not in {'PFR', 'PV'}:
        safe_pfr_or_pv = 'PFR'

    program.title = (title or '').strip()
    program.program_code = (program_code or '').strip().upper()

    program.program_type = safe_program_type
    program.service_medium = safe_service_medium
    program.service_factor = service_factor

    program.nominal_size_in = nominal_size_in
    program.npr_mpa = npr_mpa
    program.maot_c = maot_c
    program.laot_c = laot_c

    program.intended_service = (intended_service or '').strip()
    program.product_family = (product_family or '').strip()
    program.reinforcement_type = (reinforcement_type or 'NONMETALLIC').strip().upper()
    program.liner_material = (liner_material or '').strip()
    program.reinforcement_material = (reinforcement_material or '').strip()
    program.cover_material = (cover_material or '').strip()

    program.custom_requirements = (custom_requirements or '').strip()
    program.custom_acceptance_criteria = (custom_acceptance_criteria or '').strip()
    program.notes = (notes or '').strip()

    if safe_program_type == 'OTHER':
        program.qualification_standard = (qualification_standard or 'OTHER QUALIFICATION').strip()
        program.pfr_or_pv = 'PFR'
        program.parent_program_id = None
        program.pfr_reference_code = ''
    else:
        program.qualification_standard = (qualification_standard or 'API 15S R3').strip()
        program.pfr_or_pv = safe_pfr_or_pv
        program.parent_program_id = parent_program_id

        if parent_program_id:
            parent = session.get(RndQualificationProgram, parent_program_id)
            if parent:
                program.pfr_reference_code = parent.program_code
            else:
                program.pfr_reference_code = ''
        else:
            program.pfr_reference_code = ''

    _touch_program(program)
    session.add(program)
    session.commit()

    # If switching into API_15S, ensure the default API matrix exists
    if old_program_type != 'API_15S' and program.program_type == 'API_15S':
        _seed_test_matrix(session, program)
        _ensure_complete_test_matrix(session, program)

    return RedirectResponse(url=f'/rnd/qualifications/{program.id}', status_code=303)
    

@router.get('/qualifications/{program_id}/tests/{test_id}/guidance-edit')
def rnd_test_guidance_edit_page(
    program_id: int,
    test_id: int,
    request: Request,
    session: Session = Depends(get_session),
    user: User = Depends(_require_user),
):
    program = session.get(RndQualificationProgram, program_id)
    if not program:
        raise HTTPException(404, 'Program not found')

    test = session.get(RndQualificationTest, test_id)
    if not test or test.program_id != program_id:
        raise HTTPException(404, 'Test not found')

    guidance = get_test_guidance(test.code, test)

    return TEMPLATES.TemplateResponse(
        request=request,
        name='rnd_test_guidance_edit.html',
        context={
            'request': request,
            'user': user,
            'program': program,
            'test': test,
            'guidance': guidance,
        },
    )


@router.post('/qualifications/{program_id}/tests/{test_id}/guidance-edit')
def rnd_test_guidance_edit_save(
    program_id: int,
    test_id: int,
    session: Session = Depends(get_session),
    user: User = Depends(_require_user),

    specimen_count: Optional[int] = Form(None),
    when_required: str = Form(''),
    api_clause: str = Form(''),
    external_standard: str = Form(''),
    conditioning_required: str = Form(''),
    conditioning_steps: str = Form(''),
    core_process: str = Form(''),
    acceptance: str = Form(''),
    retest_logic: str = Form(''),
    practical_notes: str = Form(''),
):
    program = session.get(RndQualificationProgram, program_id)
    if not program:
        raise HTTPException(404, 'Program not found')

    test = session.get(RndQualificationTest, test_id)
    if not test or test.program_id != program_id:
        raise HTTPException(404, 'Test not found')

    test.specimen_count = specimen_count
    if specimen_count is not None and specimen_count > 0:
        unit = 'specimen' if specimen_count == 1 else 'specimens'
        test.specimen_requirement = f'{specimen_count} {unit}'

    test.guidance_when_required_override = (when_required or '').strip()
    test.guidance_specimen_count_override = test.specimen_requirement or ''
    test.guidance_api_clause_override = (api_clause or '').strip()
    test.guidance_external_standard_override = (external_standard or '').strip()
    test.guidance_conditioning_required_override = (conditioning_required or '').strip()
    test.guidance_conditioning_steps_override = (conditioning_steps or '').strip()
    test.guidance_core_process_override = (core_process or '').strip()
    test.guidance_acceptance_override = (acceptance or '').strip()
    test.guidance_retest_logic_override = (retest_logic or '').strip()
    test.guidance_practical_notes_override = (practical_notes or '').strip()
    test.updated_at = datetime.utcnow()

    session.add(test)
    session.commit()

    return RedirectResponse(url=f'/rnd/qualifications/{program_id}', status_code=303)
    

@router.post('/qualifications/{program_id}/tests/{test_id}')
def rnd_update_test(program_id: int, test_id: int, status: str = Form(...), result_summary: str = Form(''), session: Session = Depends(get_session), user: User = Depends(_require_user)):
    test = session.get(RndQualificationTest, test_id)
    if not test or test.program_id != program_id:
        raise HTTPException(404, 'Test not found')
    test.status = (status or 'PLANNED').strip().upper(); test.result_summary = result_summary or ''; _touch_row(test); session.add(test)
    program = session.get(RndQualificationProgram, program_id)
    if program:
        _touch_program(program); session.add(program)
    session.commit()
    return RedirectResponse(url=f'/rnd/qualifications/{program_id}', status_code=303)


@router.post('/qualifications/{program_id}/materials/{material_id}')
def rnd_update_material(
    program_id: int,
    material_id: int,
    material_name: str = Form(''),
    supplier_name: str = Form(''),
    manufacturer_name: str = Form(''),
    grade_name: str = Form(''),
    trade_name: str = Form(''),
    certificate_ref: str = Form(''),
    batch_ref: str = Form(''),
    lot_ref: str = Form(''),
    status: str = Form('PLANNED'),
    notes: str = Form(''),
    standard_ref: str = Form(''),
    material_family: str = Form(''),
    service_fluid_basis: str = Form(''),
    service_notes: str = Form(''),
    max_service_temp_c: Optional[float] = Form(None),
    min_service_temp_c: Optional[float] = Form(None),
    classification_basis: str = Form(''),
    pe_cell_class: str = Form(''),
    hdb_basis: str = Form(''),
    uv_class: str = Form(''),
    reinforcement_type: str = Form(''),
    reinforcement_form: str = Form(''),
    reinforcement_layer_count: Optional[int] = Form(None),
    reinforcement_layout_notes: str = Form(''),
    reinforcement_matrix_material: str = Form(''),
    steel_processing_history: str = Form(''),
    fiber_sizing_notes: str = Form(''),
    matrix_material_name: str = Form(''),
    matrix_resin_type: str = Form(''),
    cure_method: str = Form(''),
    tg_min_c: Optional[float] = Form(None),
    session: Session = Depends(get_session),
    user: User = Depends(_require_user),
):
    row = session.get(RndMaterialQualification, material_id)
    if not row or row.program_id != program_id:
        raise HTTPException(404, 'Material row not found')

    row.material_name = (material_name or '').strip()
    row.supplier_name = (supplier_name or '').strip()
    row.manufacturer_name = (manufacturer_name or '').strip()
    row.grade_name = (grade_name or '').strip()
    row.trade_name = (trade_name or '').strip()
    row.certificate_ref = (certificate_ref or '').strip()
    row.batch_ref = (batch_ref or '').strip()
    row.lot_ref = (lot_ref or '').strip()
    row.status = (status or 'PLANNED').strip().upper()
    row.notes = (notes or '').strip()

    row.standard_ref = (standard_ref or '').strip()
    row.material_family = (material_family or '').strip().upper()
    row.service_fluid_basis = (service_fluid_basis or '').strip()
    row.service_notes = (service_notes or '').strip()
    row.max_service_temp_c = max_service_temp_c
    row.min_service_temp_c = min_service_temp_c

    row.classification_basis = (classification_basis or '').strip()
    row.pe_cell_class = (pe_cell_class or '').strip().upper()
    row.hdb_basis = (hdb_basis or '').strip()
    row.uv_class = (uv_class or '').strip().upper()

    row.reinforcement_type = (reinforcement_type or '').strip().upper()
    row.reinforcement_form = (reinforcement_form or '').strip()
    row.reinforcement_layer_count = reinforcement_layer_count
    row.reinforcement_layout_notes = (reinforcement_layout_notes or '').strip()
    row.reinforcement_matrix_material = (reinforcement_matrix_material or '').strip()
    row.steel_processing_history = (steel_processing_history or '').strip()
    row.fiber_sizing_notes = (fiber_sizing_notes or '').strip()

    row.matrix_material_name = (matrix_material_name or '').strip()
    row.matrix_resin_type = (matrix_resin_type or '').strip().upper()
    row.cure_method = (cure_method or '').strip()
    row.tg_min_c = tg_min_c

    _touch_row(row)
    session.add(row)

    program = session.get(RndQualificationProgram, program_id)
    _refresh_material_review(session, row, program)

    if program:
        _touch_program(program)
        session.add(program)

    session.commit()
    return RedirectResponse(url=f'/rnd/qualifications/{program_id}?tab=materials', status_code=303)



@router.post('/qualifications/{program_id}/materials/{material_id}/tests/new')
def rnd_add_material_test_record(
    program_id: int,
    material_id: int,
    test_type: str = Form(...),
    standard_ref: str = Form(''),
    specimen_ref: str = Form(''),
    report_ref: str = Form(''),
    lab_name: str = Form(''),
    test_date: Optional[date] = Form(None),
    result_value: str = Form(''),
    result_unit: str = Form(''),
    acceptance_basis: str = Form(''),
    decision: str = Form('PENDING'),
    notes: str = Form(''),
    session: Session = Depends(get_session),
    user: User = Depends(_require_user),
):
    material = session.get(RndMaterialQualification, material_id)
    if not material or material.program_id != program_id:
        raise HTTPException(404, 'Material row not found')

    row = RndMaterialTestRecord(
        program_id=program_id,
        material_id=material_id,
        test_type=(test_type or '').strip().upper(),
        standard_ref=(standard_ref or '').strip(),
        specimen_ref=(specimen_ref or '').strip(),
        report_ref=(report_ref or '').strip(),
        lab_name=(lab_name or '').strip(),
        test_date=test_date,
        result_value=(result_value or '').strip(),
        result_unit=(result_unit or '').strip(),
        acceptance_basis=(acceptance_basis or '').strip(),
        decision=(decision or 'PENDING').strip().upper(),
        notes=(notes or '').strip(),
    )
    session.add(row)

    program = session.get(RndQualificationProgram, program_id)
    _refresh_material_review(session, material, program)

    if program:
        _touch_program(program)
        session.add(program)

    session.commit()
    return RedirectResponse(url=f'/rnd/qualifications/{program_id}?tab=materials', status_code=303)


@router.post('/qualifications/{program_id}/materials/tests/{test_row_id}/delete')
def rnd_delete_material_test_record(
    program_id: int,
    test_row_id: int,
    session: Session = Depends(get_session),
    user: User = Depends(_require_user),
):
    row = session.get(RndMaterialTestRecord, test_row_id)
    if not row or row.program_id != program_id:
        raise HTTPException(404, 'Material test record not found')

    material = session.get(RndMaterialQualification, row.material_id)
    session.delete(row)

    program = session.get(RndQualificationProgram, program_id)
    if material:
        _refresh_material_review(session, material, program)

    if program:
        _touch_program(program)
        session.add(program)

    session.commit()
    return RedirectResponse(url=f'/rnd/qualifications/{program_id}?tab=materials', status_code=303)


@router.post('/qualifications/{program_id}/attachments/new')
def rnd_add_attachment_register(program_id: int, category: str = Form('REPORT'), title: str = Form(...), reference_no: str = Form(''), file_note: str = Form(''), session: Session = Depends(get_session), user: User = Depends(_require_user)):
    program = session.get(RndQualificationProgram, program_id)
    if not program:
        raise HTTPException(404, 'Program not found')
    session.add(RndAttachmentRegister(program_id=program_id, category=(category or 'REPORT').strip().upper(), title=title.strip(), reference_no=reference_no.strip(), file_note=file_note.strip()))
    _touch_program(program); session.add(program); session.commit()
    return RedirectResponse(url=f'/rnd/qualifications/{program_id}', status_code=303)


@router.post('/qualifications/{program_id}/specimens/new')
def rnd_add_specimen(program_id: int, session: Session = Depends(get_session), user: User = Depends(_require_user), specimen_id: str = Form(...), test_type: str = Form(...), test_id: Optional[int] = Form(None), sample_date: date = Form(...), material_ref: str = Form(''), nominal_size_in: float = Form(0.0), pressure_mpa: float = Form(0.0), temperature_c: float = Form(0.0), failure_hours: Optional[float] = Form(None), failure_cycles: Optional[float] = Form(None), failure_mode: str = Form(''), permissible_failure: Optional[str] = Form(None), is_runout: Optional[str] = Form(None), include_in_regression: Optional[str] = Form(None), fitting_type: str = Form('Field fitting'), lab_name: str = Form(''), witness_name: str = Form(''), notes: str = Form('')):
    program = session.get(RndQualificationProgram, program_id)
    if not program:
        raise HTTPException(404, 'Program not found')
    materials = session.exec(select(RndMaterialQualification).where(RndMaterialQualification.program_id == program_id).order_by(RndMaterialQualification.component.asc(), RndMaterialQualification.id.asc())).all()
    specimen = RndQualificationSpecimen(program_id=program_id, test_id=test_id, specimen_id=(specimen_id or '').strip().upper(), test_type=(test_type or 'STATIC_REGRESSION').strip().upper(), sample_date=sample_date, material_ref=_coalesce_material_ref(material_ref, '', '', materials), nominal_size_in=nominal_size_in or program.nominal_size_in, pressure_mpa=pressure_mpa, temperature_c=temperature_c, failure_hours=failure_hours, failure_cycles=failure_cycles, failure_mode=failure_mode, permissible_failure=bool(permissible_failure), is_runout=bool(is_runout), include_in_regression=bool(include_in_regression), fitting_type=fitting_type, lab_name=lab_name, witness_name=witness_name, notes=notes)
    session.add(specimen); _touch_program(program); session.add(program); session.commit()
    return RedirectResponse(url=f'/rnd/qualifications/{program_id}', status_code=303)


@router.post('/qualifications/{program_id}/specimens/{specimen_id}/delete')
def rnd_delete_specimen(program_id: int, specimen_id: int, session: Session = Depends(get_session), user: User = Depends(_require_user)):
    specimen = session.get(RndQualificationSpecimen, specimen_id)
    if not specimen or specimen.program_id != program_id:
        raise HTTPException(404, 'Specimen not found')
    session.delete(specimen)
    program = session.get(RndQualificationProgram, program_id)
    if program:
        _touch_program(program); session.add(program)
    session.commit()
    return RedirectResponse(url=f'/rnd/qualifications/{program_id}', status_code=303)



@router.post('/qualifications/{program_id}/tests/custom')
def rnd_add_custom_test(program_id: int, title: str = Form(...), code: str = Form('OTHER'), description: str = Form(''), specimen_requirement: str = Form('As required'), clause_ref: str = Form('CUSTOM'), session: Session = Depends(get_session), user: User = Depends(_require_user)):
    program = session.get(RndQualificationProgram, program_id)
    if not program:
        raise HTTPException(404, 'Program not found')
    existing = session.exec(select(RndQualificationTest).where(RndQualificationTest.program_id == program_id).order_by(RndQualificationTest.sort_order.desc(), RndQualificationTest.id.desc())).first()
    next_order = (existing.sort_order if existing else 0) + 1
    norm_code = (code or 'OTHER').strip().upper()
    if norm_code == 'OTHER':
        norm_code = f"OTHER_{next_order}"
    row = RndQualificationTest(program_id=program_id, sort_order=next_order, clause_ref=(clause_ref or 'CUSTOM').strip(), code=norm_code, title=(title or '').strip(), description=(description or '').strip(), specimen_requirement=(specimen_requirement or 'As required').strip(), applicability='CUSTOM', status='PLANNED')
    session.add(row)
    _touch_program(program)
    session.add(program)
    session.commit()
    return RedirectResponse(url=f'/rnd/qualifications/{program_id}', status_code=303)


@router.post('/qualifications/{program_id}/add-pv-extension')
def rnd_add_pv_extension(
    program_id: int,
    session: Session = Depends(get_session),
    user: User = Depends(_require_user),
):
    program = session.get(RndQualificationProgram, program_id)
    if not program:
        raise HTTPException(404, 'Program not found')

    if (program.program_type or 'API_15S').strip().upper() != 'API_15S':
        raise HTTPException(400, 'PV extension is only available for API 15S programs.')

    result = _add_pv_extension_to_program(session, program)
    return RedirectResponse(
        url=f"/rnd/qualifications/{program_id}?pv_extension_added=1&created={len(result['created_codes'])}&updated={len(result['updated_codes'])}",
        status_code=303,
    )
    

@router.get('/qualifications/{program_id}/regression')
def rnd_regression_view(program_id: int, request: Request, session: Session = Depends(get_session), user: User = Depends(_require_user)):
    from sqlmodel import text

    program_row = session.exec(
        text("""
            SELECT
                id,
                program_code,
                title,
                program_type,
                qualification_standard,
                nominal_size_in,
                npr_mpa,
                maot_c,
                laot_c,
                pfr_or_pv,
                parent_program_id,
                pfr_reference_code,
                service_medium,
                service_factor,
                intended_service,
                product_family,
                reinforcement_type,
                liner_material,
                reinforcement_material,
                cover_material,
                custom_requirements,
                custom_acceptance_criteria,
                status,
                notes,
                created_by_name,
                created_at,
                updated_at,
                COALESCE(is_archived, FALSE) AS is_archived
            FROM rndqualificationprogram
            WHERE id = :program_id
        """).bindparams(program_id=program_id)
    ).first()

    if not program_row:
        raise HTTPException(404, 'Program not found')

    class RowObj:
        def __init__(self, r):
            self.__dict__.update(dict(r._mapping))

    program = RowObj(program_row)

    specimen_rows = session.exec(
        text("""
            SELECT *
            FROM rndqualificationspecimen
            WHERE program_id = :program_id
            ORDER BY created_at ASC
        """).bindparams(program_id=program_id)
    ).all()
    specimens = [RowObj(r) for r in specimen_rows]

    static_reg = _regression_from_specimens(specimens, 'STATIC_REGRESSION', program.npr_mpa, getattr(program, 'service_factor', 1.0))
    cyclic_reg = _regression_from_specimens(specimens, 'CYCLIC_REGRESSION', program.npr_mpa)

    pv_formula = None

    return TEMPLATES.TemplateResponse(
        request,
        'rnd_regression_view.html',
        {
            'request': request,
            'user': user,
            'program': program,
            'specimens': specimens,
            'static_reg': static_reg,
            'cyclic_reg': cyclic_reg,
            'pv_formula': pv_formula,
            'guide': _qualification_guide(program),
            'design_factor_nonmetallic': DESIGN_FACTOR_NONMETALLIC,
            'rcrt_hours': RCRT_HOURS,
        }
    )

@router.get('/qualifications/{program_id}/tests/{test_id}')
def rnd_test_detail(program_id: int, test_id: int, request: Request, session: Session = Depends(get_session), user: User = Depends(_require_user)):
    test = session.get(RndQualificationTest, test_id)
    if not test:
        raise HTTPException(404, 'Test not found')
    
    program = session.get(RndQualificationProgram, test.program_id)
    mpr_test = session.exec(
        select(RndQualificationTest)
        .where(RndQualificationTest.program_id == program_id)
        .where(RndQualificationTest.code == 'MPR_REG')
    ).first()
    if not program:
        raise HTTPException(404, 'Program not found')
    
    if int(program_id) != int(test.program_id):
        raise HTTPException(400, 'Test does not belong to this program')

    prep = get_specimen_prep(test.code)
    guidance = get_test_guidance(test.code, test)

    specimens = session.exec(
        select(RndQualificationSpecimen)
        .where(RndQualificationSpecimen.program_id == program_id)
        .where(RndQualificationSpecimen.test_id == test_id)
        .order_by(RndQualificationSpecimen.created_at.desc())
    ).all()

    attachments = session.exec(
        select(RndAttachmentRegister)
        .where(RndAttachmentRegister.program_id == program_id)
        .where(RndAttachmentRegister.test_id == test_id)
        .order_by(RndAttachmentRegister.created_at.desc())
    ).all()

    materials = session.exec(
        select(RndMaterialQualification)
        .where(RndMaterialQualification.program_id == program_id)
        .order_by(RndMaterialQualification.component.asc(), RndMaterialQualification.id.asc())
    ).all()
    material_options = _material_reference_options(materials)

    execution = _execution_requirements(test.code)
    acceptance = _acceptance_criteria(test.code)
    evidence = _evidence_status(test.code, attachments)
    specimen_state = _specimen_readiness(specimens, prep)
    specimen_lifecycle = _specimen_lifecycle_summary(specimens)
    progress = _test_progress_snapshot(test, specimens, attachments)
    generated_report_document_type = _preferred_generated_document_type(test.code)
    generated_report_options = _generated_report_options(test.code)

    return TEMPLATES.TemplateResponse(
        request,
        'rnd_test_detail.html',
        {
            'request': request,
            'user': user,
            'program': program,
            'test': test,
            'prep': prep,
            'specimens': specimens,
            'attachments': attachments,
            'execution': execution,
            'acceptance': acceptance,
            'evidence': evidence,
            'specimen_state': specimen_state,
            'specimen_lifecycle': specimen_lifecycle,
            'progress': progress,
            'guidance': guidance,
            'materials': materials,
            'material_options': material_options,
            'mpr_test': mpr_test,
            'generated_report_document_type': generated_report_document_type,
            'generated_report_options': generated_report_options,
        }
    )

@router.post('/qualifications/{program_id}/tests/{test_id}/specimens/new')
def rnd_add_test_specimen(
    request: Request,
    program_id: int,
    test_id: int,
    specimen_id: str = Form(...),
    test_type: str = Form(''),
    sample_date: str = Form(''),
    nominal_size_in: str = Form(''),
    confirmed_od_mm: str = Form(''),
    preparation_rule_basis: str = Form(''),
    pressure_mpa: str = Form(''),
    temperature_c: str = Form(''),
    failure_hours: str = Form(''),
    failure_cycles: str = Form(''),
    failure_mode: str = Form(''),
    permissible_failure: str = Form(''),
    is_runout: Optional[str] = Form(None),
    include_in_regression: Optional[str] = Form(None),
    fitting_type: str = Form(''),
    lab_name: str = Form(''),
    witness_name: str = Form(''),
    notes: str = Form(''),
    batch_ref: str = Form(''),
    source_pipe_ref: str = Form(''),
    cut_by: str = Form(''),
    total_cut_length_mm: str = Form(''),
    effective_length_mm: str = Form(''),
    end_allowance_each_side_mm: str = Form(''),
    trimming_margin_mm: str = Form(''),
    conditioning_complete: Optional[str] = Form(None),
    pretest_visual_ok: Optional[str] = Form(None),
    released_for_test: Optional[str] = Form(None),
    planned_pressure_mpa: str = Form(''),
    actual_pressure_at_failure_mpa: str = Form(''),
    pressure_at_hold_mpa: str = Form(''),
    failure_time_sec: str = Form(''),
    pre_failure_condition: str = Form(''),
    pre_failure_visual: str = Form(''),
    post_failure_visual: str = Form(''),
    failure_location: str = Form(''),
    failure_description: str = Form(''),
    leak_observation: str = Form(''),
    result_status: str = Form('PENDING'),
    qa_review_status: str = Form('PENDING'),
    session: Session = Depends(get_session),
):
    def as_float(value):
        if value is None:
            return None
        value = str(value).strip()
        if value == '':
            return None
        try:
            return float(value)
        except Exception:
            return None

    def as_bool(value):
        return value in (True, 'true', 'True', 'on', '1', 1)

    def as_date(value):
        value = (value or '').strip()
        if not value:
            return date.today()
        try:
            return datetime.strptime(value, '%Y-%m-%d').date()
        except Exception:
            return date.today()

    # First try to recover the test, even if program_id is stale
    test = session.get(RndQualificationTest, test_id)

    if not test:
        desired_code = (test_type or '').strip().upper()
        if desired_code:
            test = session.exec(
                select(RndQualificationTest).where(
                    RndQualificationTest.program_id == program_id,
                    RndQualificationTest.code == desired_code,
                )
            ).first()

    if not test:
        raise HTTPException(404, 'Test not found')

    # Use the real program from the recovered test, not blindly from URL
    program = session.get(RndQualificationProgram, test.program_id)
    if not program:
        raise HTTPException(404, 'Program not found')

    existing_specimen = session.exec(
        select(RndQualificationSpecimen)
        .where(RndQualificationSpecimen.program_id == program.id)
        .where(RndQualificationSpecimen.test_id == test.id)
        .where(RndQualificationSpecimen.specimen_id == (specimen_id or '').strip())
    ).first()
    
    if existing_specimen:
        return RedirectResponse(
            url=f'/rnd/qualifications/{program.id}/tests/{test.id}?tab=specimens&error=duplicate_specimen_id',
            status_code=303,
        )

    safe_batch_ref = (batch_ref or '').strip()
    safe_source_pipe_ref = (source_pipe_ref or '').strip()

    material_ref_parts = []
    if safe_batch_ref:
        material_ref_parts.append(f'Batch {safe_batch_ref}')
    if safe_source_pipe_ref:
        material_ref_parts.append(f'Pipe {safe_source_pipe_ref}')
    safe_material_ref = ' | '.join(material_ref_parts) or 'FINAL_PRODUCT'

    specimen = RndQualificationSpecimen(
        program_id=program.id,
        test_id=test.id,
        specimen_id=(specimen_id or '').strip(),
        test_type=(test.code or test_type or 'MPR_REG').strip().upper(),
        sample_date=as_date(sample_date),
        material_ref=safe_material_ref,
        conditioning_required=_conditioning_required_flag(
            get_test_guidance(test.code).get('conditioning_required')
        ),
        nominal_size_in=as_float(nominal_size_in) or program.nominal_size_in,
        confirmed_od_mm=as_float(confirmed_od_mm),
        preparation_rule_basis=(preparation_rule_basis or '').strip(),
        pressure_mpa=(
            as_float(actual_pressure_at_failure_mpa)
            if as_float(actual_pressure_at_failure_mpa) is not None and (test.code or '').upper() in {'MPR_REG', 'CYCLIC_REG'}
            else (as_float(pressure_mpa) or 0.0)
        ),
        temperature_c=as_float(temperature_c) or program.maot_c,
        failure_hours=as_float(failure_hours),
        failure_cycles=as_float(failure_cycles),
        failure_mode=(failure_mode or '').strip(),
        permissible_failure=as_bool(permissible_failure),
        is_runout=as_bool(is_runout),
        include_in_regression=True if include_in_regression is None else as_bool(include_in_regression),
        fitting_type=(fitting_type or 'Field fitting').strip(),
        lab_name=(lab_name or '').strip(),
        witness_name=(witness_name or '').strip(),
        notes=(notes or '').strip(),
        batch_ref=safe_batch_ref,
        source_pipe_ref=safe_source_pipe_ref,
        cut_by=(cut_by or '').strip(),
        total_cut_length_mm=as_float(total_cut_length_mm),
        effective_length_mm=as_float(effective_length_mm),
        end_allowance_each_side_mm=as_float(end_allowance_each_side_mm),
        trimming_margin_mm=as_float(trimming_margin_mm),
        conditioning_complete=as_bool(conditioning_complete),
        pretest_visual_ok=as_bool(pretest_visual_ok),
        released_for_test=as_bool(released_for_test),
        planned_pressure_mpa=as_float(planned_pressure_mpa),
        actual_pressure_at_failure_mpa=as_float(actual_pressure_at_failure_mpa),
        pressure_at_hold_mpa=as_float(pressure_at_hold_mpa),
        failure_time_sec=as_float(failure_time_sec),
        pre_failure_condition=(pre_failure_condition or '').strip(),
        pre_failure_visual=(pre_failure_visual or '').strip(),
        post_failure_visual=(post_failure_visual or '').strip(),
        failure_location=(failure_location or '').strip(),
        failure_description=(failure_description or '').strip(),
        leak_observation=(leak_observation or '').strip(),
        result_status=(result_status or 'PENDING').strip().upper(),
        qa_review_status=(qa_review_status or 'PENDING').strip().upper(),
    )

    session.add(specimen)

    test.status = 'IN_PROGRESS'
    _touch_row(test)
    session.add(test)

    _touch_program(program)
    session.add(program)

    session.commit()

    return RedirectResponse(
        url=f'/rnd/qualifications/{program.id}/tests/{test.id}?tab=specimens&specimen_id={specimen.id}',
        status_code=303,
    )

@router.post('/qualifications/{program_id}/tests/{test_id}/specimens/{specimen_row_id}/update')
def rnd_update_test_specimen(
    program_id: int,
    test_id: int,
    specimen_row_id: int,
    session: Session = Depends(get_session),
    user: User = Depends(_require_user),
    material_ref: str = Form(''),
    planned_pressure_mpa: Optional[float] = Form(None),
    actual_pressure_at_failure_mpa: Optional[float] = Form(None),
    pressure_at_hold_mpa: Optional[float] = Form(None),
    failure_time_sec: Optional[float] = Form(None),
    failure_hours: Optional[float] = Form(None),
    failure_cycles: Optional[float] = Form(None),
    failure_mode: str = Form(''),
    failure_location: str = Form(''),
    failure_description: str = Form(''),
    leak_observation: str = Form(''),
    pre_failure_condition: str = Form(''),
    pre_failure_visual: str = Form(''),
    post_failure_visual: str = Form(''),
    result_status: str = Form('PENDING'),
    qa_review_status: str = Form('PENDING'),
    permissible_failure: Optional[str] = Form(None),
    is_runout: Optional[str] = Form(None),
    include_in_regression: Optional[str] = Form(None),
    notes: str = Form(''),
):
    specimen = session.get(RndQualificationSpecimen, specimen_row_id)
    if not specimen or specimen.program_id != program_id or specimen.test_id != test_id:
        raise HTTPException(404, 'Specimen not found')

    test = session.get(RndQualificationTest, test_id)
    if not test or test.program_id != program_id:
        raise HTTPException(404, 'Test not found')

    test_code = (test.code or '').strip().upper()

    specimen.material_ref = (material_ref or specimen.material_ref or 'FINAL_PRODUCT').strip()

    if specimen.conditioning_required is None:
        specimen.conditioning_required = _conditioning_required_flag(
            get_test_guidance(test.code).get('conditioning_required')
        )

    specimen.result_status = (result_status or 'PENDING').strip().upper()
    specimen.qa_review_status = (qa_review_status or 'PENDING').strip().upper()
    specimen.permissible_failure = bool(permissible_failure)
    specimen.is_runout = bool(is_runout)
    specimen.include_in_regression = bool(include_in_regression)

    # Shared fields kept for all test types
    specimen.planned_pressure_mpa = planned_pressure_mpa
    specimen.pressure_at_hold_mpa = pressure_at_hold_mpa
    specimen.failure_time_sec = failure_time_sec
    specimen.failure_hours = failure_hours
    specimen.failure_cycles = failure_cycles
    specimen.failure_mode = (failure_mode or '').strip()
    specimen.failure_location = (failure_location or '').strip()
    specimen.failure_description = (failure_description or '').strip()
    specimen.leak_observation = (leak_observation or '').strip()
    specimen.pre_failure_condition = (pre_failure_condition or '').strip()
    specimen.pre_failure_visual = (pre_failure_visual or '').strip()
    specimen.post_failure_visual = (post_failure_visual or '').strip()

    # Test-specific handling
    if test_code in {'MPR_REG', 'STATIC_REGRESSION'}:
        specimen.actual_pressure_at_failure_mpa = actual_pressure_at_failure_mpa

        # Keep regression pressure aligned with actual failure pressure
        if actual_pressure_at_failure_mpa is not None:
            specimen.pressure_mpa = actual_pressure_at_failure_mpa

        specimen.notes = (notes or '').strip()

    elif test_code == 'PV_1000H':
        # PV confirmation is about hold pressure + survival duration
        specimen.actual_pressure_at_failure_mpa = actual_pressure_at_failure_mpa
        specimen.pressure_mpa = planned_pressure_mpa or specimen.pressure_mpa

        pv_lines = []
        if pre_failure_condition:
            pv_lines.append(f"Setup: {pre_failure_condition.strip()}")
        if leak_observation:
            pv_lines.append(f"End observation: {leak_observation.strip()}")
        if notes:
            pv_lines.append(f"Notes: {notes.strip()}")
        specimen.notes = "\n".join(pv_lines).strip()

    elif test_code in {'TEMP_CYCLE', 'THERMAL'}:
        specimen.actual_pressure_at_failure_mpa = actual_pressure_at_failure_mpa

        cycle_lines = [
            "TEMP_CYCLE RECORD",
            f"High temp placeholder: {planned_pressure_mpa if planned_pressure_mpa is not None else '-'}",
            f"Low temp placeholder: {pressure_at_hold_mpa if pressure_at_hold_mpa is not None else '-'}",
            f"Cycle count: {failure_cycles if failure_cycles is not None else '-'}",
            f"Leak/proof pressure: {actual_pressure_at_failure_mpa if actual_pressure_at_failure_mpa is not None else '-'}",
            f"Mode: {(failure_mode or '').strip() or '-'}",
            f"Location: {(failure_location or '').strip() or '-'}",
            f"Pre-cycle setup: {(pre_failure_condition or '').strip() or '-'}",
            f"Post-cycle observation: {(leak_observation or '').strip() or '-'}",
        ]
        if notes:
            cycle_lines.append(f"Notes: {notes.strip()}")
        specimen.notes = "\n".join(cycle_lines)

    elif test_code == 'RAPID_DECOMP':
        specimen.actual_pressure_at_failure_mpa = actual_pressure_at_failure_mpa

        decomp_lines = [
            "RAPID_DECOMP RECORD",
            f"Soak pressure: {planned_pressure_mpa if planned_pressure_mpa is not None else '-'}",
            f"Decompression stop pressure: {pressure_at_hold_mpa if pressure_at_hold_mpa is not None else '-'}",
            f"Soak duration hours: {failure_hours if failure_hours is not None else '-'}",
            f"Cycle/rate placeholder: {failure_cycles if failure_cycles is not None else '-'}",
            f"Observation mode: {(failure_mode or '').strip() or '-'}",
            f"Location: {(failure_location or '').strip() or '-'}",
            f"Gas/setup: {(pre_failure_condition or '').strip() or '-'}",
            f"Post-test visual: {(post_failure_visual or '').strip() or '-'}",
            f"Leak/decompression observation: {(leak_observation or '').strip() or '-'}",
            f"Failure description: {(failure_description or '').strip() or '-'}",
        ]
        if notes:
            decomp_lines.append(f"Notes: {notes.strip()}")
        specimen.notes = "\n".join(decomp_lines)

    elif test_code == 'IMPACT':
        specimen.actual_pressure_at_failure_mpa = actual_pressure_at_failure_mpa

        impact_lines = [
            "IMPACT RECORD",
            f"Impact setup: {(pre_failure_condition or '').strip() or '-'}",
            f"Follow-up pressure: {planned_pressure_mpa if planned_pressure_mpa is not None else '-'}",
            f"Result mode: {(failure_mode or '').strip() or '-'}",
            f"Location: {(failure_location or '').strip() or '-'}",
            f"Post-impact visual: {(post_failure_visual or '').strip() or '-'}",
        ]
        if notes:
            impact_lines.append(f"Notes: {notes.strip()}")
        specimen.notes = "\n".join(impact_lines)

    elif test_code in {'AXIAL_LOAD', 'AXIAL'}:
        specimen.actual_pressure_at_failure_mpa = actual_pressure_at_failure_mpa

        axial_lines = [
            "AXIAL LOAD RECORD",
            f"Applied axial load: {(pre_failure_condition or '').strip() or '-'}",
            f"Hold duration sec: {failure_time_sec if failure_time_sec is not None else '-'}",
            f"Follow-up pressure: {planned_pressure_mpa if planned_pressure_mpa is not None else '-'}",
            f"Result mode: {(failure_mode or '').strip() or '-'}",
            f"Location: {(failure_location or '').strip() or '-'}",
            f"Post-load observation: {(leak_observation or '').strip() or '-'}",
        ]
        if notes:
            axial_lines.append(f"Notes: {notes.strip()}")
        specimen.notes = "\n".join(axial_lines)

    else:
        specimen.actual_pressure_at_failure_mpa = actual_pressure_at_failure_mpa
        specimen.notes = (notes or '').strip()

    _touch_row(specimen)
    session.add(specimen)

    all_specimens = session.exec(
        select(RndQualificationSpecimen)
        .where(RndQualificationSpecimen.program_id == program_id)
        .where(RndQualificationSpecimen.test_id == test_id)
        .order_by(RndQualificationSpecimen.id.asc())
    ).all()

    executed = sum(1 for s in all_specimens if (s.result_status or 'PENDING') != 'PENDING')
    accepted = sum(1 for s in all_specimens if (s.qa_review_status or 'PENDING') in {'ACCEPTED', 'APPROVED'})

    test.result_summary = f"{executed} specimen(s) executed, {accepted} QA accepted"
    if all_specimens and accepted == len(all_specimens):
        test.status = 'COMPLETE'
    elif executed:
        test.status = 'IN_PROGRESS'

    _touch_row(test)
    session.add(test)

    program = session.get(RndQualificationProgram, program_id)
    if program:
        _touch_program(program)
        session.add(program)

    session.commit()
    return RedirectResponse(
        url=f'/rnd/qualifications/{program_id}/tests/{test_id}?tab=specimens&specimen_id={specimen.id}',
        status_code=303,
    )

@router.post('/qualifications/{program_id}/tests/{test_id}/generate-report')
def rnd_generate_test_report(
    program_id: int,
    test_id: int,
    session: Session = Depends(get_session),
    user: User = Depends(_require_user),
    title: str = Form(''),
    reference_no: str = Form(''),
    document_type: str = Form(''),
    approval_status: str = Form('PENDING'),
):
    program = session.get(RndQualificationProgram, program_id)
    if not program:
        raise HTTPException(404, 'Program not found')

    test = session.get(RndQualificationTest, test_id)
    if not test or test.program_id != program_id:
        raise HTTPException(404, 'Test not found')

    specimens = session.exec(
        select(RndQualificationSpecimen)
        .where(RndQualificationSpecimen.program_id == program_id)
        .where(RndQualificationSpecimen.test_id == test_id)
        .order_by(RndQualificationSpecimen.created_at.asc())
    ).all()

    attachments = session.exec(
        select(RndAttachmentRegister)
        .where(RndAttachmentRegister.program_id == program_id)
        .where(RndAttachmentRegister.test_id == test_id)
        .order_by(RndAttachmentRegister.created_at.asc())
    ).all()

    materials = session.exec(
        select(RndMaterialQualification)
        .where(RndMaterialQualification.program_id == program_id)
        .order_by(RndMaterialQualification.component.asc(), RndMaterialQualification.id.asc())
    ).all()

    guidance = get_test_guidance(test.code, test)
    evidence = _evidence_status(test.code, attachments)

    safe_document_type = (document_type or '').strip().upper() or _preferred_generated_document_type(test.code)

    saved = _save_generated_test_report_docx(
        program=program,
        test=test,
        guidance=guidance,
        specimens=specimens,
        attachments=attachments,
        evidence=evidence,
        materials=materials,
        document_type=safe_document_type,
    )

    safe_title = (title or '').strip() or _auto_report_title(test, safe_document_type)
    safe_reference_no = (reference_no or '').strip() or _auto_report_reference(program, test)
    safe_title = (title or '').strip() or f"{test.title or test.code or 'Qualification Test'} Report"

    row = RndAttachmentRegister(
        program_id=program_id,
        test_id=test_id,
        category='REPORT',
        document_type=safe_document_type,
        title=safe_title,
        reference_no=safe_reference_no,
        file_note='System-generated technical report based on current program, specimen, and evidence records.',
        is_mandatory=False,
        uploaded_by_name=(getattr(user, 'display_name', '') or getattr(user, 'username', '') or ''),
        approval_status=(approval_status or 'PENDING').strip().upper(),
        original_filename=saved['original_filename'],
        stored_filename=saved['stored_filename'],
        file_path=saved['file_path'],
        content_type=saved['content_type'],
        file_size_bytes=saved['file_size_bytes'],
        source_mode='GENERATED',
        is_signed_copy=False,
    )
    session.add(row)

    test.status = 'IN_PROGRESS' if (test.status or '').upper() == 'PLANNED' else test.status
    _touch_row(test)
    session.add(test)

    _touch_program(program)
    session.add(program)

    session.commit()
    return RedirectResponse(url=f'/rnd/qualifications/{program_id}/tests/{test_id}?tab=evidence', status_code=303)

@router.post('/qualifications/{program_id}/tests/{test_id}/attachments/new')
def rnd_add_test_attachment(
    program_id: int,
    test_id: int,
    session: Session = Depends(get_session),
    user: User = Depends(_require_user),
    category: str = Form('REPORT'),
    document_type: str = Form(...),
    title: str = Form(...),
    reference_no: str = Form(''),
    file_note: str = Form(''),
    is_mandatory: Optional[str] = Form(None),
    approval_status: str = Form('PENDING'),
    source_mode: str = Form('UPLOAD'),
    is_signed_copy: Optional[str] = Form(None),
    uploaded_file: Optional[UploadFile] = File(None),
):
    program = session.get(RndQualificationProgram, program_id)
    if not program:
        raise HTTPException(404, 'Program not found')

    test = session.get(RndQualificationTest, test_id)
    if not test or test.program_id != program_id:
        raise HTTPException(404, 'Test not found')

    safe_source_mode = (source_mode or 'UPLOAD').strip().upper()

    saved = None

    if safe_source_mode == 'UPLOAD':
        if uploaded_file is None or not getattr(uploaded_file, 'filename', None):
            raise HTTPException(400, 'Please choose a file to upload.')
        saved = _save_rnd_upload(program_id, uploaded_file)

    elif safe_source_mode == 'GENERATED':
        generated_dir = _program_upload_dir(program_id)
        generated_name = f"{uuid.uuid4().hex}.txt"
        generated_path = generated_dir / generated_name

        generated_content = []
        generated_content.append(f"Qualification Program: {program.program_code} - {program.title}")
        generated_content.append(f"Test: {test.code} - {test.title}")
        generated_content.append(f"Clause: {test.clause_ref or ''}")
        generated_content.append(f"Generated At: {datetime.utcnow().isoformat()} UTC")
        generated_content.append("")
        generated_content.append("Result Summary:")
        generated_content.append(test.result_summary or "No result summary yet.")
        generated_content.append("")
        generated_content.append("Guidance:")
        generated_content.append(get_test_guidance(test.code, test).get("when_required", ""))

        generated_path.write_text("\n".join(generated_content), encoding="utf-8")

        saved = {
            "original_filename": f"{(test.code or 'report').lower()}_generated_report.txt",
            "stored_filename": generated_name,
            "file_path": str(generated_path),
            "content_type": "text/plain",
            "file_size_bytes": generated_path.stat().st_size,
        }

    else:
        raise HTTPException(400, f'Unsupported source mode: {safe_source_mode}')

    row = RndAttachmentRegister(
        program_id=program_id,
        test_id=test_id,
        category=(category or 'REPORT').strip().upper(),
        document_type=(document_type or '').strip().upper(),
        title=(title or '').strip(),
        reference_no=(reference_no or '').strip(),
        file_note=(file_note or '').strip(),
        is_mandatory=bool(is_mandatory),
        uploaded_by_name=(getattr(user, 'display_name', '') or getattr(user, 'username', '') or ''),
        approval_status=(approval_status or 'PENDING').strip().upper(),
        original_filename=saved["original_filename"],
        stored_filename=saved["stored_filename"],
        file_path=saved["file_path"],
        content_type=saved["content_type"],
        file_size_bytes=saved["file_size_bytes"],
        source_mode=safe_source_mode,
        is_signed_copy=bool(is_signed_copy),
    )
    session.add(row)

    test.status = 'IN_PROGRESS' if (test.status or '').upper() == 'PLANNED' else test.status
    _touch_row(test)
    session.add(test)

    _touch_program(program)
    session.add(program)

    session.commit()
    return RedirectResponse(url=f'/rnd/qualifications/{program_id}/tests/{test_id}?tab=evidence', status_code=303)

@router.get('/qualifications/{program_id}/attachments/{attachment_id}/download')
def rnd_download_attachment(
    program_id: int,
    attachment_id: int,
    session: Session = Depends(get_session),
    user: User = Depends(_require_user),
):
    row = session.get(RndAttachmentRegister, attachment_id)
    if not row or row.program_id != program_id:
        raise HTTPException(404, 'Attachment not found')

    if not row.file_path or not os.path.exists(row.file_path):
        raise HTTPException(404, 'Stored file not found')

    filename = row.original_filename or row.title or 'attachment'
    return FileResponse(
        path=row.file_path,
        media_type=row.content_type or 'application/octet-stream',
        filename=filename,
    )


@router.get('/qualifications/{program_id}/attachments/{attachment_id}/view')
def rnd_view_attachment(
    program_id: int,
    attachment_id: int,
    session: Session = Depends(get_session),
    user: User = Depends(_require_user),
):
    row = session.get(RndAttachmentRegister, attachment_id)
    if not row or row.program_id != program_id:
        raise HTTPException(404, 'Attachment not found')

    if not row.file_path or not os.path.exists(row.file_path):
        raise HTTPException(404, 'Stored file not found')

    media_type = row.content_type or 'application/octet-stream'

    return FileResponse(
        path=row.file_path,
        media_type=media_type,
        filename=row.original_filename or row.title or 'attachment',
    )
